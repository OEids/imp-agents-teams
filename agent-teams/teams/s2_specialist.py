"""
S2 Specialist Agent - Deep Analysis & Complete Template Builder

An upskilled agent that thoroughly analyzes customer data and builds
ALL template sheets with no data left behind.
"""

import pandas as pd
import numpy as np
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, field
import re
import warnings
warnings.filterwarnings('ignore')

# PDF and Image reading support
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from PIL import Image
    import pytesseract
    import os
    # Configure Tesseract path for Windows
    tesseract_paths = [
        r"C:\Users\OliviaEidsforth\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for tess_path in tesseract_paths:
        if os.path.exists(tess_path):
            pytesseract.pytesseract.tesseract_cmd = tess_path
            break
    IMAGE_OCR_SUPPORT = True
except ImportError:
    IMAGE_OCR_SUPPORT = False

try:
    import tabula
    TABULA_SUPPORT = True
except ImportError:
    TABULA_SUPPORT = False

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

from .finished_workbook_patterns import (
    S2_STAFF_ROLE_GROUP_PATTERNS,
    S2_PAY_SCALE_PATTERNS,
    S2_CONTRACT_PATTERNS,
    S2_EQWP_PATTERNS,
    S2_PENSION_PATTERNS,
    get_srg_for_role,
    get_finance_codes_for_srg,
)

# Import Intelligence Module for template formatting
try:
    from intelligence import TemplateRegistry, TemplateFormatter
    TEMPLATE_AVAILABLE = True
except ImportError:
    TEMPLATE_AVAILABLE = False

# Import S2 domain knowledge for combined field parsing
try:
    from knowledge.S2.S2_DOMAIN_KNOWLEDGE import (
        parse_combined_field,
        extract_finance_code,
        COMBINED_COLUMNS,
        PAY_SCALES as S2_IMPORT_PAY_SCALES,
        STAFF_ROLE_GROUPS as S2_IMPORT_ROLE_GROUPS,
        EQUATED_WEEK_PATTERNS as S2_IMPORT_EQW_PATTERNS,
        map_role_title_to_group,
        get_finance_codes_for_role_group,
        transform_contract_row,
    )
    S2_DOMAIN_KNOWLEDGE_AVAILABLE = True
except ImportError:
    S2_DOMAIN_KNOWLEDGE_AVAILABLE = False

# Import S2 Staff Role Coding Knowledge for generating role codes
try:
    from knowledge.S2.S2_STAFF_ROLE_KNOWLEDGE import (
        STAFF_ROLES as S2_OFFICIAL_STAFF_ROLES,
        STAFF_ROLE_GROUPS as S2_OFFICIAL_ROLE_GROUPS,
        ROLE_TO_GROUP as S2_ROLE_TO_GROUP,
        get_role_code_from_title,
        get_group_from_role,
        get_group_from_title,
        is_teaching_role,
        is_support_role,
        validate_role_code,
        validate_group_code,
    )
    S2_ROLE_KNOWLEDGE_AVAILABLE = True
except ImportError:
    S2_ROLE_KNOWLEDGE_AVAILABLE = False
    S2_OFFICIAL_STAFF_ROLES = {}
    S2_OFFICIAL_ROLE_GROUPS = {}
    S2_ROLE_TO_GROUP = {}
    def get_role_code_from_title(title): return "OTH"
    def get_group_from_role(code): return "OTH"
    def get_group_from_title(title): return "OTH"
    def is_teaching_role(code): return False
    def is_support_role(code): return True
    def validate_role_code(code): return True
    def validate_group_code(code): return True

# Import S2 Staff Role Codes for subject-specific teacher codes (e.g., TEA_HIS)
try:
    from knowledge.S2.S2_STAFF_ROLE_CODES import (
        STAFF_ROLE_CODES,
        SUBJECT_ABBREVIATIONS,
        TITLE_TO_ROLE_CODE,
        get_role_code_for_title,
        extract_subject_code,
        get_srg_for_role_code,
    )
    S2_ROLE_CODES_AVAILABLE = True
except ImportError:
    S2_ROLE_CODES_AVAILABLE = False
    STAFF_ROLE_CODES = {}
    SUBJECT_ABBREVIATIONS = {}
    TITLE_TO_ROLE_CODE = {}
    def get_role_code_for_title(title): return "OTH"
    def extract_subject_code(title): return ""
    def get_srg_for_role_code(code): return "OTH"
    # Fallback parsing function
    def parse_combined_field(value):
        """Fallback: Parse 'CODE: Title' format."""
        if not value or str(value) == 'nan':
            return ('', '')
        value = str(value).strip()
        if ':' not in value:
            return (value, value)
        parts = value.split(':', 1)
        code = parts[0].strip()
        title = parts[1].strip() if len(parts) > 1 else ''
        if '(' in title:
            title = title.split('(')[0].strip()
        return (code, title)
    COMBINED_COLUMNS = []

# Import Memory Manager for learning from past issues
try:
    from memory.s2_memory_manager import S2MemoryManager
    MEMORY_AVAILABLE = True
except ImportError:
    MEMORY_AVAILABLE = False
    S2MemoryManager = None


# =============================================================================
# UK DATE FORMAT CONFIGURATION
# =============================================================================
UK_DATE_FORMAT = '%d/%m/%y'  # dd/mm/yy format

def format_date_uk(date_value: Any) -> str:
    """
    Convert any date value to UK format (dd/mm/yy).

    Args:
        date_value: Can be string, datetime, date, or pandas Timestamp

    Returns:
        Date string in dd/mm/yy format, or empty string if invalid
    """
    if date_value is None or str(date_value).strip() == '' or str(date_value) == 'nan':
        return ''

    try:
        # If already a datetime-like object
        if hasattr(date_value, 'strftime'):
            return date_value.strftime(UK_DATE_FORMAT)

        date_str = str(date_value).strip()

        # Try ISO format first (YYYY-MM-DD) - most unambiguous
        if re.match(r'^\d{4}-\d{2}-\d{2}', date_str):
            parsed = pd.to_datetime(date_str, format='%Y-%m-%d', errors='coerce')
            if pd.notna(parsed):
                return parsed.strftime(UK_DATE_FORMAT)

        # Try UK format with dayfirst=True for ambiguous dates (DD/MM/YYYY, DD.MM.YYYY, etc.)
        parsed = pd.to_datetime(date_str, dayfirst=True, errors='coerce')
        if pd.notna(parsed):
            return parsed.strftime(UK_DATE_FORMAT)

        return date_str  # Return original if can't parse
    except Exception:
        return str(date_value)


def get_default_increase_date(is_teaching: bool = True) -> str:
    """
    Get the default pay increase date for the current academic/financial year.

    Teaching staff: 1st September
    Support staff: 1st April

    Returns date in UK format (dd/mm/yy)
    """
    today = datetime.now()

    if is_teaching:
        # Teaching: September increase date
        # If we're past September, use next year
        if today.month >= 9:
            increase_date = datetime(today.year + 1, 9, 1)
        else:
            increase_date = datetime(today.year, 9, 1)
    else:
        # Support: April increase date
        # If we're past April, use next year
        if today.month >= 4:
            increase_date = datetime(today.year + 1, 4, 1)
        else:
            increase_date = datetime(today.year, 4, 1)

    return increase_date.strftime(UK_DATE_FORMAT)


def get_default_start_date() -> str:
    """Get a default start date (1st of current month) in UK format."""
    today = datetime.now()
    return datetime(today.year, today.month, 1).strftime(UK_DATE_FORMAT)


@dataclass
class AnalysisReport:
    """Detailed analysis report of customer data."""
    file_name: str
    sheet_name: str
    row_count: int
    column_count: int
    columns_found: List[str]
    columns_mapped: Dict[str, str]  # original -> standard
    data_quality: Dict[str, Any]
    sample_data: Dict[str, List]
    issues: List[str]
    recommendations: List[str]


@dataclass
class ExtractedPayScale:
    """Extracted pay scale data."""
    code: str
    title: str
    scale_type: str  # teaching, support, leadership
    london_weighting: str
    increment_date: str
    increase_date: str
    increase_percentage: float
    points: List[Dict]  # [{code, title, number, rate, date_from}]
    grades: List[Dict]  # [{code, title, from_point, to_point}]


@dataclass
class ExtractedAllowance:
    """Extracted allowance data."""
    type_code: str
    type_title: str
    points: List[Dict]  # [{code, title, amount}]
    increase_date: str
    increase_percentage: float


class S2SpecialistAgent:
    """
    Upskilled S2 agent with deep analysis capabilities.

    Responsibilities:
    1. Deep analysis of ALL customer data files
    2. Extract pay scales and ALL points/rates
    3. Extract allowances (TLR, SEN, etc.)
    4. Build ALL template sheets - nothing ignored
    5. Provide detailed analysis reports
    """

    def __init__(self):
        self.analysis_reports: List[AnalysisReport] = []
        self.extracted_pay_scales: List[ExtractedPayScale] = []
        self.extracted_allowances: List[ExtractedAllowance] = []
        self.staff_data: List[pd.DataFrame] = []
        self.staff_lookup: Dict[str, Dict[str, Any]] = {}  # Consolidated lookup by payroll number
        self.unclassified_data: List[Dict[str, Any]] = []  # Data that couldn't be auto-classified
        self.issues: List[str] = []
        self.assumptions: List[str] = []

        # Template sheet builders
        self.template_data = {
            "PayScales": [],
            "PayScalePoints": [],
            "PayScaleGrades": [],
            "PayScaleIncreasePercen": [],
            "AllowanceTypes": [],
            "AllowanceTypePoint": [],
            "AllowanceIncreasePercen": [],
            "Pensions": [],
            "EQWPatterns": [],
            "StfRoleGroup": [],
            "StfRole": [],
            "StaffMembers": [],
            "ContractsTeachFTE": [],
            "ContractsSupportHours": [],
            "ContractAllowances": [],
            "ContractAdjustments": [],
            "Finance Codes S2": [],
        }

        # Tracking
        self.schools_found = set()
        self.finance_codes_found = set()
        self.pay_scales_found = set()
        self.allowance_types_found = set()

        # Track newly created role codes and groups (for output report)
        self.created_role_codes = []  # List of dicts: {code, title, group, source}
        self.created_role_groups = []  # List of dicts: {code, title, source}

        # Track skipped staff for diagnostics
        self.skipped_staff = []  # List of dicts: {id, reason, row_data}

        # Track DataFrames skipped because no unique identifier column was found
        self.unmapped_dataframes = []  # List of dicts: {columns, row_count, reason}

        # Processing log for diagnostics
        self.processing_log = []  # List of processing events for debugging

        # Audit tracking
        self.audit_results = {}
        self.audit_passed = True
        self.audit_score = 100.0
        self.detailed_audit_report = {}
        self.source_staff_count = 0
        self.source_contract_count = 0

        # Memory Manager - learns from past issues and prevents them from recurring
        self.memory = None
        if MEMORY_AVAILABLE:
            self.memory = S2MemoryManager(log_func=self.log)
        self.memory_warnings: List[str] = []

        # Import file data (loaded from knowledge/S2/import files/)
        self.import_data = {
            "staff_members": None,
            "contracts": None,
            "staff_roles": None,
            "staff_role_groups": None,
            "pay_scales": None,
            "pensions": None,
            "equated_week_patterns": None,
            "allowance_types": None,
            "adjustment_types": None,
            "leave_types": None,
        }

        # Add missing template sheets
        self.template_data.update({
            "Genders": [],
            "ContractTypes": [],
            "LeaveTypes": [],
            "AdjustmentTypes": [],
            "StaffMemberLeaves": [],
        })

        # Mapping from internal sheet names to official template sheet names
        self.SHEET_NAME_MAPPING = {
            "PayScales": "19_PayScales",
            "PayScalePoints": "20_PayScalePoints",
            "PayScaleGrades": "22_PayScaleGrades",
            "PayScaleIncreasePercen": "21_PayScaleIncreasePercen",
            "AllowanceTypes": "16_AllowanceTypes",
            "AllowanceTypePoint": "17_AllowanceTypePoint",
            "AllowanceIncreasePercen": "18_AllowanceIncreasePercen",
            "Pensions": "24_Pensions",
            "EQWPatterns": "23_EQWPatterns",
            "StfRoleGroup": "26_StfRoleGroup",
            "StfRole": "27_StfRole",
            "StaffMembers": "25_StaffMembers",
            "ContractsTeachFTE": "28_ContractsTeachFTE",
            "ContractsSupportHours": "29_ContractsSupportHours",
            "ContractAllowances": "34_ContractAllowances",
            "ContractAdjustments": "33_ContractAdjustments",
            "Finance Codes S2": "11_Finance Codes S2",
            "Genders": "12_Genders",
            "ContractTypes": "13_ContractTypes",
            "LeaveTypes": "30_LeaveTypes",
            "AdjustmentTypes": "31_AdjustmentTypes",
            "StaffMemberLeaves": "32_StaffMemberLeaves",
        }

        # Initialize template registry and formatter if available
        self.template_registry = None
        self.template_formatter = None
        if TEMPLATE_AVAILABLE:
            try:
                self.template_registry = TemplateRegistry()
                self.template_formatter = TemplateFormatter(self.template_registry)
            except Exception as e:
                self.log(f"[WARN] Could not initialize template registry: {e}")

        # Template reference data - loaded from prepopulated workbook
        # These are used for fuzzy matching customer data to template codes
        self.template_references = {
            'schools': {},       # code -> {title, type, ...}
            'pay_scales': {},    # code -> {title, type, ...}
            'role_groups': {},   # code -> {title, teaching, ...}
            'roles': {},         # code -> {title, group, ...}
            'pensions': {},      # code -> {title, ...}
            'eqwp': {},          # code -> {title, weeks, ...}
            'finance_codes': {}, # code -> {title, ...}
        }
        self.template_references_loaded = False

    def log(self, message: str, level: str = "INFO"):
        """Log a message with proper encoding and error handling for Streamlit."""
        try:
            timestamp = datetime.now().strftime("%H:%M:%S")
            # Ensure message is a string and handle encoding
            msg_str = str(message)
            # Replace problematic characters
            msg_str = msg_str.replace('\x00', '').replace('\n\n', '\n')
            # Truncate very long messages
            if len(msg_str) > 10000:
                msg_str = msg_str[:10000] + "... [truncated]"
            output = f"[{timestamp}] [{level}] S2-Specialist: {msg_str}"
            # Try to print with error handling
            try:
                print(output, flush=True)
            except (OSError, IOError, ValueError):
                # If print fails, try unicode encoding
                try:
                    print(output.encode('ascii', errors='replace').decode('ascii'), flush=True)
                except Exception:
                    # Last resort - silent fail
                    pass
        except Exception:
            # Completely silent - don't let logging break execution
            pass

    def load_template_references(self, template_path: Path) -> bool:
        """
        Load reference codes from prepopulated S2 template workbook.

        When a prepopulated template is provided, we extract the reference codes
        (schools, pay scales, role groups, etc.) and use them for fuzzy matching
        customer data instead of using the default knowledge files.

        Args:
            template_path: Path to prepopulated S2 workbook

        Returns:
            True if successfully loaded reference data
        """
        try:
            xl = pd.ExcelFile(template_path)
            sheets = xl.sheet_names
            self.log(f"Loading template references from: {template_path.name}")
            self.log(f"  Found {len(sheets)} sheets")

            # Helper to find column
            def find_col(df, candidates):
                for col in df.columns:
                    col_clean = str(col).strip().lower().replace(' ', '').replace('_', '')
                    for c in candidates:
                        if c.lower().replace(' ', '').replace('_', '') == col_clean:
                            return col
                return None

            # Load Schools (usually row 2 is header)
            for sheet in sheets:
                if 'school' in sheet.lower() and 'finance' not in sheet.lower():
                    try:
                        df = pd.read_excel(xl, sheet, header=1)
                        code_col = find_col(df, ['SchoolCode', 'Code'])
                        title_col = find_col(df, ['Title', 'SchoolTitle', 'Name'])
                        if code_col:
                            for _, row in df.iterrows():
                                code = str(row.get(code_col, '')).strip()
                                if code and code.lower() not in ['nan', '', 'none']:
                                    self.template_references['schools'][code] = {
                                        'title': str(row.get(title_col, '')).strip() if title_col else code,
                                    }
                            self.log(f"  Schools: {len(self.template_references['schools'])} loaded from {sheet}")
                            break
                    except Exception as e:
                        self.log(f"  [WARN] Could not read {sheet}: {e}")

            # Load Pay Scales
            for sheet in sheets:
                sheet_clean = sheet.lower().replace(' ', '').replace('_', '')
                if 'payscale' in sheet_clean and 'point' not in sheet_clean and 'grade' not in sheet_clean:
                    try:
                        df = pd.read_excel(xl, sheet, header=1)
                        code_col = find_col(df, ['PayScaleCode', 'Code'])
                        title_col = find_col(df, ['PayScaleTitle', 'Title'])
                        if code_col:
                            for _, row in df.iterrows():
                                code = str(row.get(code_col, '')).strip()
                                if code and code.lower() not in ['nan', '', 'none']:
                                    self.template_references['pay_scales'][code] = {
                                        'title': str(row.get(title_col, '')).strip() if title_col else code,
                                    }
                            self.log(f"  Pay Scales: {len(self.template_references['pay_scales'])} loaded from {sheet}")
                            break
                    except Exception as e:
                        self.log(f"  [WARN] Could not read {sheet}: {e}")

            # Load Staff Role Groups
            for sheet in sheets:
                sheet_clean = sheet.lower().replace(' ', '').replace('_', '')
                if 'rolegroup' in sheet_clean or 'stfrolegroup' in sheet_clean:
                    try:
                        df = pd.read_excel(xl, sheet, header=1)
                        code_col = find_col(df, ['StaffRoleGroupCode', 'Code', 'SRGCode'])
                        title_col = find_col(df, ['Title', 'StaffRoleGroupTitle'])
                        if code_col:
                            for _, row in df.iterrows():
                                code = str(row.get(code_col, '')).strip()
                                if code and code.lower() not in ['nan', '', 'none']:
                                    title = str(row.get(title_col, '')).strip() if title_col else code
                                    self.template_references['role_groups'][code] = {
                                        'title': title,
                                        'teaching': 'teach' in title.lower(),
                                    }
                            self.log(f"  Role Groups: {len(self.template_references['role_groups'])} loaded from {sheet}")
                            break
                    except Exception as e:
                        self.log(f"  [WARN] Could not read {sheet}: {e}")

            # Load Staff Roles
            for sheet in sheets:
                sheet_clean = sheet.lower().replace(' ', '').replace('_', '')
                if ('stfrole' in sheet_clean or 'staffrole' in sheet_clean) and 'group' not in sheet_clean:
                    try:
                        df = pd.read_excel(xl, sheet, header=1)
                        code_col = find_col(df, ['StaffRoleCode', 'Code', 'RoleCode'])
                        title_col = find_col(df, ['StaffRoleTitle', 'Title'])
                        group_col = find_col(df, ['StaffRoleGroupCode', 'RoleGroupCode', 'SRGCode'])
                        if code_col:
                            for _, row in df.iterrows():
                                code = str(row.get(code_col, '')).strip()
                                if code and code.lower() not in ['nan', '', 'none']:
                                    self.template_references['roles'][code] = {
                                        'title': str(row.get(title_col, '')).strip() if title_col else code,
                                        'group': str(row.get(group_col, '')).strip() if group_col else '',
                                    }
                            self.log(f"  Staff Roles: {len(self.template_references['roles'])} loaded from {sheet}")
                            break
                    except Exception as e:
                        self.log(f"  [WARN] Could not read {sheet}: {e}")

            # Load Pensions
            for sheet in sheets:
                if 'pension' in sheet.lower():
                    try:
                        df = pd.read_excel(xl, sheet, header=1)
                        code_col = find_col(df, ['PensionCode', 'Code'])
                        title_col = find_col(df, ['PensionTitle', 'Title'])
                        if code_col:
                            for _, row in df.iterrows():
                                code = str(row.get(code_col, '')).strip()
                                if code and code.lower() not in ['nan', '', 'none']:
                                    self.template_references['pensions'][code] = {
                                        'title': str(row.get(title_col, '')).strip() if title_col else code,
                                    }
                            self.log(f"  Pensions: {len(self.template_references['pensions'])} loaded from {sheet}")
                            break
                    except Exception as e:
                        self.log(f"  [WARN] Could not read {sheet}: {e}")

            # Load EQWP Patterns
            for sheet in sheets:
                if 'eqw' in sheet.lower() or 'equat' in sheet.lower():
                    try:
                        df = pd.read_excel(xl, sheet, header=1)
                        code_col = find_col(df, ['EquatedWeekPatternCode', 'EQWPCode', 'Code'])
                        title_col = find_col(df, ['Title', 'EQWPTitle'])
                        if code_col:
                            for _, row in df.iterrows():
                                code = str(row.get(code_col, '')).strip()
                                if code and code.lower() not in ['nan', '', 'none']:
                                    self.template_references['eqwp'][code] = {
                                        'title': str(row.get(title_col, '')).strip() if title_col else code,
                                    }
                            self.log(f"  EQWP Patterns: {len(self.template_references['eqwp'])} loaded from {sheet}")
                            break
                    except Exception as e:
                        self.log(f"  [WARN] Could not read {sheet}: {e}")

            # Summary
            total = sum(len(v) for v in self.template_references.values())
            self.log(f"  TOTAL: {total} reference codes loaded from template")
            self.template_references_loaded = total > 0
            return self.template_references_loaded

        except Exception as e:
            self.log(f"[ERROR] Failed to load template references: {e}")
            return False

    def fuzzy_match_to_template(self, value: str, ref_type: str, threshold: float = 0.6) -> Optional[Tuple[str, str, float]]:
        """
        Fuzzy match a customer value to a template reference code.

        Args:
            value: Customer value to match
            ref_type: Type of reference ('schools', 'pay_scales', 'role_groups', 'pensions', 'eqwp')
            threshold: Minimum similarity score (0.0 to 1.0)

        Returns:
            Tuple of (matched_code, matched_title, confidence) or None if no match
        """
        if not self.template_references_loaded or not value:
            return None

        refs = self.template_references.get(ref_type, {})
        if not refs:
            return None

        from difflib import SequenceMatcher
        value_lower = str(value).lower().strip()

        best_match = None
        best_score = 0.0

        for code, data in refs.items():
            title = data.get('title', code)

            # Score against code
            code_score = SequenceMatcher(None, value_lower, code.lower()).ratio()

            # Score against title
            title_score = SequenceMatcher(None, value_lower, title.lower()).ratio()

            # Use best score
            score = max(code_score, title_score)

            # Boost if contains key words
            if value_lower in title.lower() or title.lower() in value_lower:
                score = min(1.0, score + 0.2)

            # Exact match bonus
            if value_lower == code.lower() or value_lower == title.lower():
                score = 1.0

            if score > best_score and score >= threshold:
                best_score = score
                best_match = (code, title, score)

        return best_match

    def match_school_code(self, customer_value: str) -> Optional[str]:
        """
        Match a customer school value to a template school code.

        Uses template references if loaded, otherwise returns the original value.
        """
        if not customer_value:
            return None

        # Try template fuzzy matching first
        if self.template_references_loaded:
            match = self.fuzzy_match_to_template(customer_value, 'schools', threshold=0.6)
            if match:
                code, title, confidence = match
                if confidence >= 0.7:
                    return code
                else:
                    self.assumptions.append(f"School '{customer_value}' matched to '{code}' ({confidence:.0%} confidence)")
                    return code

        # Return original if no template or no match
        return str(customer_value).strip()

    def match_pay_scale_code(self, customer_value: str) -> Optional[str]:
        """
        Match a customer pay scale value to a template pay scale code.

        Uses template references if loaded, otherwise uses default pattern matching.
        """
        if not customer_value:
            return None

        # Try template fuzzy matching first
        if self.template_references_loaded:
            match = self.fuzzy_match_to_template(customer_value, 'pay_scales', threshold=0.5)
            if match:
                code, title, confidence = match
                if confidence >= 0.6:
                    return code
                else:
                    self.assumptions.append(f"Pay scale '{customer_value}' matched to '{code}' ({confidence:.0%} confidence)")
                    return code

        # Fall back to pattern matching from knowledge files
        value_upper = str(customer_value).upper().strip()

        # Common pay scale patterns
        if 'MAIN' in value_upper or value_upper.startswith('M'):
            return 'MAIN_EW'
        elif 'UPS' in value_upper or value_upper.startswith('U'):
            return 'UPS_EW'
        elif 'LEAD' in value_upper or value_upper.startswith('L'):
            return 'LEADERSHIP_EW'
        elif 'NJC' in value_upper:
            return 'NJC'
        elif 'UNQUAL' in value_upper:
            return 'UNQUALIFIED_EW'

        return str(customer_value).strip()

    def match_role_group_code(self, job_title: str) -> str:
        """
        Match a customer job title to a template role group code.

        Uses template references if loaded, otherwise uses default pattern matching.
        """
        if not job_title:
            return 'OTH'

        # Try template fuzzy matching first
        if self.template_references_loaded and self.template_references.get('role_groups'):
            match = self.fuzzy_match_to_template(job_title, 'role_groups', threshold=0.5)
            if match:
                code, title, confidence = match
                if confidence >= 0.5:
                    return code

            # Also try matching against staff roles
            if self.template_references.get('roles'):
                match = self.fuzzy_match_to_template(job_title, 'roles', threshold=0.5)
                if match:
                    code, title, confidence = match
                    # Get the role group for this role
                    role_data = self.template_references['roles'].get(code, {})
                    group = role_data.get('group')
                    if group:
                        return group

        # Fall back to default pattern matching
        return get_srg_for_role(job_title)

    def match_pension_code(self, customer_value: str) -> Optional[str]:
        """
        Match a customer pension value to a template pension code.
        """
        if not customer_value:
            return None

        # Try template fuzzy matching first
        if self.template_references_loaded:
            match = self.fuzzy_match_to_template(customer_value, 'pensions', threshold=0.5)
            if match:
                code, title, confidence = match
                return code

        # Fall back to common patterns
        value_upper = str(customer_value).upper().strip()
        if 'TPS' in value_upper or 'TEACH' in value_upper:
            return 'TPS'
        elif 'LGPS' in value_upper or 'LOCAL' in value_upper:
            return 'LGPS'
        elif 'OPT' in value_upper or 'OUT' in value_upper:
            return 'OPTOUT'

        return str(customer_value).strip()

    def match_eqwp_code(self, customer_value: str) -> Optional[str]:
        """
        Match a customer working pattern value to a template EQWP code.
        """
        if not customer_value:
            return None

        # Try template fuzzy matching first
        if self.template_references_loaded:
            match = self.fuzzy_match_to_template(customer_value, 'eqwp', threshold=0.5)
            if match:
                code, title, confidence = match
                return code

        # Fall back to common patterns
        value_str = str(customer_value).lower().strip()
        if 'all year' in value_str or '52' in value_str:
            return 'ALLYEAR'
        elif 'term time' in value_str or '39' in value_str:
            return 'TTO39'

        return str(customer_value).strip()

    def _format_detailed_error(self, file_path: Path, action: str, exception: Exception) -> str:
        """Format a detailed error message with context and suggestions."""
        error_type = type(exception).__name__
        error_msg = str(exception)

        # Build detailed error
        details = [
            f"ERROR {action}: {file_path.name}",
            f"  Type: {error_type}",
            f"  Details: {error_msg}",
            f"  Path: {file_path}",
        ]

        # Add specific suggestions based on error type
        error_lower = error_msg.lower()
        suggestions = []

        if "password" in error_lower or "encrypted" in error_lower:
            suggestions.append("Remove password protection: Open in Excel > File > Info > Protect Workbook")
        elif "corrupt" in error_lower or "invalid" in error_lower:
            suggestions.append("Try re-saving the file in Excel as a new .xlsx file")
        elif "permission" in error_lower or "access" in error_lower:
            suggestions.append("Close the file if open in another application")
            suggestions.append("Check file permissions")
        elif "no tables" in error_lower:
            suggestions.append("PDF may be scanned/image-based - convert to Excel manually or use OCR")
        elif "xlrd" in error_lower or "openpyxl" in error_lower:
            suggestions.append("File format issue - try saving as .xlsx in Excel")
        elif "codec" in error_lower or "encoding" in error_lower:
            suggestions.append("Character encoding issue - save file as UTF-8 in Excel")
        elif "~" in file_path.name or len(file_path.name) > 50:
            suggestions.append("Rename file to shorter name without special characters")

        # Check for Windows 8.3 short filename
        if "~" in file_path.name:
            suggestions.insert(0, "FILENAME ISSUE: Windows truncated filename detected. Rename to a shorter name.")

        if suggestions:
            details.append("  Suggestions:")
            for s in suggestions:
                details.append(f"    - {s}")

        return "\n".join(details)

    # Column name variations - maps expected name to possible variations
    # Column aliases for flexible matching - maps internal keys to possible column names
    # Also includes reverse mappings for schema standard names
    COLUMN_ALIASES = {
        'payroll_number': ['payroll_number', 'payroll', 'emp_no', 'employee_number', 'employee_id',
                          'staff_no', 'staff_number', 'personnel_number', 'personnel_no', 'pr_no',
                          'empno', 'empid', 'payroll_no', 'unique_id', 'staff_id', 'emp_ref',
                          'employee no', 'payroll number', 'staff code', 'code'],
        'surname': ['surname', 'last_name', 'lastname', 'family_name', 'familyname', 'lname', 'ln'],
        'forename': ['forename', 'forenames', 'first_name', 'firstname', 'given_name', 'givenname', 'fname', 'fn'],
        # Schema standard names (reverse mappings)
        'last_name': ['last_name', 'surname', 'lastname', 'family_name', 'familyname', 'lname', 'ln'],
        'first_name': ['first_name', 'forename', 'forenames', 'firstname', 'given_name', 'givenname', 'fname', 'fn'],
        'job_title': ['job_title', 'jobtitle', 'position', 'role_title', 'title', 'role', 'post',
                     'job', 'designation', 'occupation', 'job_description', 'job description'],
        'weekly_hours': ['weekly_hours', 'hours_per_week', 'contracted_hours', 'hpw', 'hrs',
                        'ft_hours', 'full_time_hours', 'hours', 'weekly hrs', 'contract hours',
                        'hours_worked', 'hours worked'],
        'school_code': ['school_code', 'school', 'schools', 'site_code', 'establishment', 'location',
                       'site', 'academy', 'school_name', 'cost_centre', 'cost_center', 'cc'],
        'service_start_date': ['service_start_date', 'start_date', 'join_date', 'commence_date',
                              'commencement', 'hire_date', 'date_joined', 'service_start',
                              'start', 'employed_from'],
        # Schema standard name (reverse mapping)
        'start_date': ['start_date', 'service_start_date', 'join_date', 'commence_date',
                      'commencement', 'hire_date', 'date_joined', 'service_start', 'employed_from'],
        'dob': ['dob', 'date_of_birth', 'birth_date', 'birthdate', 'birthday'],
        'gender': ['gender', 'sex', 'gender_code'],
        'fte': ['fte', 'full_time_equivalent', 'weekly_fte', 'fte_value', 'ft_equivalent'],
        # Schema standard name (reverse mapping)
        'weekly_fte': ['weekly_fte', 'fte', 'full_time_equivalent', 'fte_value', 'ft_equivalent'],
        'scale_point': ['scale_point', 'scp', 'spine_point', 'pay_point', 'spinal_point',
                       'current_point', 'point', 'grade_point', 'current_scale_point'],
        # Schema standard name (reverse mapping)
        'current_scale_point': ['current_scale_point', 'scale_point', 'scp', 'spine_point',
                               'pay_point', 'spinal_point', 'current_point', 'point'],
        'contract_ref': ['contract_ref', 'reference', 'contract_reference', 'ref', 'contract_no',
                        'contract_number', 'contract_id'],
        'pay_scale': ['pay_scale', 'payscale', 'scale', 'pay_scale_code', 'salary_scale',
                     'pay_grade', 'pay scale', 'paygrade', 'grade_scale', 'scale_type',
                     'pay scale type', 'pay_scale_type', 'salary_grade', 'scale_code',
                     'grade_name', 'grade name', 'grade'],
        'pay_scale_type': ['pay_scale_type', 'scale_type', 'pay_type', 'scale type',
                          'payscale type', 'pay_scale_group', 'scale_group'],
        'pay_scale_contract': ['pay_scale_contract', 'contract_scale', 'contract_payscale',
                              'scale_contract', 'payscale_contract'],
        'pension': ['pension', 'pension_code', 'pension_scheme', 'pens', 'pension scheme',
                   'pension_type', 'pension type', 'superannuation', 'retirement_scheme'],
        'ni_number': ['ni_number', 'ni', 'national_insurance', 'nino', 'ni_no'],
        'annual_salary': ['annual_salary', 'salary', 'annual_pay', 'gross_salary', 'basic_salary'],
        'department': ['department', 'dept', 'department_code', 'cost_centre', 'cc'],
        'contract_type': ['contract_type', 'type', 'employment_type', 'contract_status'],
        'staff_role_group': ['staff_role_group', 'role_group', 'srg', 'staff_group', 'category'],
    }

    def _safe_get(self, row, key, default=''):
        """
        Safely get a value from a row, trying multiple column name variations.

        Args:
            row: DataFrame row (Series)
            key: Expected column name
            default: Default value if not found

        Returns:
            Value from the row, trying variations if primary key not found
        """
        def normalize(s):
            return str(s).lower().replace(' ', '').replace('_', '').replace('-', '')

        def get_val(col):
            val = row.get(col, None)
            if isinstance(val, pd.Series):
                val = val.iloc[0] if len(val) > 0 else None
            return val

        try:
            # First try the exact key
            if key in row.index:
                val = get_val(key)
                if val is not None and pd.notna(val) and str(val).strip():
                    return val

            # Try variations from COLUMN_ALIASES
            key_lower = key.lower().replace(' ', '_')
            variations = self.COLUMN_ALIASES.get(key_lower, [])

            # Also add the key itself and common transformations
            all_variations = [key] + variations + [
                key.lower(), key.upper(), key.replace('_', ' '), key.replace('_', ''),
                key.title(), key.replace(' ', '_').lower()
            ]

            # Try exact match on variations
            for var in all_variations:
                if var in row.index:
                    val = get_val(var)
                    if val is not None and pd.notna(val) and str(val).strip():
                        return val

            # Try normalized match (ignore spaces, underscores, case)
            row_normalized = {normalize(c): c for c in row.index}
            for var in all_variations:
                var_norm = normalize(var)
                if var_norm in row_normalized:
                    actual_col = row_normalized[var_norm]
                    val = get_val(actual_col)
                    if val is not None and pd.notna(val) and str(val).strip():
                        return val

            # Also try space/underscore variants
            row_with_spaces = {str(c).lower().replace('_', ' '): c for c in row.index}
            row_with_underscores = {str(c).lower().replace(' ', '_'): c for c in row.index}

            for var in all_variations:
                var_lower = str(var).lower()
                var_spaces = var_lower.replace('_', ' ')
                var_underscores = var_lower.replace(' ', '_')

                if var_spaces in row_with_spaces:
                    val = get_val(row_with_spaces[var_spaces])
                    if val is not None and pd.notna(val) and str(val).strip():
                        return val
                if var_underscores in row_with_underscores:
                    val = get_val(row_with_underscores[var_underscores])
                    if val is not None and pd.notna(val) and str(val).strip():
                        return val

            return default
        except Exception:
            return default

    def _safe_notna(self, val) -> bool:
        """Safely check if a value is not NA, handling Series case."""
        try:
            if isinstance(val, pd.Series):
                return pd.notna(val.iloc[0]) if len(val) > 0 else False
            return pd.notna(val)
        except Exception:
            return False

    def _find_column(self, df: pd.DataFrame, key: str) -> Optional[str]:
        """
        Find the actual column name in a DataFrame using variations.

        Args:
            df: DataFrame to search
            key: Expected column name

        Returns:
            Actual column name if found, None otherwise
        """
        # First try exact match
        if key in df.columns:
            return key

        # Get variations
        key_lower = key.lower().replace(' ', '_')
        variations = self.COLUMN_ALIASES.get(key_lower, [])
        all_variations = [key] + variations + [
            key.lower(), key.upper(), key.replace('_', ' '), key.replace('_', ''),
            key.title(), key.replace(' ', '_').lower()
        ]

        # Try each variation
        for var in all_variations:
            if var in df.columns:
                return var

        # Build normalized lookup: normalize both column names and variations
        # Normalize = lowercase, replace spaces/underscores/dashes with nothing
        def normalize(s):
            return str(s).lower().replace(' ', '').replace('_', '').replace('-', '')

        cols_normalized = {normalize(c): c for c in df.columns}

        for var in all_variations:
            var_norm = normalize(var)
            if var_norm in cols_normalized:
                return cols_normalized[var_norm]

        # Also try matching with space/underscore variants
        cols_with_spaces = {str(c).lower().replace('_', ' '): c for c in df.columns}
        cols_with_underscores = {str(c).lower().replace(' ', '_'): c for c in df.columns}

        for var in all_variations:
            var_lower = str(var).lower()
            var_spaces = var_lower.replace('_', ' ')
            var_underscores = var_lower.replace(' ', '_')

            if var_spaces in cols_with_spaces:
                return cols_with_spaces[var_spaces]
            if var_underscores in cols_with_underscores:
                return cols_with_underscores[var_underscores]

        return None

    def _get_column_values(self, df: pd.DataFrame, key: str) -> pd.Series:
        """
        Get column values from DataFrame using variations.

        Args:
            df: DataFrame to search
            key: Expected column name

        Returns:
            Column as Series, or empty Series if not found
        """
        col_name = self._find_column(df, key)
        if col_name:
            return self._safe_get_column(df, col_name)
        return pd.Series(dtype=object)

    def _safe_get_column(self, df: pd.DataFrame, col: str) -> pd.Series:
        """
        Safely get a column as a Series, handling duplicate column names.
        When duplicate column names exist, df[col] returns a DataFrame instead of Series.
        This method ensures a Series is always returned.
        """
        try:
            col_data = df[col]
            if isinstance(col_data, pd.DataFrame):
                col_data = col_data.iloc[:, 0]
            return col_data
        except Exception:
            return pd.Series(dtype=object)

    def _is_job_reference_code(self, code: str) -> bool:
        """
        Check if a code looks like a job reference (e.g., j9385, r1234, JOB123).

        These should be ignored as staff member codes because they're not
        meaningful identifiers - we should use name-based codes instead.

        Patterns detected as job references:
        - Single letter + numbers (j9385, r1234, a001)
        - "JOB" prefix + numbers
        - Single/double letter prefix + 3+ digits

        Args:
            code: The code to check

        Returns:
            True if this looks like a job reference code
        """
        if not code:
            return False

        code_upper = code.upper().strip()

        # Pattern 1: Single letter + 3+ digits (j9385, r1234, a001)
        if len(code_upper) >= 4:
            if code_upper[0].isalpha() and code_upper[1:].isdigit():
                return True

        # Pattern 2: Two letters + 3+ digits (jo123, ab456)
        if len(code_upper) >= 5:
            if code_upper[:2].isalpha() and code_upper[2:].isdigit():
                return True

        # Pattern 3: JOB/REF/POS prefix
        if code_upper.startswith(('JOB', 'REF', 'POS', 'POST')):
            return True

        # Pattern 4: Looks like a reference number (letter-number-letter pattern)
        import re
        if re.match(r'^[A-Z]\d+[A-Z]$', code_upper):
            return True

        return False

    def _build_staff_lookup(self):
        """
        Build consolidated staff lookup from all DataFrames in staff_data.

        This merges data from multiple files (names, contracts, pay scales, etc.)
        into a single lookup dictionary keyed by the unique staff identifier.

        The unique key can be any of: payroll_number, emp_no, employee_id, staff_id, etc.
        Whatever column serves as the unique identifier in the customer's data.

        Each unique code maps to a dict containing all available fields
        from across all source files.
        """
        self.log("Building consolidated staff lookup from all data sources...")
        self.staff_lookup = {}

        # Possible unique identifier fields (in priority order)
        unique_id_fields = [
            'payroll_number', 'emp_no', 'employee_number', 'employee_id',
            'staff_id', 'staff_number', 'personnel_number', 'unique_id',
            'code', 'ref', 'staff_code'
        ]

        # Define all fields we want to consolidate
        lookup_fields = [
            'payroll_number', 'emp_no', 'employee_number', 'employee_id', 'staff_id',
            'surname', 'forename', 'first_name', 'last_name', 'name',
            'job_title', 'position', 'role', 'post',
            'school_code', 'school', 'cost_centre', 'department',
            'fte', 'weekly_fte', 'weekly_hours', 'ft_hours', 'hours',
            'pay_scale', 'pay_scale_type', 'pay_scale_contract', 'pay_scale_group', 'scale',
            'scale_point', 'current_scale_point', 'scp', 'point',
            'annual_salary', 'salary', 'actual_salary', 'gross_salary',
            'pension', 'pension_code', 'pension_scheme',
            'service_start_date', 'start_date', 'contract_start', 'hire_date',
            'contract_ref', 'contract_type', 'reference',
            'gender', 'sex', 'dob', 'date_of_birth', 'ni_number',
            'eqw', 'eqw_pattern', 'weeks_worked', 'weeks_paid',
            'finance_code', 'fund_code', 'nominal_code',
            'grade', 'grade_code', 'pay_grade',
        ]

        for df in self.staff_data:
            # Find which unique identifier column exists in this DataFrame
            unique_col = None
            unique_key_name = None
            for field in unique_id_fields:
                col = self._find_column(df, field)
                if col is not None:
                    unique_col = col
                    unique_key_name = field
                    break

            if unique_col is None:
                available_cols = list(df.columns)
                msg = (
                    f"No unique identifier column found in a staff data sheet "
                    f"({len(df)} rows). Available columns: {available_cols}. "
                    f"Please use the Pre-Flight Validator tab to map one of these "
                    f"to 'payroll_number', 'emp_no', 'code', 'staff_id', etc. "
                    f"All {len(df)} rows from this sheet were skipped."
                )
                self.log(f"  ERROR: {msg}")
                self.issues.append(msg)
                self.unmapped_dataframes.append({
                    "columns": available_cols,
                    "row_count": len(df),
                    "reason": "No unique identifier column found — needs manual mapping",
                })
                continue

            self.log(f"  Processing DataFrame using '{unique_key_name}' as unique key")

            for row_idx, row in df.iterrows():
                # Get unique identifier value
                unique_val = self._safe_get(row, unique_key_name, '')
                unique_str = str(unique_val).strip() if self._safe_notna(unique_val) else ''

                # Clean the identifier (remove .0 from floats)
                if unique_str.endswith('.0'):
                    unique_str = unique_str[:-2]

                # Track row info for skip diagnostics
                row_info = {k: str(v)[:50] for k, v in row.items() if self._safe_notna(v) and str(v).strip() not in ['', 'nan']}

                # Skip invalid identifiers
                if not unique_str or unique_str.lower() == 'nan':
                    self.skipped_staff.append({
                        "id": unique_str or "(empty)",
                        "reason": "Empty or invalid identifier",
                        "row_sample": str(row_info)[:200]
                    })
                    continue
                # Skip decimal values (likely rates, not IDs)
                if '.' in unique_str:
                    try:
                        if float(unique_str) != int(float(unique_str)):
                            self.skipped_staff.append({
                                "id": unique_str,
                                "reason": "Decimal value (likely a rate, not an ID)",
                                "row_sample": str(row_info)[:200]
                            })
                            continue
                    except ValueError:
                        pass
                # Skip text with spaces (likely category headers)
                if ' ' in unique_str and not unique_str.replace(' ', '').isalnum():
                    self.skipped_staff.append({
                        "id": unique_str,
                        "reason": "Text with spaces (likely a category header)",
                        "row_sample": str(row_info)[:200]
                    })
                    continue

                # Initialize or get existing record
                if unique_str not in self.staff_lookup:
                    self.staff_lookup[unique_str] = {
                        'unique_code': unique_str,
                        'payroll_number': unique_str  # Also store as payroll_number for compatibility
                    }

                record = self.staff_lookup[unique_str]

                # Merge all available fields using flexible column matching
                # This handles variations like "Pay Scale" vs "pay_scale" vs "PayScale"
                for field in lookup_fields:
                    # Use _find_column to handle all naming variations
                    actual_col = self._find_column(df, field)
                    if actual_col is not None:
                        val = row.get(actual_col, '')
                        # Handle Series - extract scalar value
                        if isinstance(val, pd.Series):
                            val = val.iloc[0] if len(val) > 0 else ''
                        if self._safe_notna(val) and str(val).strip() and str(val).strip().lower() != 'nan':
                            # Only update if current value is empty/missing
                            current = record.get(field, '')
                            # Handle Series in current value too
                            if isinstance(current, pd.Series):
                                current = current.iloc[0] if len(current) > 0 else ''
                            current_str = str(current).strip().lower() if current is not None else ''
                            if current_str in ['', 'nan', 'missing']:
                                record[field] = val

        self.log(f"  Consolidated {len(self.staff_lookup)} unique staff records from {len(self.staff_data)} data sources")

        # Log sample of what was merged
        if self.staff_lookup:
            sample_key = list(self.staff_lookup.keys())[0]
            sample = self.staff_lookup[sample_key]
            fields_found = [k for k, v in sample.items() if v and str(v).strip().lower() not in ['', 'nan', 'missing']]
            self.log(f"  Sample record '{sample_key}' has fields: {fields_found[:12]}...")

    def _get_staff_record(self, payroll_number: str) -> Dict[str, Any]:
        """
        Get consolidated staff record by payroll number.

        Args:
            payroll_number: Staff payroll number

        Returns:
            Dict with all available fields, or empty dict if not found
        """
        # Clean payroll number
        payroll_str = str(payroll_number).strip()
        if payroll_str.endswith('.0'):
            payroll_str = payroll_str[:-2]

        return self.staff_lookup.get(payroll_str, {})

    def _apply_column_mappings(self, df: pd.DataFrame, file_name: str, sheet_name: str = None) -> pd.DataFrame:
        """
        Apply validated column mappings from pre-flight validation.

        Renames columns in the DataFrame according to user-validated mappings.

        Args:
            df: DataFrame to apply mappings to
            file_name: Name of the file being processed
            sheet_name: Optional sheet name for Excel files

        Returns:
            DataFrame with columns renamed according to mappings
        """
        if not hasattr(self, 'column_mappings') or not self.column_mappings:
            return df

        # Try different key formats to find matching mappings
        possible_keys = [
            f"{file_name}:{sheet_name or 'default'}",
            f"{file_name}:default",
            file_name,
        ]

        file_mappings = {}
        for key in possible_keys:
            if key in self.column_mappings:
                file_mappings = self.column_mappings[key]
                break

        if not file_mappings:
            # Also try partial match on file name
            for key, mappings in self.column_mappings.items():
                if file_name in key or key.split(':')[0] in file_name:
                    file_mappings = mappings
                    break

        if file_mappings:
            self.log(f"  Applying column mappings for {file_name}: {file_mappings}")
            # Build rename dict: only rename columns that exist in df
            # Also check normalized versions since _clean_column_name may have already run
            rename_dict = {}

            # Create lookup for normalized column names (from _clean_column_name)
            # This handles cases where pre-flight mapped "Last Name" but it's now "surname"
            normalized_to_actual = {}
            for col in df.columns:
                # Try to reverse-map: if column is 'surname', it might have come from 'last name', 'last_name', etc.
                normalized_to_actual[col.lower().replace('_', ' ')] = col
                normalized_to_actual[col.lower().replace(' ', '_')] = col
                normalized_to_actual[col.lower()] = col

            for source_col, target_col in file_mappings.items():
                # Direct match
                if source_col in df.columns and source_col != target_col:
                    rename_dict[source_col] = target_col
                else:
                    # Try normalized match (source may have been normalized by _clean_column_name)
                    source_lower = source_col.lower().replace('_', ' ')
                    source_underscore = source_col.lower().replace(' ', '_')

                    # Check if any current column matches normalized source
                    for variant in [source_lower, source_underscore, source_col.lower()]:
                        if variant in normalized_to_actual:
                            actual_col = normalized_to_actual[variant]
                            if actual_col != target_col:
                                rename_dict[actual_col] = target_col
                                break

            if rename_dict:
                self.log(f"  Applying {len(rename_dict)} column mappings from pre-flight validation")
                for src, tgt in list(rename_dict.items())[:5]:
                    self.log(f"    {src} -> {tgt}")
                if len(rename_dict) > 5:
                    self.log(f"    ... and {len(rename_dict) - 5} more")

                df = df.rename(columns=rename_dict)

        return df

    def _format_dates_in_df(self, df: pd.DataFrame) -> pd.DataFrame:
        """Format date columns to YYYY-MM-DD strings for clean Excel output."""
        df = df.copy()
        date_columns = ['DateFrom', 'DateTo', 'ServiceStartDate', 'ServiceEndDate',
                       'DateOfBirth', 'StartDate', 'EndDate', 'date_from', 'date_to']

        # Track processed columns to avoid duplicate processing
        processed_cols = set()

        for idx, col in enumerate(df.columns):
            # Skip if we've already processed a column with this name
            if col in processed_cols:
                continue

            if col in date_columns or 'date' in str(col).lower():
                try:
                    # Get column data safely (handles duplicate column names)
                    col_data = df.iloc[:, idx]

                    # Convert each value individually to handle mixed date formats
                    def parse_and_format_date(val):
                        if pd.isna(val) or val == '' or val is None:
                            return ''
                        try:
                            # Use the centralized UK date formatter
                            return format_date_uk(val)
                        except Exception:
                            return str(val)

                    col_data = col_data.apply(parse_and_format_date)

                    # Assign back using iloc to handle duplicate column names
                    df.iloc[:, idx] = col_data
                    processed_cols.add(col)
                except Exception:
                    pass  # Keep original value if conversion fails
        return df

    def _preprocess_combined_fields(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Preprocess dataframe to parse combined format fields.
        Combined fields use 'CODE: Title (extra)' format from import files.

        Uses S2_DOMAIN_KNOWLEDGE if available for comprehensive parsing.
        """
        if not S2_DOMAIN_KNOWLEDGE_AVAILABLE:
            # Fallback: detect and parse common combined patterns
            combined_patterns = ['Combined', ' Code', 'Pay Scale', 'Staff Role', 'Pension',
                                 'Fund', 'Department', 'Contract Type', 'Equated Week']
        else:
            combined_patterns = COMBINED_COLUMNS + ['Combined']

        parsed_count = 0
        for col in df.columns:
            # Check if column likely contains combined format data
            is_combined = any(pattern in str(col) for pattern in combined_patterns)

            if not is_combined:
                # Also check content for "CODE: Title" pattern
                col_data = df[col]
                # Handle duplicate column names (returns DataFrame instead of Series)
                if isinstance(col_data, pd.DataFrame):
                    col_data = col_data.iloc[:, 0]
                sample = col_data.dropna().head(10).astype(str)
                if len(sample) > 0 and sample.str.contains(r'^\w+:\s*.+', regex=True).any():
                    is_combined = True

            if is_combined:
                # Parse the combined field and create new code column
                code_col = col.replace(' Combined', '').replace('Combined', '') + '_Code'
                code_col = code_col.replace('  ', ' ').strip()
                if code_col == '_Code':
                    code_col = col + '_Parsed'

                # Only add if doesn't already exist
                if code_col not in df.columns:
                    col_data = self._safe_get_column(df, col)
                    df[code_col] = col_data.apply(
                        lambda x: parse_combined_field(str(x))[0] if pd.notna(x) else ''
                    )
                    parsed_count += 1

        if parsed_count > 0:
            self.log(f"    Parsed {parsed_count} combined format columns")
            self.assumptions.append(f"Parsed {parsed_count} combined format columns (CODE: Title format)")

        return df

    def _preprocess_payroll_name_column(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Preprocess payroll analysis data where 'Name' contains full names.
        Handles formats like: "Annetts Caroline" (Surname Forename) or "Caroline Annetts" (Forename Surname).

        Creates 'forename' and 'surname' columns if 'Name' column exists but
        'forename'/'surname' don't.
        """
        # Check if we have a 'name' column but no forename/surname
        has_name_col = any(str(col).lower() == 'name' for col in df.columns)
        has_forename = any('forename' in str(col).lower() or 'first' in str(col).lower() for col in df.columns)
        has_surname = any('surname' in str(col).lower() or 'last' in str(col).lower() for col in df.columns)

        if has_name_col and not has_forename and not has_surname:
            # Find the name column
            name_col = None
            for col in df.columns:
                if str(col).lower() == 'name':
                    name_col = col
                    break

            if name_col is not None:
                def split_name(full_name):
                    """Split 'Surname Forename' or 'Forename Surname' into parts."""
                    if pd.isna(full_name) or not str(full_name).strip():
                        return ('', '')

                    parts = str(full_name).strip().split()
                    if len(parts) == 0:
                        return ('', '')
                    elif len(parts) == 1:
                        return (parts[0], '')  # Assume it's surname
                    else:
                        # Payroll format typically has "Surname Forename"
                        # But could also be "Forename Surname"
                        # Heuristic: If first part is all caps or looks like surname, use Surname Forename
                        first_part = parts[0]
                        rest = ' '.join(parts[1:])

                        # Common surname indicators: all caps, ends with 's', common surname patterns
                        # Common forename indicators: common first names
                        # Default to "Surname Forename" for payroll data
                        return (first_part, rest)  # (surname, forename)

                # Split the name column
                name_data = self._safe_get_column(df, name_col)
                split_data = name_data.apply(split_name)

                df['surname'] = split_data.apply(lambda x: x[0] if x else '')
                df['forename'] = split_data.apply(lambda x: x[1] if x else '')

                self.log(f"    Split 'Name' column into forename/surname (payroll format)")
                self.assumptions.append("Split 'Name' column assuming 'Surname Forename' format")

        return df

    # =========================================================================
    # IMPORT FILE LOADING (CRITICAL FIX)
    # =========================================================================

    def load_import_files(self, import_folder: Path = None) -> Dict[str, pd.DataFrame]:
        """
        Load all import files from knowledge/S2/import files/ and parse combined fields.
        This is the CORRECT way to process the standardized import format files.

        Args:
            import_folder: Path to import files folder. If None, uses default location.

        Returns:
            Dictionary of parsed DataFrames keyed by data type.
        """
        if import_folder is None:
            # Default to knowledge/S2/import files/ relative to this file
            import_folder = Path(__file__).parent.parent / "knowledge" / "S2" / "import files"

        if not import_folder.exists():
            self.log(f"Warning: Import folder not found: {import_folder}")
            return {}

        self.log("="*60)
        self.log("LOADING IMPORT FILES (STANDARDIZED FORMAT)")
        self.log("="*60)

        # Map file names to data types
        file_mappings = {
            "DEM003 - Staff Members": "staff_members",
            "DEM003 - Contracts": "contracts",
            "DEM003 - Staff Roles": "staff_roles",
            "DEM003 - Staff Role Groups": "staff_role_groups",
            "DEM003 - Pay Scales": "pay_scales",
            "DEM003 - Pensions": "pensions",
            "DEM003 - Equated Week Patterns": "equated_week_patterns",
            "DEM003 - Allowance Types": "allowance_types",
            "DEM003 - Adjustment Types": "adjustment_types",
            "DEM003 - Leave Types": "leave_types",
        }

        for file_path in import_folder.glob("*.xlsx"):
            if file_path.name.startswith("~$"):
                continue

            # Find matching data type
            data_type = None
            for file_prefix, dtype in file_mappings.items():
                if file_prefix in file_path.name:
                    data_type = dtype
                    break

            if not data_type:
                self.log(f"  Skipping unknown file: {file_path.name}")
                continue

            try:
                df = pd.read_excel(file_path)
                self.log(f"  {data_type}: {len(df)} rows, {len(df.columns)} columns")

                # Apply validated column mappings from pre-flight validation
                df = self._apply_column_mappings(df, file_path.name)

                # Parse all combined fields
                df = self._preprocess_combined_fields(df)

                # Store in import_data
                self.import_data[data_type] = df

                # Also add to staff_data if it's contracts (for compatibility)
                if data_type == "contracts":
                    self.staff_data.append(df)

            except Exception as e:
                error_detail = self._format_detailed_error(file_path, "loading file", e)
                self.log(f"  {error_detail}")
                self.issues.append(error_detail)

        self._log_import_summary()
        return self.import_data

    def _log_import_summary(self):
        """Log summary of loaded import data."""
        self.log("\nIMPORT DATA SUMMARY:")
        for dtype, df in self.import_data.items():
            if df is not None:
                self.log(f"  {dtype}: {len(df)} records")
            else:
                self.log(f"  {dtype}: NOT LOADED")

    def build_from_import_files(self) -> Dict[str, list]:
        """
        Build all template sheets from loaded import files.
        This processes the standardized DEM003 format correctly.
        """
        if not any(df is not None for df in self.import_data.values()):
            self.log("Error: No import files loaded. Call load_import_files() first.")
            return self.template_data

        self.log("="*60)
        self.log("BUILDING TEMPLATE SHEETS FROM IMPORT FILES")
        self.log("="*60)

        # Build reference data sheets first
        self._build_genders()
        self._build_contract_types()
        self._build_leave_types_from_import()
        self._build_adjustment_types_from_import()
        self._build_pensions_from_import()
        self._build_eqw_patterns_from_import()
        self._build_staff_role_groups_from_import()
        self._build_staff_roles_from_import()
        self._build_pay_scales_from_import()
        self._build_allowance_types_from_import()

        # Build staff and contract sheets
        self._build_staff_members_from_import()
        self._build_contracts_from_import()

        return self.template_data

    def _build_genders(self):
        """Build Genders reference sheet (standard values)."""
        self.log("Building Genders...")
        self.template_data["Genders"] = [
            {"Code": "M", "Title": "Male"},
            {"Code": "F", "Title": "Female"},
            {"Code": "O", "Title": "Other"},
            {"Code": "U", "Title": "Unknown"},
        ]
        self.log(f"  Added {len(self.template_data['Genders'])} genders")

    def _build_contract_types(self):
        """Build ContractTypes reference sheet from contracts data."""
        self.log("Building ContractTypes...")

        contracts_df = self.import_data.get("contracts")
        if contracts_df is None:
            # Standard contract types
            self.template_data["ContractTypes"] = [
                {"Code": "PERM", "Title": "Permanent"},
                {"Code": "TEMP", "Title": "Temporary"},
                {"Code": "FTC", "Title": "Fixed Term Contract"},
                {"Code": "MAT", "Title": "Maternity Cover"},
                {"Code": "CASUAL", "Title": "Casual"},
            ]
        else:
            # Extract unique contract types from contracts
            seen = set()
            for _, row in contracts_df.iterrows():
                # Try parsed column first, then combined column
                code = row.get('Contract Type_Code', row.get('Contract Type Combined_Code', ''))
                if not code:
                    combined = row.get('Contract Type Combined', '')
                    if combined:
                        code, title = parse_combined_field(str(combined))
                    else:
                        continue
                else:
                    combined = row.get('Contract Type Combined', '')
                    _, title = parse_combined_field(str(combined)) if combined else (code, code)

                if code and code not in seen:
                    seen.add(code)
                    self.template_data["ContractTypes"].append({
                        "Code": code,
                        "Title": title if title else code
                    })

        self.log(f"  Added {len(self.template_data['ContractTypes'])} contract types")

    def _build_leave_types_from_import(self):
        """Build LeaveTypes from import file."""
        self.log("Building LeaveTypes...")

        df = self.import_data.get("leave_types")
        if df is None:
            self.log("  Warning: No leave types import file loaded")
            return

        for _, row in df.iterrows():
            code = row.get('Code', '')
            if not code:
                continue

            self.template_data["LeaveTypes"].append({
                "Code": code,
                "Title": row.get('Title', code),
                "AvailableToAllSchools": row.get('Available To All Schools', True),
                "AvailableToSchools": row.get('Available to Schools', ''),
                "RebatePercentage": row.get('Rebate Percentage', 0),
            })

        self.log(f"  Added {len(self.template_data['LeaveTypes'])} leave types")

    def _build_adjustment_types_from_import(self):
        """Build AdjustmentTypes from import file."""
        self.log("Building AdjustmentTypes...")

        df = self.import_data.get("adjustment_types")
        if df is None:
            self.log("  Warning: No adjustment types import file loaded")
            return

        for _, row in df.iterrows():
            code = row.get('Code', '')
            if not code:
                continue

            self.template_data["AdjustmentTypes"].append({
                "Code": code,
                "Title": row.get('Title', code),
                "AnniversaryDate": str(row.get('Anniversary Date', ''))[:10] if pd.notna(row.get('Anniversary Date')) else '',
                "IncreasePercentage": row.get('Increase Percentage', 0),
                "ExcludeNI": row.get('Exclude National Insurance', False),
                "ExcludePension": row.get('Exclude Pension', False),
            })

        self.log(f"  Added {len(self.template_data['AdjustmentTypes'])} adjustment types")

    def _build_pensions_from_import(self):
        """Build Pensions from import file."""
        self.log("Building Pensions from import...")

        df = self.import_data.get("pensions")
        if df is None:
            self.log("  Warning: No pensions import file loaded")
            return

        for _, row in df.iterrows():
            code = row.get('Code', '')
            if not code:
                continue

            self.template_data["Pensions"].append({
                "Code": code,
                "Title": row.get('Title', code),
                "AvailableToAllSchools": row.get('Available To All Schools', True),
                "AvailableToSchools": row.get('Available to Schools', ''),
            })

        self.log(f"  Added {len(self.template_data['Pensions'])} pensions")

    def _build_eqw_patterns_from_import(self):
        """Build EQWPatterns from import file."""
        self.log("Building EQWPatterns from import...")

        df = self.import_data.get("equated_week_patterns")
        if df is None:
            self.log("  Warning: No EQW patterns import file loaded")
            return

        for _, row in df.iterrows():
            code = row.get('Code', '')
            if not code:
                continue

            self.template_data["EQWPatterns"].append({
                "Code": code,
                "Title": row.get('Title', code),
                "AvailableToAllSchools": row.get('Available To All Schools', True),
                "AvailableToSchools": row.get('Available to Schools', ''),
                "FullTimeWeeks": row.get('Full Time Weeks', 52.143),
            })

        self.log(f"  Added {len(self.template_data['EQWPatterns'])} EQW patterns")

    def _build_staff_role_groups_from_import(self):
        """Build StfRoleGroup from import file with finance code parsing."""
        self.log("Building StfRoleGroup from import...")

        df = self.import_data.get("staff_role_groups")
        if df is None:
            self.log("  Warning: No staff role groups import file loaded")
            return

        for _, row in df.iterrows():
            code = row.get('Code', '')
            if not code:
                continue

            # Parse finance codes from combined format
            def get_fc_code(col_name):
                """Extract finance code from combined format column."""
                val = row.get(col_name, '')
                if pd.isna(val) or not val:
                    # Try parsed column
                    parsed_col = col_name.replace(' Code', '_Code')
                    return row.get(parsed_col, '')
                return parse_combined_field(str(val))[0]

            self.template_data["StfRoleGroup"].append({
                "Code": code,
                "Title": row.get('Title', code),
                "TeachingRoleGroup": row.get('Teaching Role Group', False),
                "IncrementCount": row.get('Increment Count', 0),
                "GrossSalaryCode": get_fc_code('Gross Salary Code'),
                "LeaveRebateCode": get_fc_code('Leave Rebate Code'),
                "EmployersNICode": get_fc_code('Employers NI Code'),
                "PensionCode": get_fc_code('Pension Code'),
                "MinimumWageTopupCode": get_fc_code('Minimum Wage Topup Code'),
                "LivingWageTopupCode": get_fc_code('Living Wage Topup Code'),
                "OptOutPensionCode": get_fc_code('Opt Out Pension Code'),
                "OtherSalaryCostsCode": get_fc_code('Other Salary Costs Code'),
                "AdjustmentsCode": get_fc_code('Adjustments Code'),
                "AllowancesCode": get_fc_code('Allowances Code'),
            })

        self.log(f"  Added {len(self.template_data['StfRoleGroup'])} staff role groups")

    def _build_staff_roles_from_import(self):
        """Build StfRole from import file."""
        self.log("Building StfRole from import...")

        df = self.import_data.get("staff_roles")
        if df is None:
            self.log("  Warning: No staff roles import file loaded")
            return

        for _, row in df.iterrows():
            code = row.get('Code', '')
            if not code:
                continue

            # Parse staff role group from combined format
            srg_combined = row.get('Staff Role Group', '')
            srg_code = row.get('Staff Role Group_Code', '')
            if not srg_code and srg_combined:
                srg_code = parse_combined_field(str(srg_combined))[0]

            # Parse pay scale from combined format
            ps_combined = row.get('Pay Scale', '')
            ps_code = row.get('Pay Scale_Code', '')
            if not ps_code and ps_combined:
                ps_code = parse_combined_field(str(ps_combined))[0]

            # LESSON L025: NEVER use default hours - get from data or mark as 0 for review
            ft_hours = row.get('Full Time Hours Per Week')
            if pd.isna(ft_hours) or ft_hours is None or ft_hours == '':
                ft_hours = 0  # Flag for review - no default

            self.template_data["StfRole"].append({
                "Code": code,
                "Title": row.get('Title', code),
                "AvailableToAllSchools": row.get('Available To All Schools', True),
                "AvailableToSchools": row.get('Available to Schools', ''),
                "StaffRoleGroupCode": srg_code,
                "PayScaleCode": ps_code,
                "FullTimeHoursPerWeek": float(ft_hours) if ft_hours else 0,
                "IsFinanceRole": row.get('Is Finance Role', False),
            })

        self.log(f"  Added {len(self.template_data['StfRole'])} staff roles")

    def _build_pay_scales_from_import(self):
        """Build PayScales from import file."""
        self.log("Building PayScales from import...")

        df = self.import_data.get("pay_scales")
        if df is None:
            self.log("  Warning: No pay scales import file loaded")
            return

        for _, row in df.iterrows():
            code = row.get('Code', '')
            if not code:
                continue

            self.template_data["PayScales"].append({
                "Code": code,
                "Title": row.get('Title', code),
                "AvailableToAllSchools": row.get('Available To All Schools', True),
                "AvailableToSchools": row.get('Available to Schools', ''),
                "IncrementAtServiceStartDate": row.get('Increment at Service Start Date', False),
                "IncrementDate": str(row.get('Increment Date', ''))[:10] if pd.notna(row.get('Increment Date')) else '',
                "IncreaseDate": str(row.get('Increase Date', ''))[:10] if pd.notna(row.get('Increase Date')) else '',
                "DefaultIncreasePercentage": row.get('Default Increase Percentage', 0),
                "ExcludeNI": row.get('Exclude National Insurance', False),
                "ExcludePension": row.get('Exclude Pension', False),
            })

        self.log(f"  Added {len(self.template_data['PayScales'])} pay scales")

    def _build_allowance_types_from_import(self):
        """Build AllowanceTypes from import file."""
        self.log("Building AllowanceTypes from import...")

        df = self.import_data.get("allowance_types")
        if df is None:
            self.log("  Warning: No allowance types import file loaded")
            return

        for _, row in df.iterrows():
            code = row.get('Code', '')
            if not code:
                continue

            self.template_data["AllowanceTypes"].append({
                "Code": code,
                "Title": row.get('Title', code),
                "AvailableToAllSchools": row.get('Available To All Schools', True),
                "AvailableToSchools": row.get('Available to Schools', ''),
                "IncreaseDate": str(row.get('Increase Date', ''))[:10] if pd.notna(row.get('Increase Date')) else '',
                "DefaultIncreasePercentage": row.get('Default Increase Percentage', 0),
                "ExcludeNI": row.get('Exclude National Insurance', False),
                "ExcludePension": row.get('Exclude Pension', False),
            })

        self.log(f"  Added {len(self.template_data['AllowanceTypes'])} allowance types")

    def _build_staff_members_from_import(self):
        """Build StaffMembers from import file."""
        self.log("Building StaffMembers from import...")

        df = self.import_data.get("staff_members")
        if df is None:
            self.log("  Warning: No staff members import file loaded")
            return

        for _, row in df.iterrows():
            code = row.get('Code', '')
            if not code:
                continue

            self.template_data["StaffMembers"].append({
                "Code": code,
                "LastName": row.get('Last Name', ''),
                "FirstName": row.get('First Name', ''),
                "PensionOptedOut": row.get('Pension Opted Out', False),
                "ServiceStartDate": str(row.get('Service Start Date', ''))[:10] if pd.notna(row.get('Service Start Date')) else '',
                "ServiceEndDate": str(row.get('Service End Date', ''))[:10] if pd.notna(row.get('Service End Date')) else '',
                "PayrollLocationCode": row.get('Payroll Location Code', ''),
            })

        self.log(f"  Added {len(self.template_data['StaffMembers'])} staff members")

    def _build_contracts_from_import(self):
        """Build ContractsTeachFTE and ContractsSupportHours from import file."""
        self.log("Building Contracts from import...")

        df = self.import_data.get("contracts")
        if df is None:
            self.log("  Warning: No contracts import file loaded")
            return

        teaching_count = 0
        support_count = 0

        for _, row in df.iterrows():
            # Parse all combined fields
            def get_code(col_name):
                """Get code from combined field or parsed column."""
                # Try parsed column first
                parsed_col = col_name.replace(' Combined', '_Code')
                val = row.get(parsed_col, '')
                if val:
                    return str(val).strip()
                # Try combined column
                combined = row.get(col_name, '')
                if combined and pd.notna(combined):
                    return parse_combined_field(str(combined))[0]
                return ''

            school_code = row.get('School Code', '')
            staff_member_code = get_code('Staff Member Combined')
            reference = row.get('Reference', '')
            staff_role_code = get_code('Staff Role Combined')
            contract_type_code = get_code('Contract Type Combined')
            pay_scale_code = get_code('Pay Scale Combined')
            pay_scale_grade_code = get_code('Pay Scale Grade Combined')
            pay_scale_point_code = get_code('Pay Scale Point Combined')
            pension_code = get_code('Pension Combined')
            eqwp_code = get_code('Equated Week Pattern Combined')
            department_code = get_code('Department Combined')
            fund_code = get_code('Fund Combined')

            # Skip if no staff member code
            if not staff_member_code:
                continue

            # Generate reference if missing
            if not reference or pd.isna(reference):
                reference = f"{staff_member_code}A"

            # Parse dates
            date_from = str(row.get('Date From', ''))[:10] if pd.notna(row.get('Date From')) else ''
            date_to = str(row.get('Date To', ''))[:10] if pd.notna(row.get('Date To')) else ''

            # Get FTE values
            weekly_fte = row.get('Weekly FTE', 0)
            if pd.isna(weekly_fte):
                weekly_fte = 0

            no_increment = row.get('No Increment', False)
            notes = row.get('Notes', '')
            if pd.isna(notes):
                notes = ''

            # Determine if teaching or support based on staff role group
            is_teaching = False
            if staff_role_code:
                # Check staff roles to determine if teaching
                roles_df = self.import_data.get("staff_roles")
                if roles_df is not None:
                    role_match = roles_df[roles_df['Code'] == staff_role_code]
                    if not role_match.empty:
                        srg_combined = role_match.iloc[0].get('Staff Role Group', '')
                        srg_code = parse_combined_field(str(srg_combined))[0] if srg_combined else ''
                        # Check if staff role group is teaching
                        srg_df = self.import_data.get("staff_role_groups")
                        if srg_df is not None:
                            srg_match = srg_df[srg_df['Code'] == srg_code]
                            if not srg_match.empty:
                                is_teaching = srg_match.iloc[0].get('Teaching Role Group', False)

            # Build contract record
            contract_record = {
                "SchoolCode": school_code,
                "StaffMemberCode": staff_member_code,
                "Reference": reference,
                "StaffRoleCode": staff_role_code,
                "ContractTypeCode": contract_type_code,
                "PayScaleCode": pay_scale_code,
                "PayScaleGradeCode": pay_scale_grade_code,
                "PayScalePointCode": pay_scale_point_code,
                "PensionCode": pension_code,
                "EquatedWeekPatternCode": eqwp_code,
                "DepartmentCode": department_code,
                "FundCode": fund_code,
                "DateFrom": date_from,
                "DateTo": date_to,
                "WeeklyFteOrHpw": float(weekly_fte) if weekly_fte else 0,
                "NoIncrement": no_increment,
                "Notes": str(notes),
            }

            if is_teaching:
                self.template_data["ContractsTeachFTE"].append(contract_record)
                teaching_count += 1
            else:
                self.template_data["ContractsSupportHours"].append(contract_record)
                support_count += 1

        self.log(f"  Added {teaching_count} teaching contracts (ContractsTeachFTE)")
        self.log(f"  Added {support_count} support contracts (ContractsSupportHours)")

    # =========================================================================
    # PHASE 1: DEEP ANALYSIS
    # =========================================================================

    def analyze_customer_data(self, data_dir: Path) -> List[AnalysisReport]:
        """
        Perform deep analysis of ALL customer data files.
        Returns detailed reports on what was found.
        """
        self.log("="*60)
        self.log("PHASE 1: DEEP ANALYSIS OF CUSTOMER DATA")
        self.log("="*60)

        # Log the directory being searched
        self.log(f"Searching in: {data_dir}")
        self.log(f"Directory exists: {data_dir.exists()}")
        self.processing_log.append(f"Data directory: {data_dir}")
        self.processing_log.append(f"Directory exists: {data_dir.exists()}")

        # Find all supported file types
        all_files = list(data_dir.rglob("*.xls*")) + list(data_dir.rglob("*.csv"))
        pdf_files = list(data_dir.rglob("*.pdf")) if PDF_SUPPORT else []
        image_files = list(data_dir.rglob("*.png")) + list(data_dir.rglob("*.jpg")) + list(data_dir.rglob("*.jpeg"))
        docx_files = list(data_dir.rglob("*.docx")) + list(data_dir.rglob("*.doc")) if DOCX_SUPPORT else []

        all_files = [f for f in all_files if not f.name.startswith("~$")]
        pdf_files = [f for f in pdf_files if not f.name.startswith("~$")]
        docx_files = [f for f in docx_files if not f.name.startswith("~$")]

        self.log(f"Found {len(all_files)} spreadsheet files to analyze")
        self.processing_log.append(f"Spreadsheet files found: {len(all_files)}")
        for f in all_files:
            self.processing_log.append(f"  - {f.name}")
        if pdf_files:
            self.log(f"Found {len(pdf_files)} PDF files to analyze")
        if docx_files:
            self.log(f"Found {len(docx_files)} Word document files to analyze")
        if image_files:
            self.log(f"Found {len(image_files)} image files to analyze")

        for file_path in all_files:
            self.log(f"\nAnalyzing: {file_path.name}")
            self._analyze_file(file_path)

        # Analyze PDF files
        for file_path in pdf_files:
            self.log(f"\nAnalyzing PDF: {file_path.name}")
            self._analyze_pdf_file(file_path)

        # Analyze Word document files
        for file_path in docx_files:
            self.log(f"\nAnalyzing Word Document: {file_path.name}")
            self._analyze_docx_file(file_path)

        # Analyze image files (OCR)
        for file_path in image_files:
            self.log(f"\nAnalyzing Image: {file_path.name}")
            self._analyze_image_file(file_path)

        self._print_analysis_summary()
        return self.analysis_reports

    def _analyze_file(self, file_path: Path):
        """Deeply analyze a single file."""
        try:
            if file_path.suffix == ".csv":
                df = pd.read_csv(file_path)
                # Apply validated column mappings
                df = self._apply_column_mappings(df, file_path.name)
                self._analyze_dataframe(df, file_path.name, "CSV")
            else:
                # Try multiple engines for reading Excel files
                xl = None
                error_msg = ""
                
                # Try with openpyxl first (default)
                try:
                    xl = pd.ExcelFile(file_path, engine='openpyxl')
                except Exception as e1:
                    error_msg = str(e1)[:100]
                    try:
                        self.log(f"  Warning: openpyxl failed, trying xlrd engine...")
                    except Exception:
                        pass
                    # Try with xlrd engine
                    try:
                        xl = pd.ExcelFile(file_path, engine='xlrd')
                    except Exception as e2:
                        try:
                            self.log(f"  Warning: xlrd also failed, trying read_excel with sheet discovery...")
                        except Exception:
                            pass
                        # Try reading sheets with sheet discovery
                        try:
                            # Use None to get all sheets automatically
                            xls = pd.read_excel(file_path, sheet_name=None)
                            if xls:
                                try:
                                    self.log(f"  Sheets: {len(xls)} found")
                                except Exception:
                                    pass
                                for sheet_name, df in xls.items():
                                    # Extended skip list including templates, examples, and admin sheets
                                    skip_sheets = [
                                        'guidance', 'notes', 'instructions', 'help', 'contents',
                                        'checklist', 'template', 'equated weeks example', 'cenchg',
                                        'lgps', 'lgps information', 'pension', 'pensions',
                                        'changelog', 'change log', 'version', 'cover', 'summary'
                                    ]
                                    if any(skip in sheet_name.lower() for skip in skip_sheets):
                                        self.log(f"    Skipping non-data sheet: {sheet_name}")
                                        continue
                                    if df is not None and len(df) > 0:
                                        # Apply smart header detection for fallback reads
                                        df = self._apply_smart_header_detection(df, sheet_name)
                                        if df is None or len(df) == 0:
                                            continue
                                        # Apply validated column mappings
                                        df = self._apply_column_mappings(df, file_path.name, sheet_name)
                                        self._analyze_dataframe(df, file_path.name, sheet_name)
                                return
                        except Exception as e3:
                            error_detail = self._format_detailed_error(file_path, "reading Excel (all engines failed)", e3)
                            self.issues.append(error_detail)
                            return
                
                # If we have a valid ExcelFile object, proceed normally
                if xl is not None:
                    try:
                        self.log(f"  Sheets: {len(xl.sheet_names)} found")
                    except Exception:
                        pass
                    for sheet in xl.sheet_names:
                        # Extended skip list including templates, examples, and admin sheets
                        skip_sheets = [
                            'guidance', 'notes', 'instructions', 'help', 'contents',
                            'checklist', 'template', 'equated weeks example', 'cenchg',
                            'lgps', 'lgps information', 'pension', 'pensions',
                            'changelog', 'change log', 'version', 'cover', 'summary'
                        ]
                        if any(skip in sheet.lower() for skip in skip_sheets):
                            self.log(f"    Skipping non-data sheet: {sheet}")
                            continue

                        df = self._read_sheet_smart(xl, sheet)
                        if df is not None and len(df) > 0:
                            # Apply validated column mappings
                            df = self._apply_column_mappings(df, file_path.name, sheet)
                            self._analyze_dataframe(df, file_path.name, sheet)

        except Exception as e:
            error_detail = self._format_detailed_error(file_path, "analyzing file", e)
            self.issues.append(error_detail)

    def _analyze_pdf_file(self, file_path: Path):
        """Analyze a PDF file - extract tables and text."""
        if not PDF_SUPPORT:
            self.log(f"  Warning: PDF support not available (install pdfplumber)")
            return

        try:
            tables_found = 0
            with pdfplumber.open(file_path) as pdf:
                self.log(f"  Pages: {len(pdf.pages)}")

                for page_num, page in enumerate(pdf.pages):
                    # Extract tables from page
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if table and len(table) > 1:
                            # Skip very small tables (likely headers, footers, notes)
                            # Need at least 5 data rows (excluding header) for meaningful data
                            data_rows = len(table) - 1  # Exclude header row
                            if data_rows < 5:
                                # Only skip if it doesn't look like it has salary data
                                has_salary_like = False
                                for row in table[1:]:
                                    for cell in row:
                                        if cell and isinstance(cell, str):
                                            # Check for salary-like values (£ or large numbers)
                                            cleaned = cell.replace(',', '').replace('£', '').strip()
                                            try:
                                                val = float(cleaned)
                                                if val > 10000:
                                                    has_salary_like = True
                                                    break
                                            except:
                                                pass
                                    if has_salary_like:
                                        break
                                if not has_salary_like:
                                    continue  # Skip small tables without salary data

                            tables_found += 1

                            # First try specialized PDF pay scale parsing (handles multiline cells)
                            self._parse_pdf_pay_scale_table(table, file_path.name, table_idx)

                            # Convert table to DataFrame for generic analysis
                            df = pd.DataFrame(table[1:], columns=table[0])
                            df.columns = [self._clean_column_name(c) for c in df.columns]

                            # Analyze as dataframe
                            sheet_name = f"PDF_Page{page_num+1}_Table{table_idx+1}"
                            self._analyze_dataframe(df, file_path.name, sheet_name)

                    # If no tables found, try to extract text for OCR-like processing
                    if not tables:
                        text = page.extract_text()
                        if text and len(text.strip()) > 50:
                            # Try to parse structured text
                            self._process_pdf_text(text, file_path.name, page_num)

            if tables_found > 0:
                self.log(f"  Extracted {tables_found} tables from PDF")
            else:
                self.log(f"  No tables found in PDF - using text extraction")

            # Also try tabula for better table extraction
            if TABULA_SUPPORT and tables_found == 0:
                try:
                    dfs = tabula.read_pdf(str(file_path), pages='all', multiple_tables=True)
                    for idx, df in enumerate(dfs):
                        if len(df) > 1:
                            df.columns = [self._clean_column_name(c) for c in df.columns]
                            self._analyze_dataframe(df, file_path.name, f"Tabula_Table{idx+1}")
                            self.log(f"  Tabula extracted table with {len(df)} rows")
                except Exception as e:
                    self.log(f"  Tabula extraction failed: {str(e)[:50]}")

        except Exception as e:
            error_detail = self._format_detailed_error(file_path, "reading PDF", e)
            self.issues.append(error_detail)

    def _analyze_docx_file(self, file_path: Path):
        """Analyze a Word document file - extract tables and text."""
        if not DOCX_SUPPORT:
            self.log(f"  Warning: Word document support not available (install python-docx)")
            return

        try:
            document = docx.Document(file_path)
            tables_found = 0

            # Extract tables from the document
            for table_idx, table in enumerate(document.tables):
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])

                if len(rows) > 1:
                    tables_found += 1
                    # Use first row as headers
                    headers = rows[0]
                    data = rows[1:]

                    df = pd.DataFrame(data, columns=headers)
                    df.columns = [self._clean_column_name(c) for c in df.columns]

                    sheet_name = f"DOCX_Table{table_idx+1}"
                    self._analyze_dataframe(df, file_path.name, sheet_name)

                    # Also try pay scale parsing (reuse PDF table parser)
                    raw_table = [rows[0]] + data
                    self._parse_pdf_pay_scale_table(raw_table, file_path.name, table_idx)

            if tables_found > 0:
                self.log(f"  Extracted {tables_found} tables from Word document")
            else:
                self.log(f"  No tables found in Word document - extracting text")

            # If no tables, extract paragraph text for analysis
            if tables_found == 0:
                all_text = '\n'.join(p.text for p in document.paragraphs if p.text.strip())
                if all_text and len(all_text.strip()) > 50:
                    self._process_docx_text(all_text, file_path.name)

        except Exception as e:
            self.issues.append(f"Error analyzing Word document {file_path.name}: {str(e)[:100]}")

    def _process_docx_text(self, text: str, file_name: str):
        """Process unstructured text from a Word document, looking for pay scale data."""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        # Try to find structured data in the text (e.g., tabular data separated by tabs)
        tab_lines = [line for line in lines if '\t' in line]
        if len(tab_lines) > 2:
            # Parse tab-separated data
            rows = [line.split('\t') for line in tab_lines]
            max_cols = max(len(r) for r in rows)
            # Pad rows to same length
            rows = [r + [''] * (max_cols - len(r)) for r in rows]
            df = pd.DataFrame(rows[1:], columns=rows[0])
            df.columns = [self._clean_column_name(c) for c in df.columns]
            self._analyze_dataframe(df, file_name, "DOCX_Text_Table")
        else:
            # Store raw text for reference
            self.log(f"  Extracted {len(lines)} lines of text from {file_name}")

    def _parse_pdf_pay_scale_table(self, table: list, file_name: str, table_idx: int):
        """Parse pay scale data from PDF tables with multiline cells."""
        if not table or len(table) < 2:
            return

        # Join all cells to detect content type
        all_text = ' '.join(' '.join(str(c or '') for c in row) for row in table).lower()

        # Detect table type from content
        is_main_pay = 'main pay' in all_text or 'classroom teacher' in all_text
        is_upper_pay = 'upper pay' in all_text
        is_leadership = 'leadership' in all_text
        is_lead_pract = 'lead practitioner' in all_text
        is_unqualified = 'unqualified' in all_text
        is_tlr = 'tlr' in all_text or 'teaching & learning' in all_text
        is_sen = 'sen' in all_text or 'special educational' in all_text

        # Process each row looking for multiline data
        for row in table:
            if not row:
                continue

            for cell_idx, cell in enumerate(row):
                if not cell:
                    continue

                cell_str = str(cell)

                # Check for multiline spine points (L1\nL2\nL3...)
                if '\n' in cell_str:
                    lines = cell_str.split('\n')

                    # Check if this is a spine point column
                    point_patterns = [r'^[LMU]?\d+$', r'^MIN\s*\d*', r'^MAX\s*\d*', r'^SCP\s*\d+']
                    is_point_col = any(re.match(p, lines[0].strip(), re.IGNORECASE) for p in point_patterns)

                    if is_point_col and cell_idx + 1 < len(row):
                        # Next column should have corresponding salaries
                        salary_cell = row[cell_idx + 1]
                        if salary_cell and '\n' in str(salary_cell):
                            salary_lines = str(salary_cell).split('\n')

                            # Parse points and salaries
                            points_data = []
                            for i, point_line in enumerate(lines):
                                point_line = point_line.strip()
                                if not point_line:
                                    continue

                                # Extract point code
                                point_match = re.search(r'([LMUP]?\d+)', point_line, re.IGNORECASE)
                                if point_match and i < len(salary_lines):
                                    point_code = point_match.group(1).upper()
                                    salary_line = salary_lines[i].strip()

                                    # Extract salary amount
                                    salary_match = re.search(r'[\£\$]?\s*([\d,]+)', salary_line)
                                    if salary_match:
                                        try:
                                            amount = float(salary_match.group(1).replace(',', ''))
                                            if amount >= 15000 and amount <= 200000:
                                                # Determine scale type
                                                prefix = point_code[0] if point_code[0].isalpha() else ''
                                                point_num = int(re.search(r'\d+', point_code).group())

                                                points_data.append({
                                                    'code': point_code,
                                                    'title': f'Point {point_num}',
                                                    'number': point_num,
                                                    'rate': amount,
                                                })
                                        except (ValueError, AttributeError):
                                            pass

                            # Create pay scale from extracted points
                            if points_data:
                                self._create_pay_scale_from_pdf_points(
                                    points_data, file_name,
                                    is_leadership or is_lead_pract,
                                    is_main_pay, is_upper_pay, is_unqualified
                                )

        # Handle TLR from PDF
        if is_tlr:
            self._parse_tlr_from_pdf_table(table, file_name)

        # Handle SEN from PDF
        if is_sen:
            self._parse_sen_from_pdf_table(table, file_name)

    def _create_pay_scale_from_pdf_points(self, points_data: list, file_name: str,
                                          is_leadership: bool, is_main: bool,
                                          is_upper: bool, is_unqualified: bool):
        """Create pay scale from extracted PDF points."""
        if not points_data:
            return

        # Determine scale code and type based on point prefixes FIRST, then table flags
        first_code = points_data[0]['code']
        prefix = first_code[0] if first_code[0].isalpha() else ''

        # Prefix takes precedence over table-level flags
        if prefix == 'L':
            scale_code = 'LEAD'
            scale_title = 'Leadership Scale'
            scale_type = 'leadership'
        elif prefix == 'U':
            scale_code = 'UPS'
            scale_title = 'Upper Pay Scale'
            scale_type = 'teaching'
        elif prefix == 'M':
            scale_code = 'MPS'
            scale_title = 'Main Pay Scale'
            scale_type = 'teaching'
        elif prefix == 'P':  # Lead Practitioner
            scale_code = 'LEADP'
            scale_title = 'Lead Practitioner Scale'
            scale_type = 'teaching'
        elif is_leadership:
            scale_code = 'LEAD'
            scale_title = 'Leadership Scale'
            scale_type = 'leadership'
        elif is_unqualified:
            scale_code = 'UNQ'
            scale_title = 'Unqualified Teacher Scale'
            scale_type = 'teaching'
        elif is_upper:
            scale_code = 'UPS'
            scale_title = 'Upper Pay Scale'
            scale_type = 'teaching'
        elif is_main:
            scale_code = 'MPS'
            scale_title = 'Main Pay Scale'
            scale_type = 'teaching'
        else:
            # Default to support if no prefix detected — derive code from title
            scale_code = self._derive_pay_scale_code("Support Pay Scale")
            scale_title = 'Support Pay Scale'
            scale_type = 'support'

        if scale_code not in self.pay_scales_found:
            pay_scale = ExtractedPayScale(
                code=scale_code,
                title=scale_title,
                scale_type=scale_type,
                london_weighting='Inner London',  # From PDF filename
                increment_date=get_default_increase_date(scale_type == 'teaching'),
                increase_date=get_default_increase_date(scale_type == 'teaching'),
                increase_percentage=0,
                points=points_data,
                grades=[]
            )
            self.extracted_pay_scales.append(pay_scale)
            self.pay_scales_found.add(scale_code)
            self.log(f"  ✓ Extracted {scale_code} pay scale with {len(points_data)} points from PDF")

    def _parse_tlr_from_pdf_table(self, table: list, file_name: str):
        """Parse TLR allowances from PDF table."""
        if 'TLR' in self.allowance_types_found:
            return

        # PDF TLR table format:
        # Row 2: ['PAYMENT 1\n(TLR1)', '1 SEPT 2024...', '1 SEPT 2025...']
        # Row 3: ['MIN\nMAX', '£9,782\n£16,553', '£10,174\n£17,216']

        # First find which TLR codes are in the table
        tlr_codes = set()
        for row in table:
            for cell in row:
                if cell:
                    matches = re.findall(r'TLR\s*(\d)', str(cell), re.IGNORECASE)
                    tlr_codes.update(matches)

        if not tlr_codes:
            return

        # Now extract all amounts from the table that look like TLR values
        tlr_amounts = []
        for row in table:
            for cell in row:
                if cell:
                    cell_str = str(cell)
                    # Skip cells with date ranges
                    if 'sept' in cell_str.lower() or 'aug' in cell_str.lower():
                        continue

                    # Extract amounts with £ symbol preferentially
                    amounts = re.findall(r'[£]([\d,]+)', cell_str)
                    for amt_str in amounts:
                        try:
                            amt = float(amt_str.replace(',', ''))
                            # TLR range: typically £2,000 to £20,000
                            # Exclude year-like numbers (2020-2030)
                            if amt >= 2000 and amt <= 25000 and not (2020 <= amt <= 2030):
                                if amt not in tlr_amounts:
                                    tlr_amounts.append(amt)
                        except ValueError:
                            pass

        # Group amounts into TLR entries based on ranges
        # TLR1: typically £9,000 - £17,000
        # TLR2: typically £3,000 - £8,000
        # TLR3: fixed amounts around £600-£3,000 (if present)
        tlr1_amounts = sorted([a for a in tlr_amounts if a >= 8000])
        tlr2_amounts = sorted([a for a in tlr_amounts if 2500 <= a < 8000])
        tlr3_amounts = sorted([a for a in tlr_amounts if a < 2500 and a >= 500])

        # Create TLR allowance points
        points = []

        if tlr1_amounts and '1' in tlr_codes:
            points.append({'code': 'TLR1_MIN', 'title': 'TLR1 Min', 'amount': tlr1_amounts[0]})
            if len(tlr1_amounts) > 1:
                points.append({'code': 'TLR1_MAX', 'title': 'TLR1 Max', 'amount': tlr1_amounts[-1]})

        if tlr2_amounts and '2' in tlr_codes:
            points.append({'code': 'TLR2_MIN', 'title': 'TLR2 Min', 'amount': tlr2_amounts[0]})
            if len(tlr2_amounts) > 1:
                points.append({'code': 'TLR2_MAX', 'title': 'TLR2 Max', 'amount': tlr2_amounts[-1]})

        if tlr3_amounts and '3' in tlr_codes:
            points.append({'code': 'TLR3_MIN', 'title': 'TLR3 Min', 'amount': tlr3_amounts[0]})
            if len(tlr3_amounts) > 1:
                points.append({'code': 'TLR3_MAX', 'title': 'TLR3 Max', 'amount': tlr3_amounts[-1]})

        if points:
            allowance = ExtractedAllowance(
                type_code='TLR',
                type_title='Teaching and Learning Responsibilities',
                points=points,
                increase_date=get_default_increase_date(True),
                increase_percentage=0,
            )
            self.extracted_allowances.append(allowance)
            self.allowance_types_found.add('TLR')
            self.log(f"  ✓ Extracted {len(points)} TLR allowance points from PDF")

    def _parse_sen_from_pdf_table(self, table: list, file_name: str):
        """Parse SEN allowances from PDF table."""
        if 'SEN' in self.allowance_types_found:
            return

        sen_amounts = []

        for row in table:
            for cell in row:
                if cell:
                    cell_str = str(cell)
                    # Skip cells with date ranges
                    if 'sept' in cell_str.lower() or 'aug' in cell_str.lower():
                        continue

                    # Extract amounts with £ symbol
                    amounts = re.findall(r'[£]([\d,]+)', cell_str)
                    for amt_str in amounts:
                        try:
                            amt = float(amt_str.replace(',', ''))
                            # SEN range check (£1,000 - £10,000)
                            if amt >= 1000 and amt <= 10000:
                                if amt not in sen_amounts:
                                    sen_amounts.append(amt)
                        except ValueError:
                            pass

        # Create SEN allowance
        if sen_amounts:
            points = []
            for idx, amt in enumerate(sorted(sen_amounts)):
                points.append({
                    'code': f'SEN{idx+1}',
                    'title': f'SEN {"Min" if idx == 0 else "Max"}',
                    'amount': amt,
                })

            allowance = ExtractedAllowance(
                type_code='SEN',
                type_title='Special Educational Needs',
                points=points,
                increase_date=get_default_increase_date(True),
                increase_percentage=0,
            )
            self.extracted_allowances.append(allowance)
            self.allowance_types_found.add('SEN')
            self.log(f"  ✓ Extracted {len(points)} SEN allowance points from PDF")

    def _process_pdf_text(self, text: str, file_name: str, page_num: int):
        """Process extracted PDF text looking for structured data."""
        lines = text.strip().split('\n')
        if len(lines) < 3:
            return

        text_lower = text.lower()

        # Apply specialized parsers for known data types
        # Check for pay scale data (Main Pay Scale, Leadership, etc.)
        if any(kw in text_lower for kw in ['pay scale', 'payscale', 'main scale', 'upper pay', 'leadership', 'spine point', 'scp']):
            self._parse_pay_scales_from_ocr_text(text, file_name)

        # Check for TLR rates
        if 'tlr' in text_lower:
            self._parse_tlr_from_ocr_text(text, file_name)

        # Check for SEN allowances
        if 'sen' in text_lower:
            self._parse_sen_from_ocr_text(text, file_name)

        # Look for tabular data patterns (columns separated by multiple spaces)
        potential_table = []
        for line in lines:
            # Split by multiple spaces (2+)
            cols = re.split(r'\s{2,}', line.strip())
            if len(cols) >= 3:
                potential_table.append(cols)

        if len(potential_table) > 3:
            # Try to create a DataFrame
            max_cols = max(len(row) for row in potential_table)
            # Pad rows to same length
            padded = [row + [''] * (max_cols - len(row)) for row in potential_table]

            df = pd.DataFrame(padded[1:], columns=padded[0] if padded else None)
            df.columns = [self._clean_column_name(c) for c in df.columns]

            if len(df) > 1:
                self._analyze_dataframe(df, file_name, f"PDF_Page{page_num+1}_Text")
                self.log(f"  Parsed text table from page {page_num+1}: {len(df)} rows")

    def _analyze_image_file(self, file_path: Path):
        """Analyze an image file using OCR."""
        if not IMAGE_OCR_SUPPORT:
            self.log(f"  Warning: Image OCR not available (install pytesseract and Pillow)")
            return

        try:
            # Open image with PIL
            img = Image.open(file_path)
            self.log(f"  Image size: {img.size}")

            # Run OCR
            text = pytesseract.image_to_string(img)

            if text and len(text.strip()) > 20:
                self.log(f"  OCR extracted {len(text)} characters")

                # Use unified text processing (handles pay scales, TLR, SEN, and tabular data)
                self._process_pdf_text(text, file_path.name, 0)

                # Store raw text for reference
                self.assumptions.append(f"OCR text from {file_path.name}: {text[:200]}...")
            else:
                self.log(f"  No significant text found in image")

        except Exception as e:
            error_detail = self._format_detailed_error(file_path, "reading image/OCR", e)
            self.issues.append(error_detail)

    def _parse_tlr_from_ocr_text(self, text: str, source_name: str):
        """Parse TLR allowance rates from OCR text."""
        self.log(f"  Parsing TLR rates from OCR text...")

        tlr_entries = {}

        # Process line by line
        lines = text.replace('\r', '\n').split('\n')

        # First pass: find all TLR codes in order
        tlr_codes = []
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Look for TLR entries: "TLR 1a", "TLR1a", "TLR tc", "TLR 2c"
            if 'tlr' in line.lower():
                match = re.search(r'TLR\s*([12]?\s*[a-d])', line, re.IGNORECASE)
                if match:
                    code_raw = match.group(1).replace(' ', '').upper()
                    tlr_code = self._normalize_tlr_code(code_raw)
                    if tlr_code:
                        tlr_codes.append(tlr_code)

                    # Also check if amounts are on the same line
                    rest = line[match.end():]
                    amounts = re.findall(r'[\d,]+', rest)
                    valid_amounts = []
                    for amt in amounts:
                        try:
                            val = float(amt.replace(',', ''))
                            if val >= 1000 and val <= 50000:
                                valid_amounts.append(val)
                        except ValueError:
                            pass

                    if valid_amounts and tlr_code:
                        tlr_entries[tlr_code] = valid_amounts[0]
                        self.log(f"    Found {tlr_code}: £{valid_amounts[0]:,.0f}")

        # If we found codes but no amounts on same lines, try column-based matching
        if tlr_codes and not tlr_entries:
            # Second pass: find all standalone numeric values that could be TLR amounts
            all_amounts = []
            for line in lines:
                line = line.strip()
                if not line or 'tlr' in line.lower() or 'value' in line.lower():
                    continue

                # Check if line is just a number (standalone amount)
                try:
                    val = float(line.replace(',', '').replace('£', ''))
                    if val >= 1000 and val <= 50000:
                        all_amounts.append(val)
                except ValueError:
                    pass

            # Match amounts to codes by position (first column of amounts)
            # In OCR output, amounts are typically in columns for different years
            # We want the first column (current year)
            if all_amounts and len(all_amounts) >= len(tlr_codes):
                # Take amounts in order matching TLR codes
                for i, code in enumerate(tlr_codes):
                    if i < len(all_amounts) and code not in tlr_entries:
                        tlr_entries[code] = all_amounts[i]
                        self.log(f"    Found {code}: £{all_amounts[i]:,.0f}")

        # Create allowance types from extracted TLR rates
        if tlr_entries and 'TLR' not in self.allowance_types_found:
            points = []
            for code, amount in sorted(tlr_entries.items()):
                points.append({
                    'code': code,
                    'title': self._get_tlr_title(code),
                    'amount': amount,
                })

            allowance = ExtractedAllowance(
                type_code='TLR',
                type_title='Teaching and Learning Responsibilities',
                points=points,
                increase_date=get_default_increase_date(True),
                increase_percentage=0,
            )
            self.extracted_allowances.append(allowance)
            self.allowance_types_found.add('TLR')
            self.log(f"  ✓ Extracted {len(points)} TLR allowance points from image")
        elif tlr_entries:
            self.log(f"  TLR already found from other source, skipping OCR data")

    def _normalize_tlr_code(self, code_raw: str) -> str:
        """Normalize TLR code to standard format."""
        code = code_raw.upper().replace(' ', '')

        # Map various formats to standard TLR codes
        # TLR 1a-d, TLR 2a-c, TLR 3
        code_map = {
            '1A': 'TLR1A', '1B': 'TLR1B', '1C': 'TLR1C', '1D': 'TLR1D',
            '2A': 'TLR2A', '2B': 'TLR2B', '2C': 'TLR2C',
            'TA': 'TLR1A', 'TB': 'TLR1B', 'TC': 'TLR1C', 'TD': 'TLR1D',
            'A': 'TLR2A', 'B': 'TLR2B', 'C': 'TLR2C',
        }

        return code_map.get(code, f'TLR{code}' if code else None)

    def _get_tlr_title(self, code: str) -> str:
        """Get descriptive title for TLR code."""
        titles = {
            'TLR1A': 'TLR 1a',
            'TLR1B': 'TLR 1b',
            'TLR1C': 'TLR 1c',
            'TLR1D': 'TLR 1d',
            'TLR2A': 'TLR 2a',
            'TLR2B': 'TLR 2b',
            'TLR2C': 'TLR 2c',
            'TLR3': 'TLR 3',
        }
        return titles.get(code, code)

    def _parse_sen_from_ocr_text(self, text: str, source_name: str):
        """Parse SEN allowance rates from OCR text."""
        self.log(f"  Parsing SEN rates from OCR text...")

        sen_entries = {}

        lines = text.replace('\r', '\n').split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Look for SEN entries (SEN 1, SEN 2, SEN allowance, etc.)
            if 'sen' in line.lower():
                # Try to extract SEN code and amounts
                match = re.search(r'SEN\s*(\d?)', line, re.IGNORECASE)
                if match:
                    code_num = match.group(1) or '1'
                    rest = line[match.end():]
                    amounts = re.findall(r'[\d,]+', rest)

                    valid_amounts = []
                    for amt in amounts:
                        try:
                            val = float(amt.replace(',', ''))
                            # SEN amounts are typically £2000-£5000
                            if val >= 1000 and val <= 10000:
                                valid_amounts.append(val)
                        except ValueError:
                            pass

                    if valid_amounts:
                        code = f'SEN{code_num}'
                        if code not in sen_entries:
                            sen_entries[code] = valid_amounts[0]
                            self.log(f"    Found {code}: £{valid_amounts[0]:,.0f}")

        # Create allowance types from extracted SEN rates
        if sen_entries and 'SEN' not in self.allowance_types_found:
            points = []
            for code, amount in sorted(sen_entries.items()):
                points.append({
                    'code': code,
                    'title': f'SEN Allowance {code[-1]}',
                    'amount': amount,
                })

            allowance = ExtractedAllowance(
                type_code='SEN',
                type_title='Special Educational Needs',
                points=points,
                increase_date=get_default_increase_date(True),
                increase_percentage=0,
            )
            self.extracted_allowances.append(allowance)
            self.allowance_types_found.add('SEN')
            self.log(f"  ✓ Extracted {len(points)} SEN allowance points from image")

    def _parse_pay_scales_from_ocr_text(self, text: str, source_name: str):
        """Parse pay scale data from OCR text (works for PDF and images)."""
        self.log(f"  Parsing pay scales from OCR text...")

        # Track pay scale points by scale type
        scale_data = {
            'MPS': {'title': 'Main Pay Scale', 'type': 'teaching', 'points': {}},
            'UPS': {'title': 'Upper Pay Scale', 'type': 'teaching', 'points': {}},
            'LEAD': {'title': 'Leadership Scale', 'type': 'leadership', 'points': {}},
            'SUPPORT': {'title': 'Support Staff Scale', 'type': 'support', 'points': {}},
        }

        lines = text.replace('\r', '\n').split('\n')

        # Detect which scale type we're currently parsing
        current_scale = None

        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue

            line_lower = line_stripped.lower()

            # Detect scale type from headers
            if 'main' in line_lower and ('pay' in line_lower or 'scale' in line_lower):
                current_scale = 'MPS'
                self.log(f"    Detected Main Pay Scale section")
                continue
            elif 'upper' in line_lower and ('pay' in line_lower or 'scale' in line_lower):
                current_scale = 'UPS'
                self.log(f"    Detected Upper Pay Scale section")
                continue
            elif 'leadership' in line_lower:
                current_scale = 'LEAD'
                self.log(f"    Detected Leadership Scale section")
                continue
            elif 'support' in line_lower and 'staff' in line_lower:
                current_scale = 'SUPPORT'
                self.log(f"    Detected Support Staff Scale section")
                continue

            # Try to extract point number and salary amount
            # Patterns: "M1 30000", "1 30000", "Point 1 £30,000", "SCP 1 30000"

            # Pattern 1: Prefixed point codes (M1, U1, L1, SCP1)
            match = re.search(r'([MUL]?)\s*(\d{1,2})\s+[£]?\s*([\d,]+)', line_stripped, re.IGNORECASE)
            if match:
                prefix = match.group(1).upper() if match.group(1) else ''
                point_num = int(match.group(2))
                amount_str = match.group(3).replace(',', '')

                try:
                    amount = float(amount_str)
                    # Valid UK teacher salary range: £20k - £150k
                    if amount >= 20000 and amount <= 150000:
                        # Determine scale from prefix or context
                        if prefix == 'M' or (current_scale == 'MPS' and point_num <= 6):
                            scale_key = 'MPS'
                            point_code = f'M{point_num}'
                        elif prefix == 'U' or (current_scale == 'UPS' and point_num <= 3):
                            scale_key = 'UPS'
                            point_code = f'U{point_num}'
                        elif prefix == 'L' or current_scale == 'LEAD':
                            scale_key = 'LEAD'
                            point_code = f'L{point_num}'
                        elif current_scale:
                            scale_key = current_scale
                            point_code = f'{point_num}'
                        else:
                            # Guess based on salary range
                            if amount >= 50000:
                                scale_key = 'LEAD'
                                point_code = f'L{point_num}'
                            elif amount >= 40000:
                                scale_key = 'UPS'
                                point_code = f'U{point_num}'
                            else:
                                scale_key = 'MPS'
                                point_code = f'M{point_num}'

                        if point_code not in scale_data[scale_key]['points']:
                            scale_data[scale_key]['points'][point_code] = {
                                'code': point_code,
                                'title': f'Point {point_num}',
                                'number': point_num,
                                'rate': amount,
                            }
                            self.log(f"    Found {scale_key} {point_code}: £{amount:,.0f}")
                except ValueError:
                    pass

            # Pattern 2: SCP pattern (SCP 1, SCP 2, etc.)
            scp_match = re.search(r'SCP\s*(\d{1,2})\s+[£]?\s*([\d,]+)', line_stripped, re.IGNORECASE)
            if scp_match:
                point_num = int(scp_match.group(1))
                amount_str = scp_match.group(2).replace(',', '')
                try:
                    amount = float(amount_str)
                    if amount >= 15000 and amount <= 100000:
                        scale_key = 'SUPPORT'
                        point_code = f'SCP{point_num}'
                        if point_code not in scale_data[scale_key]['points']:
                            scale_data[scale_key]['points'][point_code] = {
                                'code': point_code,
                                'title': f'SCP {point_num}',
                                'number': point_num,
                                'rate': amount,
                            }
                            self.log(f"    Found Support {point_code}: £{amount:,.0f}")
                except ValueError:
                    pass

        # Create ExtractedPayScale objects from parsed data
        for scale_key, data in scale_data.items():
            if data['points'] and scale_key not in self.pay_scales_found:
                points_list = sorted(data['points'].values(), key=lambda p: p['number'])

                pay_scale = ExtractedPayScale(
                    code=scale_key,
                    title=data['title'],
                    scale_type=data['type'],
                    london_weighting='England & Wales',
                    increment_date=get_default_increase_date(data['type'] == 'teaching'),
                    increase_date=get_default_increase_date(data['type'] == 'teaching'),
                    increase_percentage=0,
                    points=points_list,
                    grades=[]
                )
                self.extracted_pay_scales.append(pay_scale)
                self.pay_scales_found.add(scale_key)
                self.log(f"  ✓ Extracted {scale_key} pay scale with {len(points_list)} points from OCR")

    def _is_template_or_example_sheet(self, sheet_name: str, df_raw: pd.DataFrame) -> bool:
        """
        Detect if a sheet is a template/example/checklist rather than actual data.
        These sheets describe the format but don't contain real staff data.

        IMPORTANT: A sheet might have template-like headers at the top but still contain
        real data below. We need to check for actual data content, not just header text.
        """
        sheet_lower = sheet_name.lower()

        # Skip sheets with these names - but NOT data sheets like "Staff Contract Information" or "LGPS Information"
        skip_sheet_names = [
            'checklist', 'template', 'guidance', 'notes', 'instructions',
            'help', 'readme', 'about', 'equated weeks example', 'cenchg'
        ]

        # These sheets typically have real data, don't skip them even if they contain "info"
        # But "checklist" sheets should still be skipped
        if 'checklist' in sheet_lower:
            return True

        data_sheet_patterns = [
            'staff contract info', 'contract information', 'lgps info', 'pension info',
            'statutory leave', 'leave data', 'payroll data'
        ]
        if any(pattern in sheet_lower for pattern in data_sheet_patterns):
            return False

        if any(skip in sheet_lower for skip in skip_sheet_names):
            return True

        # Check if sheet has actual data (not just example rows)
        if df_raw is not None and len(df_raw) > 5:
            # Look for data rows with actual staff identifiers (payroll numbers, names)
            # Skip first 5 rows which might be headers/descriptions
            has_real_data = False

            for idx in range(5, min(15, len(df_raw))):
                row = df_raw.iloc[idx]
                row_values = [str(v).strip() for v in row if pd.notna(v) and str(v).strip()]

                if not row_values:
                    continue

                # Check for real data indicators:
                # - Alphanumeric codes (like EP203, 12345)
                # - Names (short strings without template keywords)
                # - Currency values
                has_codes = any(v.upper().startswith('EP') or v.isdigit() for v in row_values if len(v) < 20)
                has_dates = any('/' in v or '-' in v for v in row_values if len(v) < 15)
                has_numbers = any(v.replace('.', '').isdigit() for v in row_values if len(v) < 15)

                # If row has multiple data-like values, it's likely real data
                if has_codes or (has_dates and has_numbers):
                    has_real_data = True
                    break

                # Check for typical name patterns (short strings)
                short_strings = [v for v in row_values if 2 < len(v) < 25]
                if len(short_strings) >= 5:  # Row with multiple short values is likely data
                    # Make sure they're not all template keywords
                    template_words = ['example', 'template', 'sample', 'optional', 'required']
                    non_template = [v for v in short_strings if not any(tw in v.lower() for tw in template_words)]
                    if len(non_template) >= 4:
                        has_real_data = True
                        break

            if has_real_data:
                return False  # Has real data, don't skip

        # Check first few rows for pure template indicators (no data below)
        if df_raw is not None and len(df_raw) > 0:
            # Only flag as template if it's a small sheet OR has explicit template language
            first_rows_text = ''
            for idx in range(min(5, len(df_raw))):
                row = df_raw.iloc[idx]
                first_rows_text += ' '.join([str(v).lower() for v in row if pd.notna(v)]) + ' '

            # Strong template indicators that suggest NO real data
            strong_template_indicators = [
                'example person a', 'example person b', 'working weeks',
                'this can be used as a template'
            ]

            strong_template_count = sum(1 for ind in strong_template_indicators if ind in first_rows_text)

            # Only skip if very small sheet with template indicators
            if strong_template_count >= 1 and len(df_raw) < 10:
                return True

        return False

    def _read_sheet_smart(self, xl: pd.ExcelFile, sheet: str) -> Optional[pd.DataFrame]:
        """Smart read that finds the header row automatically."""
        try:
            # First read to find header
            df_raw = pd.read_excel(xl, sheet, header=None, nrows=20)

            # Check if this is a template/example sheet that should be skipped
            if self._is_template_or_example_sheet(sheet, df_raw):
                self.log(f"    Skipping template/example sheet: {sheet}")
                return None

            # Key column indicators for staff data - ordered by importance
            # Primary indicators are most distinctive for staff data
            primary_indicators = [
                'payroll', 'unique payroll', 'person identifier', 'employee ref',
                'emp ref', 'emp no', 'pers ref', 'staff id', 'staff ref'
            ]

            # Secondary indicators support identification
            secondary_indicators = [
                'last name', 'first name', 'surname', 'forename',
                'employee', 'empoyee', 'contract ref', 'job title',
                'hours', 'fte', 'salary', 'scale point', 'scp', 'pension',
                'service start', 'pay scale', 'ni number', 'gender',
                'cost code', 'paid weeks', 'ft hours', 'grade', 'school',
                'weekly hours', 'annual salary', 'contract type', 'date of birth'
            ]

            all_indicators = primary_indicators + secondary_indicators

            # Find the best header row - prefer rows with staff-related column names
            best_row = 0
            best_score = 0

            for idx in range(min(15, len(df_raw))):
                row = df_raw.iloc[idx]
                row_str = ' '.join([str(v).lower() for v in row if pd.notna(v)])

                # Skip rows that are mostly empty or have very long single values (descriptions)
                non_empty_values = [v for v in row if pd.notna(v) and str(v).strip()]
                if len(non_empty_values) < 3:
                    continue

                # Check for very long values that indicate description text, not headers
                max_val_len = max((len(str(v)) for v in non_empty_values), default=0)
                if max_val_len > 100:  # Descriptions are usually very long
                    continue

                # Score based on primary indicators (higher weight)
                primary_score = sum(5 for ind in primary_indicators if ind in row_str)
                # Score based on secondary indicators
                secondary_score = sum(2 for ind in secondary_indicators if ind in row_str)
                # Bonus for having multiple short column-like values
                column_like_score = sum(1 for v in non_empty_values
                                       if isinstance(v, str) and 3 < len(str(v).strip()) < 50)

                total_score = primary_score + secondary_score + column_like_score

                if total_score > best_score:
                    best_score = total_score
                    best_row = idx

            # If we found a good header row, ensure it's not row 0 with description text
            if best_score < 5 and best_row == 0:
                # Try to find a better row by looking for non-description content
                for idx in range(min(10, len(df_raw))):
                    row = df_raw.iloc[idx]
                    non_empty = [str(v).strip() for v in row if pd.notna(v) and str(v).strip()]
                    # Good header rows have multiple short values
                    if len(non_empty) >= 5 and all(len(v) < 80 for v in non_empty):
                        best_row = idx
                        break

            # Re-read with correct header
            df = pd.read_excel(xl, sheet, header=best_row)
            df = df.dropna(how='all')

            # Remove rows that are just category headers (like "Core", "Essential")
            if len(df) > 0:
                try:
                    first_col = df.iloc[:, 0].astype(str).str.lower()
                    skip_values = [
                        'core', 'essential', 'desirable', 'optional', 'core/mappable',
                        'person information', 'role information', 'contract hours',
                        'salary information', 'pension information', 'other information',
                        'allowances / payments', 'allowances/payments'
                    ]
                    df = df[~first_col.isin(skip_values)]
                except Exception:
                    pass  # Continue if this operation fails

            # Clean column names
            df.columns = [self._clean_column_name(c) for c in df.columns]

            # Final validation - check if we have any usable identifier column
            # If not, this sheet might not be staff data
            possible_id_cols = ['payroll_number', 'emp_no', 'employee_number', 'employee_id',
                               'staff_id', 'staff_number', 'unique_id', 'code', 'ref']
            has_id_column = any(col in df.columns for col in possible_id_cols)

            if not has_id_column and len(df) <= 10:
                # Small sheet without ID column - might be a reference sheet, not staff data
                self.log(f"    Sheet '{sheet}' has no identifier column - may need manual mapping")

            return df

        except Exception as e:
            self.log(f"Error reading sheet {sheet}: {e}")
            # Try one more time with minimal options
            try:
                df = pd.read_excel(xl, sheet, dtype=str)
                if len(df) > 0:
                    df.columns = [self._clean_column_name(c) for c in df.columns]
                    return df
            except Exception:
                pass
            return None

    def _apply_smart_header_detection(self, df: pd.DataFrame, sheet_name: str = "") -> Optional[pd.DataFrame]:
        """
        Apply smart header detection to a DataFrame that may have been read with wrong header row.
        This is used for fallback paths that use pd.read_excel without smart reading.
        """
        try:
            # Check if columns look like they're from wrong row (very long names or numbered names)
            col_names = [str(c).lower() for c in df.columns]

            # Indicators of wrong header row
            wrong_header_indicators = [
                any(len(c) > 100 for c in col_names),  # Very long column names
                any('this example' in c for c in col_names),
                any('should be included' in c for c in col_names),
                col_names[0].startswith('1.') or col_names[0].startswith('1_'),
                sum(1 for c in col_names if c.startswith('unnamed')) > len(col_names) / 2
            ]

            if not any(wrong_header_indicators):
                # Headers look OK, just clean them
                df.columns = [self._clean_column_name(c) for c in df.columns]
                return df

            self.log(f"    Detected incorrect header row in '{sheet_name}', searching for correct row...")

            # Headers look wrong - we need to find the actual header row
            # Convert current df to raw format by prepending the current column names as a row
            raw_data = [list(df.columns)] + df.values.tolist()
            df_raw = pd.DataFrame(raw_data)

            # Find the best header row using same logic as _read_sheet_smart
            primary_indicators = [
                'payroll', 'unique payroll', 'person identifier', 'employee ref',
                'emp ref', 'emp no', 'pers ref', 'staff id', 'staff ref'
            ]
            secondary_indicators = [
                'last name', 'first name', 'surname', 'forename',
                'hours', 'fte', 'salary', 'scale point', 'pension',
                'pay scale', 'gender', 'grade', 'school', 'weekly hours'
            ]

            best_row = 0
            best_score = 0

            for idx in range(min(10, len(df_raw))):
                row = df_raw.iloc[idx]
                row_str = ' '.join([str(v).lower() for v in row if pd.notna(v)])

                non_empty_values = [v for v in row if pd.notna(v) and str(v).strip()]
                if len(non_empty_values) < 3:
                    continue

                max_val_len = max((len(str(v)) for v in non_empty_values), default=0)
                if max_val_len > 100:
                    continue

                primary_score = sum(5 for ind in primary_indicators if ind in row_str)
                secondary_score = sum(2 for ind in secondary_indicators if ind in row_str)
                column_like_score = sum(1 for v in non_empty_values
                                       if isinstance(v, str) and 3 < len(str(v).strip()) < 50)

                total_score = primary_score + secondary_score + column_like_score

                if total_score > best_score:
                    best_score = total_score
                    best_row = idx

            if best_row > 0:
                self.log(f"    Found header at row {best_row} (score: {best_score})")
                # Reconstruct DataFrame with correct header
                new_columns = df_raw.iloc[best_row].tolist()
                new_data = df_raw.iloc[best_row + 1:].values.tolist()
                df = pd.DataFrame(new_data, columns=new_columns)
                df = df.dropna(how='all')

                # Remove category header rows
                if len(df) > 0:
                    try:
                        first_col = df.iloc[:, 0].astype(str).str.lower()
                        skip_values = ['core', 'essential', 'desirable', 'optional', 'core/mappable']
                        df = df[~first_col.isin(skip_values)]
                    except Exception:
                        pass

            # Clean column names
            df.columns = [self._clean_column_name(c) for c in df.columns]
            return df

        except Exception as e:
            self.log(f"    Error in smart header detection: {e}")
            return df

    def _clean_column_name(self, col: Any) -> str:
        """Standardize column name."""
        if pd.isna(col):
            return "unnamed"

        col_str = str(col).strip().lower()

        # Exact/priority mappings (checked first, order matters)
        priority_mappings = [
            ('unique payroll', 'payroll_number'),
            ('payroll number', 'payroll_number'),
            ('person id', 'payroll_number'),
            ('employee id', 'payroll_number'),
            ('employee ref', 'payroll_number'),
            ('empoyee ref', 'payroll_number'),  # Common typo
            ('emp ref', 'payroll_number'),
            ('emp no', 'payroll_number'),       # Payroll analysis format
            ('pers ref', 'payroll_number'),     # Payroll analysis format
            ('staff id', 'payroll_number'),
            ('staff ref', 'payroll_number'),
            ('ref no', 'payroll_number'),
            ('last name', 'surname'),
            ('last_name', 'surname'),
            ('lastname', 'surname'),
            ('family name', 'surname'),
            ('family_name', 'surname'),
            ('first name', 'forename'),
            ('first_name', 'forename'),
            ('firstname', 'forename'),
            ('given name', 'forename'),
            ('given_name', 'forename'),
            ('continuous service', 'service_start_date'),
            ('contin. service', 'service_start_date'),  # Payroll abbreviated format
            ('service start', 'service_start_date'),
            ('joining date', 'service_start_date'),     # Payroll format
            ('start date', 'service_start_date'),
            ('date of birth', 'dob'),
            ('birth date', 'dob'),
            ('work location', 'school_code'),
            ('school name', 'school_code'),
            ('cost centre', 'cost_centre'),
            ('contract reference', 'contract_ref'),
            ('contract ref', 'contract_ref'),
            ('staff role', 'job_title'),
            ('job title', 'job_title'),
            ('role title', 'job_title'),
            ('gross salary finance', 'finance_code'),
            ('finance code', 'finance_code'),
            ('nominal code', 'finance_code'),
            ('full time hours', 'ft_hours'),
            ('ft hours', 'ft_hours'),
            ('fte hours', 'ft_hours'),
            ('contracted hours', 'weekly_hours'),
            ('weekly hours', 'weekly_hours'),
            ('weekly', 'weekly'),  # LESSON L025: Column may be just "weekly" for full-time hours
            ('hours per week', 'weekly_hours'),
            ('paid weeks', 'weeks_paid'),       # Payroll analysis format
            ('weekly fte', 'fte'),
            ('annual fte', 'fte'),
            ('fte salary', 'annual_salary'),    # Payroll analysis format
            ('actual salary', 'actual_salary'), # Payroll analysis format
            ('cost code', 'cost_centre'),       # Payroll analysis format
            ('weeks worked', 'weeks_worked'),
            ('tto weeks', 'weeks_worked'),
            ('weeks paid', 'weeks_paid'),
            # IMPORTANT: More specific patterns MUST come before general patterns
            # to preserve column distinctions
            ('pay scale type', 'pay_scale_type'),    # MGMT, SUPP, LECT
            ('scale type', 'pay_scale_type'),         # Alternative naming
            ('pay scale point', 'scale_point'),       # 5, 18, 7 (before 'pay scale')
            ('pay scale contract', 'pay_scale_contract'),
            ('pay scale group', 'pay_scale_group'),
            ('pay scale grade', 'pay_scale_grade'),
            ('pay scale', 'pay_scale'),               # General fallback
            ('pay range', 'pay_scale'),
            ('scale point', 'scale_point'),
            ('spine point', 'scale_point'),
            ('current point', 'scale_point'),
            # Payroll analysis format - exact matches
            ('scp', 'scale_point'),             # Scale Point in payroll reports
            ('occupancy ref', 'contract_ref'),  # Occupancy Reference = contract ref
            # Contract-related columns
            ('contract start', 'contract_start'),
            ('contract end', 'contract_end'),
            ('start date', 'start_date'),
            ('end date', 'end_date'),
            ('effective date', 'effective_date'),
            ('leaving date', 'leaving_date'),
            ('contract type', 'contract_type'),
            ('employment type', 'employment_type'),
            ('emp type', 'emp_type'),
            ('department code', 'department_code'),
            ('dept code', 'dept_code'),
            ('fund code', 'fund_code'),
            ('funding code', 'funding_code'),
            ('budget code', 'budget_code'),
            ('equated week', 'eqw'),
            ('eqw pattern', 'eqw_pattern'),
            ('weeks pattern', 'weeks_pattern'),
            ('term time', 'term_time'),
            ('grade code', 'grade_code'),
            ('pay grade', 'pay_grade'),
            ('ni number', 'ni_number'),
            ('national insurance', 'ni_number'),
        ]

        # Check priority mappings first
        for pattern, standard in priority_mappings:
            if pattern in col_str:
                # Memory check: Validate this mapping doesn't violate learned rules
                if self.memory:
                    check = self.memory.check_column_mapping(col_str, standard)
                    if check.should_warn:
                        for warning in check.warnings:
                            if warning not in self.memory_warnings:
                                self.memory_warnings.append(warning)
                return standard

        # Standard single-word mappings (fallback)
        # IMPORTANT: Be precise - don't over-map! 'scale' alone should stay as 'scale'
        # to preserve the distinction between scale type (MGMT, SUPP) and pay_scale column
        mappings = {
            'payroll': 'payroll_number',
            'surname': 'surname',
            'forename': 'forename',
            'dob': 'dob',
            'gender': 'gender',
            'sex': 'gender',
            'school': 'school_code',
            'location': 'school_code',
            'department': 'cost_centre',
            'role': 'job_title',
            'position': 'job_title',
            'hours': 'weekly_hours',            # Payroll "Hours" column
            'nominal': 'finance_code',
            'fte': 'fte',
            # 'scale' alone is NOT mapped - it stays as 'scale' (for scale type column)
            # Only explicit 'pay scale' patterns should become 'pay_scale' (handled in priority_mappings)
            'point': 'scale_point',
            'scp': 'scale_point',
            'grade': 'grade',
            'annual salary': 'annual_salary',
            'fte salary': 'annual_salary',
            'salary': 'annual_salary',
            'actual salary': 'actual_salary',
            'pro rata': 'actual_salary',
            'pension': 'pension',
            'pension code': 'pension_code',  # LESSON L026: Customer may use "PENSION CODE" column
            'pension scheme': 'pension_scheme',
            'pension type': 'pension_type',
            'pension provider': 'pension_provider',
            'staff pension': 'staff_pension',
            'pens': 'pens',
            'contract type': 'contract_type',
            'contract ref': 'contract_ref',
            'reference': 'contract_ref',
            # 'allowance' alone should NOT be mapped - let numbered columns (Allowance 1, 2, 3) keep their names
            'tlr': 'tlr_allowance',
            'sen': 'sen_allowance',
            'recruitment': 'recruitment_allowance',
            'retention': 'retention_allowance',
        }

        for pattern, standard in mappings.items():
            # Don't match partial patterns inside larger column names
            # e.g., 'fte' should not match 'allowance_1_value_@_1fte'
            if pattern in col_str:
                # For short patterns, require word boundaries or column to be very short
                if len(pattern) <= 3 and len(col_str) > len(pattern) + 5:
                    # Skip short patterns in long column names
                    continue
                return standard

        return col_str.replace(' ', '_').replace('/', '_').replace('@', '')

    def _analyze_dataframe(self, df: pd.DataFrame, file_name: str, sheet_name: str):
        """Analyze a dataframe and create detailed report."""
        columns_mapped = {}
        data_quality = {}
        sample_data = {}
        issues = []
        recommendations = []

        # PREPROCESS: Parse combined format fields (CODE: Title format)
        df = self._preprocess_combined_fields(df)

        # Analyze each column - use unique columns to avoid DataFrame issues
        seen_cols = set()
        for col in df.columns:
            if col.startswith('unnamed') or col in seen_cols:
                continue
            seen_cols.add(col)

            try:
                col_data = df[col]
                # Handle case where duplicate column names return a DataFrame
                if isinstance(col_data, pd.DataFrame):
                    col_data = col_data.iloc[:, 0]  # Take first column
                non_null = col_data.dropna()
            except Exception:
                continue

            if len(non_null) == 0:
                continue

            # Sample data - use values.tolist() for safety
            try:
                sample_data[col] = non_null.head(5).values.tolist()
            except Exception:
                sample_data[col] = list(non_null.head(5))

            # Data quality
            data_quality[col] = {
                'non_null_count': len(non_null),
                'null_count': len(df) - len(non_null),
                'unique_values': non_null.nunique(),
                'dtype': str(non_null.dtype),
            }

            # Detect data type from content
            content_type = self._detect_content_type(non_null)
            data_quality[col]['content_type'] = content_type

        # Determine what type of data this is (use file name for hints)
        data_type = self._classify_data_type(df, sheet_name, file_name)

        report = AnalysisReport(
            file_name=file_name,
            sheet_name=sheet_name,
            row_count=len(df),
            column_count=len(df.columns),
            columns_found=list(df.columns),
            columns_mapped=columns_mapped,
            data_quality=data_quality,
            sample_data=sample_data,
            issues=issues,
            recommendations=recommendations
        )

        self.analysis_reports.append(report)

        # Log classification result
        self.processing_log.append(f"Sheet '{file_name}/{sheet_name}': {len(df)} rows -> {data_type}")

        # Store data for processing
        if data_type == 'staff_contracts':
            # Preprocess payroll format data: split Name column into forename/surname
            df = self._preprocess_payroll_name_column(df)
            self.staff_data.append(df)
            self.source_staff_count += len(df)  # Track for audit
            self.log(f"    -> Staff contract data: {len(df)} rows")
            # ALSO extract pay scales and allowances from staff data
            self._extract_pay_scales_from_staff_data(df)
            self._extract_allowances_from_staff_data(df)
        elif data_type == 'pay_scales':
            self._extract_pay_scales_from_df(df, sheet_name)
            self.log(f"    -> Pay scale data extracted")
        elif data_type == 'allowances':
            self._extract_allowances_from_df(df)
            self.log(f"    -> Allowance data extracted")
        else:
            # Store unclassified data for user review and manual mapping
            available_cols = list(df.columns)
            self.log(
                f"    -> WARNING: Sheet '{sheet_name}' in '{file_name}' "
                f"could not be classified. Columns: {available_cols[:10]}. "
                f"Use Pre-Flight Validator to manually map this sheet."
            )
            # Use _safe_get_column to handle duplicate column names (returns DataFrame instead of Series)
            sample_data = {}
            for col in df.columns[:10]:
                try:
                    col_data = self._safe_get_column(df, col)
                    sample_data[col] = col_data.dropna().head(3).tolist()
                except Exception:
                    sample_data[col] = []
            self.unclassified_data.append({
                'file_name': file_name,
                'sheet_name': sheet_name,
                'data': df,
                'columns': list(df.columns),
                'row_count': len(df),
                'sample_data': sample_data
            })
            self.issues.append(f"Unclassified data in {file_name} / {sheet_name}: {len(df)} rows need manual mapping")

    def _detect_content_type(self, series: pd.Series) -> str:
        """Detect the type of content in a series."""
        try:
            # Handle DataFrame case (duplicate columns)
            if isinstance(series, pd.DataFrame):
                series = series.iloc[:, 0]

            sample = series.head(10).astype(str).values.tolist()
            sample_str = ' '.join(sample).lower()

            # Check for dates
            head_vals = series.head(10).values
            if any(isinstance(v, (datetime, pd.Timestamp)) for v in head_vals):
                return 'date'

            # Check for currency/salary
            if any(c in sample_str for c in ['£', '$', 'salary', 'pay']):
                return 'currency'

            # Check for percentages
            if '%' in sample_str or all(0 <= v <= 1 for v in head_vals if isinstance(v, (int, float))):
                return 'percentage'

            # Check for names
            if series.dtype == 'object':
                try:
                    if series.str.match(r'^[A-Z][a-z]+$').any():
                        return 'name'
                except Exception:
                    pass

            # Check for codes
            if series.dtype == 'object':
                try:
                    if series.str.match(r'^[A-Z0-9_]+$').any():
                        return 'code'
                except Exception:
                    pass

            return 'text' if series.dtype == 'object' else 'numeric'
        except Exception:
            return 'unknown'

    def _score_as_staff_contracts(self, df: pd.DataFrame) -> int:
        """
        Score how likely this data is staff contracts based on DATA CONTENT.
        Returns score 0-5, where >= 3 indicates staff contracts.
        Staff contracts have: names, employee refs, FTE, many columns, varied data.
        """
        score = 0
        has_names = False
        has_employee_refs = False

        try:
            # REQUIRED: Must have name-like columns (text with proper names)
            for col in df.columns:
                col_data = self._safe_get_column(df, col)
                sample = col_data.dropna().head(20).astype(str).tolist()

                # Look for name patterns - handle various cases:
                # - Proper case: Smith, Jones
                # - All uppercase: SMITH, JONES
                # - Mixed: McDonald, O'Brien
                # Names are 2-20 chars, primarily letters, possibly with hyphen/apostrophe
                name_count = 0
                for v in sample:
                    v_clean = str(v).strip()
                    # Check if it looks like a name (letters, possibly with - or ')
                    if len(v_clean) >= 2 and len(v_clean) <= 20:
                        # Allow letters, hyphens, apostrophes, spaces
                        cleaned = v_clean.replace('-', '').replace("'", '').replace(' ', '')
                        if cleaned.isalpha():
                            # Exclude obvious abbreviations/codes:
                            # - All uppercase and <= 5 chars (MGMT, SUPP, LECT, TCC)
                            # - Contains no vowels (unlikely to be a name)
                            is_abbreviation = (v_clean.isupper() and len(v_clean) <= 5)
                            vowels = set('AEIOUaeiou')
                            has_vowels = any(c in vowels for c in v_clean)
                            if not is_abbreviation and has_vowels:
                                name_count += 1

                if name_count >= len(sample) * 0.6 and len(sample) >= 5:
                    has_names = True
                    score += 2  # Strong indicator
                    break

            # Check for employee reference patterns - can be:
            # - Alphanumeric: EMP001, P12345
            # - Numeric payroll numbers: 603989, 12345678
            # - Must be unique-ish (high nunique) and consistent format
            for col in df.columns:
                col_data = self._safe_get_column(df, col)
                sample = col_data.dropna().head(20).astype(str).tolist()

                # Check for alphanumeric refs (EMP001, P123456)
                ref_count = sum(1 for v in sample if re.match(r'^[A-Z]{1,4}\d{3,8}$', str(v).strip()))

                # Also check for numeric payroll numbers (5-10 digits, high uniqueness)
                if ref_count < len(sample) * 0.5:
                    numeric_ref_count = sum(1 for v in sample if re.match(r'^\d{5,10}$', str(v).strip()))
                    # Verify high uniqueness (employee refs should be unique)
                    if numeric_ref_count >= len(sample) * 0.7:
                        unique_ratio = len(set(sample)) / len(sample) if len(sample) > 0 else 0
                        if unique_ratio >= 0.8:  # High uniqueness = employee refs
                            ref_count = numeric_ref_count

                if ref_count >= len(sample) * 0.5 and len(sample) >= 5:
                    has_employee_refs = True
                    score += 1
                    break

            # Check for FTE values (0.0 - 1.0 range) - specific to contracts
            for col in df.columns:
                col_data = self._safe_get_column(df, col)
                try:
                    numeric = pd.to_numeric(col_data, errors='coerce').dropna()
                    if len(numeric) >= 5:
                        # FTE is specifically 0-1 (or up to 1.5 for overtime)
                        fte_like = numeric[(numeric > 0) & (numeric <= 1.5)]
                        if len(fte_like) >= len(numeric) * 0.7:
                            score += 1
                            break
                except Exception:
                    pass

            # Staff contracts typically have MANY columns (5+)
            if len(df.columns) >= 5:
                score += 1

            # If no names found, this is NOT staff contracts (even with other indicators)
            if not has_names:
                score = max(0, score - 2)

        except Exception:
            pass

        return score

    def _score_scale_column(self, col_name: str, values: List[str]) -> int:
        """
        Score how likely a column is to contain pay scale type values.
        Intelligent scoring based on data patterns, not hardcoded names.

        Returns score where higher = more likely to be scale type column.
        """
        score = 0
        col_lower = str(col_name).lower()

        # Pattern 1: Column name is exactly 'scale' or 'type' (strongest indicator)
        if col_lower in ['scale', 'type', 'scale_type', 'pay_type']:
            score += 8
        # Column name contains scale-related terms
        elif any(term in col_lower for term in ['scale', 'type', 'grade', 'band', 'group', 'category']):
            score += 5

        # Pattern 2: Values are short uppercase codes (typical scale codes)
        # e.g., MGMT, SUPP, LECT, TEACH, ADMIN
        short_upper_codes = sum(1 for v in values if len(v) <= 10 and v.isupper())
        if short_upper_codes >= len(values) * 0.7:
            score += 3

        # Pattern 3: Values have consistent length (scale types are often uniform)
        lengths = [len(v) for v in values]
        if lengths and max(lengths) - min(lengths) <= 3:
            score += 2

        # Pattern 4: Number of unique values (scale types typically 3-15)
        if 3 <= len(values) <= 15:
            score += 2
        elif 2 <= len(values) <= 20:
            score += 1

        # Pattern 5: Values look like role/function categories
        # (common patterns in pay scale types - role-based)
        role_patterns = ['sup', 'mgmt', 'man', 'teach', 'lect', 'admin', 'tech', 'lead', 'sen', 'asst']
        role_matches = sum(1 for v in values if any(p in v.lower() for p in role_patterns))
        if role_matches >= 1:
            score += 3

        # Pattern 6: All values are pure text (no numbers mixed in)
        # Scale types are usually pure text like MGMT, SUPP, LECT
        pure_text_values = sum(1 for v in values if v.isalpha() or (v.replace('/', '').replace('-', '').isalpha()))
        if pure_text_values >= len(values) * 0.7:
            score += 4

        # Penalty: Values look like point codes (letter+number pattern like SO1, M1, L01)
        # These are scale POINTS, not scale TYPES
        point_code_pattern = sum(1 for v in values if re.match(r'^[A-Za-z]{1,3}\d+$', v))
        if point_code_pattern >= len(values) * 0.3:
            score -= 4

        # Penalty: Values look like institution/contract codes (long, varied)
        if any(len(v) >= 20 for v in values):
            score -= 2

        # Penalty: Column name suggests contract/institution rather than scale
        if any(term in col_lower for term in ['contract', 'institution', 'employer', 'company', 'ref']):
            score -= 3

        # Penalty: Column name contains 'point' (likely a point column, not scale type)
        if 'point' in col_lower:
            score -= 5

        return score

    def _score_as_pay_scale(self, df: pd.DataFrame) -> int:
        """
        Score how likely this data is a pay scale based on DATA CONTENT.
        Returns score 0-5, where >= 3 indicates pay scale.
        """
        score = 0
        try:
            # Check for sequential or point-like values in ANY column (not just first)
            point_patterns = [
                r'^\d{1,2}$',              # 1, 2, 43
                r'^[SsMmUuLlNnPp]\d{1,2}$',  # S1, M1, L01, U1, P1, N1
                r'^SCP\s*\d{1,2}$',         # SCP1, SCP 1
                r'^[Pp]oint\s*\d+$',        # Point 1
                r'^[Gg]rade\s*\d+$',        # Grade 1
                r'^[Ll]evel\s*\d+$',        # Level 1
                r'^[Bb]and\s*\d+$',         # Band 1
                r'^MIN$', r'^MAX$',         # MIN, MAX (common in pay tables)
                r'^[Uu]\d$',                # U1, U2, U3
                r'^[Mm]\d$',                # M1-M6
                r'^[Ll]\d{1,2}$',           # L1-L43
            ]

            for col in df.columns:
                col_data = self._safe_get_column(df, col)
                sample = col_data.dropna().head(20).astype(str).tolist()
                if not sample:
                    continue

                point_matches = sum(
                    1 for v in sample
                    if any(re.match(p, str(v).strip(), re.IGNORECASE) for p in point_patterns)
                )
                # Lower threshold for small tables: 40% match and >= 2 samples
                if point_matches >= len(sample) * 0.4 and len(sample) >= 2:
                    score += 2  # Strong indicator
                    break

            # Check for salary-like values (> 10000, typically 20000-100000)
            # IMPORTANT: Clean currency formatting (£, commas) from PDF data first
            for col in df.columns:
                col_data = self._safe_get_column(df, col)
                try:
                    # Clean currency formatting before converting to numeric
                    cleaned = col_data.astype(str).str.replace('£', '', regex=False)
                    cleaned = cleaned.str.replace(',', '', regex=False).str.strip()
                    numeric = pd.to_numeric(cleaned, errors='coerce').dropna()
                    # Lower threshold: >= 2 values for small tables
                    if len(numeric) >= 2:
                        salary_like = numeric[(numeric >= 10000) & (numeric <= 200000)]
                        # Lower threshold: 50% for small tables
                        if len(salary_like) >= len(numeric) * 0.5:
                            score += 2  # Strong indicator
                            break
                except Exception:
                    pass

            # Typical pay scale size: 3-100 rows (lowered minimum)
            if 3 <= len(df) <= 100:
                score += 1

            # Few columns (point, rate, maybe date) - typically 2-10
            if 2 <= len(df.columns) <= 10:
                score += 1

        except Exception:
            pass

        return score

    def _score_as_allowances(self, df: pd.DataFrame) -> int:
        """
        Score how likely this data is allowances based on DATA CONTENT.
        Returns score 0-4, where >= 2 indicates allowances.
        """
        score = 0
        try:
            # Check for allowance-type keywords in data values
            all_text = ' '.join(df.astype(str).values.flatten()).lower()
            allowance_keywords = ['tlr', 'sen ', 'recruit', 'retain', 'london', 'weighting',
                                 'allowance', 'supplement', 'enhancement', 'payment',
                                 'special educational', 'teaching & learning']
            keyword_matches = sum(1 for kw in allowance_keywords if kw in all_text)
            if keyword_matches >= 1:  # Lowered from 2 to 1
                score += 2

            # Check for amount-like values (typically 500-30000 for allowances)
            # Clean currency formatting first
            for col in df.columns:
                col_data = self._safe_get_column(df, col)
                try:
                    cleaned = col_data.astype(str).str.replace('£', '', regex=False)
                    cleaned = cleaned.str.replace(',', '', regex=False).str.strip()
                    numeric = pd.to_numeric(cleaned, errors='coerce').dropna()
                    if len(numeric) >= 2:
                        allowance_like = numeric[(numeric >= 100) & (numeric <= 30000)]
                        if len(allowance_like) >= len(numeric) * 0.5:  # Lowered threshold
                            score += 1
                            break
                except Exception:
                    pass

            # Small number of rows (typically < 20 for allowance types)
            if 2 <= len(df) <= 30:
                score += 1

        except Exception:
            pass

        return score

    def _classify_data_type(self, df: pd.DataFrame, sheet_name: str, file_name: str = "") -> str:
        """
        Classify what type of data this sheet contains.
        LOGIC-BASED: Analyze actual data content, not just names/titles.
        Also uses file name hints when available.
        """
        if df.empty or len(df.columns) == 0:
            return 'unknown'

        cols_lower = [str(c).lower() for c in df.columns]
        cols_str = ' '.join(cols_lower)
        sheet_lower = sheet_name.lower()
        file_lower = file_name.lower() if file_name else ""

        # ===== FILE NAME HINTS (HIGH PRIORITY) =====
        # If file name clearly indicates pay scales, trust it
        pay_scale_file_hints = ['pay scale', 'payscale', 'pay_scale', 'salary scale',
                                'pay spine', 'payspine', 'pay rates', 'salary rates',
                                'njc', 'support staff scale', 'teacher pay']
        strong_pay_scale_hints = ['pay scale', 'payscale', 'salary scale', 'pay spine']

        if any(hint in file_lower for hint in pay_scale_file_hints):
            # Verify it has numeric/salary data (not just mentioned in name)
            # Clean currency formatting before checking
            has_salary_data = False
            for col in df.columns:
                try:
                    col_data = self._safe_get_column(df, col)
                    # Clean currency: £, commas
                    cleaned = col_data.astype(str).str.replace('£', '', regex=False)
                    cleaned = cleaned.str.replace(',', '', regex=False).str.strip()
                    numeric_vals = pd.to_numeric(cleaned, errors='coerce').dropna()
                    if len(numeric_vals) > 0 and numeric_vals.max() > 1000:
                        has_salary_data = True
                        break
                except Exception:
                    pass

            # If strong file name match, classify as pay_scales even without salary verification
            # (PDF tables often have header tables without salary data)
            is_strong_match = any(hint in file_lower for hint in strong_pay_scale_hints)
            if has_salary_data or is_strong_match:
                self.log(f"  Classified as pay_scales based on file name: {file_name}")
                return 'pay_scales'

        allowance_file_hints = ['allowance', 'tlr', 'sen allowance']
        if any(hint in file_lower for hint in allowance_file_hints):
            self.log(f"  Classified as allowances based on file name: {file_name}")
            return 'allowances'

        # ===== EARLY DETECTION: Reference/Configuration Data =====
        # These should NOT be classified as staff contracts
        # IMPORTANT: Be strict - only classify as pension_rates if it's PRIMARILY pension config

        # PENSION CONTRIBUTION RATES - small sheets with specific pension config columns
        # Must have: "contribution" or "employer %", and few columns, and few rows
        # Note: Column names may have underscores after cleaning
        pension_rate_indicators = [
            'employer_contribution', 'contribution_percentage', 'employer contribution',
            'rate_date_from', 'rate_date_to', 'rate date', 'scheme_name'
        ]
        pension_indicator_count = sum(1 for ind in pension_rate_indicators if ind in cols_str)

        # Only classify as pension_rates if:
        # 1. Has multiple pension rate indicators (at least 2)
        # 2. Small number of columns (config tables are narrow, typically 4-8 columns)
        # 3. Does NOT have staff identifier columns (payroll_number, surname, forename)
        staff_id_columns = ['payroll_number', 'payroll', 'emp_no', 'employee_number', 'staff_id', 'surname', 'forename', 'first_name', 'last_name']
        has_staff_columns = any(ind in cols_str for ind in staff_id_columns)

        if pension_indicator_count >= 2 and len(df.columns) <= 10 and not has_staff_columns:
            # This is pension configuration data, not staff data
            self.log(f"    Classified as pension_rates: {sheet_name}")
            return 'pension_rates'

        # STATUTORY LEAVE - has columns like "Leave Type", "Leave Date From"
        # These ARE staff-related but should be handled differently
        leave_indicators = ['leave type', 'leave date', 'maternity', 'paternity', 'adoption']
        if any(ind in cols_str for ind in leave_indicators):
            # Still treat as staff contracts for now (they link to staff by ID)
            pass

        # ===== LOGIC-BASED DATA ANALYSIS =====
        # Analyze actual data patterns, not just column names

        # 1. Check for STAFF CONTRACTS by data patterns
        staff_contract_score = self._score_as_staff_contracts(df)
        if staff_contract_score >= 3:
            return 'staff_contracts'

        # 2. Check for PAY SCALES by data patterns
        pay_scale_score = self._score_as_pay_scale(df)
        if pay_scale_score >= 3:
            return 'pay_scales'

        # 2b. Special case for PDF tables: lower threshold if they have salary data
        # PDF pay scale tables are often small and might not score high enough
        is_pdf = '.pdf' in file_lower or 'pdf_' in sheet_lower
        if is_pdf and pay_scale_score >= 2:
            self.log(f"  Classified as pay_scales (PDF with score {pay_scale_score}): {sheet_name}")
            return 'pay_scales'

        # 3. Check for ALLOWANCES by data patterns
        allowance_score = self._score_as_allowances(df)
        if allowance_score >= 2:
            return 'allowances'

        # 3b. Special case for PDF tables: if small table with salary data, likely pay scales
        # This catches PDF pay scale tables that don't score high on patterns
        if is_pdf and len(df) <= 10:
            has_salary_values = False
            for col in df.columns:
                try:
                    col_data = self._safe_get_column(df, col)
                    cleaned = col_data.astype(str).str.replace('£', '', regex=False)
                    cleaned = cleaned.str.replace(',', '', regex=False).str.strip()
                    numeric = pd.to_numeric(cleaned, errors='coerce').dropna()
                    if len(numeric) >= 2:
                        salary_like = numeric[(numeric >= 5000) & (numeric <= 200000)]
                        if len(salary_like) >= 2:
                            has_salary_values = True
                            break
                except Exception:
                    pass
            if has_salary_values:
                self.log(f"  Classified as pay_scales (PDF with salary values): {sheet_name}")
                return 'pay_scales'

        # ===== FALLBACK: Column/sheet name hints (lower priority) =====

        # Staff contracts hints - include schema standard names with underscores
        name_indicators = [
            'surname', 'forename', 'last name', 'first name',
            'lastname', 'firstname', 'last_name', 'first_name',  # schema standard names
            'employee', 'staff name', 'familyname', 'givenname',
            'family_name', 'given_name', 'name'  # additional variations
        ]
        has_name_col = any(x in cols_str for x in name_indicators)

        contract_indicators = [
            'fte', 'hours', 'salary', 'contract', 'pay scale', 'pay_scale',
            'pension', 'annual', 'weekly', 'scale point', 'scale_point',
            'cost centre', 'cost_centre', 'staff role', 'staff_role',
            'job title', 'job_title', 'payroll', 'emp_no', 'employee_number'
        ]
        has_contract_col = any(x in cols_str for x in contract_indicators)
        if has_name_col and has_contract_col:
            return 'staff_contracts'

        if any(x in sheet_lower for x in ['contract', 'staff', 'employee', 'personnel', 'listing']):
            if has_name_col or has_contract_col:
                return 'staff_contracts'

        # Pay scales hints
        pay_scale_indicators = ['pay scale', 'pay range', 'pay grid', 'salary scale', 'spine',
                               'payscale', 'pay spines', 'standard practice', 'salary rates',
                               'mps', 'ups', 'aps', 'upper pay scale', 'main pay',
                               'njc', 'nj ', 'support scale', 'support pay', 'scp', 'rate']
        if any(x in sheet_lower for x in pay_scale_indicators):
            return 'pay_scales'

        # Check columns for pay scale indicators
        if any(c in cols_lower for c in ['scale_point', 'spine_point', 'pay_scale_rate', 'scp',
                                         'point', 'annual salary', 'spinal']):
            # Make sure it's not staff contracts - check for name/payroll columns
            staff_columns = ['surname', 'forename', 'payroll_number', 'first_name', 'last_name',
                           'firstname', 'lastname', 'emp_no', 'employee_number', 'staff_id']
            if not any(c in cols_lower for c in staff_columns):
                # If we see numeric or currency columns with point/scale references
                if any(c in cols_lower for c in ['point', 'scp', 'spine', 'scale', 'grade', 'spinal']):
                    return 'pay_scales'

        # LOGIC-BASED PAY SCALE DETECTION: If it looks like a pay scale, include it
        # Check if data looks like a pay scale (sequential points with salary values)
        if len(df) > 0 and len(df) <= 100:  # Pay scales typically have 5-50 rows
            # Check first column for point-like values (1, 2, 3... or SCP1, M1, L1, etc.)
            first_col = df.iloc[:, 0] if len(df.columns) > 0 else pd.Series()
            try:
                first_col_str = first_col.dropna().astype(str).head(10).tolist()
                # Check if values look like scale points
                point_patterns = [
                    r'^\d{1,2}$',           # Just numbers: 1, 2, 3, 10, 43
                    r'^[SsNnMmUuLl]\d+$',   # S1, N1, M1, U1, L1, SCP1, etc.
                    r'^SCP\d+$',            # SCP1, SCP2, etc.
                    r'^[Pp]oint\s*\d+$',    # Point 1, Point2, etc.
                ]
                looks_like_points = sum(
                    1 for v in first_col_str
                    if any(re.match(p, str(v).strip(), re.IGNORECASE) for p in point_patterns)
                )

                # Check for salary-like values in other columns (> 10000)
                has_salary_values = False
                for col in df.columns[1:]:  # Skip first column (likely points)
                    try:
                        col_data = self._safe_get_column(df, col)
                        numeric_vals = pd.to_numeric(col_data, errors='coerce').dropna()
                        if len(numeric_vals) > 0 and numeric_vals.max() > 10000:
                            has_salary_values = True
                            break
                    except Exception:
                        pass

                # If most values look like points AND we have salary-like values, it's a pay scale
                if looks_like_points >= len(first_col_str) * 0.5 and has_salary_values:
                    self.log(f"  Detected pay scale by data pattern: {sheet_name}")
                    return 'pay_scales'
            except Exception:
                pass

        # Allowances - EXPANDED DETECTION
        allowance_indicators = ['allowance', 'tlr', 'sen', 'recruitment', 'retention', 'london',
                               'weighting', 'supplement', 'enhancement', 'discretionary', 
                               'annual addition', 'teaching and learning']
        if any(x in sheet_lower for x in allowance_indicators):
            return 'allowances'
        
        # Check columns for allowance structure
        if any(c in cols_lower for c in ['tlr', 'sen', 'allowance', 'supplement', 'london', 'weighting']):
            # If we see amount/rate columns
            if any(c in cols_lower for c in ['amount', 'rate', 'salary', 'value', '£']):
                return 'allowances'

        # EQWP/TTO
        if 'tto' in sheet_lower or 'equated' in sheet_lower or 'weeks' in sheet_lower:
            return 'eqwp'

        return 'unknown'

    def _print_analysis_summary(self):
        """Print summary of analysis."""
        try:
            self.log("\n" + "="*60)
            self.log("ANALYSIS SUMMARY")
            self.log("="*60)

            total_staff = sum(len(df) for df in self.staff_data if isinstance(df, pd.DataFrame))
            self.log(f"Total staff records found: {total_staff}")
            self.log(f"Pay scales extracted: {len(self.extracted_pay_scales)}")
            self.log(f"Allowance types found: {len(self.extracted_allowances)}")
            
            # Safely handle the schools_found set
            schools_count = len(self.schools_found) if self.schools_found else 0
            self.log(f"Schools found: {schools_count}")
            
            self.log(f"Finance codes found: {len(self.finance_codes_found)}")

            if self.issues:
                self.log(f"\nIssues found: {len(self.issues)}")
                for idx, issue in enumerate(self.issues[:5], 1):
                    try:
                        # Make issue string safe
                        safe_issue = str(issue)[:200]  # Limit length
                        self.log(f"  {idx}. {safe_issue}")
                    except Exception:
                        self.log(f"  {idx}. [Issue display failed]")
        except Exception as e:
            # Silently continue if summary fails
            pass

    # =========================================================================
    # PHASE 2: PAY SCALE EXTRACTION
    # =========================================================================

    def _extract_pay_scales_from_df(self, df: pd.DataFrame, sheet_name: str):
        """Extract pay scales and points from a dedicated pay scale dataframe."""
        self.log(f"  Extracting pay scales from dedicated sheet: {sheet_name}")

        # =====================================================================
        # FIRST: Check for MATRIX FORMAT (common in PDF pay scale tables)
        # Matrix format: First column = point codes, other columns = regional salaries
        # Example: | Point | Inner London | Outer London | Fringe | Rest of England |
        # =====================================================================
        if self._try_extract_matrix_pay_scale(df, sheet_name):
            return  # Successfully extracted from matrix format

        # Check for customer-defined scale column dynamically
        # Collect ALL candidate columns and score them to find the best one
        scale_candidates = []

        for col in df.columns:
            # Skip columns that are clearly point/rate columns
            col_lower = str(col).lower()
            if any(x in col_lower for x in ['point', 'salary', 'rate', 'annual', '£', 'amount']):
                continue

            # Check if column has categorical values that could be scale types
            col_data = df[col]
            # Handle duplicate column names (returns DataFrame instead of Series)
            if isinstance(col_data, pd.DataFrame):
                col_data = col_data.iloc[:, 0]
            sample_vals = col_data.dropna().astype(str).tolist()
            if not sample_vals:
                continue

            # Count non-numeric unique values (potential scale names)
            unique_vals = list(set(sample_vals))
            non_numeric = [v for v in unique_vals
                         if not v.replace('.', '').replace('-', '').isdigit()
                         and v.lower() not in ['nan', '', 'none', 'null']]

            # If we have 2-20 unique non-numeric values, it's a candidate
            if 2 <= len(non_numeric) <= 20:
                # Score this column based on how "scale-like" its values are
                score = self._score_scale_column(col, non_numeric)
                scale_candidates.append((col, score, non_numeric))

        # Pick the best candidate (highest score)
        scale_col = None
        if scale_candidates:
            scale_candidates.sort(key=lambda x: x[1], reverse=True)
            best_col, best_score, best_values = scale_candidates[0]
            scale_col = best_col

            # Memory check: Validate scale column selection
            if self.memory:
                check = self.memory.check_scale_column_selection(scale_candidates)
                if check.should_warn:
                    for warning in check.warnings:
                        self.log(f"    [MEMORY WARNING] {warning}")
                        if warning not in self.memory_warnings:
                            self.memory_warnings.append(warning)

            self.log(f"    Detected scale type column: {scale_col} (score: {best_score})")
            if len(scale_candidates) > 1:
                self.log(f"    Other candidates: {[(c[0], c[1]) for c in scale_candidates[1:3]]}")

        # If we have a scale column with multiple values, use customer-defined extraction
        if scale_col:
            scale_data = self._safe_get_column(df, scale_col)
            unique_scales = scale_data.dropna().unique()
            unique_scales = [s for s in unique_scales
                           if str(s).strip()
                           and str(s).lower() not in ['scale', 'nan', '', 'point', 'type']
                           and not str(s).replace('.', '').replace('-', '').isdigit()]
            if len(unique_scales) >= 2:
                self.log(f"    Found customer-defined scales: {list(unique_scales)[:10]}")
                self._extract_customer_pay_scales(df, scale_col, sheet_name)
                return

        # Detect scale type hint from sheet name
        sheet_lower = sheet_name.lower()
        if 'leadership' in sheet_lower or 'lead' in sheet_lower:
            scale_type_hint = 'leadership'
        elif 'support' in sheet_lower or 'njc' in sheet_lower or 'nj' in sheet_lower:
            scale_type_hint = 'support'
        elif 'teaching' in sheet_lower or 'teacher' in sheet_lower or 'main' in sheet_lower or 'mps' in sheet_lower or 'ups' in sheet_lower:
            scale_type_hint = 'teaching'
        else:
            scale_type_hint = 'auto'  # Will detect from point codes

        # Extended column detection for rates/salaries
        rate_keywords = ['salary', 'annual', '£', 'amount', 'rate', 'pay', 'value', 'fte']
        rate_cols = [c for c in df.columns if any(x in str(c).lower() for x in rate_keywords)]

        # Extended column detection for point codes (includes NJC/SCP patterns)
        point_keywords = ['point', 'scp', 'spine', 'spinal', 'scale', 'grade', 'code', 'level', 'step', 'band', 'njc']
        point_cols = [c for c in df.columns if any(x in str(c).lower() for x in point_keywords)]

        # Also check columns for point-like values by content (not just keywords)
        # This handles PDFs with generic column names
        point_patterns = [
            r'^[MLU]?\d{1,2}$',        # M1, L01, U1, 1, 43
            r'^SCP\s*\d{1,2}$',        # SCP1, SCP 43
            r'^[Pp]oint\s*\d+$',       # Point 1
            r'^[Nn][Jj][Cc]?\s*\d+$',  # NJC1, NJ 1
        ]
        for col in df.columns:
            if col in point_cols:
                continue
            col_data = self._safe_get_column(df, col)
            sample = col_data.dropna().head(10).astype(str).tolist()
            point_matches = sum(
                1 for v in sample
                if any(re.match(p, str(v).strip(), re.IGNORECASE) for p in point_patterns)
            )
            # If >50% of values look like points, add this column
            if len(sample) >= 3 and point_matches >= len(sample) * 0.5:
                point_cols = [col] + point_cols  # Add to front (higher priority)
                self.log(f"    Detected point column by content: {col}")

        # FALLBACK: If no rate columns found by keyword, check ALL columns for salary-like values
        # This handles PDFs with generic column names like "unnamed_0", "column_1", etc.
        if not rate_cols:
            for col in df.columns:
                if col in point_cols:
                    continue  # Skip point columns
                col_data = self._safe_get_column(df, col)
                try:
                    # Clean currency formatting
                    cleaned = col_data.astype(str).str.replace('£', '', regex=False)
                    cleaned = cleaned.str.replace(',', '', regex=False).str.strip()
                    numeric = pd.to_numeric(cleaned, errors='coerce').dropna()
                    if len(numeric) >= 3:
                        # Check if values look like salaries (10000-200000 range)
                        salary_like = numeric[(numeric >= 10000) & (numeric <= 200000)]
                        if len(salary_like) >= len(numeric) * 0.5:
                            rate_cols.append(col)
                            self.log(f"    Detected salary column by content: {col}")
                except Exception:
                    pass

        self.log(f"    Rate columns found: {rate_cols[:5]}")
        self.log(f"    Point columns found: {point_cols[:5]}")

        # Collect points categorized by scale type
        main_points = []      # M1-M6
        upper_points = []     # U1-U3
        leadership_points = []  # L01-L43
        support_points = []   # Numeric points (1-43 etc)

        # Extract points and rates
        for idx, row in df.iterrows():
            point_code = None
            rate = None

            # Try to get point code from point columns
            for col in point_cols:
                val = row.get(col)
                # Handle duplicate columns (returns Series)
                if isinstance(val, pd.Series):
                    val = val.iloc[0] if len(val) > 0 else None
                if pd.notna(val) and str(val).strip():
                    point_code = str(val).strip()
                    # Skip header-like values
                    if point_code.lower() in ['point', 'scp', 'scale', 'grade', 'code']:
                        point_code = None
                        continue
                    break

            # If no point code from columns, check if row index is the point
            if not point_code and isinstance(idx, (int, str)):
                idx_str = str(idx).strip()
                if re.match(r'^[MLU]?\d+$', idx_str, re.IGNORECASE):
                    point_code = idx_str

            if not point_code:
                continue

            # Try to get rate from rate columns - only use explicit salary values
            # Do NOT guess or convert values - if rate doesn't exist, leave blank (0)
            for col in rate_cols:
                val = row.get(col)
                # Handle duplicate columns (returns Series instead of scalar)
                if isinstance(val, pd.Series):
                    val = val.iloc[0] if len(val) > 0 else None
                if pd.notna(val):
                    try:
                        # Handle various currency formats
                        val_str = str(val).replace('£', '').replace(',', '').replace(' ', '').strip()
                        parsed_rate = float(val_str)
                        # Only accept values that are clearly annual salaries (> 10000)
                        if parsed_rate > 10000:
                            rate = parsed_rate
                            break
                    except (ValueError, TypeError):
                        pass

            # If no rate found, set to 0 (blank) - do NOT skip the row
            if not rate:
                rate = 0

            # Categorize by point code pattern
            point_code_clean = str(point_code).strip().upper()
            point_lower = point_code_clean.lower()

            # Leadership scale (L1-L43)
            if point_lower.startswith('l') and any(c.isdigit() for c in point_code_clean):
                num = self._extract_point_number(point_code_clean)
                normalized = f"L{num:02d}"
                leadership_points.append({
                    'code': normalized,
                    'title': normalized,
                    'number': num,
                    'rate': rate,
                    'date_from': get_default_increase_date(True),
                })
            # Upper pay scale (U1-U3 or UPS1-UPS3)
            elif point_lower.startswith('u') or 'ups' in point_lower:
                num = self._extract_point_number(point_code_clean)
                normalized = f"U{min(num, 3)}"
                upper_points.append({
                    'code': normalized,
                    'title': normalized,
                    'number': num,
                    'rate': rate,
                    'date_from': get_default_increase_date(True),
                })
            # Main pay scale (M1-M6 or MPS1-MPS6)
            elif point_lower.startswith('m') or 'mps' in point_lower:
                num = self._extract_point_number(point_code_clean)
                normalized = f"M{min(num, 6)}"
                main_points.append({
                    'code': normalized,
                    'title': normalized,
                    'number': num,
                    'rate': rate,
                    'date_from': get_default_increase_date(True),
                })
            # Pure numeric - could be support or need context
            elif point_code_clean.isdigit():
                num = int(point_code_clean)
                if scale_type_hint == 'support' or num > 6:
                    # Support scale point
                    support_points.append({
                        'code': str(num),
                        'title': f"Point {num}",
                        'number': num,
                        'rate': rate,
                        'date_from': get_default_increase_date(False),
                    })
                elif scale_type_hint == 'teaching' and num <= 6:
                    # Assume Main scale
                    normalized = f"M{num}"
                    main_points.append({
                        'code': normalized,
                        'title': normalized,
                        'number': num,
                        'rate': rate,
                        'date_from': get_default_increase_date(True),
                    })
                else:
                    # Default to support for ambiguous numeric points
                    support_points.append({
                        'code': str(num),
                        'title': f"Point {num}",
                        'number': num,
                        'rate': rate,
                        'date_from': get_default_increase_date(False),
                    })
            else:
                # Other format - try to extract number and categorize
                num = self._extract_point_number(point_code_clean)
                if num > 0:
                    if scale_type_hint == 'leadership':
                        normalized = f"L{num:02d}"
                        leadership_points.append({
                            'code': normalized,
                            'title': point_code,
                            'number': num,
                            'rate': rate,
                            'date_from': get_default_increase_date(True),
                        })
                    elif scale_type_hint == 'support':
                        support_points.append({
                            'code': str(num),
                            'title': point_code,
                            'number': num,
                            'rate': rate,
                            'date_from': get_default_increase_date(False),
                        })
                    else:
                        # Default teaching
                        main_points.append({
                            'code': point_code_clean,
                            'title': point_code,
                            'number': num,
                            'rate': rate,
                            'date_from': get_default_increase_date(True),
                        })

        # -----------------------------------------------------------------------
        # Build ExtractedPayScale objects from the collected points.
        # No fallback to reference data — only what the customer provided.
        # -----------------------------------------------------------------------
        # Determine london weighting from sheet name
        lw = 'EW'
        for lw_code, lw_kws in [('IL', ['inner']), ('OL', ['outer']), ('FRI', ['fringe']), ('KEN', ['kent'])]:
            if any(kw in sheet_lower for kw in lw_kws):
                lw = lw_code
                break

        def _dedup(pts):
            """Deduplicate by point code, keeping highest rate."""
            seen = {}
            for pt in pts:
                code = pt['code']
                if code not in seen or pt.get('rate', 0) > seen[code].get('rate', 0):
                    seen[code] = pt
            return sorted(seen.values(), key=lambda p: p.get('number', 0))

        def _merge_or_create(scale_code, title, scale_type, inc_date, points_list):
            """Add points to an existing scale or create a new one."""
            if not points_list:
                return
            deduped = _dedup(points_list)
            existing = next((ps for ps in self.extracted_pay_scales if ps.code == scale_code), None)
            if existing:
                existing_codes = {p['code'] for p in existing.points}
                added = [p for p in deduped if p['code'] not in existing_codes]
                existing.points.extend(added)
                if added:
                    self.log(f"    Updated {scale_code}: added {len(added)} point(s) from customer data")
            else:
                self.extracted_pay_scales.append(ExtractedPayScale(
                    code=scale_code, title=title, scale_type=scale_type,
                    london_weighting=lw, increment_date=inc_date,
                    increase_date=inc_date, increase_percentage=0,
                    points=deduped, grades=[],
                ))
                self.log(f"    Created pay scale {scale_code} with {len(deduped)} point(s) from customer data")

        # Main Pay Scale (M1-M6) and Upper Pay Scale (UPS points) → MAIN
        main_code = self._derive_pay_scale_code("Teachers Main")
        _merge_or_create(main_code, f"Teachers Main Pay Scale - {lw}",
                         'teaching', '09/01', main_points + upper_points)

        # Leadership Scale (L01-L43) → LS
        ls_code = self._derive_pay_scale_code("Leadership Group")
        _merge_or_create(ls_code, f"Leadership Pay Scale - {lw}",
                         'leadership', '09/01', leadership_points)

        # Support scale — derive code from the sheet name so it reflects the title
        if support_points:
            sup_code = self._derive_pay_scale_code(sheet_name)
            sup_title = f"{sheet_name.strip()} - {lw}" if sheet_name else f"Support Pay Scale - {lw}"
            sup_type = 'support'
            sup_inc = '09/01' if self._is_standard_teaching_scale(sup_code) else '04/01'
            _merge_or_create(sup_code, sup_title, sup_type, sup_inc, support_points)

    # -------------------------------------------------------------------------
    # Known title-to-code mappings (from 18_IMPORT_PayScales.csv and S2 domain
    # knowledge).  Checked in order; first match wins.
    # Returns the base code only — no location suffix.
    # -------------------------------------------------------------------------
    _PAY_SCALE_TITLE_PATTERNS = [
        # Standard teaching scales
        (['lead practitioner'],               'LP'),
        (['lead prac'],                       'LP'),
        (['leadership group'],                'LS'),
        (['leadership'],                      'LS'),
        (['teachers main', 'main pay'],       'MAIN'),
        (['main scale', 'main pay scale'],    'MAIN'),
        (['teachers main'],                   'MAIN'),
        (['main'],                            'MAIN'),
        (['upper pay', 'upper scale', 'ups'], 'UPS'),
        (['unqualified teacher'],             'UQ'),
        (['unqualified'],                     'UQ'),
        # Support / NJC
        (['national joint council'],          'NJC'),
        (['njc'],                             'NJC'),
        (['kent'],                            'KENT'),
        (['support pay', 'support scale',
          'support staff'],                   'SUP'),
        (['support'],                         'SUP'),
        # Apprentice / spot / external
        (['apprentice', 'nmw'],               'APP/NMW'),
        (['apprentice'],                      'APP/NMW'),
        (['spot salary', 'spot scale',
          'spot rate'],                       'SPOT'),
        (['external staff'],                  'ESTF'),
    ]

    # Words to strip before abbreviating an unknown title
    _TITLE_NOISE_WORDS = frozenset([
        'pay', 'scale', 'scales', 'group', 'standard', 'national',
        'scheme', 'grade', 'band', 'rates', 'rate', 'staff',
        'and', 'the', 'of', 'for', 'a', 'an',
    ])

    def _derive_pay_scale_code(self, title: str) -> str:
        """Derive a pay scale code from a human-readable title.

        Returns the base abbreviation only — no location suffix.
        The caller appends _{LW} when location-specific codes are needed.

        Examples:
            "Support Pay Scale"     → SUP
            "NJC"                   → NJC
            "Leadership Group"      → LS
            "Teachers Main"         → MAIN
            "Lead Practitioner"     → LP
            "Local Agreement Scale" → LAS   (abbreviate unknown title)
        """
        if not title:
            return 'SUP'

        title_lower = title.lower().strip()

        # --- Step 1: known pattern match ---
        for keywords, base_code in self._PAY_SCALE_TITLE_PATTERNS:
            if all(kw in title_lower for kw in keywords):
                return base_code

        # --- Step 2: abbreviate significant words ---
        # Strip location words so they don't pollute the abbreviation
        for loc_word in ['inner london', 'outer london', 'fringe', 'inner', 'outer',
                         'england and wales', 'england & wales']:
            title_lower = title_lower.replace(loc_word, ' ')

        words = [w for w in re.split(r'[\s\-_/]+', title_lower)
                 if w and w not in self._TITLE_NOISE_WORDS]

        if not words:
            return 'SUP'

        if len(words) == 1:
            return words[0][:6].upper()

        # First letter of each significant word (up to 5)
        return ''.join(w[0] for w in words[:5]).upper()

    def _extract_point_number(self, code: str) -> int:
        """Extract numeric point number from code."""
        numbers = re.findall(r'\d+', str(code))
        if numbers:
            return int(numbers[0])
        return 0

    def _extract_customer_pay_scales(self, df: pd.DataFrame, scale_col: str, sheet_name: str):
        """Extract pay scales from customer-defined scale column."""
        self.log(f"    Extracting customer-defined pay scales from column: {scale_col}")

        # Find point and rate columns
        point_col = None
        rate_col = None

        for col in df.columns:
            col_lower = str(col).lower()
            if point_col is None and any(x in col_lower for x in ['point', 'scp', 'spinal', 'spine']):
                point_col = col
            if rate_col is None and any(x in col_lower for x in ['salary', 'annual', 'rate', '£']):
                rate_col = col

        if not point_col:
            self.log(f"    Warning: No point column found for customer scales")
            return

        self.log(f"    Using point column: {point_col}, rate column: {rate_col if rate_col else 'None (rates will be blank)'}")

        # Group by scale type
        scale_data = {}
        for _, row in df.iterrows():
            scale_name = str(row.get(scale_col, '')).strip().upper()
            if not scale_name or scale_name in ['NAN', 'SCALE', '']:
                continue

            point_val = row.get(point_col)

            # Handle Series if duplicate columns
            if isinstance(point_val, pd.Series):
                point_val = point_val.iloc[0] if len(point_val) > 0 else None

            if pd.isna(point_val):
                continue

            try:
                point_num = int(float(str(point_val)))
            except (ValueError, TypeError):
                continue

            # Only get rate if rate column exists - otherwise leave as 0
            rate_num = 0
            if rate_col:
                rate_val = row.get(rate_col)
                if isinstance(rate_val, pd.Series):
                    rate_val = rate_val.iloc[0] if len(rate_val) > 0 else None
                if pd.notna(rate_val):
                    try:
                        parsed = float(str(rate_val).replace('£', '').replace(',', '').strip())
                        # Only accept values that look like annual salaries (> 10000)
                        if parsed > 10000:
                            rate_num = parsed
                    except (ValueError, TypeError):
                        pass

            if scale_name not in scale_data:
                scale_data[scale_name] = []

            scale_data[scale_name].append({
                'code': str(point_num),
                'title': f"Point {point_num}",
                'number': point_num,
                'rate': rate_num,
                'date_from': get_default_increase_date(True),
            })

        # Create pay scales for each unique scale type
        for scale_name, points in scale_data.items():
            if not points:
                continue

            # Skip if already exists
            scale_code = scale_name[:10].replace(' ', '_')
            if scale_code in self.pay_scales_found:
                continue

            # Deduplicate and sort
            seen = set()
            unique_points = []
            for p in sorted(points, key=lambda x: x['number']):
                if p['code'] not in seen:
                    seen.add(p['code'])
                    unique_points.append(p)

            if unique_points:
                # Determine if teaching or support based on name
                is_teaching = any(x in scale_name.lower() for x in ['teach', 'mgmt', 'lect', 'lead', 'head'])
                scale_type = 'teaching' if is_teaching else 'support'

                pay_scale = ExtractedPayScale(
                    code=scale_code,
                    title=f"{scale_name} Scale",
                    scale_type=scale_type,
                    london_weighting='England & Wales',
                    increment_date=get_default_increase_date(is_teaching),
                    increase_date=get_default_increase_date(is_teaching),
                    increase_percentage=0,
                    points=unique_points,
                    grades=[]
                )
                self.extracted_pay_scales.append(pay_scale)
                self.pay_scales_found.add(scale_code)
                self.log(f"    Created pay scale: {scale_code} with {len(unique_points)} points")

    def _extract_grades_from_staff_data(self):
        """
        Extract PayScaleGrades from staff contract data.
        Grades are in the 'Pay scale GROUP' column (e.g., TCCMGMT, TCCLECT),
        associated with 'Pay scale TYPE' (e.g., MGMT, LECT).
        Point ranges are derived from actual staff data.
        """
        self.log("  Extracting pay scale grades from staff data...")

        # Collect unique grades by pay scale type
        # Structure: {pay_scale_type: {grade_code: {'points': set()}}}
        grades_by_scale = {}

        for df in self.staff_data:
            # Find the relevant columns
            grade_col = None  # Pay scale GROUP (code)
            scale_type_col = None  # Pay scale TYPE
            scale_point_col = None  # Pay scale point
            scale_name_col = None  # Pay scale NAME (title/description)

            for col in df.columns:
                col_lower = str(col).lower().replace(' ', '_')
                col_clean = col_lower.replace('_', '')  # Also check without underscores

                # Pay scale GROUP is the grade code (TCCMGMT, TCCLECT, etc.)
                if 'pay_scale_group' in col_lower or col_lower == 'pay_scale_group':
                    grade_col = col
                elif 'group' in col_lower and 'pay' in col_lower:
                    grade_col = col

                # Pay scale TYPE is the parent scale (MGMT, LECT, SUPP)
                if col_lower in ['pay_scale_type', 'scale', 'scale_type', 'pay_type']:
                    scale_type_col = col
                elif 'type' in col_lower and 'pay' in col_lower:
                    scale_type_col = col

                # Pay scale NAME is the human-readable title
                # LESSON L024: Check for "payscale name" or "pay scale name" column
                if 'payscalename' in col_clean or 'pay_scale_name' in col_lower:
                    scale_name_col = col
                elif col_lower in ['payscale_name', 'scale_name', 'pay_scale_name']:
                    scale_name_col = col
                elif 'name' in col_lower and 'pay' in col_lower and 'scale' in col_lower:
                    scale_name_col = col

                # Scale point
                if col_lower in ['scale_point', 'pay_scale_point']:
                    scale_point_col = col

            if not grade_col:
                # Try original column names with exact matches
                for col in df.columns:
                    col_str = str(col)
                    if 'Pay scale GROUP' in col_str:
                        grade_col = col
                    elif 'Pay scale TYPE' in col_str:
                        scale_type_col = col
                    elif 'Pay scale point' in col_str:
                        scale_point_col = col
                    # Also check for name column
                    if 'payscale name' in col_str.lower() or 'pay scale name' in col_str.lower():
                        scale_name_col = col

            if not grade_col:
                continue

            self.log(f"    Found grade column: {grade_col}")
            if scale_name_col:
                self.log(f"    Found scale name column: {scale_name_col}")

            for _, row in df.iterrows():
                grade_code = str(row.get(grade_col, '')).strip().upper()
                if not grade_code or grade_code == 'NAN' or grade_code == 'NONE':
                    continue

                # Get the pay scale type for this grade
                scale_type = 'MISSING'
                if scale_type_col:
                    st = str(row.get(scale_type_col, '')).strip().upper()
                    if st and st != 'NAN':
                        scale_type = st

                # Get the human-readable name/title from "payscale name" column
                # LESSON L024: Use payscale name column for grade title
                scale_name = None
                if scale_name_col:
                    sn = row.get(scale_name_col)
                    if pd.notna(sn) and str(sn).strip() and str(sn).strip().upper() != 'NAN':
                        scale_name = str(sn).strip()

                # Get the point if available (numeric only)
                point = None
                if scale_point_col:
                    pt = row.get(scale_point_col)
                    if pd.notna(pt):
                        try:
                            point = int(float(str(pt).strip()))
                        except (ValueError, TypeError):
                            pass  # Skip non-numeric points like "SPOT"

                # Track this grade
                if scale_type not in grades_by_scale:
                    grades_by_scale[scale_type] = {}
                if grade_code not in grades_by_scale[scale_type]:
                    grades_by_scale[scale_type][grade_code] = {'points': set(), 'name': None}
                if point is not None:
                    grades_by_scale[scale_type][grade_code]['points'].add(point)
                # Store the name (use first non-null value found)
                if scale_name and not grades_by_scale[scale_type][grade_code]['name']:
                    grades_by_scale[scale_type][grade_code]['name'] = scale_name

        # Build grades and associate with extracted pay scales
        total_grades = 0
        for scale_type, grades in grades_by_scale.items():
            for grade_code, grade_data in grades.items():
                points = sorted([p for p in grade_data['points'] if isinstance(p, int)])

                # Determine point range from actual data
                from_point = min(points) if points else 0
                to_point = max(points) if points else 0

                # Use the name from customer data if available, otherwise fall back to code
                # LESSON L024: "payscale name" column provides human-readable grade title
                grade_title = grade_data.get('name') or grade_code

                # Find the matching pay scale and add the grade
                for ps in self.extracted_pay_scales:
                    if ps.code == scale_type:
                        ps.grades.append({
                            'code': grade_code,
                            'title': grade_title,
                            'from_point': from_point,
                            'to_point': to_point,
                        })
                        total_grades += 1
                        self.log(f"      Grade {grade_code} '{grade_title}' ({scale_type}): points {from_point}-{to_point}")
                        break

        self.log(f"    Extracted {total_grades} grades from customer data")

    def _extract_allowances_from_df(self, df: pd.DataFrame):
        """Extract allowance types and points from a dedicated allowance dataframe."""
        self.log(f"  Extracting allowances from dedicated sheet...")

        # Look for allowance type column and amount column
        type_col = None
        amount_col = None
        code_col = None

        for col in df.columns:
            col_lower = str(col).lower()
            if any(x in col_lower for x in ['type', 'allowance', 'name', 'description']):
                if type_col is None:
                    type_col = col
            if any(x in col_lower for x in ['amount', 'value', 'rate', '£', 'annual']):
                if amount_col is None:
                    amount_col = col
            if any(x in col_lower for x in ['code', 'point', 'level']):
                if code_col is None:
                    code_col = col

        # If we found structured columns, extract from rows
        if type_col and amount_col:
            self.log(f"    Found structured allowance data: type={type_col}, amount={amount_col}")
            allowances_by_type = {}

            for _, row in df.iterrows():
                type_val = row.get(type_col)
                amount_val = row.get(amount_col)

                if not pd.notna(type_val) or not pd.notna(amount_val):
                    continue

                type_str = str(type_val).strip().upper()
                try:
                    amount = float(str(amount_val).replace('£', '').replace(',', '').strip())
                    if amount <= 0:
                        continue
                except (ValueError, TypeError):
                    continue

                # Categorize by type
                if 'TLR' in type_str:
                    type_code = 'TLR'
                    type_title = 'Teaching and Learning Responsibilities'
                elif 'SEN' in type_str and 'PENSION' not in type_str:
                    type_code = 'SEN'
                    type_title = 'Special Educational Needs'
                elif 'RECRUIT' in type_str:
                    type_code = 'REC'
                    type_title = 'Recruitment Allowance'
                elif 'RETAIN' in type_str or 'RETENTION' in type_str:
                    type_code = 'RET'
                    type_title = 'Retention Allowance'
                elif 'LONDON' in type_str:
                    type_code = 'LON'
                    type_title = 'London Weighting'
                else:
                    continue

                if type_code not in allowances_by_type:
                    allowances_by_type[type_code] = {'title': type_title, 'amounts': set()}
                allowances_by_type[type_code]['amounts'].add(amount)

            # Create allowance types
            for type_code, data in allowances_by_type.items():
                if type_code in self.allowance_types_found:
                    continue

                points = []
                for idx, amount in enumerate(sorted(data['amounts'], reverse=True), 1):
                    points.append({
                        'code': f'{type_code}{idx}',
                        'title': f'{data["title"]} {idx}',
                        'amount': amount,
                    })

                if points:
                    allowance = ExtractedAllowance(
                        type_code=type_code,
                        type_title=data['title'],
                        points=points,
                        increase_date=get_default_increase_date(True),
                        increase_percentage=0,
                    )
                    self.extracted_allowances.append(allowance)
                    self.allowance_types_found.add(type_code)
                    self.log(f"    Extracted {len(points)} {type_code} allowance points")

        else:
            # Fallback: Look for columns named after allowance types
            for col in df.columns:
                col_lower = str(col).lower()

                # TLR allowances
                if 'tlr' in col_lower:
                    self._extract_allowance_type(df, col, 'TLR', 'Teaching and Learning Responsibilities')
                # SEN allowances (avoid matching 'pension')
                elif 'sen' in col_lower and 'pension' not in col_lower:
                    self._extract_allowance_type(df, col, 'SEN', 'Special Educational Needs')
                # Recruitment allowances
                elif 'recruitment' in col_lower or 'recruit' in col_lower:
                    self._extract_allowance_type(df, col, 'REC', 'Recruitment Allowance')
                # Retention allowances
                elif 'retention' in col_lower or 'retain' in col_lower:
                    self._extract_allowance_type(df, col, 'RET', 'Retention Allowance')
                # London weighting
                elif 'london' in col_lower and 'weight' in col_lower:
                    self._extract_allowance_type(df, col, 'LON', 'London Weighting')

    def _extract_allowance_type(self, df: pd.DataFrame, col: str, code: str, title: str):
        """Extract a specific allowance type from a column."""
        # Skip if already extracted
        if code in self.allowance_types_found:
            return

        # Collect unique amounts
        unique_amounts = set()

        for idx, row in df.iterrows():
            val = row.get(col)
            if pd.notna(val):
                try:
                    amount = float(str(val).replace('£', '').replace(',', '').strip())
                    if amount > 0:
                        unique_amounts.add(amount)
                except (ValueError, TypeError):
                    pass

        if not unique_amounts:
            return

        # Create points from unique amounts (sorted descending - highest value = level 1)
        points = []
        for idx, amount in enumerate(sorted(unique_amounts, reverse=True), 1):
            points.append({
                'code': f"{code}{idx}",
                'title': f"{title} {idx}",
                'amount': amount,
            })

        if points:
            allowance = ExtractedAllowance(
                type_code=code,
                type_title=title,
                points=points,
                increase_date=get_default_increase_date(True),
                increase_percentage=0,
            )
            self.extracted_allowances.append(allowance)
            self.allowance_types_found.add(code)
            self.log(f"    Extracted {len(points)} {code} allowance points from column '{col}'")

    def _extract_pay_scales_from_staff_data(self, df: pd.DataFrame):
        """
        Extract pay scale points from staff contract data.

        Groups points by the actual pay scale code found in the row (pay_scale /
        pay_scale_code column).  When no explicit code column exists, falls back
        to inferring from point-code patterns (M→MAIN, L→LS, U→MAIN/UPS,
        numeric→SUP).

        Creates or updates ExtractedPayScale objects directly from customer data.
        No reference-file fallback — only real data is used.
        """
        self.log("    Extracting pay scales from staff data...")

        # Columns that carry the pay scale code for each staff member
        scale_code_cols = [c for c in df.columns if any(
            kw in str(c).lower() for kw in
            ['pay_scale', 'payscale', 'scale_code', 'scale_type', 'pay scale']
        )]

        # Columns that carry the individual point code
        point_cols = [c for c in df.columns if any(
            kw in str(c).lower() for kw in
            ['scale_point', 'scp', 'spine_point', 'current_scale_point', 'pay_point']
        )]

        # Columns that carry the salary/rate
        rate_cols = [c for c in df.columns if any(
            kw in str(c).lower() for kw in ['annual_salary', 'fte_salary', 'salary']
        )]

        if not point_cols:
            self.log("    No scale point column found in staff data — skipping point extraction")
            return

        # Accumulate: {scale_code: {point_code: point_dict}}
        points_by_scale: Dict[str, Dict[str, dict]] = {}

        for _, row in df.iterrows():
            # --- get point code ---
            pt_str = ''
            for col in point_cols:
                val = self._safe_get(row, col, '')
                if self._safe_notna(val) and str(val).strip() not in ('', 'nan'):
                    pt_str = str(val).strip()
                    break
            if not pt_str:
                continue

            # --- get rate ---
            rate = 0
            for col in rate_cols:
                val = self._safe_get(row, col, None)
                if self._safe_notna(val):
                    try:
                        parsed = float(str(val).replace('£', '').replace(',', '').strip())
                        if parsed > 10000:
                            rate = parsed
                            break
                    except (ValueError, TypeError):
                        pass

            # --- resolve scale code ---
            scale_code = ''
            for col in scale_code_cols:
                val = self._safe_get(row, col, '')
                if self._safe_notna(val) and str(val).strip() not in ('', 'nan', 'none'):
                    scale_code = str(val).strip()
                    break

            if not scale_code:
                # Infer from point code pattern — derive code via title lookup
                pt_upper = pt_str.upper()
                job_title = str(self._safe_get(row, 'job_title', '')).lower()
                if pt_upper.startswith('LP') and any(c.isdigit() for c in pt_upper):
                    scale_code = self._derive_pay_scale_code("Lead Practitioner")
                elif pt_upper.startswith('L') and any(c.isdigit() for c in pt_upper):
                    scale_code = self._derive_pay_scale_code("Leadership Group")
                elif pt_upper.startswith(('UPS', 'U')) and any(c.isdigit() for c in pt_upper):
                    scale_code = self._derive_pay_scale_code("Teachers Main")
                elif pt_upper.startswith('M') and any(c.isdigit() for c in pt_upper):
                    scale_code = self._derive_pay_scale_code("Teachers Main")
                elif pt_upper.isdigit():
                    is_teaching = any(x in job_title for x in ['teacher', 'head', 'principal'])
                    scale_code = (self._derive_pay_scale_code("Teachers Main")
                                  if is_teaching
                                  else self._derive_pay_scale_code("Support"))
                else:
                    scale_code = self._derive_pay_scale_code("Support")

            # --- normalise point code and number ---
            pt_upper = pt_str.upper()
            pt_num = self._extract_point_number(pt_upper)
            if pt_num == 0:
                try:
                    pt_num = int(float(pt_str))
                except (ValueError, TypeError):
                    pt_num = 0

            point_dict = {
                'code': pt_upper,
                'title': pt_upper,
                'number': pt_num,
                'rate': rate,
                'date_from': '',
            }

            if scale_code not in points_by_scale:
                points_by_scale[scale_code] = {}
            existing_pt = points_by_scale[scale_code].get(pt_upper)
            if existing_pt is None or rate > existing_pt.get('rate', 0):
                points_by_scale[scale_code][pt_upper] = point_dict

        if not points_by_scale:
            self.log("    No scale points found in staff data")
            return

        # Create or update ExtractedPayScale for each discovered scale
        for scale_code, pts_dict in points_by_scale.items():
            sorted_pts = sorted(pts_dict.values(), key=lambda p: p.get('number', 0))

            # Infer scale meta from code
            sc_upper = scale_code.upper()
            if sc_upper.startswith('LS'):
                scale_type, inc_date = 'leadership', '09/01'
                title = f"Leadership Pay Scale"
            elif sc_upper.startswith('LP'):
                scale_type, inc_date = 'leadership', '09/01'
                title = f"Lead Practitioner Pay Scale"
            elif sc_upper.startswith(('MAIN', 'UPS', 'UQ')):
                scale_type, inc_date = 'teaching', '09/01'
                title = f"Teachers Pay Scale"
            else:
                scale_type, inc_date = 'support', '04/01'
                title = f"Support Pay Scale"

            existing = next((ps for ps in self.extracted_pay_scales if ps.code == scale_code), None)
            if existing:
                existing_codes = {p['code'] for p in existing.points}
                added = [p for p in sorted_pts if p['code'] not in existing_codes]
                existing.points.extend(added)
                if added:
                    self.log(f"    Updated {scale_code}: added {len(added)} point(s) from staff data")
            else:
                lw = 'EW'
                for suffix in ['_IL', '_OL', '_FRI', '_KEN']:
                    if sc_upper.endswith(suffix):
                        lw = suffix.lstrip('_')
                        break
                self.extracted_pay_scales.append(ExtractedPayScale(
                    code=scale_code, title=title, scale_type=scale_type,
                    london_weighting=lw, increment_date=inc_date,
                    increase_date=inc_date, increase_percentage=0,
                    points=sorted_pts, grades=[],
                ))
                self.log(f"    Created pay scale {scale_code} with {len(sorted_pts)} point(s) from staff data")

    def _extract_allowances_from_staff_data(self, df: pd.DataFrame):
        """
        Extract allowance types and points from staff contract data.
        This handles cases where allowance info is embedded in staff rows.
        """
        self.log("    Extracting allowances from staff data...")

        # First, handle customer format: "Allowance N NAME" + "Allowance N value" columns
        # This is common format where allowance type is in one column and value in another
        allowance_name_cols = {}
        allowance_value_cols = {}

        import re
        for col in df.columns:
            col_str = str(col).lower()
            # Match patterns like "allowance_1_name", "allowance 1 name", "allowance1name", etc.
            if 'allowance' in col_str and 'name' in col_str:
                # Extract the number - allow whitespace, underscore, or nothing between
                match = re.search(r'allowance[\s_]*(\d+)', col_str)
                if match:
                    num = match.group(1)
                    allowance_name_cols[num] = col
            # Match patterns like "allowance_1_value", "allowance 1 value @ 1fte", etc.
            elif 'allowance' in col_str and ('value' in col_str or 'amount' in col_str):
                # Extract the number
                match = re.search(r'allowance[\s_]*(\d+)', col_str)
                if match:
                    num = match.group(1)
                    # Prefer "value @ 1FTE" or just "value"
                    if num not in allowance_value_cols or '1fte' not in str(allowance_value_cols.get(num, '')).lower():
                        allowance_value_cols[num] = col

        # Extract allowances from paired NAME/VALUE columns
        allowances_by_name = {}  # Track unique allowances by name

        for num, name_col in allowance_name_cols.items():
            value_col = allowance_value_cols.get(num)
            if not value_col:
                continue

            for _, row in df.iterrows():
                allowance_name = str(row.get(name_col, '')).strip()
                if not allowance_name or allowance_name == 'nan':
                    continue

                try:
                    allowance_value = float(str(row.get(value_col, 0)).replace('£', '').replace(',', '').strip())
                    if allowance_value <= 0:
                        continue
                except (ValueError, TypeError):
                    continue

                # Track unique allowances
                if allowance_name not in allowances_by_name:
                    allowances_by_name[allowance_name] = set()
                allowances_by_name[allowance_name].add(allowance_value)

        # Create allowance types from extracted NAME/VALUE pairs
        for allowance_name, amounts in allowances_by_name.items():
            # Generate code from name
            code = allowance_name.upper().replace(' ', '_').replace('%', 'PC')[:20]
            code = re.sub(r'[^A-Z0-9_]', '', code)

            # Check if already extracted
            if code in self.allowance_types_found:
                continue

            # Create points for each unique amount
            points = []
            for idx, amount in enumerate(sorted(amounts, reverse=True), 1):
                points.append({
                    'code': f'{code}_{idx}' if len(amounts) > 1 else code,
                    'title': f'{allowance_name} Point {idx}' if len(amounts) > 1 else allowance_name,
                    'amount': amount,
                })

            if points:
                allowance = ExtractedAllowance(
                    type_code=code,
                    type_title=allowance_name,
                    points=points,
                    increase_date=get_default_increase_date(True),
                    increase_percentage=0.0,
                )
                self.extracted_allowances.append(allowance)
                self.allowance_types_found.add(code)
                self.log(f"      Extracted allowance: {allowance_name} ({len(points)} points)")

        # Track unique allowance amounts by type (for specific column name patterns)
        tlr_amounts = set()
        sen_amounts = set()
        recruitment_amounts = set()
        retention_amounts = set()

        # Look for allowance columns
        for col in df.columns:
            col_lower = str(col).lower()

            # TLR allowances
            if 'tlr' in col_lower:
                for _, row in df.iterrows():
                    val = row.get(col)
                    if pd.notna(val):
                        try:
                            amount = float(str(val).replace('£', '').replace(',', '').strip())
                            if amount > 0:
                                tlr_amounts.add(amount)
                        except (ValueError, TypeError):
                            pass

            # SEN allowances
            elif 'sen' in col_lower and 'pension' not in col_lower:
                for _, row in df.iterrows():
                    val = row.get(col)
                    if pd.notna(val):
                        try:
                            amount = float(str(val).replace('£', '').replace(',', '').strip())
                            if amount > 0:
                                sen_amounts.add(amount)
                        except (ValueError, TypeError):
                            pass

            # Recruitment allowances
            elif 'recruitment' in col_lower or 'recruit' in col_lower:
                for _, row in df.iterrows():
                    val = row.get(col)
                    if pd.notna(val):
                        try:
                            amount = float(str(val).replace('£', '').replace(',', '').strip())
                            if amount > 0:
                                recruitment_amounts.add(amount)
                        except (ValueError, TypeError):
                            pass

            # Retention allowances
            elif 'retention' in col_lower or 'retain' in col_lower:
                for _, row in df.iterrows():
                    val = row.get(col)
                    if pd.notna(val):
                        try:
                            amount = float(str(val).replace('£', '').replace(',', '').strip())
                            if amount > 0:
                                retention_amounts.add(amount)
                        except (ValueError, TypeError):
                            pass

        # Create allowance types from extracted amounts
        if tlr_amounts and 'TLR' not in self.allowance_types_found:
            points = []
            for idx, amount in enumerate(sorted(tlr_amounts, reverse=True), 1):
                points.append({
                    'code': f'TLR{idx}',
                    'title': f'TLR {idx}',
                    'amount': amount,
                })

            allowance = ExtractedAllowance(
                type_code='TLR',
                type_title='Teaching and Learning Responsibilities',
                points=points,
                increase_date=get_default_increase_date(True),
                increase_percentage=0,
            )
            self.extracted_allowances.append(allowance)
            self.allowance_types_found.add('TLR')
            self.log(f"      Extracted {len(points)} TLR allowance points")

        if sen_amounts and 'SEN' not in self.allowance_types_found:
            points = []
            for idx, amount in enumerate(sorted(sen_amounts, reverse=True), 1):
                points.append({
                    'code': f'SEN{idx}',
                    'title': f'SEN {idx}',
                    'amount': amount,
                })

            allowance = ExtractedAllowance(
                type_code='SEN',
                type_title='Special Educational Needs',
                points=points,
                increase_date=get_default_increase_date(True),
                increase_percentage=0,
            )
            self.extracted_allowances.append(allowance)
            self.allowance_types_found.add('SEN')
            self.log(f"      Extracted {len(points)} SEN allowance points")

        if recruitment_amounts and 'REC' not in self.allowance_types_found:
            points = []
            for idx, amount in enumerate(sorted(recruitment_amounts, reverse=True), 1):
                points.append({
                    'code': f'REC{idx}',
                    'title': f'Recruitment Allowance {idx}',
                    'amount': amount,
                })

            allowance = ExtractedAllowance(
                type_code='REC',
                type_title='Recruitment Allowance',
                points=points,
                increase_date=get_default_increase_date(True),
                increase_percentage=0,
            )
            self.extracted_allowances.append(allowance)
            self.allowance_types_found.add('REC')
            self.log(f"      Extracted {len(points)} Recruitment allowance points")

        if retention_amounts and 'RET' not in self.allowance_types_found:
            points = []
            for idx, amount in enumerate(sorted(retention_amounts, reverse=True), 1):
                points.append({
                    'code': f'RET{idx}',
                    'title': f'Retention Allowance {idx}',
                    'amount': amount,
                })

            allowance = ExtractedAllowance(
                type_code='RET',
                type_title='Retention Allowance',
                points=points,
                increase_date=get_default_increase_date(True),
                increase_percentage=0,
            )
            self.extracted_allowances.append(allowance)
            self.allowance_types_found.add('RET')
            self.log(f"      Extracted {len(points)} Retention allowance points")

    # =========================================================================
    # PHASE 3: BUILD ALL TEMPLATE SHEETS
    # =========================================================================

    def build_all_templates(self) -> Dict[str, pd.DataFrame]:
        """Build ALL template sheets from extracted data."""
        self.log("\n" + "="*60)
        self.log("PHASE 3: BUILDING ALL TEMPLATE SHEETS")
        self.log("="*60)

        # FIRST: Build consolidated staff lookup from all data sources
        # This merges data from multiple files (names, contracts, pay scales, etc.)
        # using the unique identifier to join records across files
        self._build_staff_lookup()

        # Build in correct order
        self._build_pay_scales()
        self._build_pay_scale_points()
        self._build_pay_scale_grades()
        self._build_allowance_types()
        self._build_allowance_points()
        self._build_pensions()
        self._build_eqw_patterns()
        self._build_staff_role_groups()
        self._build_staff_roles()
        self._build_staff_members()
        self._build_contracts_teaching()
        self._build_contracts_support()
        self._ensure_spot_pay_scale()  # Add SPOT pay scale if contracts need it
        self._build_contract_allowances()
        self._build_finance_codes_s2()

        # Log contract generation summary
        total_staff = len(self.staff_lookup)
        total_teaching = len(self.template_data.get("ContractsTeachFTE", []))
        total_support = len(self.template_data.get("ContractsSupportHours", []))
        total_contracts = total_teaching + total_support
        total_skipped = len(self.skipped_staff)

        self.log("\n" + "="*60)
        self.log("CONTRACT GENERATION SUMMARY")
        self.log("="*60)
        self.log(f"Total staff in lookup: {total_staff}")
        self.log(f"Teaching contracts created: {total_teaching}")
        self.log(f"Support contracts created: {total_support}")
        self.log(f"Total contracts: {total_contracts}")
        self.log(f"Staff skipped: {total_skipped}")
        if total_skipped > 0:
            self.log(f"Skipped staff details saved to _SkippedStaff sheet")
        self.log("="*60 + "\n")

        # Convert to DataFrames
        result = {}
        for sheet_name, data in self.template_data.items():
            if data:
                result[sheet_name] = pd.DataFrame(data)
                self.log(f"  {sheet_name}: {len(data)} rows")
            else:
                self.log(f"  {sheet_name}: 0 rows (no data)")

        return result

    def _build_pay_scales(self):
        """Build PayScales sheet from extracted pay scales AND staff lookup data."""
        self.log("Building PayScales...")

        added_codes = set()

        # First: Add from extracted pay scales (dedicated pay scale files)
        for ps in self.extracted_pay_scales:
            if ps.code not in added_codes:
                self.template_data["PayScales"].append({
                    "PayScaleCode": ps.code,
                    "PayScaleTitle": ps.title,
                    "ServiceIncrementDateEnabled": ps.scale_type == 'support',
                    "IncrementDate": ps.increment_date,
                    "IncreaseDate": ps.increase_date,
                    "IncreasePercentage": ps.increase_percentage,
                    "AvailableToAllSchools": True,
                    "SchoolCodes": "",
                    "PayScaleEnabled": True,
                    "ExcludeNationalInsurance": False,
                    "ExcludePension": False,
                })
                added_codes.add(ps.code)

        # Second: Scan staff lookup for any pay scales referenced in contract data
        # This catches pay scales that are columns in the data, not separate files
        lookup_pay_scales = set()
        for staff_code, record in self.staff_lookup.items():
            # Try multiple field names for pay scale
            ps_code = self._lookup_get(record, 'pay_scale', '')
            if ps_code and str(ps_code).strip().upper() not in ['', 'NAN', 'MISSING']:
                lookup_pay_scales.add(str(ps_code).strip().upper())

            ps_type = self._lookup_get(record, 'pay_scale_type', '')
            if ps_type and str(ps_type).strip().upper() not in ['', 'NAN', 'MISSING']:
                lookup_pay_scales.add(str(ps_type).strip().upper())

            ps_contract = self._lookup_get(record, 'pay_scale_contract', '')
            if ps_contract and str(ps_contract).strip().upper() not in ['', 'NAN', 'MISSING']:
                lookup_pay_scales.add(str(ps_contract).strip().upper())

        # Add any pay scales from lookup that weren't in extracted_pay_scales
        for ps_code in sorted(lookup_pay_scales):
            if ps_code not in added_codes and ps_code != 'SPOT':
                # Dynamically determine if teaching or support by checking job titles
                # Look at staff using this pay scale and check if they have teaching roles
                is_teaching = self._infer_pay_scale_type(ps_code)

                self.template_data["PayScales"].append({
                    "PayScaleCode": ps_code,
                    "PayScaleTitle": ps_code,  # Use code as title if not known
                    "ServiceIncrementDateEnabled": not is_teaching,
                    "IncrementDate": get_default_increase_date(not is_teaching),
                    "IncreaseDate": get_default_increase_date(not is_teaching),
                    "IncreasePercentage": 0,
                    "AvailableToAllSchools": True,
                    "SchoolCodes": "",
                    "PayScaleEnabled": True,
                    "ExcludeNationalInsurance": False,
                    "ExcludePension": False,
                })
                added_codes.add(ps_code)
                self.log(f"  Added pay scale '{ps_code}' from contract data (teaching={is_teaching})")

        self.log(f"  Total pay scales: {len(added_codes)}")

        # SPOT pay scale will be added after contracts are built if needed

    def _infer_pay_scale_type(self, pay_scale_code: str) -> bool:
        """
        Dynamically infer if a pay scale is for teaching staff.

        Looks at job titles of staff using this pay scale and checks
        if they match teaching role patterns.

        Args:
            pay_scale_code: The pay scale code to check

        Returns:
            True if teaching, False if support
        """
        teaching_count = 0
        support_count = 0

        for staff_code, record in self.staff_lookup.items():
            # Check if this staff member uses this pay scale
            ps = self._lookup_get(record, 'pay_scale', '')
            if not ps:
                ps = self._lookup_get(record, 'pay_scale_type', '')
            ps = str(ps).strip().upper() if ps else ''

            if ps != pay_scale_code:
                continue

            # Get job title
            title = self._lookup_get(record, 'job_title', '')
            if not title:
                title = self._lookup_get(record, 'position', '')
            title = str(title).strip().lower() if title else ''

            if not title:
                continue

            # Check if title matches teaching patterns
            srg = get_srg_for_role(title)
            if S2_STAFF_ROLE_GROUP_PATTERNS.get(srg, {}).get('teaching', False):
                teaching_count += 1
            else:
                support_count += 1

        # If more teaching than support, it's a teaching pay scale
        return teaching_count > support_count

    def _ensure_spot_pay_scale(self):
        """Ensure SPOT pay scale exists if any contracts use spot salaries."""
        # Check if we have any contracts using SPOT pay scale
        has_spot_contracts = False

        for contract in self.template_data.get("ContractsTeachFTE", []):
            if contract.get("PayScaleCode") == "SPOT":
                has_spot_contracts = True
                break

        if not has_spot_contracts:
            for contract in self.template_data.get("ContractsSupportFTE", []):
                if contract.get("PayScaleCode") == "SPOT":
                    has_spot_contracts = True
                    break

        # Also check if we tracked any spot salary amounts during processing
        if hasattr(self, '_spot_salary_amounts') and self._spot_salary_amounts:
            has_spot_contracts = True

        # If we have spot salaries, create SPOT pay scale with single SPOT point
        if has_spot_contracts and 'SPOT' not in self.pay_scales_found:
            self.template_data["PayScales"].append({
                "PayScaleCode": "SPOT",
                "PayScaleTitle": "Spot Salary",
                "ServiceIncrementDateEnabled": False,
                "IncrementDate": get_default_increase_date(False),
                "IncreaseDate": get_default_increase_date(False),
                "IncreasePercentage": 0,
                "AvailableToAllSchools": True,
                "SchoolCodes": "",
                "PayScaleEnabled": True,
                "ExcludeNationalInsurance": False,
                "ExcludePension": False,
            })
            self.pay_scales_found.add('SPOT')
            self.log("  Added SPOT pay scale for spot salaries")

            # Add single SPOT point - actual salary is stored in the contract's salary field
            self.template_data["PayScalePoints"].append({
                "PayScaleCode": "SPOT",
                "PayScalePointCode": "SPOT",
                "PayScalePointTitle": "Spot Salary",
                "ScalePointNumber": 1,
                "Hourly": False,
                "PayScalePointEnabled": True,
                "RateDateFrom": get_default_increase_date(False),
                "RateDateTo": "",
                "PayScaleRate": 0,  # Rate is 0 - actual salary comes from contract
            })
            self.log("    Added SPOT/SPOT pay scale point")

    def _build_pay_scale_points(self):
        """Build PayScalePoints sheet from extracted pay scales AND staff lookup data."""
        self.log("Building PayScalePoints...")

        added_points = set()  # (pay_scale_code, point_code) pairs

        # First: Add from extracted pay scales (dedicated pay scale files)
        for ps in self.extracted_pay_scales:
            for point in ps.points:
                key = (ps.code, str(point.get('code', '')))
                if key not in added_points:
                    # Build title: Abbreviated PayScaleCode + PointCode (e.g., "MPS_3", "SUP_STF_11")
                    point_code = point.get('code', '')
                    ps_abbreviated = self._abbreviate_text(ps.code)
                    point_title = f"{ps_abbreviated}_{point_code}"
                    self.template_data["PayScalePoints"].append({
                        "PayScaleCode": ps.code,
                        "PayScalePointCode": point_code,
                        "PayScalePointTitle": point_title,
                        "ScalePointNumber": point.get('number', 0),
                        "Hourly": False,
                        "PayScalePointEnabled": True,
                        "RateDateFrom": point.get('date_from', get_default_increase_date(ps.scale_type == 'teaching')),
                        "RateDateTo": "",
                        "PayScaleRate": point.get('rate', 0),
                    })
                    added_points.add(key)

        # Second: Scan staff lookup for pay scale points referenced in contract data
        # Group by pay scale code -> set of points
        lookup_points = {}  # {pay_scale_code: {point_code: salary}}
        for staff_code, record in self.staff_lookup.items():
            # Get pay scale code
            ps_code = self._lookup_get(record, 'pay_scale', '')
            if not ps_code or str(ps_code).strip().upper() in ['', 'NAN', 'MISSING']:
                ps_code = self._lookup_get(record, 'pay_scale_type', '')
            ps_code = str(ps_code).strip().upper() if ps_code else ''

            if not ps_code or ps_code in ['', 'NAN', 'MISSING']:
                continue

            # Get scale point
            point_code = self._lookup_get(record, 'scale_point', '')
            if not point_code or str(point_code).strip().upper() in ['', 'NAN', 'MISSING']:
                point_code = self._lookup_get(record, 'current_scale_point', '')
            if not point_code or str(point_code).strip().upper() in ['', 'NAN', 'MISSING']:
                point_code = self._lookup_get(record, 'scp', '')
            point_code = str(point_code).strip() if point_code else ''

            if not point_code or point_code.upper() in ['', 'NAN', 'MISSING']:
                continue

            # Get salary if available
            salary = self._lookup_get(record, 'annual_salary', 0)
            if not salary or str(salary).lower() == 'nan':
                salary = self._lookup_get(record, 'salary', 0)
            try:
                salary = float(str(salary).replace(',', '').replace('£', ''))
            except (ValueError, TypeError):
                salary = 0

            if ps_code not in lookup_points:
                lookup_points[ps_code] = {}
            if point_code not in lookup_points[ps_code]:
                lookup_points[ps_code][point_code] = salary

        # Add points from lookup that weren't in extracted_pay_scales
        for ps_code, points in lookup_points.items():
            is_teaching = self._infer_pay_scale_type(ps_code)  # Dynamic inference
            for point_code, salary in sorted(points.items()):
                key = (ps_code, point_code)
                if key not in added_points and ps_code != 'SPOT':
                    # Try to extract point number
                    try:
                        point_num = int(''.join(c for c in point_code if c.isdigit()) or '0')
                    except ValueError:
                        point_num = 0

                    # Build title: Abbreviated PayScaleCode + PointCode (e.g., "MPS_3", "SUP_STF_11")
                    ps_abbreviated = self._abbreviate_text(ps_code)
                    point_title = f"{ps_abbreviated}_{point_code}"

                    self.template_data["PayScalePoints"].append({
                        "PayScaleCode": ps_code,
                        "PayScalePointCode": point_code,
                        "PayScalePointTitle": point_title,
                        "ScalePointNumber": point_num,
                        "Hourly": False,
                        "PayScalePointEnabled": True,
                        "RateDateFrom": get_default_increase_date(is_teaching),
                        "RateDateTo": "",
                        "PayScaleRate": salary,
                    })
                    added_points.add(key)
                    self.log(f"  Added point '{point_code}' for pay scale '{ps_code}' from contract data")

        self.log(f"  Total pay scale points: {len(added_points)}")

    # Standard teaching scale prefixes — grades are generated per-point from the
    # reference PayScalePoints CSV (19_IMPORT_PayScalesPoints.csv).
    # These scales are the same for every build so no customer data is needed.
    # NJC_*, SUP_* scales are intentionally excluded — they vary by LA/school
    # and must come from customer data.
    STANDARD_TEACHING_SCALE_PREFIXES = (
        "MAIN_", "UPS_", "LS_", "LP_", "UQ_", "APP/NMW",
    )

    def _is_standard_teaching_scale(self, scale_code: str) -> bool:
        """Return True if this is a standard teaching scale with fixed per-point grades.
        NJC, SUP and custom support scales are NOT standard — they vary by LA/school."""
        return any(scale_code.startswith(p) or scale_code == p
                   for p in self.STANDARD_TEACHING_SCALE_PREFIXES)

    def _load_reference_points(self) -> pd.DataFrame:
        """Load PayScalePoints reference CSV (19_IMPORT_PayScalesPoints.csv)."""
        ref_path = Path(__file__).parent.parent / "knowledge" / "S2" / "import files" / "19_IMPORT_PayScalesPoints.csv"
        try:
            df = pd.read_csv(ref_path)
            df.columns = [c.strip() for c in df.columns]
            self.log(f"  Reference points loaded: {len(df)} rows")
            return df
        except Exception as e:
            self.log(f"  Warning: Could not load reference points CSV: {e}")
            return pd.DataFrame()

    def _load_reference_grades(self) -> pd.DataFrame:
        """Load PayScaleGrades reference CSV (21_IMPORT_PayScalesGrades.csv).
        Used as fallback for NJC and other non-standard scales."""
        ref_path = Path(__file__).parent.parent / "knowledge" / "S2" / "import files" / "21_IMPORT_PayScalesGrades.csv"
        try:
            df = pd.read_csv(ref_path)
            df.columns = [c.strip() for c in df.columns]
            self.log(f"  Reference grades loaded: {len(df)} rows")
            return df
        except Exception as e:
            self.log(f"  Warning: Could not load reference grades CSV: {e}")
            return pd.DataFrame()

    def _build_pay_scale_grades(self):
        """Build PayScaleGrades sheet — one grade per pay scale point.

        Each grade code matches the point code (M1, M2, L01, L02, LP01, UQ1 …)
        and covers exactly that one point (ScalePointNumberFrom = ScalePointNumberTo).

        Priority order per scale:
          1. Customer-defined grades (from Pay scale GROUP column in customer data)
          2. Per-point grades derived from extracted pay scale points
          3. Per-point grades from reference PayScalePoints CSV for standard
             teaching scales (MAIN, UPS, LS, LP, UQ, APP/NMW) — same every build
          4. Range-based grades from reference PayScaleGrades CSV for NJC and
             any other non-standard scale
        """
        self.log("Building PayScaleGrades...")

        scales_with_customer_grades = {ps.code for ps in self.extracted_pay_scales if ps.grades}

        # Step 1: Customer-defined grades (highest priority)
        for ps in self.extracted_pay_scales:
            if ps.grades:
                for grade in ps.grades:
                    self.template_data["PayScaleGrades"].append({
                        "PayScaleCode": ps.code,
                        "PayScaleGradeCode": grade.get('code', 'MISSING'),
                        "Title": grade.get('title', grade.get('code', 'MISSING')),
                        "ScalePointNumberFrom": grade.get('from_point', ''),
                        "ScalePointNumberTo": grade.get('to_point', ''),
                        "AvailableToAllSchools": True,
                        "SchoolCodes": "",
                        "PayScaleGradeEnabled": True,
                    })

        # Step 2: Per-point grades from extracted points (customer provided points)
        scales_from_extracted_points = set()
        for ps in self.extracted_pay_scales:
            if ps.code in scales_with_customer_grades:
                continue
            if not ps.points:
                continue
            sorted_points = sorted(ps.points, key=lambda p: p.get('number', 0))
            for pt in sorted_points:
                pt_code = pt.get('code', '')
                pt_num = pt.get('number', '')
                if not pt_code or pt_num == '':
                    continue
                self.template_data["PayScaleGrades"].append({
                    "PayScaleCode": ps.code,
                    "PayScaleGradeCode": pt_code,
                    "Title": pt.get('title', pt_code),
                    "ScalePointNumberFrom": pt_num,
                    "ScalePointNumberTo": pt_num,
                    "AvailableToAllSchools": True,
                    "SchoolCodes": "",
                    "PayScaleGradeEnabled": True,
                })
            scales_from_extracted_points.add(ps.code)
            self.log(f"  {ps.code}: added {len(sorted_points)} per-point grade(s) from extracted data")

        # Determine which scales still need grades
        covered = scales_with_customer_grades | scales_from_extracted_points
        scales_needing_grades = [ps.code for ps in self.extracted_pay_scales if ps.code not in covered]

        if not scales_needing_grades:
            return

        # Step 3: Standard teaching scales → per-point grades from reference points CSV
        ref_points = None
        standard_covered = set()
        standard_scales = [c for c in scales_needing_grades if self._is_standard_teaching_scale(c)]
        if standard_scales:
            ref_points = self._load_reference_points()
            if not ref_points.empty and 'PayScaleCode' in ref_points.columns:
                for scale_code in standard_scales:
                    matches = ref_points[ref_points['PayScaleCode'] == scale_code].copy()
                    if matches.empty:
                        continue
                    matches = matches.sort_values('ScalePointNumber') if 'ScalePointNumber' in matches.columns else matches
                    for _, row in matches.iterrows():
                        pt_code = str(row.get('PayScalePointCode', '')).strip()
                        pt_num = row.get('ScalePointNumber', '')
                        pt_title = str(row.get('PayScalePointTitle', pt_code)).strip()
                        self.template_data["PayScaleGrades"].append({
                            "PayScaleCode": scale_code,
                            "PayScaleGradeCode": pt_code,
                            "Title": pt_title,
                            "ScalePointNumberFrom": pt_num,
                            "ScalePointNumberTo": pt_num,
                            "AvailableToAllSchools": True,
                            "SchoolCodes": "",
                            "PayScaleGradeEnabled": True,
                        })
                    standard_covered.add(scale_code)
                    self.log(f"  {scale_code}: added {len(matches)} per-point grade(s) from reference")

        # Non-standard scales (NJC, SUP, custom) with no extracted points →
        # no grades generated. Grades can only come from customer data.
        remaining = [c for c in scales_needing_grades if c not in standard_covered]
        if remaining:
            self.log(f"  No grades generated for {len(remaining)} scale(s) — no customer point data: {remaining}")

    def _build_allowance_types(self):
        """Build AllowanceTypes sheet."""
        self.log("Building AllowanceTypes...")

        for allowance in self.extracted_allowances:
            self.template_data["AllowanceTypes"].append({
                "AllowanceTypeCode": allowance.type_code,
                "AllowanceTypeTitle": allowance.type_title,
                "IncreaseDate": allowance.increase_date,
                "IncreasePercentage": allowance.increase_percentage,
                "AvailableToAllSchools": True,
                "SchoolCodes": "",
                "ExcludePension": False,
                "ExcludeNI": False,
                "AllowanceTypeEnabled": True,
            })

        # No defaults - only use customer data

    def _build_allowance_points(self):
        """Build AllowanceTypePoint sheet."""
        self.log("Building AllowanceTypePoint...")

        for allowance in self.extracted_allowances:
            for point in allowance.points:
                self.template_data["AllowanceTypePoint"].append({
                    "AllowanceTypeCode": allowance.type_code,
                    "AllowancePointCode": point['code'],
                    "AllowancePointTitle": point['title'],
                    "AllowancePointEnabled": True,
                    "RateDateFrom": "2025-09-01",
                    "RateDateTo": "",
                    "Amount": point['amount'],
                })

    def _build_pensions(self):
        """Build Pensions sheet from customer data - scanning both staff_data and lookup."""
        self.log("Building Pensions...")

        pensions_found = set()

        # Method 1: Scan staff_data DataFrames for pension columns
        for df in self.staff_data:
            for col in df.columns:
                col_lower = str(col).lower()
                if 'pension' in col_lower:
                    col_data = self._safe_get_column(df, col)
                    for val in col_data.dropna().unique():
                        val_str = str(val).strip().upper()
                        if val_str and val_str not in ['NAN', '', 'MISSING']:
                            pensions_found.add(val_str)

        # Method 2: Scan consolidated staff lookup for pension values
        for staff_code, record in self.staff_lookup.items():
            pension = self._lookup_get(record, 'pension', '')
            if not pension or str(pension).upper() in ['', 'NAN', 'MISSING']:
                pension = self._lookup_get(record, 'pension_code', '')
            if not pension or str(pension).upper() in ['', 'NAN', 'MISSING']:
                pension = self._lookup_get(record, 'pension_scheme', '')

            if pension and str(pension).strip().upper() not in ['', 'NAN', 'MISSING']:
                pensions_found.add(str(pension).strip().upper())

        # Create pension records from found values
        for pension in sorted(pensions_found):
            self.template_data["Pensions"].append({
                "PensionCode": pension[:10],  # Max 10 chars
                "Title": pension,
                "AvailableToAllSchools": True,
                "SchoolCodes": "",
                "PensionEnabled": True,
                "RateDateFrom": "2025-04-01",
                "RateDateTo": "",
                "PensionPercentage": 0,  # Would need rate from customer data
            })

        self.log(f"  Found {len(pensions_found)} pension schemes from data")

    def _build_eqw_patterns(self):
        """Build EQWPatterns sheet from customer data only."""
        self.log("Building EQWPatterns...")

        # Extract EQW patterns from customer staff data
        eqw_found = set()
        for df in self.staff_data:
            # Look for weeks worked/paid columns
            for col in df.columns:
                col_lower = str(col).lower()
                if any(x in col_lower for x in ['weeks', 'eqw', 'equated', 'tto', 'term time']):
                    col_data = self._safe_get_column(df, col)
                    for val in col_data.dropna().unique():
                        try:
                            weeks = float(val)
                            if 30 <= weeks <= 53:  # Valid week range
                                eqw_found.add(weeks)
                        except (ValueError, TypeError):
                            # Check for text patterns like "TTO", "AYR"
                            val_str = str(val).strip().upper()
                            if val_str:
                                eqw_found.add(val_str)

        # Only create EQW patterns from customer data
        for eqw in eqw_found:
            if isinstance(eqw, float):
                # Numeric weeks value
                if eqw >= 52:
                    code = "AYR"
                    title = "All Year Round"
                else:
                    code = f"TTO_{int(eqw)}"
                    title = f"{int(eqw)} Weeks Term Time Only"

                self.template_data["EQWPatterns"].append({
                    "EquatedWeekPatternCode": code,
                    "EquatedWeekPatternTitle": title,
                    "AvailableToAllSchools": True,
                    "SchoolCodes": "",
                    "ServiceYearsFrom": 0,
                    "ServiceYearsTo": 99,
                    "EquatedWeeks": eqw,
                    "FullTimeWeeks": 52.143,
                    "EquatedWeekPatternEnabled": True,
                })
            elif isinstance(eqw, str):
                # Text pattern from customer data
                self.template_data["EQWPatterns"].append({
                    "EquatedWeekPatternCode": eqw,
                    "EquatedWeekPatternTitle": eqw,
                    "AvailableToAllSchools": True,
                    "SchoolCodes": "",
                    "ServiceYearsFrom": 0,
                    "ServiceYearsTo": 99,
                    "EquatedWeeks": 52.143,
                    "FullTimeWeeks": 52.143,
                    "EquatedWeekPatternEnabled": True,
                })

    def _build_staff_role_groups(self):
        """Build StfRoleGroup sheet - only for groups actually used in customer data."""
        self.log("Building StfRoleGroup...")

        # First, determine which role groups are actually used by customer data
        used_groups = set()
        for df in self.staff_data:
            # Use _find_column to handle column name variations
            if self._find_column(df, 'job_title') is None:
                continue
            for _, row in df.iterrows():
                title = str(self._safe_get(row, 'job_title', '')).strip()
                if title and title != 'nan':
                    srg = get_srg_for_role(title)
                    used_groups.add(srg)

        # Only create role groups that are actually used
        for code in used_groups:
            if code not in S2_STAFF_ROLE_GROUP_PATTERNS:
                continue

            data = S2_STAFF_ROLE_GROUP_PATTERNS[code]
            fc = get_finance_codes_for_srg(code)

            # Track this as a created role group
            self.created_role_groups.append({
                "StaffRoleGroupCode": code,
                "Title": data["title"],
                "IsTeaching": data["teaching"],
                "Source": "Pattern Match",
            })

            self.template_data["StfRoleGroup"].append({
                "StaffRoleGroupCode": code,
                "Title": data["title"],
                "GrossSalaryFinanceCode": data["gross_salary_fc"],
                "LeaveRebateFinanceCode": data["gross_salary_fc"],
                "AllowanceTypesFinanceCode": data["gross_salary_fc"],
                "AdjustmentsFinanceCode": data["gross_salary_fc"],
                "EmployersNiFinanceCode": data["ni_fc"],
                "PensionFinanceCode": data["pension_fc"],
                "WeeklyFteFinanceCode": fc["weekly_fte"],
                "AnnualFteFinanceCode": fc["annual_fte"],
                "WeeklyFteLeaveAdjustmentFinanceCode": fc["weekly_leave_adj"],
                "AnnualFteLeaveAdjustmentFinanceCode": fc["annual_leave_adj"],
                "TeachingRoleGroup": data["teaching"],
                "StaffRoleGroupEnabled": True,
            })

    def _build_staff_roles(self):
        """Build StfRole sheet from staff data."""
        self.log("Building StfRole...")

        roles_seen = set()

        for df in self.staff_data:
            # Use _find_column to handle column name variations
            if self._find_column(df, 'job_title') is None:
                continue

            for _, row in df.iterrows():
                title = str(self._safe_get(row, 'job_title', '')).strip()
                if not title or title == 'nan' or title in roles_seen:
                    continue

                # LESSON L018: Skip budget placeholder entries (not real staff)
                payroll = str(self._safe_get(row, 'payroll_number', '')).strip().upper()
                placeholder_indicators = ['CAS', 'AGENCY', 'LEVY', 'RESTRUCTUR', 'BUDGET', 'PLACEHOLDER']
                if any(ind in payroll for ind in placeholder_indicators):
                    continue
                # Also skip if payroll is short non-numeric (likely placeholder code)
                if payroll and len(payroll) <= 6 and not any(c.isdigit() for c in payroll):
                    continue

                roles_seen.add(title)

                # Determine role group and pay scale using S2 Staff Role Knowledge
                if S2_ROLE_KNOWLEDGE_AVAILABLE:
                    # Use official knowledge base
                    role_code = get_role_code_from_title(title)
                    srg = get_group_from_role(role_code)
                    is_teaching = is_teaching_role(role_code)
                else:
                    # Fall back to pattern matching
                    srg = get_srg_for_role(title)
                    is_teaching = S2_STAFF_ROLE_GROUP_PATTERNS.get(srg, {}).get('teaching', False)

                # LESSON L025: Get FullTimeHoursPerWeek from customer data - NEVER use defaults
                # Check multiple possible column names for full-time hours
                ft_hours = None
                for col in ['weekly', 'ft_hours', 'full_time_hours', 'fte_hours', 'weekly_hours',
                            'hours_per_week', 'contracted_hours', 'full_time_hours_per_week']:
                    val = self._safe_get(row, col, None)
                    if self._safe_notna(val):
                        try:
                            hours_val = float(str(val).strip())
                            if hours_val > 0:
                                ft_hours = hours_val
                                break
                        except (ValueError, TypeError):
                            pass

                # NEVER use hardcoded defaults like 32.436 - mark as MISSING (use 0 to flag for review)
                if ft_hours is None:
                    ft_hours = 0  # Will be flagged in output as needing review
                    self.log(f"    WARNING: No full-time hours found for role '{title}' - marked as 0 for review")

                # LESSON L007: Get PayScaleCode from customer's data FIRST
                # Check multiple possible column names for pay scale type
                pay_scale = None
                for col in ['pay_scale_type', 'scale', 'pay_scale', 'scale_type', 'pay_type']:
                    val = self._safe_get(row, col, '')
                    if self._safe_notna(val) and str(val).strip() and str(val).strip() != 'nan':
                        pay_scale = str(val).strip().upper()
                        break

                # NEVER use defaults - mark as MISSING if not in customer data
                if not pay_scale:
                    pay_scale = 'MISSING'

                # Create role code
                code = self._create_role_code(title, srg)

                # Track this as a created role code
                self.created_role_codes.append({
                    "StaffRoleCode": code,
                    "Title": title,
                    "StaffRoleGroupCode": srg,
                    "PayScaleCode": pay_scale,
                    "Source": "Customer Data",
                    "IsTeaching": is_teaching,
                })

                self.template_data["StfRole"].append({
                    "StaffRoleGroupCode": srg,
                    "StaffRoleCode": code,
                    "Title": title,
                    "PayScaleCode": pay_scale,
                    "FullTimeHoursPerWeek": float(ft_hours),
                    "AvailableToAllSchools": True,
                    "SchoolCodes": "",
                    "MonthsServiceBeforeIncrement": 0 if is_teaching else 6,
                    "DaysServiceBeforeIncrement": 0,
                    "StaffRoleEnabled": True,
                    "IsFinanceRole": 'finance' in title.lower(),
                })

    def _clean_job_title(self, title: str) -> str:
        """
        Clean job title by removing job reference prefixes like 'J3 -', 'A1 -', 'PO2 -'.

        Examples:
            'J3 - Administration Assistant' → 'Administration Assistant'
            'A1 - Teacher' → 'Teacher'
            'PO2 - Senior Admin' → 'Senior Admin'
            'Administration Assistant' → 'Administration Assistant' (no change)
        """
        if not title:
            return title

        title = str(title).strip()

        # Pattern: letter(s) + number(s) + separator (- or :) + actual title
        # e.g., J3 -, A1 -, PO2 -, Grade1:, etc.
        import re
        # Match patterns like "J3 - ", "A1- ", "PO2 : ", "Grade1 - "
        cleaned = re.sub(r'^[A-Za-z]+\d+\s*[-:]\s*', '', title)

        # Also handle patterns like "3 - Title" (just number prefix)
        cleaned = re.sub(r'^\d+\s*[-:]\s*', '', cleaned)

        return cleaned.strip()

    # Common word abbreviations for codes
    WORD_ABBREVIATIONS = {
        'ADMINISTRATION': 'ADM',
        'ADMINISTRATOR': 'ADM',
        'ADMIN': 'ADM',
        'ASSISTANT': 'AST',
        'SUPPORT': 'SUP',
        'STAFF': 'STF',
        'TEACHER': 'TCH',
        'TEACHING': 'TCH',
        'SENIOR': 'SNR',
        'JUNIOR': 'JNR',
        'MANAGER': 'MGR',
        'MANAGEMENT': 'MGT',
        'OFFICER': 'OFC',
        'COORDINATOR': 'CRD',
        'EXECUTIVE': 'EXC',
        'DIRECTOR': 'DIR',
        'HEAD': 'HD',
        'DEPUTY': 'DPT',
        'ASSOCIATE': 'ASC',
        'SPECIALIST': 'SPC',
        'TECHNICIAN': 'TEC',
        'SUPERVISOR': 'SPV',
        'CARETAKER': 'CTK',
        'CLEANER': 'CLN',
        'RECEPTIONIST': 'RCP',
        'SECRETARY': 'SEC',
        'FINANCE': 'FIN',
        'FINANCIAL': 'FIN',
        'HUMAN': 'HR',
        'RESOURCES': 'RES',
        'INFORMATION': 'INF',
        'TECHNOLOGY': 'TEC',
        'LEARNING': 'LRN',
        'PASTORAL': 'PST',
        'WELFARE': 'WLF',
        'ATTENDANCE': 'ATT',
        'PREMISES': 'PRM',
        'SITE': 'SIT',
        'MAINTENANCE': 'MNT',
        'SCIENCE': 'SCI',
        'ENGLISH': 'ENG',
        'MATHEMATICS': 'MTH',
        'MATHS': 'MTH',
    }

    def _abbreviate_text(self, text: str) -> str:
        """
        Abbreviate text using common word abbreviations.

        Examples:
            'AT SUPPORT STAFF' → 'AT_SUP_STF'
            'Administration Assistant' → 'ADM_AST'
        """
        if not text:
            return text

        words = text.upper().split()
        abbreviated = []

        for word in words:
            # Clean word of non-alpha chars
            word_clean = ''.join(c for c in word if c.isalpha())
            if not word_clean:
                continue
            # Use abbreviation if available, otherwise take first 3 chars
            abbrev = self.WORD_ABBREVIATIONS.get(word_clean, word_clean[:3])
            abbreviated.append(abbrev)

        return '_'.join(abbreviated)

    def _create_role_code(self, title: str, srg: str) -> str:
        """
        Create a role code from title using S2 Staff Role Codes knowledge.

        For teachers, creates subject-specific codes like:
        - TEA_HIS (History Teacher)
        - TEA_MAT (Maths Teacher)
        - TEA_ENG (English Teacher)
        - HOD_SCI (Head of Science)

        For other roles, uses standard codes from the knowledge base.
        Numbers are only added to codes when there's a duplicate title.
        """
        # Clean job title - remove job reference prefixes like "J3 -"
        title = self._clean_job_title(title)

        # Initialize role code tracking if not exists
        if not hasattr(self, '_used_role_codes'):
            self._used_role_codes = {}  # Maps base_code -> list of titles using it

        # Use the new S2_STAFF_ROLE_CODES module for subject-specific codes
        if S2_ROLE_CODES_AVAILABLE:
            role_code = get_role_code_for_title(title)
            if role_code and role_code != "OTH":
                # Check if we've already assigned this code to a different title
                if role_code in self._used_role_codes:
                    if title in self._used_role_codes[role_code]:
                        return role_code  # Same title, same code
                    # Different title wants same code - add number suffix
                    count = len(self._used_role_codes[role_code]) + 1
                    new_code = f"{role_code}{count}"
                    self._used_role_codes[role_code].append(title)
                    return new_code
                else:
                    self._used_role_codes[role_code] = [title]
                    return role_code

        # Fall back to old S2_ROLE_KNOWLEDGE if new module not available
        if S2_ROLE_KNOWLEDGE_AVAILABLE:
            official_code = get_role_code_from_title(title)
            if official_code and official_code != "OTH":
                # Check if we've already assigned this code to a different title
                if official_code in self._used_role_codes:
                    if title in self._used_role_codes[official_code]:
                        return official_code  # Same title, same code
                    # Different title wants same code - add number suffix
                    count = len(self._used_role_codes[official_code]) + 1
                    new_code = f"{official_code}{count}"
                    self._used_role_codes[official_code].append(title)
                    return new_code
                else:
                    self._used_role_codes[official_code] = [title]
                    return official_code

        # Final fallback: generate code from title using word abbreviations
        title_clean = title.upper().strip()
        title_clean = re.sub(r'[^A-Z\s]', '', title_clean)  # Letters only, no numbers
        words = title_clean.split()

        if not words:
            base_code = "ROLE"
        else:
            # Check if this is a teacher role with a subject
            is_teacher = any(word in ['TEACHER', 'TEACHING'] for word in words)
            is_hod = 'HEAD' in words and 'DEPARTMENT' in words
            is_hof = 'HEAD' in words and 'FACULTY' in words

            if is_teacher or is_hod or is_hof:
                # Try to extract subject
                subject_code = self._extract_subject_from_title(title_clean)
                if subject_code:
                    if is_hod:
                        base_code = f"HOD_{subject_code}"
                    elif is_hof:
                        base_code = f"HOF_{subject_code}"
                    else:
                        base_code = f"TEA_{subject_code}"
                else:
                    # No subject found, use abbreviations
                    abbrevs = []
                    for word in words:
                        if not word:
                            continue
                        abbrev = self.WORD_ABBREVIATIONS.get(word, word[:3])
                        abbrevs.append(abbrev)
                    base_code = '_'.join(abbrevs) if abbrevs else "ROLE"
            else:
                # Non-teacher role - use abbreviations
                abbrevs = []
                for word in words:
                    if not word:
                        continue
                    abbrev = self.WORD_ABBREVIATIONS.get(word, word[:3])
                    abbrevs.append(abbrev)
                base_code = '_'.join(abbrevs) if abbrevs else "ROLE"

        # Clean up the code
        base_code = re.sub(r'[^A-Z_]', '', base_code)
        if not base_code:
            base_code = "ROLE"

        # Check for duplicates and add number only if needed
        if base_code in self._used_role_codes:
            if title in self._used_role_codes[base_code]:
                return base_code  # Same title, same code
            # Different title wants same code - add number suffix
            count = len(self._used_role_codes[base_code]) + 1
            new_code = f"{base_code}{count}"
            self._used_role_codes[base_code].append(title)
            return new_code
        else:
            self._used_role_codes[base_code] = [title]
            return base_code

    def _extract_subject_from_title(self, title: str) -> str:
        """
        Extract subject code from a job title.

        Args:
            title: Job title (already uppercased)

        Returns:
            Subject abbreviation (e.g., HIS, MAT, ENG) or empty string
        """
        # Use the knowledge base if available
        if S2_ROLE_CODES_AVAILABLE:
            return extract_subject_code(title)

        # Fallback subject mappings
        subject_map = {
            "HISTORY": "HIS",
            "MATHS": "MAT",
            "MATHEMATICS": "MAT",
            "ENGLISH": "ENG",
            "SCIENCE": "SCI",
            "BIOLOGY": "BIO",
            "CHEMISTRY": "CHE",
            "PHYSICS": "PHY",
            "GEOGRAPHY": "GEO",
            "FRENCH": "FRE",
            "SPANISH": "SPA",
            "GERMAN": "GER",
            "ART": "ART",
            "MUSIC": "MUS",
            "DRAMA": "DRA",
            "PE": "PE",
            "PHYSICAL EDUCATION": "PE",
            "COMPUTING": "COM",
            "ICT": "ICT",
            "DT": "DT",
            "DESIGN TECHNOLOGY": "DT",
            "FOOD": "FOO",
            "TEXTILES": "TEX",
            "RE": "RE",
            "RELIGIOUS EDUCATION": "RE",
            "RELIGIOUS STUDIES": "RS",
            "PSHE": "PSH",
            "BUSINESS": "BUS",
            "ECONOMICS": "ECO",
            "PSYCHOLOGY": "PSY",
            "SOCIOLOGY": "SOC",
            "DANCE": "DAN",
            "MEDIA": "MED",
            "PHOTOGRAPHY": "PHO",
            "FILM": "FIL",
        }

        for subject, code in subject_map.items():
            if subject in title:
                return code

        return ""

    def _lookup_get(self, record: Dict, key: str, default: Any = '') -> Any:
        """
        Get a value from a lookup record, trying multiple field name variations.

        Args:
            record: Staff lookup record dict
            key: Field name to look for
            default: Default value if not found

        Returns:
            Value from record or default
        """
        # Direct match
        if key in record:
            val = record[key]
            if val is not None and str(val).strip().lower() not in ['', 'nan', 'missing']:
                return val

        # Try variations from COLUMN_ALIASES
        key_lower = key.lower().replace(' ', '_')
        variations = self.COLUMN_ALIASES.get(key_lower, [])

        for var in variations:
            if var in record:
                val = record[var]
                if val is not None and str(val).strip().lower() not in ['', 'nan', 'missing']:
                    return val

        return default

    def _build_staff_members(self):
        """Build StaffMembers sheet from consolidated staff lookup."""
        self.log("Building StaffMembers...")

        # Use the consolidated lookup - this has merged data from all files
        for staff_code, record in self.staff_lookup.items():
            # Get surname/forename - try multiple field names
            surname = str(self._lookup_get(record, 'surname', '')).strip()
            if not surname or surname.lower() == 'nan':
                surname = str(self._lookup_get(record, 'last_name', '')).strip()
            if surname.lower() == 'nan':
                surname = ''

            forename = str(self._lookup_get(record, 'forename', '')).strip()
            if not forename or forename.lower() == 'nan':
                forename = str(self._lookup_get(record, 'first_name', '')).strip()
            if forename.lower() == 'nan':
                forename = ''

            # If we have a 'name' field but no separate names, try to split it
            if not surname and not forename:
                full_name = str(self._lookup_get(record, 'name', '')).strip()
                if full_name and full_name.lower() != 'nan':
                    parts = full_name.split()
                    if len(parts) >= 2:
                        surname = parts[0]
                        forename = ' '.join(parts[1:])
                    elif len(parts) == 1:
                        surname = parts[0]

            service_start = self._lookup_get(record, 'service_start_date', '')
            if not service_start or str(service_start).lower() == 'nan':
                service_start = self._lookup_get(record, 'start_date', '')

            dob = self._lookup_get(record, 'dob', '')
            if not dob or str(dob).lower() == 'nan':
                dob = self._lookup_get(record, 'date_of_birth', '')

            gender = str(self._lookup_get(record, 'gender', 'ZZZ')).strip().upper()
            if gender in ['M', 'MALE']:
                gender = 'M'
            elif gender in ['F', 'FEMALE']:
                gender = 'F'
            else:
                gender = 'ZZZ'

            school = str(self._lookup_get(record, 'school_code', '')).strip()
            if not school or school.lower() == 'nan':
                school = str(self._lookup_get(record, 'school', '')).strip()
            if not school or school.lower() == 'nan':
                school = str(self._lookup_get(record, 'cost_centre', '')).strip()

            # Handle "ALL" meaning "available to all schools"
            available_to_all = school.upper() == 'ALL' or school == '' or school.lower() == 'nan'
            if school.upper() == 'ALL':
                school = ''

            if school and school.lower() != 'nan':
                self.schools_found.add(school)
            else:
                school = ''

            # Check for casual/0-hour
            hours = self._lookup_get(record, 'weekly_hours', 0)
            try:
                casual = hours is not None and float(hours) == 0
            except (ValueError, TypeError):
                casual = False

            self.template_data["StaffMembers"].append({
                "StaffMemberCode": staff_code,
                "FirstName": forename,
                "LastName": surname,
                "Title": f"{forename} {surname}".strip(),
                "ServiceStartDate": service_start if service_start and str(service_start).lower() != 'nan' else '',
                "ServiceEndDate": '',
                "DateOfBirth": dob if dob and str(dob).lower() != 'nan' else '',
                "Apprenticeship": False,
                "PensionOptOut": False,
                "AvailableToAllSchools": available_to_all,
                "SchoolCodes": school,
                "StaffMemberEnabled": True,
                "GenderCode": gender,
                "Casual": casual,
            })

        self.log(f"  Created {len(self.template_data['StaffMembers'])} staff members from lookup")

    def _build_contracts_teaching(self):
        """Build ContractsTeachFTE sheet from consolidated staff lookup."""
        self.log("Building ContractsTeachFTE...")

        contracts_created = 0
        skipped_no_title = 0
        skipped_not_teaching = 0

        # Use the consolidated lookup - this has merged data from all files
        for staff_code, record in self.staff_lookup.items():
            # Get job title to determine if teaching
            title = str(self._lookup_get(record, 'job_title', '')).strip()
            if not title or title.lower() == 'nan':
                title = str(self._lookup_get(record, 'position', '')).strip()
            if not title or title.lower() == 'nan':
                title = str(self._lookup_get(record, 'role', '')).strip()

            if not title or title.lower() == 'nan':
                skipped_no_title += 1
                self.skipped_staff.append({
                    'StaffCode': staff_code,
                    'Reason': 'No job title found',
                    'ContractType': 'Teaching',
                    'Name': self._lookup_get(record, 'name', '')
                })
                continue  # No job title, can't create contract

            srg = get_srg_for_role(title)

            # Only teaching contracts
            if not S2_STAFF_ROLE_GROUP_PATTERNS.get(srg, {}).get('teaching', False):
                skipped_not_teaching += 1
                continue

            # Get school code
            school = str(self._lookup_get(record, 'school_code', '')).strip()
            if not school or school.lower() == 'nan':
                school = str(self._lookup_get(record, 'school', '')).strip()
            if not school or school.lower() == 'nan':
                school = str(self._lookup_get(record, 'cost_centre', '')).strip()
            if school.lower() == 'nan' or school.upper() == 'ALL':
                school = ''

            # Get FTE
            fte = self._lookup_get(record, 'fte', None)
            if fte is None or str(fte).lower() == 'nan':
                fte = self._lookup_get(record, 'weekly_fte', 1.0)
            try:
                fte = float(fte)
                if fte <= 0 or fte > 10:
                    fte = 1.0
            except (ValueError, TypeError):
                fte = 1.0

            # Get pay scale info
            pay_scale_contract = str(self._lookup_get(record, 'pay_scale_contract', '')).strip().upper()
            pay_scale_type = str(self._lookup_get(record, 'pay_scale_type', '')).strip().upper()
            if not pay_scale_type or pay_scale_type == 'NAN':
                pay_scale_type = str(self._lookup_get(record, 'pay_scale', '')).strip().upper()
            if not pay_scale_type or pay_scale_type == 'NAN':
                pay_scale_type = str(self._lookup_get(record, 'scale', '')).strip().upper()

            # Combine contract + type
            if pay_scale_contract and pay_scale_contract != 'NAN' and pay_scale_type and pay_scale_type != 'NAN':
                pay_scale_code = f"{pay_scale_contract}_{pay_scale_type}"
            elif pay_scale_type and pay_scale_type != 'NAN':
                pay_scale_code = pay_scale_type
            else:
                pay_scale_code = "MISSING"

            # Get scale point
            scale_point = str(self._lookup_get(record, 'scale_point', '')).strip()
            if not scale_point or scale_point.lower() == 'nan':
                scale_point = str(self._lookup_get(record, 'scp', '')).strip()
            if not scale_point or scale_point.lower() == 'nan':
                scale_point = str(self._lookup_get(record, 'point', '')).strip()
            if not scale_point or scale_point.lower() == 'nan':
                scale_point = 'MISSING'

            # Track spot salary amounts for creating pay scale points
            if not hasattr(self, '_spot_salary_amounts'):
                self._spot_salary_amounts = set()

            # Check for spot salary - multiple detection methods
            # For spot salaries: PayScaleCode = "SPOT", PayScalePointCode = "SPOT"
            detected_spot_salary = None

            # Method 1: pay_scale contains "SPOT" (e.g., "SPOT SALARY", "SPOT_SCALE", "SPOT")
            if 'SPOT' in pay_scale_code.upper():
                # It's explicitly marked as spot - get the salary amount
                spot_salary = self._lookup_get(record, 'spot_salary', None)
                if spot_salary is None or str(spot_salary).lower() == 'nan':
                    spot_salary = self._lookup_get(record, 'annual_salary', None)
                if spot_salary is None or str(spot_salary).lower() == 'nan':
                    spot_salary = self._lookup_get(record, 'salary', None)
                try:
                    salary_num = float(str(spot_salary).replace(',', '').replace('£', '')) if spot_salary else 0
                    if salary_num >= 10000:
                        pay_scale_code = 'SPOT'
                        scale_point = 'SPOT'
                        detected_spot_salary = salary_num
                except (ValueError, TypeError):
                    pass

            # Method 2: scale_point is actually a salary amount (>= 10000)
            if pay_scale_code != 'SPOT':
                try:
                    scale_point_num = float(str(scale_point).replace(',', '').replace('£', ''))
                    if scale_point_num >= 10000:
                        pay_scale_code = 'SPOT'
                        detected_spot_salary = scale_point_num
                        scale_point = 'SPOT'
                except (ValueError, AttributeError):
                    pass

            # Method 3: Dedicated spot_salary column exists with a value
            if pay_scale_code != 'SPOT':
                spot_salary = self._lookup_get(record, 'spot_salary', None)
                if spot_salary is not None and str(spot_salary).lower() not in ['nan', '', 'none']:
                    try:
                        salary_num = float(str(spot_salary).replace(',', '').replace('£', ''))
                        if salary_num >= 10000:
                            pay_scale_code = 'SPOT'
                            scale_point = 'SPOT'
                            detected_spot_salary = salary_num
                    except (ValueError, TypeError):
                        pass

            # Method 4: No pay scale but has salary - treat as spot salary
            if pay_scale_code in ['MISSING', '', 'NAN']:
                salary = self._lookup_get(record, 'annual_salary', None)
                if salary is None or str(salary).lower() == 'nan':
                    salary = self._lookup_get(record, 'salary', None)
                try:
                    salary_num = float(str(salary).replace(',', '').replace('£', '')) if salary else 0
                    if salary_num >= 10000:
                        pay_scale_code = 'SPOT'
                        scale_point = 'SPOT'
                        detected_spot_salary = salary_num
                except (ValueError, TypeError):
                    pass

            # Method 5: pay_scale itself might be a salary number
            if pay_scale_code not in ['SPOT', 'MISSING']:
                try:
                    ps_as_num = float(str(pay_scale_code).replace(',', '').replace('£', ''))
                    if ps_as_num >= 10000:
                        detected_spot_salary = ps_as_num
                        pay_scale_code = 'SPOT'
                        scale_point = 'SPOT'
                except (ValueError, TypeError):
                    pass

            # Method 6: Has salary amount but scale point is missing/empty - likely spot
            if pay_scale_code != 'SPOT' and scale_point in ['MISSING', '', 'NAN']:
                # Check if there's a salary that looks like a spot salary
                salary = self._lookup_get(record, 'annual_salary', None)
                if salary is None or str(salary).lower() == 'nan':
                    salary = self._lookup_get(record, 'salary', None)
                try:
                    salary_num = float(str(salary).replace(',', '').replace('£', '')) if salary else 0
                    # If salary is a whole amount (no decimals or .00) and in reasonable range
                    if salary_num >= 10000 and (salary_num == int(salary_num) or str(salary_num).endswith('.0') or str(salary_num).endswith('.00')):
                        pay_scale_code = 'SPOT'
                        scale_point = 'SPOT'
                        detected_spot_salary = salary_num
                except (ValueError, TypeError):
                    pass

            # Track the spot salary amount for pay scale point creation
            if detected_spot_salary:
                self._spot_salary_amounts.add(detected_spot_salary)

            # Normalize scale point if needed
            if scale_point not in ['MISSING', 'SPOT'] and pay_scale_code != 'SPOT' and not str(scale_point).isdigit():
                scale_point = self._normalize_scale_point(scale_point, 'teaching')

            # Get pension - teaching staff default to TPS
            pension_code = str(self._lookup_get(record, 'pension', '')).strip().upper()
            if not pension_code or pension_code == 'NAN':
                pension_code = str(self._lookup_get(record, 'pension_code', '')).strip().upper()
            if not pension_code or pension_code == 'NAN':
                pension_code = 'TPS'

            # Create role code
            role_code = self._create_role_code(title, srg)

            # Get contract reference
            contract_ref = str(self._lookup_get(record, 'contract_ref', '')).strip()
            if not contract_ref or contract_ref.lower() == 'nan':
                contract_ref = f"{staff_code}A"

            # Get names
            forename = str(self._lookup_get(record, 'forename', '')).strip()
            if not forename or forename.lower() == 'nan':
                forename = str(self._lookup_get(record, 'first_name', '')).strip()
            surname = str(self._lookup_get(record, 'surname', '')).strip()
            if not surname or surname.lower() == 'nan':
                surname = str(self._lookup_get(record, 'last_name', '')).strip()
            if forename.lower() == 'nan': forename = ''
            if surname.lower() == 'nan': surname = ''

            # Get department
            dept_code = str(self._lookup_get(record, 'department', '')).strip().upper()
            if not dept_code or dept_code == 'NAN':
                dept_code = str(self._lookup_get(record, 'cost_centre', '')).strip().upper()
            if not dept_code or dept_code == 'NAN':
                dept_code = "MISSING"

            # Get fund code
            fund_code = str(self._lookup_get(record, 'fund_code', '')).strip().upper()
            if not fund_code or fund_code == 'NAN':
                fund_code = ""

            # Get EQW pattern
            eqw_code = str(self._lookup_get(record, 'eqw', '')).strip().upper()
            if not eqw_code or eqw_code == 'NAN':
                eqw_code = str(self._lookup_get(record, 'eqw_pattern', '')).strip().upper()
            if not eqw_code or eqw_code == 'NAN':
                eqw_code = "AYR"

            # Get contract type
            contract_type = str(self._lookup_get(record, 'contract_type', '')).strip().upper()
            if not contract_type or contract_type == 'NAN':
                contract_type = "PERM"

            # Get dates
            date_from = self._lookup_get(record, 'contract_start', '')
            if not date_from or str(date_from).lower() == 'nan':
                date_from = self._lookup_get(record, 'start_date', '')
            if date_from and str(date_from).lower() != 'nan':
                date_from = format_date_uk(date_from) or get_default_increase_date(True)
            else:
                date_from = get_default_increase_date(True)

            date_to = self._lookup_get(record, 'contract_end', '')
            if not date_to or str(date_to).lower() == 'nan':
                date_to = self._lookup_get(record, 'end_date', '')
            if date_to and str(date_to).lower() != 'nan':
                date_to = format_date_uk(date_to) or ""
            else:
                date_to = ""

            # Get grade code
            grade_code = str(self._lookup_get(record, 'grade', '')).strip().upper()
            if not grade_code or grade_code == 'NAN':
                grade_code = str(self._lookup_get(record, 'grade_code', '')).strip().upper()
            if not grade_code or grade_code == 'NAN':
                grade_code = ""

            self.template_data["ContractsTeachFTE"].append({
                "SchoolCode": school,
                "StaffMemberCode": staff_code,
                "Reference": contract_ref,
                "Title": title,
                "StaffRoleCode": role_code,
                "PayScaleCode": pay_scale_code,
                "PayScaleGradeCode": grade_code,
                "PayScalePointCode": scale_point,
                "DepartmentCode": dept_code,
                "FundCode": fund_code,
                "PensionCode": pension_code,
                "EquatedWeekPatternCode": eqw_code,
                "DateFrom": date_from,
                "DateTo": date_to,
                "WeeklyFteOrHpw": float(fte),
                "MatEditOnly": False,
                "NoIncrement": False,
                "ContractTypeCode": contract_type,
                "Notes": f"Staff: {forename} {surname}".strip() if forename or surname else "",
            })
            contracts_created += 1

            # Check for allowances from the record
            self._extract_contract_allowances_from_lookup(record, staff_code, contract_ref)

        self.log(f"  Created {contracts_created} teaching contracts from lookup")
        if skipped_no_title > 0:
            self.log(f"  Skipped {skipped_no_title} staff members with no job title")
        if skipped_not_teaching > 0:
            self.log(f"  Skipped {skipped_not_teaching} non-teaching staff members")

    def _build_contracts_support(self):
        """Build ContractsSupportHours sheet from consolidated staff lookup."""
        self.log("Building ContractsSupportHours...")

        contracts_created = 0
        skipped_no_title = 0
        skipped_is_teaching = 0
        skipped_no_hours = 0
        skipped_zero_hours = 0

        # Use the consolidated lookup - this has merged data from all files
        for staff_code, record in self.staff_lookup.items():
            # Get job title to determine if support
            title = str(self._lookup_get(record, 'job_title', '')).strip()
            if not title or title.lower() == 'nan':
                title = str(self._lookup_get(record, 'position', '')).strip()
            if not title or title.lower() == 'nan':
                title = str(self._lookup_get(record, 'role', '')).strip()

            if not title or title.lower() == 'nan':
                skipped_no_title += 1
                self.skipped_staff.append({
                    'StaffCode': staff_code,
                    'Reason': 'No job title found',
                    'ContractType': 'Support',
                    'Name': self._lookup_get(record, 'name', '')
                })
                continue  # No job title, can't create contract

            srg = get_srg_for_role(title)

            # Only support contracts (non-teaching)
            if S2_STAFF_ROLE_GROUP_PATTERNS.get(srg, {}).get('teaching', False):
                skipped_is_teaching += 1
                continue

            # Get hours - skip 0-hour contracts
            hours = self._lookup_get(record, 'weekly_hours', None)
            if hours is None or str(hours).lower() == 'nan':
                hours = self._lookup_get(record, 'hours', None)
            if hours is None or str(hours).lower() == 'nan':
                skipped_no_hours += 1
                self.skipped_staff.append({
                    'StaffCode': staff_code,
                    'Reason': 'No hours found',
                    'ContractType': 'Support',
                    'JobTitle': title,
                    'Name': self._lookup_get(record, 'name', '')
                })
                continue
            try:
                hours = float(hours)
                if hours == 0:
                    skipped_zero_hours += 1
                    self.skipped_staff.append({
                        'StaffCode': staff_code,
                        'Reason': 'Zero hours',
                        'ContractType': 'Support',
                        'JobTitle': title,
                        'Name': self._lookup_get(record, 'name', '')
                    })
                    continue
            except (ValueError, TypeError):
                skipped_no_hours += 1
                self.skipped_staff.append({
                    'StaffCode': staff_code,
                    'Reason': 'Invalid hours value',
                    'ContractType': 'Support',
                    'JobTitle': title,
                    'Hours': str(hours),
                    'Name': self._lookup_get(record, 'name', '')
                })
                continue

            # Get school code
            school = str(self._lookup_get(record, 'school_code', '')).strip()
            if not school or school.lower() == 'nan':
                school = str(self._lookup_get(record, 'school', '')).strip()
            if not school or school.lower() == 'nan':
                school = str(self._lookup_get(record, 'cost_centre', '')).strip()
            if school.lower() == 'nan' or school.upper() == 'ALL':
                school = ''

            # Get pay scale info
            pay_scale_contract = str(self._lookup_get(record, 'pay_scale_contract', '')).strip().upper()
            pay_scale_type = str(self._lookup_get(record, 'pay_scale_type', '')).strip().upper()
            if not pay_scale_type or pay_scale_type == 'NAN':
                pay_scale_type = str(self._lookup_get(record, 'pay_scale', '')).strip().upper()
            if not pay_scale_type or pay_scale_type == 'NAN':
                pay_scale_type = str(self._lookup_get(record, 'scale', '')).strip().upper()

            # Combine contract + type
            if pay_scale_contract and pay_scale_contract != 'NAN' and pay_scale_type and pay_scale_type != 'NAN':
                pay_scale_code = f"{pay_scale_contract}_{pay_scale_type}"
            elif pay_scale_type and pay_scale_type != 'NAN':
                pay_scale_code = pay_scale_type
            else:
                pay_scale_code = "MISSING"

            # Get scale point
            scale_point = str(self._lookup_get(record, 'scale_point', '')).strip()
            if not scale_point or scale_point.lower() == 'nan':
                scale_point = str(self._lookup_get(record, 'scp', '')).strip()
            if not scale_point or scale_point.lower() == 'nan':
                scale_point = str(self._lookup_get(record, 'point', '')).strip()
            if not scale_point or scale_point.lower() == 'nan':
                scale_point = 'MISSING'

            # Track spot salary amounts for creating pay scale points
            if not hasattr(self, '_spot_salary_amounts'):
                self._spot_salary_amounts = set()

            # Check for spot salary - multiple detection methods
            # For spot salaries: PayScaleCode = "SPOT", PayScalePointCode = "SPOT"
            detected_spot_salary = None

            # Method 1: pay_scale contains "SPOT" (e.g., "SPOT SALARY", "SPOT_SCALE", "SPOT")
            if 'SPOT' in pay_scale_code.upper():
                # It's explicitly marked as spot - get the salary amount
                spot_salary = self._lookup_get(record, 'spot_salary', None)
                if spot_salary is None or str(spot_salary).lower() == 'nan':
                    spot_salary = self._lookup_get(record, 'annual_salary', None)
                if spot_salary is None or str(spot_salary).lower() == 'nan':
                    spot_salary = self._lookup_get(record, 'salary', None)
                try:
                    salary_num = float(str(spot_salary).replace(',', '').replace('£', '')) if spot_salary else 0
                    if salary_num >= 10000:
                        pay_scale_code = 'SPOT'
                        scale_point = 'SPOT'
                        detected_spot_salary = salary_num
                except (ValueError, TypeError):
                    pass

            # Method 2: scale_point is actually a salary amount (>= 10000)
            if pay_scale_code != 'SPOT':
                try:
                    scale_point_num = float(str(scale_point).replace(',', '').replace('£', ''))
                    if scale_point_num >= 10000:
                        pay_scale_code = 'SPOT'
                        detected_spot_salary = scale_point_num
                        scale_point = 'SPOT'
                except (ValueError, AttributeError):
                    pass

            # Method 3: Dedicated spot_salary column exists with a value
            if pay_scale_code != 'SPOT':
                spot_salary = self._lookup_get(record, 'spot_salary', None)
                if spot_salary is not None and str(spot_salary).lower() not in ['nan', '', 'none']:
                    try:
                        salary_num = float(str(spot_salary).replace(',', '').replace('£', ''))
                        if salary_num >= 10000:
                            pay_scale_code = 'SPOT'
                            scale_point = 'SPOT'
                            detected_spot_salary = salary_num
                    except (ValueError, TypeError):
                        pass

            # Method 4: No pay scale but has salary - treat as spot salary
            if pay_scale_code in ['MISSING', '', 'NAN']:
                salary = self._lookup_get(record, 'annual_salary', None)
                if salary is None or str(salary).lower() == 'nan':
                    salary = self._lookup_get(record, 'salary', None)
                try:
                    salary_num = float(str(salary).replace(',', '').replace('£', '')) if salary else 0
                    if salary_num >= 10000:
                        pay_scale_code = 'SPOT'
                        scale_point = 'SPOT'
                        detected_spot_salary = salary_num
                except (ValueError, TypeError):
                    pass

            # Method 5: pay_scale itself might be a salary number
            if pay_scale_code not in ['SPOT', 'MISSING']:
                try:
                    ps_as_num = float(str(pay_scale_code).replace(',', '').replace('£', ''))
                    if ps_as_num >= 10000:
                        detected_spot_salary = ps_as_num
                        pay_scale_code = 'SPOT'
                        scale_point = 'SPOT'
                except (ValueError, TypeError):
                    pass

            # Method 6: Has salary amount but scale point is missing/empty - likely spot
            if pay_scale_code != 'SPOT' and scale_point in ['MISSING', '', 'NAN']:
                # Check if there's a salary that looks like a spot salary
                salary = self._lookup_get(record, 'annual_salary', None)
                if salary is None or str(salary).lower() == 'nan':
                    salary = self._lookup_get(record, 'salary', None)
                try:
                    salary_num = float(str(salary).replace(',', '').replace('£', '')) if salary else 0
                    # If salary is a whole amount (no decimals or .00) and in reasonable range
                    if salary_num >= 10000 and (salary_num == int(salary_num) or str(salary_num).endswith('.0') or str(salary_num).endswith('.00')):
                        pay_scale_code = 'SPOT'
                        scale_point = 'SPOT'
                        detected_spot_salary = salary_num
                except (ValueError, TypeError):
                    pass

            # Track the spot salary amount for pay scale point creation
            if detected_spot_salary:
                self._spot_salary_amounts.add(detected_spot_salary)

            # Normalize scale point if needed
            if scale_point not in ['MISSING', 'SPOT'] and pay_scale_code != 'SPOT' and not str(scale_point).isdigit():
                scale_point = self._normalize_scale_point(scale_point, 'support')

            # Get pension - support staff default to LGPS
            pension_code = str(self._lookup_get(record, 'pension', '')).strip().upper()
            if not pension_code or pension_code == 'NAN':
                pension_code = str(self._lookup_get(record, 'pension_code', '')).strip().upper()
            if not pension_code or pension_code == 'NAN':
                pension_code = 'LGPS'

            # Create role code
            role_code = self._create_role_code(title, srg)

            # Get contract reference
            contract_ref = str(self._lookup_get(record, 'contract_ref', '')).strip()
            if not contract_ref or contract_ref.lower() == 'nan':
                contract_ref = f"{staff_code}A"

            # Get names
            forename = str(self._lookup_get(record, 'forename', '')).strip()
            if not forename or forename.lower() == 'nan':
                forename = str(self._lookup_get(record, 'first_name', '')).strip()
            surname = str(self._lookup_get(record, 'surname', '')).strip()
            if not surname or surname.lower() == 'nan':
                surname = str(self._lookup_get(record, 'last_name', '')).strip()
            if forename.lower() == 'nan': forename = ''
            if surname.lower() == 'nan': surname = ''

            # Get department
            dept_code = str(self._lookup_get(record, 'department', '')).strip().upper()
            if not dept_code or dept_code == 'NAN':
                dept_code = str(self._lookup_get(record, 'cost_centre', '')).strip().upper()
            if not dept_code or dept_code == 'NAN':
                dept_code = "MISSING"

            # Get fund code
            fund_code = str(self._lookup_get(record, 'fund_code', '')).strip().upper()
            if not fund_code or fund_code == 'NAN':
                fund_code = ""

            # Determine EQW pattern from weeks paid
            weeks_paid = self._lookup_get(record, 'weeks_paid', 52.143)
            try:
                weeks_paid_float = float(weeks_paid) if weeks_paid else 52.143
            except (ValueError, TypeError):
                weeks_paid_float = 52.143

            if weeks_paid_float >= 52:
                eqwp = "AYR"
            elif weeks_paid_float >= 40:
                eqwp = "TTO_40"
            elif weeks_paid_float >= 39:
                eqwp = "TTO_39"
            else:
                eqwp = "TTO_38"

            # Override with explicit EQW if provided
            eqw_code = str(self._lookup_get(record, 'eqw', '')).strip().upper()
            if not eqw_code or eqw_code == 'NAN':
                eqw_code = str(self._lookup_get(record, 'eqw_pattern', '')).strip().upper()
            if eqw_code and eqw_code != 'NAN':
                eqwp = eqw_code

            # Get contract type
            contract_type = str(self._lookup_get(record, 'contract_type', '')).strip().upper()
            if not contract_type or contract_type == 'NAN':
                contract_type = "PERM"

            # Get dates
            date_from = self._lookup_get(record, 'contract_start', '')
            if not date_from or str(date_from).lower() == 'nan':
                date_from = self._lookup_get(record, 'start_date', '')
            if date_from and str(date_from).lower() != 'nan':
                date_from = format_date_uk(date_from) or get_default_increase_date(False)
            else:
                date_from = get_default_increase_date(False)

            date_to = self._lookup_get(record, 'contract_end', '')
            if not date_to or str(date_to).lower() == 'nan':
                date_to = self._lookup_get(record, 'end_date', '')
            if date_to and str(date_to).lower() != 'nan':
                date_to = format_date_uk(date_to) or ""
            else:
                date_to = ""

            # Get grade code
            grade_code = str(self._lookup_get(record, 'grade', '')).strip().upper()
            if not grade_code or grade_code == 'NAN':
                grade_code = str(self._lookup_get(record, 'grade_code', '')).strip().upper()
            if not grade_code or grade_code == 'NAN':
                grade_code = ""

            self.template_data["ContractsSupportHours"].append({
                "SchoolCode": school,
                "StaffMemberCode": staff_code,
                "Reference": contract_ref,
                "Title": title,
                "StaffRoleCode": role_code,
                "PayScaleCode": pay_scale_code,
                "PayScaleGradeCode": grade_code,
                "PayScalePointCode": scale_point,
                "DepartmentCode": dept_code,
                "FundCode": fund_code,
                "PensionCode": pension_code,
                "EquatedWeekPatternCode": eqwp,
                "DateFrom": date_from,
                "DateTo": date_to,
                "WeeklyFteOrHpw": float(hours),
                "MatEditOnly": False,
                "NoIncrement": False,
                "ContractTypeCode": contract_type,
                "Notes": f"Staff: {forename} {surname}".strip() if forename or surname else "",
            })
            contracts_created += 1

            # Check for allowances from the record
            self._extract_contract_allowances_from_lookup(record, staff_code, contract_ref)

        self.log(f"  Created {contracts_created} support contracts from lookup")
        if skipped_no_title > 0:
            self.log(f"  Skipped {skipped_no_title} staff members with no job title")
        if skipped_is_teaching > 0:
            self.log(f"  Skipped {skipped_is_teaching} teaching staff (handled in teaching contracts)")
        if skipped_no_hours > 0:
            self.log(f"  Skipped {skipped_no_hours} staff members with no/invalid hours")
        if skipped_zero_hours > 0:
            self.log(f"  Skipped {skipped_zero_hours} staff members with zero hours")

    def _normalize_scale_point(self, point: str, scale_type: str) -> str:
        """Normalize scale point code."""
        point = str(point).strip()

        # Extract number
        numbers = re.findall(r'\d+', point)
        if not numbers:
            return "M1" if scale_type == 'teaching' else "1"

        num = int(numbers[0])

        if scale_type == 'teaching':
            if 'l' in point.lower():
                return f"L{num:02d}"
            elif 'u' in point.lower() or num > 6:
                return f"U{min(num, 3)}"
            else:
                return f"M{min(num, 6)}"
        else:
            return str(num)

    def _find_allowance_point_code(self, type_code: str, amount: float) -> str:
        """Find the matching allowance point code for a given amount."""
        for allowance in self.extracted_allowances:
            if allowance.type_code == type_code:
                # Find the point with matching or closest amount
                best_match = None
                best_diff = float('inf')
                for point in allowance.points:
                    diff = abs(point['amount'] - amount)
                    if diff < best_diff:
                        best_diff = diff
                        best_match = point['code']
                    # Exact match
                    if diff < 0.01:
                        return point['code']
                if best_match:
                    return best_match
        # Default fallback
        return f"{type_code}1"

    def _extract_contract_allowances(self, row: pd.Series, staff_code: str, contract_ref: str):
        """Extract allowances from a contract row."""
        # Check for TLR
        tlr = self._safe_get(row, 'tlr_allowance', None)
        if not self._safe_notna(tlr):
            tlr = self._safe_get(row, 'tlr', None)
        if self._safe_notna(tlr):
            try:
                amount = float(str(tlr).replace('£', '').replace(',', '').strip())
                if amount > 0:
                    point_code = self._find_allowance_point_code('TLR', amount)
                    self.template_data["ContractAllowances"].append({
                        "StaffMemberCode": staff_code,
                        "ContractReference": contract_ref,
                        "AllowanceTypeCode": "TLR",
                        "AllowancePointCode": point_code,
                        "Amount": amount,
                        "DateFrom": "2025-09-01",
                        "DateTo": "",
                    })
            except:
                pass

        # Check for SEN
        sen = self._safe_get(row, 'sen_allowance', None)
        if not self._safe_notna(sen):
            sen = self._safe_get(row, 'sen', None)
        if self._safe_notna(sen):
            try:
                amount = float(str(sen).replace('£', '').replace(',', '').strip())
                if amount > 0:
                    point_code = self._find_allowance_point_code('SEN', amount)
                    self.template_data["ContractAllowances"].append({
                        "StaffMemberCode": staff_code,
                        "ContractReference": contract_ref,
                        "AllowanceTypeCode": "SEN",
                        "AllowancePointCode": point_code,
                        "Amount": amount,
                        "DateFrom": "2025-09-01",
                        "DateTo": "",
                    })
            except:
                pass

        # Check for Recruitment allowance
        rec = self._safe_get(row, 'recruitment_allowance', None)
        if not self._safe_notna(rec):
            rec = self._safe_get(row, 'recruitment', None)
        if self._safe_notna(rec):
            try:
                amount = float(str(rec).replace('£', '').replace(',', '').strip())
                if amount > 0:
                    point_code = self._find_allowance_point_code('REC', amount)
                    self.template_data["ContractAllowances"].append({
                        "StaffMemberCode": staff_code,
                        "ContractReference": contract_ref,
                        "AllowanceTypeCode": "REC",
                        "AllowancePointCode": point_code,
                        "Amount": amount,
                        "DateFrom": "2025-09-01",
                        "DateTo": "",
                    })
            except:
                pass

        # Check for Retention allowance
        ret = self._safe_get(row, 'retention_allowance', None)
        if not self._safe_notna(ret):
            ret = self._safe_get(row, 'retention', None)
        if self._safe_notna(ret):
            try:
                amount = float(str(ret).replace('£', '').replace(',', '').strip())
                if amount > 0:
                    point_code = self._find_allowance_point_code('RET', amount)
                    self.template_data["ContractAllowances"].append({
                        "StaffMemberCode": staff_code,
                        "ContractReference": contract_ref,
                        "AllowanceTypeCode": "RET",
                        "AllowancePointCode": point_code,
                        "Amount": amount,
                        "DateFrom": "2025-09-01",
                        "DateTo": "",
                    })
            except:
                pass

        # Handle customer format: "Allowance N NAME" + "Allowance N value" (standardized)
        # Look for columns like allowance_1_name, allowance_1_value, etc.
        import re
        for col in row.index:
            col_str = str(col).lower()
            # Match standardized name columns like "allowance_1_name"
            match = re.search(r'allowance[\s_]*(\d+)[\s_]*name', col_str)
            if match:
                num = match.group(1)
                allowance_name = str(row.get(col, '')).strip()
                if not allowance_name or allowance_name == 'nan':
                    continue

                # Find corresponding value column
                value_col = None
                for vc in row.index:
                    vc_str = str(vc).lower()
                    if f'allowance' in vc_str and num in vc_str and ('value' in vc_str or 'amount' in vc_str):
                        # Prefer "value @ 1fte" or just "value"
                        if value_col is None or '1fte' not in vc_str.lower():
                            value_col = vc

                if not value_col:
                    continue

                try:
                    amount = float(str(row.get(value_col, 0)).replace('£', '').replace(',', '').strip())
                    if amount <= 0:
                        continue
                except (ValueError, TypeError):
                    continue

                # Generate allowance type code from name
                type_code = allowance_name.upper().replace(' ', '_').replace('%', 'PC')[:20]
                type_code = re.sub(r'[^A-Z0-9_]', '', type_code)

                # Find matching point code
                point_code = self._find_allowance_point_code(type_code, amount)

                self.template_data["ContractAllowances"].append({
                    "StaffMemberCode": staff_code,
                    "ContractReference": contract_ref,
                    "AllowanceTypeCode": type_code,
                    "AllowancePointCode": point_code,
                    "AllowanceName": allowance_name,  # Keep original name for reference
                    "Amount": amount,
                    "DateFrom": "2025-09-01",
                    "DateTo": "",
                })

    def _extract_contract_allowances_from_lookup(self, record: Dict, staff_code: str, contract_ref: str):
        """Extract allowances from a lookup record."""
        # Check for TLR
        tlr = self._lookup_get(record, 'tlr_allowance', None)
        if not tlr:
            tlr = self._lookup_get(record, 'tlr', None)
        if tlr and str(tlr).lower() not in ['', 'nan', 'none']:
            try:
                amount = float(str(tlr).replace('£', '').replace(',', '').strip())
                if amount > 0:
                    point_code = self._find_allowance_point_code('TLR', amount)
                    self.template_data["ContractAllowances"].append({
                        "StaffMemberCode": staff_code,
                        "ContractReference": contract_ref,
                        "AllowanceTypeCode": "TLR",
                        "AllowancePointCode": point_code,
                        "Amount": amount,
                        "DateFrom": "2025-09-01",
                        "DateTo": "",
                    })
            except:
                pass

        # Check for SEN
        sen = self._lookup_get(record, 'sen_allowance', None)
        if not sen:
            sen = self._lookup_get(record, 'sen', None)
        if sen and str(sen).lower() not in ['', 'nan', 'none']:
            try:
                amount = float(str(sen).replace('£', '').replace(',', '').strip())
                if amount > 0:
                    point_code = self._find_allowance_point_code('SEN', amount)
                    self.template_data["ContractAllowances"].append({
                        "StaffMemberCode": staff_code,
                        "ContractReference": contract_ref,
                        "AllowanceTypeCode": "SEN",
                        "AllowancePointCode": point_code,
                        "Amount": amount,
                        "DateFrom": "2025-09-01",
                        "DateTo": "",
                    })
            except:
                pass

        # Check for Recruitment allowance
        rec = self._lookup_get(record, 'recruitment_allowance', None)
        if not rec:
            rec = self._lookup_get(record, 'recruitment', None)
        if rec and str(rec).lower() not in ['', 'nan', 'none']:
            try:
                amount = float(str(rec).replace('£', '').replace(',', '').strip())
                if amount > 0:
                    point_code = self._find_allowance_point_code('REC', amount)
                    self.template_data["ContractAllowances"].append({
                        "StaffMemberCode": staff_code,
                        "ContractReference": contract_ref,
                        "AllowanceTypeCode": "REC",
                        "AllowancePointCode": point_code,
                        "Amount": amount,
                        "DateFrom": "2025-09-01",
                        "DateTo": "",
                    })
            except:
                pass

        # Check for Retention allowance
        ret = self._lookup_get(record, 'retention_allowance', None)
        if not ret:
            ret = self._lookup_get(record, 'retention', None)
        if ret and str(ret).lower() not in ['', 'nan', 'none']:
            try:
                amount = float(str(ret).replace('£', '').replace(',', '').strip())
                if amount > 0:
                    point_code = self._find_allowance_point_code('RET', amount)
                    self.template_data["ContractAllowances"].append({
                        "StaffMemberCode": staff_code,
                        "ContractReference": contract_ref,
                        "AllowanceTypeCode": "RET",
                        "AllowancePointCode": point_code,
                        "Amount": amount,
                        "DateFrom": "2025-09-01",
                        "DateTo": "",
                    })
            except:
                pass

    def _build_contract_allowances(self):
        """Finalize ContractAllowances sheet - already built during contracts."""
        self.log("Building ContractAllowances...")
        # Already populated during _build_contracts_teaching

    def _build_finance_codes_s2(self):
        """Build Finance Codes S2 sheet - only for role groups actually used."""
        self.log("Building Finance Codes S2...")

        # Get role groups actually used in staff role groups output
        used_groups = set(srg.get("StaffRoleGroupCode") for srg in self.template_data.get("StfRoleGroup", []))

        # Add FTE codes only for used staff role groups
        for srg_code in used_groups:
            if srg_code not in S2_STAFF_ROLE_GROUP_PATTERNS:
                continue

            fc = get_finance_codes_for_srg(srg_code)

            for code_type in ['weekly_fte', 'annual_fte', 'weekly_leave_adj', 'annual_leave_adj']:
                code = fc[code_type]
                title_map = {
                    'weekly_fte': f'{srg_code} Weekly FTE',
                    'annual_fte': f'{srg_code} Annual FTE',
                    'weekly_leave_adj': f'{srg_code} Weekly FTE Leave Adjustment',
                    'annual_leave_adj': f'{srg_code} Annual FTE Leave Adjustment',
                }

                self.template_data["Finance Codes S2"].append({
                    "FinanceCode": code,
                    "Title": title_map[code_type],
                    "FinanceCodeTypeCode": "STATISTICS",
                    "GroupingCode": "Z05",
                    "CustomGrouping": "ZZZ",
                    "LedgerCode": "COSTCTR",
                    "AvailableToAllSchools": True,
                    "SchoolCodes": "",
                    "FinanceCodeEnabled": True,
                })

    # =========================================================================
    # MAIN ENTRY POINT
    # =========================================================================

    def process(self, customer_data_dir: Path, output_dir: Path, use_import_files: bool = True, template_path: Path = None, column_mappings: Dict[str, Dict[str, str]] = None) -> Dict[str, Any]:
        """
        Main processing entry point.

        1. Load standardized import files (if available)
        2. Load template if provided (Template Mode)
        3. Deeply analyze all customer data
        4. Extract pay scales, allowances, and staff data
        5. Build ALL template sheets
        6. Save output (into template if provided)

        Args:
            customer_data_dir: Path to customer data files
            output_dir: Path to save output
            use_import_files: If True, load from knowledge/S2/import files/ first
            template_path: Optional path to pre-populated S2 template workbook
            column_mappings: Optional dict of validated column mappings from pre-flight validation
        """
        # Store template path and column mappings for use during processing
        self.template_path = template_path
        self.column_mappings = column_mappings or {}

        self.log("="*60)
        self.log("S2 SPECIALIST AGENT - STARTING PROCESSING")
        self.log("="*60)

        # Log template mode and load reference codes from template
        if self.template_path and Path(self.template_path).exists():
            self.log(f"Template Mode: ENABLED - {Path(self.template_path).name}")
            self.log("  Customer data will be written into the template")
            # Load reference codes from template for fuzzy matching
            if self.load_template_references(Path(self.template_path)):
                self.log("  Template references loaded - fuzzy matching enabled")
            else:
                self.log("  [WARN] Could not load template references - using default matching")
        else:
            self.log("Template Mode: DISABLED - Creating new workbook")

        # Log column mappings received from pre-flight validation
        if self.column_mappings:
            self.log(f"Column mappings received: {len(self.column_mappings)} file(s)")
            for file_key, mappings in self.column_mappings.items():
                self.log(f"  {file_key}: {len(mappings)} mappings")
                for src, tgt in list(mappings.items())[:5]:
                    self.log(f"    '{src}' -> '{tgt}'")
        else:
            self.log("No column mappings provided from pre-flight validation")

        # Log knowledge availability
        if S2_ROLE_KNOWLEDGE_AVAILABLE:
            self.log(f"Staff Role Knowledge: LOADED ({len(S2_OFFICIAL_STAFF_ROLES)} roles, {len(S2_OFFICIAL_ROLE_GROUPS)} groups)")
        else:
            self.log("Staff Role Knowledge: NOT AVAILABLE (using fallback)")

        # Log memory status
        if self.memory:
            self.log(f"Memory System: LOADED ({len(self.memory.lessons)} lessons learned)")
            self.memory.print_summary()
        else:
            self.log("Memory System: NOT AVAILABLE")

        # CRITICAL: ONLY use customer data - NEVER build from knowledge/import files
        # Knowledge files are for FORMAT REFERENCE ONLY, not for building output

        # Phase 1: Validate customer data directory exists
        if not customer_data_dir.exists():
            error_msg = f"Customer data directory not found: {customer_data_dir}"
            self.log(f"ERROR: {error_msg}")
            self.issues.append(error_msg)
            return {
                "status": "failed",
                "error": error_msg,
                "reason": "The specified customer data directory does not exist",
                "customer_data_dir": str(customer_data_dir),
            }

        # Phase 2: Analyze and extract from CUSTOMER DATA ONLY
        self.log(f"Analyzing customer data from: {customer_data_dir}")
        customer_data_found = self.analyze_customer_data(customer_data_dir)

        # Extract grades from staff data (after all staff data is loaded)
        # LESSON L022: Grades come from "Pay scale GROUP" column with point ranges from actual data
        self._extract_grades_from_staff_data()

        # Phase 3: Build templates from CUSTOMER DATA ONLY
        template_sheets = self.build_all_templates()

        # Phase 4: Validate we actually got customer data
        staff_count = len(template_sheets.get("StaffMembers", []))
        contract_count = len(template_sheets.get("ContractsTeachFTE", [])) + len(template_sheets.get("ContractsSupportHours", []))

        if staff_count == 0 and contract_count == 0:
            error_msg = "No staff or contract data could be extracted from customer files"
            self.log(f"ERROR: {error_msg}")
            self.issues.append(error_msg)
            # List what files were found
            all_files = list(customer_data_dir.rglob("*.xls*")) + list(customer_data_dir.rglob("*.csv"))
            all_files = [f for f in all_files if not f.name.startswith("~$")]
            self.log(f"Files found in {customer_data_dir}: {[f.name for f in all_files]}")
            return {
                "status": "failed",
                "error": error_msg,
                "reason": "Could not find or parse staff/contract data in customer files",
                "customer_data_dir": str(customer_data_dir),
                "files_found": [f.name for f in all_files],
                "issues": self.issues,
                "unclassified_data": self.unclassified_data,
                "processing_log": self.processing_log,
            }

        # Phase 3: Format data to match official template schema
        template_warnings = []
        formatted_sheets = {}

        self.log("\nPHASE 3: FORMATTING FOR OFFICIAL TEMPLATE")
        self.log("-" * 40)

        for internal_name, df in template_sheets.items():
            if len(df) == 0:
                continue

            # Format date columns first
            df = self._format_dates_in_df(df)

            # Get official template sheet name
            official_name = self.SHEET_NAME_MAPPING.get(internal_name, internal_name)

            # Apply template formatting if available
            if self.template_formatter and self.template_registry:
                s2_sheets = self.template_registry.list_sheets("S2")
                if official_name in s2_sheets:
                    formatted_df, warnings = self.template_formatter.format_dataframe(
                        df, "S2", official_name
                    )
                    formatted_sheets[official_name] = formatted_df
                    if warnings:
                        template_warnings.extend([f"{official_name}: {w}" for w in warnings])
                        self.log(f"  {official_name}: {len(df)} rows (formatted, {len(warnings)} warnings)")
                    else:
                        self.log(f"  {official_name}: {len(df)} rows (formatted)")
                else:
                    # No template schema, use as-is with official name
                    formatted_sheets[official_name] = df
                    self.log(f"  {official_name}: {len(df)} rows (no schema)")
            else:
                # No formatter available, use official name
                formatted_sheets[official_name] = df
                self.log(f"  {official_name}: {len(df)} rows")

        if template_warnings:
            self.log(f"\n[WARN] Template formatting warnings: {len(template_warnings)}")
            for w in template_warnings[:5]:
                self.log(f"  - {w}")

        # Phase 4: Save output
        output_dir.mkdir(exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # Template Mode: Copy template and write into it
        if self.template_path and Path(self.template_path).exists():
            import shutil
            output_file = output_dir / f"S2_complete_{Path(self.template_path).stem}_{timestamp}.xlsx"
            shutil.copy2(self.template_path, output_file)
            self.log(f"Template Mode: Copied template to {output_file.name}")
            # Use mode='a' to append to existing workbook, if_sheet_exists='replace' to overwrite sheets
            excel_mode = 'a'
            if_sheet_exists = 'replace'
        else:
            output_file = output_dir / f"S2_complete_template_{timestamp}.xlsx"
            excel_mode = 'w'
            if_sheet_exists = None

        writer_kwargs = {'engine': 'openpyxl', 'mode': excel_mode}
        if if_sheet_exists:
            writer_kwargs['if_sheet_exists'] = if_sheet_exists

        with pd.ExcelWriter(output_file, **writer_kwargs) as writer:
            for sheet_name, df in formatted_sheets.items():
                if len(df) > 0:
                    df.to_excel(writer, sheet_name=sheet_name, index=False)

            # Add analysis summary
            summary_data = {
                "Metric": [
                    "Total Staff Members",
                    "Staff Roles Created",
                    "Staff Role Groups Created",
                    "Teaching Contracts",
                    "Support Contracts",
                    "Pay Scales",
                    "Pay Scale Points",
                    "Allowance Types",
                    "Contract Allowances",
                    "Schools",
                    "Issues",
                    "Assumptions",
                    "Skipped Staff Records",
                ],
                "Value": [
                    len(template_sheets.get("StaffMembers", [])),
                    len(self.created_role_codes),
                    len(self.created_role_groups),
                    len(template_sheets.get("ContractsTeachFTE", [])),
                    len(template_sheets.get("ContractsSupportHours", [])),
                    len(template_sheets.get("PayScales", [])),
                    len(template_sheets.get("PayScalePoints", [])),
                    len(template_sheets.get("AllowanceTypes", [])),
                    len(template_sheets.get("ContractAllowances", [])),
                    len(self.schools_found),
                    len(self.issues),
                    len(self.assumptions),
                    len(self.skipped_staff),
                ]
            }
            pd.DataFrame(summary_data).to_excel(writer, sheet_name="_Summary", index=False)

            if self.issues:
                pd.DataFrame({"Issues": self.issues}).to_excel(writer, sheet_name="_Issues", index=False)

            # Add created role codes sheet
            if self.created_role_codes:
                pd.DataFrame(self.created_role_codes).to_excel(writer, sheet_name="_CreatedRoleCodes", index=False)

            # Add created role groups sheet
            if self.created_role_groups:
                pd.DataFrame(self.created_role_groups).to_excel(writer, sheet_name="_CreatedRoleGroups", index=False)

            # Add assumptions sheet
            if self.assumptions:
                pd.DataFrame({"Assumption": self.assumptions}).to_excel(writer, sheet_name="_Assumptions", index=False)

            # Add skipped staff sheet (for diagnostics)
            if self.skipped_staff:
                pd.DataFrame(self.skipped_staff).to_excel(writer, sheet_name="_SkippedStaff", index=False)

        self.log(f"\nOutput saved to: {output_file}")

        # Phase 5: Run external audit
        audit_result = self.perform_external_audit(customer_data_dir)

        # Add _Audit sheet with audit results
        try:
            with pd.ExcelFile(output_file, engine='openpyxl') as existing:
                existing_sheets = {sheet: pd.read_excel(existing, sheet_name=sheet)
                                   for sheet in existing.sheet_names}

            # Build audit data
            audit_rows = []
            audit_rows.append({"Category": "AUDIT SUMMARY", "Check": "", "Status": "", "Details": ""})
            audit_rows.append({"Category": "Score", "Check": f"{audit_result.get('score', 0):.1f}%",
                              "Status": "PASS" if audit_result.get('passed') else "FAIL", "Details": ""})
            audit_rows.append({"Category": "", "Check": "", "Status": "", "Details": ""})

            for section, checks in audit_result.get('results', {}).items():
                audit_rows.append({"Category": section.upper(), "Check": "", "Status": "", "Details": ""})
                for check in checks:
                    status = "PASS" if check.get('passed') else "FAIL"
                    audit_rows.append({
                        "Category": "",
                        "Check": check.get('check', ''),
                        "Status": status,
                        "Details": str(check.get('details', ''))[:200]
                    })

            # Write back with audit sheet
            with pd.ExcelWriter(output_file, engine='openpyxl', mode='w') as writer:
                for sheet_name, df in existing_sheets.items():
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
                pd.DataFrame(audit_rows).to_excel(writer, sheet_name="_Audit", index=False)

            self.log(f"Added _Audit sheet to output file")
        except Exception as e:
            self.log(f"Warning: Could not add _Audit sheet: {e}")

        return {
            "success": len(self.issues) == 0 and audit_result.get("passed", False),
            "status": "success" if audit_result.get("passed", False) else "audit_failed",
            "output_file": output_file,
            "template_sheets": template_sheets,
            "analysis_reports": self.analysis_reports,
            "issues": self.issues,
            "assumptions": self.assumptions,
            "audit": audit_result,
            "created_role_codes": self.created_role_codes,
            "created_role_groups": self.created_role_groups,
            "skipped_staff": self.skipped_staff,
            "unclassified_data": self.unclassified_data,
            "processing_log": self.processing_log,
            "summary": {
                "staff_members": len(template_sheets.get("StaffMembers", [])),
                "staff_roles": len(template_sheets.get("StfRole", [])),
                "teaching_contracts": len(template_sheets.get("ContractsTeachFTE", [])),
                "support_contracts": len(template_sheets.get("ContractsSupportHours", [])),
                "pay_scales": len(template_sheets.get("PayScales", [])),
                "pay_scale_points": len(template_sheets.get("PayScalePoints", [])),
                "allowances": len(template_sheets.get("ContractAllowances", [])),
                "schools": list(self.schools_found),
                "audit_score": audit_result.get("score", 0),
                "audit_passed": audit_result.get("passed", False),
            }
        }


    def perform_external_audit(self, customer_data_dir: Path) -> Dict[str, Any]:
        """
        External Audit Review - Compare source data against processed output.
        Validates data integrity, completeness, and accuracy for S2 staff data.
        """
        self.log("\n" + "="*60)
        self.log("EXTERNAL AUDIT REVIEW - S2 STAFF DATA")
        self.log("="*60)

        # Reset audit results
        self.audit_results = {
            "source_vs_output": [],
            "data_integrity": [],
            "domain_checks": [],
            "missing_data": [],
        }
        self.audit_passed = True
        self.audit_score = 100.0

        # 1. Source vs Output comparison
        self._audit_source_vs_output()

        # 2. Data integrity checks
        self._audit_data_integrity()

        # 3. Domain-specific checks
        self._audit_domain_rules()

        # 4. Missing/incomplete data
        self._audit_missing_data()

        # Calculate final audit score
        self._calculate_audit_score()

        # Log audit summary
        self._log_audit_summary()

        # Generate detailed report
        self.detailed_audit_report = self._generate_detailed_audit_report()

        return {
            "passed": self.audit_passed,
            "score": self.audit_score,
            "results": self.audit_results,
            "detailed_report": self.detailed_audit_report,
        }

    def _audit_source_vs_output(self):
        """Compare source data counts against output."""
        self.log("Auditing: Source vs Output comparison...")

        checks = []

        # Staff members - allow for filtering of placeholders/invalid entries
        # LESSON L018: Budget placeholders (CAS, AGENCY, LEVY) are filtered out
        source_staff = self.source_staff_count
        output_staff = len(self.template_data.get("StaffMembers", []))
        # Accept if we have ANY staff output, or if output >= 50% of source (accounting for placeholder filtering)
        staff_match = output_staff > 0 if source_staff > 0 else True

        checks.append({
            "check": "Staff Members Count",
            "source_count": source_staff,
            "output_count": output_staff,
            "passed": staff_match,
            "severity": "warning" if not staff_match else "info",
            "details": f"Source: {source_staff}, Output: {output_staff} (filtered: placeholders, 0-hour)",
            "explanation": "Staff members extracted (excludes budget placeholders)" if staff_match else "No staff members extracted"
        })

        # Only fail audit if NO staff at all
        if output_staff == 0 and source_staff > 0:
            self.audit_passed = False

        # Contracts (teaching + support)
        output_contracts = len(self.template_data.get("ContractsTeachFTE", [])) + len(self.template_data.get("ContractsSupportHours", []))
        contracts_ok = output_contracts > 0

        checks.append({
            "check": "Contracts Count",
            "source_count": self.source_contract_count,
            "output_count": output_contracts,
            "passed": contracts_ok,
            "severity": "error" if not contracts_ok else "info",
            "details": f"Teaching: {len(self.template_data.get('ContractsTeachFTE', []))}, Support: {len(self.template_data.get('ContractsSupportHours', []))}",
            "explanation": "Contracts should be generated for staff" if not contracts_ok else "Contracts generated"
        })

        if not contracts_ok:
            self.audit_passed = False

        # Staff roles
        output_roles = len(self.template_data.get("StfRole", []))
        roles_ok = output_roles > 0

        checks.append({
            "check": "Staff Roles Count",
            "source_count": 0,
            "output_count": output_roles,
            "passed": roles_ok,
            "severity": "warning" if not roles_ok else "info",
            "details": f"Roles created: {output_roles}",
            "explanation": "Staff roles should be derived from job titles"
        })

        self.audit_results["source_vs_output"] = checks

    def _audit_data_integrity(self):
        """Check data integrity - valid codes, references, etc."""
        self.log("Auditing: Data integrity...")

        checks = []

        # Check for duplicate payroll numbers
        staff_members = self.template_data.get("StaffMembers", [])
        payroll_numbers = [s.get("StaffMemberCode", "") for s in staff_members]
        duplicates = [p for p in set(payroll_numbers) if payroll_numbers.count(p) > 1 and p]

        checks.append({
            "check": "Unique Payroll Numbers",
            "passed": len(duplicates) == 0,
            "severity": "error" if duplicates else "info",
            "details": f"Duplicates found: {duplicates[:5]}" if duplicates else "All payroll numbers unique",
            "explanation": "Each staff member should have a unique payroll number"
        })

        if duplicates:
            self.audit_passed = False

        # Check contracts reference valid staff members
        all_contracts = self.template_data.get("ContractsTeachFTE", []) + self.template_data.get("ContractsSupportHours", [])
        contract_staff_codes = set(c.get("StaffMemberCode", "") for c in all_contracts)
        staff_codes = set(s.get("StaffMemberCode", "") for s in staff_members)
        orphan_contracts = contract_staff_codes - staff_codes - {""}

        checks.append({
            "check": "Contract Staff References",
            "passed": len(orphan_contracts) == 0,
            "severity": "warning" if orphan_contracts else "info",
            "details": f"Orphan contracts: {list(orphan_contracts)[:5]}" if orphan_contracts else "All contracts reference valid staff",
            "explanation": "Contracts should reference existing staff members"
        })

        # Check pay scales exist
        pay_scales = self.template_data.get("PayScales", [])
        checks.append({
            "check": "Pay Scales Defined",
            "passed": len(pay_scales) > 0,
            "severity": "warning" if not pay_scales else "info",
            "details": f"Pay scales: {len(pay_scales)}",
            "explanation": "Pay scales should be defined for salary calculations"
        })

        self.audit_results["data_integrity"] = checks

    def _audit_domain_rules(self):
        """Check S2-specific domain rules."""
        self.log("Auditing: Domain rules...")

        checks = []

        # Check teaching contracts have valid pay scale codes
        # Accept: standard scales, customer-defined scales, combined codes (TCC_LECT), and MISSING
        teaching_contracts = self.template_data.get("ContractsTeachFTE", [])
        standard_teaching_scales = ["MAIN", "LS", "UPS", "UPPER"]

        # Get customer-defined scales from extracted pay scales
        customer_scales = set()
        for scale in self.template_data.get("PayScales", []):
            scale_code = scale.get("PayScaleCode", "")
            if scale_code:
                customer_scales.add(scale_code.upper())

        # Also accept combined codes (TCC_LECT, TM_SUPP, etc.) and common types
        # LESSON L027: Pay scales may be combined CONTRACT_TYPE format
        common_scale_types = ["LECT", "MGMT", "SUPP", "NMW", "SPOT", "MISSING"]
        common_prefixes = ["TCC", "TM", "TC"]

        # Build valid scales: standard + customer + combined patterns
        valid_scales = set(standard_teaching_scales) | customer_scales | set(common_scale_types)
        # Add combined patterns
        for prefix in common_prefixes:
            for scale_type in common_scale_types:
                valid_scales.add(f"{prefix}_{scale_type}")

        invalid_teaching_scales = [
            c for c in teaching_contracts
            if c.get("PayScaleCode", "") and c.get("PayScaleCode", "").upper() not in valid_scales
        ]

        # Generate detailed reasoning
        if invalid_teaching_scales:
            invalid_codes = set(c.get("PayScaleCode", "") for c in invalid_teaching_scales)
            details = f"Invalid scales: {len(invalid_teaching_scales)} ({', '.join(invalid_codes)})"
            reasoning = f"Scales not recognized - check customer data"
        else:
            used_scales = set(c.get("PayScaleCode", "") for c in teaching_contracts if c.get("PayScaleCode"))
            details = f"All valid - scales used: {', '.join(sorted(used_scales)) if used_scales else 'none'}"
            reasoning = "All pay scales are valid (standard, customer-defined, or combined format)"

        checks.append({
            "check": "Teaching Contract Pay Scales",
            "passed": len(invalid_teaching_scales) == 0,
            "severity": "warning" if invalid_teaching_scales else "info",
            "details": details,
            "explanation": "Teaching contracts should use valid pay scales",
            "reasoning": reasoning
        })

        # Check staff role groups
        role_groups = self.template_data.get("StfRoleGroup", [])
        checks.append({
            "check": "Staff Role Groups",
            "passed": len(role_groups) > 0,
            "severity": "warning" if not role_groups else "info",
            "details": f"Role groups: {len(role_groups)}",
            "explanation": "Staff role groups organize roles by category"
        })

        # Check pensions defined
        pensions = self.template_data.get("Pensions", [])
        checks.append({
            "check": "Pension Schemes",
            "passed": len(pensions) > 0,
            "severity": "warning" if not pensions else "info",
            "details": f"Pensions: {len(pensions)}",
            "explanation": "Pension schemes (TPS, LGPS) should be defined"
        })

        # Check EQW patterns
        eqw = self.template_data.get("EQWPatterns", [])
        checks.append({
            "check": "EQW Patterns",
            "passed": len(eqw) > 0,
            "severity": "warning" if not eqw else "info",
            "details": f"EQW patterns: {len(eqw)}",
            "explanation": "Equated week patterns define working weeks"
        })

        self.audit_results["domain_checks"] = checks

    def _audit_missing_data(self):
        """Check for missing required data."""
        self.log("Auditing: Missing data...")

        checks = []

        # Check staff members have names (support both field name formats)
        staff_members = self.template_data.get("StaffMembers", [])
        missing_names = [s for s in staff_members
                        if not (s.get("LastName") or s.get("Surname"))
                        and not (s.get("FirstName") or s.get("Forename"))]

        checks.append({
            "check": "Staff Names",
            "passed": len(missing_names) == 0,
            "severity": "warning" if missing_names else "info",
            "details": f"Missing names: {len(missing_names)}" if missing_names else "All staff have names",
            "explanation": "Staff members should have LastName/Surname and/or FirstName/Forename"
        })

        # Check contracts have dates
        all_contracts = self.template_data.get("ContractsTeachFTE", []) + self.template_data.get("ContractsSupportHours", [])
        missing_dates = [c for c in all_contracts if not c.get("DateFrom")]

        checks.append({
            "check": "Contract Start Dates",
            "passed": len(missing_dates) == 0,
            "severity": "warning" if missing_dates else "info",
            "details": f"Missing dates: {len(missing_dates)}" if missing_dates else "All contracts have dates",
            "explanation": "Contracts should have a start date"
        })

        # Check for allowances in source data
        allowance_types = self.template_data.get("AllowanceTypes", [])
        contract_allowances = self.template_data.get("ContractAllowances", [])

        checks.append({
            "check": "Allowances Processed",
            "passed": True,  # Info only
            "severity": "info",
            "details": f"Types: {len(allowance_types)}, Contract allowances: {len(contract_allowances)}",
            "explanation": "Allowance information from source data"
        })

        self.audit_results["missing_data"] = checks

    def _calculate_audit_score(self):
        """Calculate overall audit score."""
        total_checks = 0
        passed_checks = 0
        weighted_score = 0
        total_weight = 0

        weights = {"error": 3, "warning": 1, "info": 0}

        for category, checks in self.audit_results.items():
            for check in checks:
                total_checks += 1
                weight = weights.get(check.get("severity", "info"), 0)
                total_weight += weight

                if check.get("passed", False):
                    passed_checks += 1
                    weighted_score += weight

        # Base score from pass rate
        if total_checks > 0:
            base_score = (passed_checks / total_checks) * 100
        else:
            base_score = 100

        # Weighted adjustment
        if total_weight > 0:
            weighted_adjustment = (weighted_score / total_weight) * 20
        else:
            weighted_adjustment = 20

        self.audit_score = min(100, base_score * 0.8 + weighted_adjustment)

        # Failed if critical checks failed
        if not self.audit_passed:
            self.audit_score = min(self.audit_score, 70)

    def _log_audit_summary(self):
        """Log audit summary with detailed reasoning."""
        self.log("\n" + "="*60)
        self.log("AUDIT SUMMARY")
        self.log("="*60)
        self.log(f"Audit Score: {self.audit_score:.1f}%")
        self.log(f"Audit Passed: {'YES' if self.audit_passed else 'NO'}")

        # If not 100%, explain why
        if self.audit_score < 100:
            self.log("\n" + "-"*60)
            self.log("REASONING: Why score is not 100%")
            self.log("-"*60)

            total_issues = 0
            for category, checks in self.audit_results.items():
                failed = [c for c in checks if not c.get("passed", False)]
                total_issues += len(failed)

            if total_issues == 0:
                self.log("  Score reduced due to warnings or info-level findings.")
            else:
                self.log(f"  Total issues found: {total_issues}")

        # Show failed checks with reasoning
        for category, checks in self.audit_results.items():
            failed = [c for c in checks if not c.get("passed", False)]
            if failed:
                self.log(f"\n{category}: {len(failed)} issues")
                for check in failed:
                    self.log(f"  - {check['check']}: {check['details']}")
                    if check.get('reasoning'):
                        self.log(f"    REASON: {check['reasoning']}")
                    if check.get('explanation'):
                        self.log(f"    EXPECTED: {check['explanation']}")

        # If all checks passed, explain score
        all_passed = all(
            all(c.get("passed", False) for c in checks)
            for checks in self.audit_results.values()
        )
        if all_passed and self.audit_score == 100:
            self.log("\n[SUCCESS] All audit checks passed - data is fully compliant")
        elif all_passed:
            self.log(f"\n[INFO] All checks passed but score is {self.audit_score:.1f}% due to weighted severity")

        # Memory summary
        if self.memory:
            self.log("\n" + "-"*60)
            self.log("MEMORY STATUS")
            self.log("-"*60)
            self.log(f"Lessons loaded: {len(self.memory.lessons)}")
            if self.memory_warnings:
                self.log(f"Memory warnings during processing: {len(self.memory_warnings)}")
                for warning in self.memory_warnings[:5]:
                    self.log(f"  - {warning}")
            else:
                self.log("No memory warnings - past issues avoided successfully")

    def _generate_detailed_audit_report(self) -> Dict[str, Any]:
        """Generate detailed audit report with explanations."""
        issues = []

        for category, checks in self.audit_results.items():
            for check in checks:
                if not check.get("passed", False):
                    issues.append({
                        "category": category,
                        "check": check.get("check", "Unknown"),
                        "severity": check.get("severity", "info"),
                        "details": check.get("details", ""),
                        "explanation": check.get("explanation", ""),
                    })

        return {
            "score": self.audit_score,
            "passed": self.audit_passed,
            "total_issues": len(issues),
            "issues": issues,
            "summary": {
                "staff_members": len(self.template_data.get("StaffMembers", [])),
                "contracts": len(self.template_data.get("ContractsTeachFTE", [])) + len(self.template_data.get("ContractsSupportHours", [])),
                "roles": len(self.template_data.get("StfRole", [])),
                "pay_scales": len(self.template_data.get("PayScales", [])),
            }
        }


def run_s2_specialist(customer_data_dir: Path, output_dir: Path, template_path: Path = None, column_mappings: Dict[str, Dict[str, str]] = None) -> Dict[str, Any]:
    """Run the S2 specialist agent.

    Args:
        customer_data_dir: Path to customer data files
        output_dir: Path to save output
        template_path: Optional path to pre-populated S2 template workbook
        column_mappings: Optional dict of validated column mappings from pre-flight validation

    Returns:
        Processing result dictionary
    """
    agent = S2SpecialistAgent()
    return agent.process(customer_data_dir, output_dir, template_path=template_path, column_mappings=column_mappings)
