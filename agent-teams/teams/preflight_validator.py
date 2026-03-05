"""
PreFlightValidator - Pre-Flight Column Validation and Mapping

Provides pre-flight analysis of uploaded data files to identify
column mappings, suggest fuzzy matches, and allow user corrections
before processing.
"""

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
import pandas as pd

# PDF support (optional)
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# Word document support (optional)
try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False


@dataclass
class ColumnMappingResult:
    """Result of column mapping analysis."""
    source_column: str
    mapped_to: Optional[str]
    confidence: float
    match_type: str  # 'exact', 'variation', 'fuzzy', 'unmapped'
    alternatives: List[Tuple[str, float]]  # (column_name, score)
    sample_values: List[Any]
    user_override: Optional[str] = None
    ignored: bool = False

    @property
    def status(self) -> str:
        """Get status for UI display.

        All mappings require user review/confirmation - never auto-accept.
        """
        if self.ignored:
            return 'ignored'
        if self.user_override:
            return 'corrected'
        # All mappings go to review - user must always confirm
        if self.mapped_to:
            return 'review'
        return 'unmapped'

    @property
    def final_mapping(self) -> Optional[str]:
        """Get final mapping - only returns value if user confirmed.

        Never auto-accept based on confidence - always require user confirmation.
        """
        if self.ignored:
            return None
        if self.user_override:
            return self.user_override
        # Do NOT auto-return based on confidence - require user confirmation
        return None

    def to_dict(self) -> Dict:
        """Convert to dictionary for serialization."""
        return {
            'source_column': self.source_column,
            'mapped_to': self.mapped_to,
            'confidence': self.confidence,
            'match_type': self.match_type,
            'alternatives': self.alternatives,
            'sample_values': [str(v) for v in self.sample_values[:5]],
            'user_override': self.user_override,
            'ignored': self.ignored,
            'status': self.status,
            'final_mapping': self.final_mapping
        }


@dataclass
class FileValidationResult:
    """Result of validating a single file."""
    file_path: Path
    file_name: str
    sheet_name: Optional[str]
    row_count: int
    column_count: int
    column_mappings: List[ColumnMappingResult]
    detected_strand: Optional[str]
    strand_confidence: float
    warnings: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    validated_at: datetime = field(default_factory=datetime.now)

    @property
    def matched_columns(self) -> List[ColumnMappingResult]:
        """Get columns that were successfully matched."""
        return [c for c in self.column_mappings if c.status == 'matched']

    @property
    def review_columns(self) -> List[ColumnMappingResult]:
        """Get columns that need review (fuzzy matches)."""
        return [c for c in self.column_mappings if c.status == 'review']

    @property
    def unmapped_columns(self) -> List[ColumnMappingResult]:
        """Get columns that couldn't be mapped."""
        return [c for c in self.column_mappings if c.status == 'unmapped']

    @property
    def mapping_summary(self) -> Dict[str, int]:
        """Get summary counts by status."""
        statuses = [c.status for c in self.column_mappings]
        return {
            'matched': statuses.count('matched'),
            'review': statuses.count('review'),
            'unmapped': statuses.count('unmapped'),
            'ignored': statuses.count('ignored'),
            'corrected': statuses.count('corrected'),
            'total': len(self.column_mappings)
        }

    def to_dict(self) -> Dict:
        """Convert to dictionary for serialization."""
        return {
            'file_name': self.file_name,
            'sheet_name': self.sheet_name,
            'row_count': self.row_count,
            'column_count': self.column_count,
            'detected_strand': self.detected_strand,
            'strand_confidence': self.strand_confidence,
            'mapping_summary': self.mapping_summary,
            'column_mappings': [c.to_dict() for c in self.column_mappings],
            'warnings': self.warnings,
            'errors': self.errors,
            'validated_at': self.validated_at.isoformat()
        }


class PreFlightValidator:
    """
    Pre-flight validation for data files.

    Analyzes uploaded files before processing to:
    1. Detect strand (S1, S2, S3)
    2. Map column names with confidence scores
    3. Identify fuzzy matches for user review
    4. Allow user corrections before processing
    """

    # Standard field names for fuzzy matching (by strand)
    STANDARD_FIELDS = {
        "S2": {
            # Staff identification
            "payroll_number": ["payroll", "payroll_number", "payroll no", "emp no", "employee number", "employee id", "staff id", "personnel number", "pay no", "pay number"],
            "surname": ["surname", "last name", "last_name", "family name", "lastname"],
            "forename": ["forename", "first name", "first_name", "given name", "firstname", "christian name"],
            "name": ["name", "full name", "fullname", "staff name", "employee name"],
            "job_title": ["job title", "job_title", "post", "position", "role", "job", "title", "designation", "post title"],
            "school_code": ["school", "school code", "school_code", "cost centre", "cost center", "cc", "location", "site", "establishment"],
            "weekly_hours": ["hours", "weekly hours", "weekly_hours", "contracted hours", "hpw", "hours per week", "ft hours", "full time hours", "weekly", "contract hours"],
            "fte": ["fte", "full time equivalent", "wte", "whole time equivalent"],
            "pay_scale": ["pay scale", "pay_scale", "scale", "payscale", "pay type", "scale type"],
            "scale_point": ["point", "scale point", "scp", "spine point", "current point", "pay point", "scale_point"],
            "annual_salary": ["salary", "annual salary", "actual salary", "gross salary", "pay", "annual pay", "fte salary"],
            "pension": ["pension", "pension scheme", "pension code", "superannuation"],
            "start_date": ["start date", "start_date", "service start", "hire date", "commencement", "date started", "contract start"],
            "date_of_birth": ["dob", "date of birth", "birth date", "birthdate"],
            "gender": ["gender", "sex", "m/f"],
            "ni_number": ["ni", "ni number", "national insurance", "nino"],
            "department": ["department", "dept", "section", "team", "division"],
            "contract_type": ["contract type", "contract_type", "employment type", "type"],
            "grade": ["grade", "pay grade", "band", "level"],
            "allowance_type": ["allowance", "tlr", "sen", "allowance type"],
            "staff_role_code": ["role code", "staff role", "job code", "post code"],
            "staff_role_group": ["role group", "srg", "category", "staff group", "job family"],
        },
        "S1": {
            "finance_code": ["finance code", "nominal", "nominal code", "account code", "gl code"],
            "school_code": ["school", "school code", "establishment", "site"],
            "department_code": ["department", "dept", "department code"],
            "fund_code": ["fund", "fund code", "funding"],
        },
        "S3": {
            "school_code": ["school", "school code", "establishment"],
            "pupil_count": ["pupils", "pupil count", "nos", "number on roll", "nor"],
        }
    }

    def __init__(self, inference_engine=None, column_learner=None):
        """
        Initialize PreFlightValidator.

        Args:
            inference_engine: Optional InferenceEngine instance
            column_learner: Optional ColumnMappingLearner for learned mappings
        """
        self.inference_engine = inference_engine
        self.column_learner = column_learner
        self._results: Dict[str, FileValidationResult] = {}
        self._corrections: Dict[str, Dict[str, str]] = {}

        # Try to get inference engine if not provided
        if self.inference_engine is None:
            try:
                from intelligence import InferenceEngine
                self.inference_engine = InferenceEngine()
            except ImportError:
                pass

        # Try to get column learner if not provided
        if self.column_learner is None:
            try:
                from memory.column_mapping_learner import ColumnMappingLearner
                self.column_learner = ColumnMappingLearner()
            except ImportError:
                pass

    def validate_file(
        self,
        file_path: Path,
        sheet_name: Optional[str] = None,
        strand: Optional[str] = None
    ) -> FileValidationResult:
        """
        Validate a single file and analyze columns.

        Args:
            file_path: Path to the file
            sheet_name: Specific sheet to validate (for Excel files)
            strand: Force specific strand (auto-detect if None)

        Returns:
            FileValidationResult with column analysis
        """
        file_path = Path(file_path)

        # Read the file
        try:
            if file_path.suffix.lower() in ['.xlsx', '.xlsm', '.xls']:
                if sheet_name:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                else:
                    xl = pd.ExcelFile(file_path)
                    # Use first sheet with data (skip empty/summary sheets)
                    sheet_name = None
                    for sn in xl.sheet_names:
                        try:
                            test_df = pd.read_excel(xl, sheet_name=sn, nrows=5)
                            if len(test_df.columns) >= 3 and len(test_df) >= 1:
                                sheet_name = sn
                                break
                        except:
                            continue
                    if not sheet_name:
                        sheet_name = xl.sheet_names[0]
                    df = pd.read_excel(xl, sheet_name=sheet_name)

                    # Skip rows if first row looks like a header/title (single value spanning columns)
                    if len(df) > 0:
                        first_row_nulls = df.iloc[0].isna().sum()
                        if first_row_nulls > len(df.columns) * 0.7:
                            # First row is mostly empty, might be a title row - try reading with header in row 1
                            df = pd.read_excel(xl, sheet_name=sheet_name, header=1)

            elif file_path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
            elif file_path.suffix.lower() == '.pdf':
                if not PDF_SUPPORT:
                    return self._create_error_result(file_path, "PDF support not available (install pdfplumber)")
                # Extract tables from PDF, fallback to text-based extraction
                df = self._extract_pdf_tables(file_path)
                if df is None or df.empty:
                    # Try text-based extraction for pay scales
                    df = self._extract_pdf_text_as_data(file_path)
                    if df is None or df.empty:
                        return self._create_error_result(file_path, "No tables or pay scale data found in PDF. The PDF may be scanned/image-based.")
                sheet_name = "PDF_Table"
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                if not DOCX_SUPPORT:
                    return self._create_error_result(file_path, "Word document support not available (install python-docx)")
                # Extract tables from Word document, fallback to text-based extraction
                df = self._extract_word_tables(file_path)
                if df is None or df.empty:
                    # Try text-based extraction for pay scales
                    df = self._extract_word_text_as_data(file_path)
                    if df is None or df.empty:
                        return self._create_error_result(file_path, "No tables or pay scale data found in Word document")
                sheet_name = "Word_Table"
            else:
                return self._create_error_result(file_path, f"Unsupported file type: {file_path.suffix}")
        except Exception as e:
            error_msg = str(e).lower()
            error_type = type(e).__name__
            full_error = f"{error_type}: {str(e)}"

            if "password" in error_msg or "encrypted" in error_msg:
                return self._create_error_result(file_path, f"File is password-protected.\nDetails: {full_error}")
            elif "could not read" in error_msg or "no engine" in error_msg:
                return self._create_error_result(file_path, f"Cannot read file - may be corrupted, password-protected, or unsupported format.\nDetails: {full_error}")
            elif "~" in file_path.name:
                return self._create_error_result(file_path, f"Windows truncated filename detected (8.3 format). Rename file to shorter name.\nDetails: {full_error}")
            else:
                return self._create_error_result(file_path, f"Error reading file.\nType: {error_type}\nDetails: {str(e)}")

        # Detect strand if not provided
        detected_strand = strand
        strand_confidence = 1.0 if strand else 0.0

        if not strand and self.inference_engine:
            try:
                strand_result = self.inference_engine.infer_strand(
                    columns=list(df.columns)
                )
                detected_strand = strand_result.decision
                strand_confidence = strand_result.confidence
            except Exception:
                detected_strand = None
                strand_confidence = 0.0

        # Analyze each column (skip blank column names)
        column_mappings = []
        for col in df.columns:
            # Skip blank/empty column names and pandas auto-generated "Unnamed:" columns
            # Handle NaN, None, empty strings, and numeric NaN values
            if col is None or (isinstance(col, float) and pd.isna(col)):
                continue
            col_str = str(col).strip()
            if not col_str or col_str.lower().startswith("unnamed") or col_str.lower() == "nan":
                continue
            # Skip columns with no data
            if df[col].dropna().empty:
                continue
            mapping = self._analyze_column(col, df[col], detected_strand)
            column_mappings.append(mapping)

        # Generate warnings
        warnings = self._generate_warnings(column_mappings, df)

        result = FileValidationResult(
            file_path=file_path,
            file_name=file_path.name,
            sheet_name=sheet_name,
            row_count=len(df),
            column_count=len(column_mappings),  # Use filtered count
            column_mappings=column_mappings,
            detected_strand=detected_strand,
            strand_confidence=strand_confidence,
            warnings=warnings
        )

        # Cache result
        cache_key = f"{file_path.name}:{sheet_name or 'default'}"
        self._results[cache_key] = result

        return result

    def _analyze_column(
        self,
        column_name: str,
        column_data: pd.Series,
        strand: Optional[str]
    ) -> ColumnMappingResult:
        """Analyze a single column and find mapping."""
        # Get sample values (non-null)
        sample_values = column_data.dropna().head(5).tolist()

        # Default result
        mapped_to = None
        confidence = 0.0
        match_type = 'unmapped'
        alternatives = []

        # FIRST: Check learned mappings from user corrections
        if self.column_learner and strand:
            # Check if this column should be ignored
            if self.column_learner.should_ignore(str(column_name), strand):
                return ColumnMappingResult(
                    source_column=str(column_name),
                    mapped_to="__IGNORE__",
                    confidence=1.0,
                    match_type='exact',
                    alternatives=[],
                    sample_values=sample_values,
                    user_override="__IGNORE__"
                )

            # Check for learned mapping
            learned_mapping = self.column_learner.get_learned_mapping(str(column_name), strand)
            if learned_mapping:
                return ColumnMappingResult(
                    source_column=str(column_name),
                    mapped_to=learned_mapping,
                    confidence=0.98,  # High confidence for learned mappings
                    match_type='exact',  # Treat as exact match
                    alternatives=[],
                    sample_values=sample_values,
                    user_override=learned_mapping  # Mark as user-overridden
                )

        # SECOND: Try inference engine if available
        if self.inference_engine and strand:
            try:
                result = self.inference_engine.infer_column_mapping(
                    source_column=str(column_name),
                    strand=strand
                )
                mapped_to = result.decision
                confidence = result.confidence

                # Determine match type
                if confidence >= 0.95:
                    match_type = 'exact'
                elif confidence >= 0.85:
                    match_type = 'variation'
                elif confidence >= 0.7:
                    match_type = 'fuzzy'
                else:
                    match_type = 'unmapped'

                # Get alternatives
                for alt in result.alternatives[:5]:
                    if isinstance(alt, dict):
                        alt_col = alt.get('column', alt.get('target', ''))
                        alt_score = alt.get('score', alt.get('confidence', 0.0))
                        if alt_col:
                            alternatives.append((alt_col, alt_score))

            except Exception:
                pass

        # Fallback: Use simple fuzzy matching if inference engine didn't provide good results
        if confidence < 0.7 and strand:
            fuzzy_result = self._fuzzy_match_column(column_name, strand)
            if fuzzy_result:
                best_match, best_score, all_matches = fuzzy_result
                if best_score > confidence:
                    mapped_to = best_match
                    confidence = best_score
                    match_type = 'exact' if best_score >= 0.95 else 'variation' if best_score >= 0.85 else 'fuzzy'
                    # Add alternatives
                    alternatives = [(m, s) for m, s in all_matches if m != best_match][:5]

        return ColumnMappingResult(
            source_column=str(column_name),  # Convert to string for datetime/numeric headers
            mapped_to=mapped_to if confidence >= 0.5 else None,
            confidence=confidence,
            match_type=match_type,
            alternatives=alternatives,
            sample_values=sample_values
        )

    def _fuzzy_match_column(self, column_name: str, strand: str) -> Optional[Tuple[str, float, List[Tuple[str, float]]]]:
        """
        Simple fuzzy matching against standard field names.

        Returns: (best_match, score, all_matches) or None
        """
        if strand not in self.STANDARD_FIELDS:
            return None

        # Normalize source column name (convert to string first for datetime/numeric headers)
        source_clean = str(column_name).lower().strip()
        source_clean = source_clean.replace('_', ' ').replace('-', ' ')

        matches = []

        for target_field, variations in self.STANDARD_FIELDS[strand].items():
            best_variation_score = 0.0

            for variation in variations:
                variation_clean = variation.lower().strip()

                # Exact match
                if source_clean == variation_clean:
                    matches.append((target_field, 1.0))
                    best_variation_score = max(best_variation_score, 1.0)
                    continue

                # Contains match
                if variation_clean in source_clean or source_clean in variation_clean:
                    # Score based on length similarity
                    len_ratio = min(len(source_clean), len(variation_clean)) / max(len(source_clean), len(variation_clean))
                    score = 0.7 + (0.25 * len_ratio)
                    best_variation_score = max(best_variation_score, score)
                    continue

                # Word overlap
                source_words = set(source_clean.split())
                variation_words = set(variation_clean.split())
                if source_words & variation_words:
                    overlap = len(source_words & variation_words) / max(len(source_words), len(variation_words))
                    score = 0.5 + (0.4 * overlap)
                    best_variation_score = max(best_variation_score, score)

            if best_variation_score > 0:
                matches.append((target_field, best_variation_score))

        if not matches:
            return None

        # Sort by score descending
        matches.sort(key=lambda x: x[1], reverse=True)

        # Remove duplicates (keep highest score for each target)
        seen = set()
        unique_matches = []
        for target, score in matches:
            if target not in seen:
                seen.add(target)
                unique_matches.append((target, score))

        if unique_matches:
            return (unique_matches[0][0], unique_matches[0][1], unique_matches)
        return None

    def _generate_warnings(
        self,
        column_mappings: List[ColumnMappingResult],
        df: pd.DataFrame
    ) -> List[str]:
        """Generate warnings based on column analysis."""
        warnings = []

        # Check unmapped column ratio
        unmapped_count = sum(1 for c in column_mappings if c.status == 'unmapped')
        total_count = len(column_mappings)

        if total_count > 0 and unmapped_count > total_count * 0.5:
            warnings.append(
                f"High proportion of unmapped columns ({unmapped_count}/{total_count}). "
                "Consider checking file format."
            )

        # Check for empty columns (with error handling for non-standard data)
        try:
            empty_cols = []
            for col in df.columns:
                try:
                    col_data = df[col]
                    if isinstance(col_data, pd.DataFrame):
                        col_data = col_data.iloc[:, 0]
                    if isinstance(col_data, pd.Series) and col_data.isna().all():
                        empty_cols.append(str(col))
                except Exception:
                    pass
            if empty_cols:
                warnings.append(f"Empty columns detected: {', '.join(empty_cols[:5])}")
        except Exception:
            pass

        # Check for duplicate columns
        try:
            if len(df.columns) != len(set(df.columns)):
                warnings.append("Duplicate column names detected")
        except Exception:
            pass

        return warnings

    def _extract_pdf_tables(self, file_path: Path) -> Optional[pd.DataFrame]:
        """Extract the largest table from a PDF file."""
        tables = self._extract_all_pdf_tables(file_path)
        if tables:
            # Return the largest table (most likely to be the main data)
            return max(tables.values(), key=lambda x: len(x) * len(x.columns))
        return None

    def _extract_all_pdf_tables(self, file_path: Path) -> Dict[str, pd.DataFrame]:
        """Extract ALL tables from a PDF file as a dict of table_name -> DataFrame."""
        if not PDF_SUPPORT:
            return {}

        all_tables = {}
        try:
            with pdfplumber.open(file_path) as pdf:
                table_idx = 0
                for page_num, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    for t_idx, table in enumerate(tables):
                        if table and len(table) > 1:
                            try:
                                # Use first row as header
                                headers = table[0]
                                data = table[1:]

                                # Clean and validate headers
                                clean_headers = []
                                for i, h in enumerate(headers):
                                    if h is None or (isinstance(h, str) and not h.strip()):
                                        clean_headers.append(f'unnamed_{i}')
                                    else:
                                        clean_h = str(h).strip().lower().replace(' ', '_').replace('\n', '_')
                                        clean_headers.append(clean_h if clean_h else f'unnamed_{i}')

                                # Create DataFrame
                                df = pd.DataFrame(data, columns=clean_headers)

                                # Convert all columns to proper types (avoid list objects)
                                for col in df.columns:
                                    df[col] = df[col].apply(lambda x: str(x) if x is not None else None)

                                if len(df) > 0 and len(df.columns) >= 2:
                                    table_name = f"PDF_Page{page_num+1}_Table{t_idx+1}"
                                    all_tables[table_name] = df
                                    table_idx += 1
                            except Exception:
                                continue
            return all_tables
        except Exception:
            return {}

    def _extract_pdf_text_as_data(self, file_path: Path) -> Optional[pd.DataFrame]:
        """Extract pay scale data from PDF text (non-table format)."""
        if not PDF_SUPPORT:
            return None

        try:
            import re
            all_text = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        all_text.append(text)

            if not all_text:
                return None

            full_text = '\n'.join(all_text)
            return self._parse_payscale_text(full_text)
        except Exception:
            return None

    def _extract_word_tables(self, file_path: Path) -> Optional[pd.DataFrame]:
        """Extract the largest table from a Word document."""
        tables = self._extract_all_word_tables(file_path)
        if tables:
            # Return the largest table
            return max(tables.values(), key=lambda x: len(x) * len(x.columns))
        return None

    def _extract_all_word_tables(self, file_path: Path) -> Dict[str, pd.DataFrame]:
        """Extract ALL tables from a Word document."""
        if not DOCX_SUPPORT:
            return {}

        all_tables = {}
        try:
            doc = DocxDocument(str(file_path))

            for t_idx, table in enumerate(doc.tables):
                if len(table.rows) < 2:
                    continue

                # Extract table data
                data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    data.append(row_data)

                if not data:
                    continue

                # Use first row as header
                headers = data[0]
                table_data = data[1:]

                # Clean headers
                clean_headers = []
                for i, h in enumerate(headers):
                    if not h or not h.strip():
                        clean_headers.append(f'unnamed_{i}')
                    else:
                        clean_h = h.strip().lower().replace(' ', '_').replace('\n', '_')
                        clean_headers.append(clean_h if clean_h else f'unnamed_{i}')

                df = pd.DataFrame(table_data, columns=clean_headers)

                if len(df) > 0 and len(df.columns) >= 2:
                    table_name = f"Word_Table_{t_idx + 1}"
                    all_tables[table_name] = df

            return all_tables
        except Exception:
            return {}

    def _extract_word_text_as_data(self, file_path: Path) -> Optional[pd.DataFrame]:
        """Extract pay scale data from Word document text (non-table format)."""
        if not DOCX_SUPPORT:
            return None

        try:
            import re
            doc = DocxDocument(str(file_path))

            # Get all text from paragraphs
            all_text = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    all_text.append(text)

            if not all_text:
                return None

            full_text = '\n'.join(all_text)
            return self._parse_payscale_text(full_text)
        except Exception:
            return None

    def _parse_payscale_text(self, text: str) -> Optional[pd.DataFrame]:
        """
        Parse pay scale data from plain text.

        Handles formats like:
        - "M1 £30,000" or "M1: £30,000"
        - "Point 1 - £25,000"
        - "Leadership L1 £45,000"
        - "UPS1: £40,000"
        - Lines with salary values
        """
        import re

        # Patterns to match pay scale points with values
        patterns = [
            # M1 £30,000 or M1: £30,000 or M1 - £30,000
            r'([MULP](?:PS|QT)?[0-9]+|L[0-9]+|SCP[0-9]+|Point\s*[0-9]+|TLR[0-9][ABC]?|SEN[0-9]?)\s*[:=-]?\s*[£$]?\s*([0-9,]+(?:\.[0-9]{2})?)',
            # Leadership L01 £45,000
            r'(Leadership\s*L?[0-9]+)\s*[:=-]?\s*[£$]?\s*([0-9,]+(?:\.[0-9]{2})?)',
            # Main Scale 1 £30,000
            r'((?:Main|Upper|Unqualified)\s*(?:Scale|Pay)?\s*[0-9]+)\s*[:=-]?\s*[£$]?\s*([0-9,]+(?:\.[0-9]{2})?)',
            # TMPS1 £30,000 or UPS1 £40,000
            r'((?:TMPS|UPS|MPS)[0-9]+)\s*[:=-]?\s*[£$]?\s*([0-9,]+(?:\.[0-9]{2})?)',
            # Grade A £25,000 or Band 1 £25,000
            r'((?:Grade|Band|Scale)\s*[A-Z0-9]+)\s*[:=-]?\s*[£$]?\s*([0-9,]+(?:\.[0-9]{2})?)',
        ]

        extracted = []
        lines = text.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            for pattern in patterns:
                matches = re.findall(pattern, line, re.IGNORECASE)
                for match in matches:
                    point_code = match[0].strip()
                    value_str = match[1].replace(',', '')
                    try:
                        value = float(value_str)
                        # Only include if it looks like a salary (> 1000) or hourly rate (< 100)
                        if value > 1000 or (value > 5 and value < 100):
                            extracted.append({
                                'point_code': point_code,
                                'value': value,
                                'original_line': line[:100]  # Keep context
                            })
                    except ValueError:
                        continue

        if not extracted:
            return None

        df = pd.DataFrame(extracted)
        return df

    def validate_file_all_sheets(
        self,
        file_path: Path,
        strand: Optional[str] = None
    ) -> List[FileValidationResult]:
        """
        Validate ALL sheets/tables from a file.

        For Excel files: validates each sheet separately
        For PDFs: validates each extracted table separately
        For CSV: validates the single file

        Returns:
            List of FileValidationResult, one per sheet/table
        """
        file_path = Path(file_path)
        results = []

        try:
            if file_path.suffix.lower() in ['.xlsx', '.xlsm', '.xls']:
                # Excel: validate each sheet
                xl = pd.ExcelFile(file_path)
                for sheet_name in xl.sheet_names:
                    # Skip common non-data sheets
                    if sheet_name.lower() in ['guidance', 'notes', 'instructions', 'help', 'contents']:
                        continue
                    try:
                        result = self.validate_file(file_path, sheet_name=sheet_name, strand=strand)
                        results.append(result)
                    except Exception:
                        continue

            elif file_path.suffix.lower() == '.pdf':
                # PDF: validate each table + text-based data
                if not PDF_SUPPORT:
                    results.append(self._create_error_result(file_path, "PDF support not available"))
                    return results

                all_tables = self._extract_all_pdf_tables(file_path)

                # Also try text-based extraction for pay scales
                text_df = self._extract_pdf_text_as_data(file_path)
                if text_df is not None and not text_df.empty:
                    all_tables["PDF_Text_Data"] = text_df

                if not all_tables:
                    results.append(self._create_error_result(file_path, "No tables or pay scale data found in PDF. The PDF may be scanned/image-based."))
                    return results

                for table_name, df in all_tables.items():
                    result = self._validate_dataframe(file_path, df, table_name, strand)
                    results.append(result)

            elif file_path.suffix.lower() == '.csv':
                # CSV: single file
                result = self.validate_file(file_path, strand=strand)
                results.append(result)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                # Word: validate each table + text-based data
                if not DOCX_SUPPORT:
                    results.append(self._create_error_result(file_path, "Word document support not available"))
                    return results

                all_tables = self._extract_all_word_tables(file_path)

                # Also try text-based extraction
                text_df = self._extract_word_text_as_data(file_path)
                if text_df is not None and not text_df.empty:
                    all_tables["Word_Text_Data"] = text_df

                if not all_tables:
                    results.append(self._create_error_result(file_path, "No tables or pay scale data found in Word document"))
                    return results

                for table_name, df in all_tables.items():
                    result = self._validate_dataframe(file_path, df, table_name, strand)
                    results.append(result)
            else:
                results.append(self._create_error_result(file_path, f"Unsupported file type: {file_path.suffix}"))

        except Exception as e:
            results.append(self._create_error_result(file_path, f"Error processing file: {e}"))

        return results if results else [self._create_error_result(file_path, "No data found")]

    def _validate_dataframe(
        self,
        file_path: Path,
        df: pd.DataFrame,
        sheet_name: str,
        strand: Optional[str] = None
    ) -> FileValidationResult:
        """Validate a DataFrame directly (used for PDF tables)."""
        # Detect strand if not provided
        detected_strand = strand
        strand_confidence = 1.0 if strand else 0.0

        if not strand and self.inference_engine:
            try:
                strand_result = self.inference_engine.infer_strand(columns=list(df.columns))
                detected_strand = strand_result.decision
                strand_confidence = strand_result.confidence
            except Exception:
                detected_strand = None
                strand_confidence = 0.0

        # Analyze each column
        column_mappings = []
        for col in df.columns:
            if col is None or (isinstance(col, float) and pd.isna(col)):
                continue
            col_str = str(col).strip()
            if not col_str or col_str.lower().startswith('unnamed'):
                continue

            # Get column data as a proper Series
            try:
                col_data = df[col]
                if isinstance(col_data, pd.DataFrame):
                    col_data = col_data.iloc[:, 0]
                # Ensure we have a Series
                if not isinstance(col_data, pd.Series):
                    col_data = pd.Series(col_data)
            except Exception:
                col_data = pd.Series(dtype=object)

            # Try to match column (pass Series, not list)
            mapping_result = self._analyze_column(col_str, col_data, detected_strand)
            column_mappings.append(mapping_result)

        # Generate warnings
        warnings = self._generate_warnings(column_mappings, df)

        # Store result
        result = FileValidationResult(
            file_path=file_path,
            file_name=file_path.name,
            sheet_name=sheet_name,
            row_count=len(df),
            column_count=len(df.columns),
            column_mappings=column_mappings,
            detected_strand=detected_strand,
            strand_confidence=strand_confidence,
            warnings=warnings
        )

        key = f"{file_path.name}:{sheet_name}"
        self._results[key] = result
        return result

    def _create_error_result(self, file_path: Path, error: str) -> FileValidationResult:
        """Create an error result for failed validation with detailed context."""
        # Build detailed error message with suggestions
        detailed_error = f"FILE: {file_path.name}\nPATH: {file_path}\nERROR: {error}"

        # Add specific suggestions based on error type
        suggestions = []
        error_lower = error.lower()

        if "password" in error_lower:
            suggestions.append("SUGGESTION: Open the file in Excel, go to File > Info > Protect Workbook > Encrypt with Password, and remove the password")
        elif "corrupt" in error_lower or "cannot read" in error_lower:
            suggestions.append("SUGGESTION: Try opening the file in Excel and re-saving it as a new .xlsx file")
            suggestions.append("SUGGESTION: Check if the file is still open in another application")
        elif "no tables found" in error_lower and "pdf" in error_lower:
            suggestions.append("SUGGESTION: This PDF may be scanned/image-based. Try using OCR software to convert it, or manually enter the data into Excel")
            suggestions.append("SUGGESTION: Check if the PDF contains actual tables or just formatted text")
        elif "pdf support" in error_lower:
            suggestions.append("SUGGESTION: Install pdfplumber with: pip install pdfplumber")
        elif "unsupported file type" in error_lower:
            suggestions.append("SUGGESTION: Convert the file to .xlsx, .xls, .csv, or .pdf format")
        elif "8.3" in str(file_path) or "~" in file_path.name:
            suggestions.append("SUGGESTION: The filename appears truncated (Windows 8.3 format). Rename the file to a shorter name without special characters")

        if suggestions:
            detailed_error += "\n" + "\n".join(suggestions)

        return FileValidationResult(
            file_path=file_path,
            file_name=file_path.name,
            sheet_name=None,
            row_count=0,
            column_count=0,
            column_mappings=[],
            detected_strand=None,
            strand_confidence=0.0,
            errors=[detailed_error]
        )

    def get_unmapped_columns(self, file_key: Optional[str] = None) -> List[ColumnMappingResult]:
        """Get all unmapped columns across validated files."""
        unmapped = []
        results = [self._results[file_key]] if file_key else self._results.values()

        for result in results:
            unmapped.extend(result.unmapped_columns)

        return unmapped

    def get_fuzzy_matches(self, file_key: Optional[str] = None) -> List[ColumnMappingResult]:
        """Get all fuzzy-matched columns that need review."""
        fuzzy = []
        results = [self._results[file_key]] if file_key else self._results.values()

        for result in results:
            fuzzy.extend(result.review_columns)

        return fuzzy

    def apply_user_corrections(self, corrections: Dict[str, str], file_key: str = None):
        """
        Apply user corrections to column mappings.

        Args:
            corrections: Dict mapping source_column -> corrected_standard_column
            file_key: Optional file key to limit corrections to
        """
        if file_key:
            self._corrections.setdefault(file_key, {}).update(corrections)

            if file_key in self._results:
                for mapping in self._results[file_key].column_mappings:
                    if mapping.source_column in corrections:
                        mapping.user_override = corrections[mapping.source_column]
        else:
            # Apply to all files
            for key, result in self._results.items():
                for mapping in result.column_mappings:
                    if mapping.source_column in corrections:
                        mapping.user_override = corrections[mapping.source_column]

    def ignore_column(self, source_column: str, file_key: str = None):
        """Mark a column as ignored (won't be processed)."""
        results = [self._results[file_key]] if file_key else self._results.values()

        for result in results:
            for mapping in result.column_mappings:
                if mapping.source_column == source_column:
                    mapping.ignored = True

    def get_final_mappings(self, file_key: str) -> Dict[str, str]:
        """
        Get final column mappings for a file after corrections.

        Returns:
            Dict mapping source_column -> standard_column
        """
        if file_key not in self._results:
            return {}

        result = self._results[file_key]
        mappings = {}

        for mapping in result.column_mappings:
            final = mapping.final_mapping
            if final:
                mappings[mapping.source_column] = final

        return mappings

    def get_sample_data(
        self,
        file_path: Path,
        column: str,
        n: int = 5,
        sheet_name: Optional[str] = None
    ) -> List[Any]:
        """Get sample data from a column."""
        try:
            if file_path.suffix.lower() in ['.xlsx', '.xlsm', '.xls']:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
            else:
                df = pd.read_csv(file_path)

            if column in df.columns:
                return df[column].dropna().head(n).tolist()
        except Exception:
            pass

        return []

    def validate_mapping_completeness(self, file_key: str, required_columns: List[str]) -> Dict:
        """
        Check if all required columns are mapped.

        Returns dict with 'complete' bool and 'missing' list.
        """
        if file_key not in self._results:
            return {'complete': False, 'missing': required_columns}

        mappings = self.get_final_mappings(file_key)
        mapped_standards = set(mappings.values())

        missing = [col for col in required_columns if col not in mapped_standards]

        return {
            'complete': len(missing) == 0,
            'missing': missing,
            'mapped': list(mapped_standards),
            'coverage': len(mapped_standards) / len(required_columns) if required_columns else 1.0
        }

    def get_all_results(self) -> Dict[str, FileValidationResult]:
        """Get all validation results."""
        return self._results.copy()

    def clear_results(self):
        """Clear all cached results and corrections."""
        self._results.clear()
        self._corrections.clear()

    def export_mappings(self, file_key: str = None) -> Dict:
        """Export all mappings for saving/reuse."""
        if file_key:
            if file_key not in self._results:
                return {}
            return {
                file_key: {
                    'mappings': self.get_final_mappings(file_key),
                    'corrections': self._corrections.get(file_key, {}),
                    'ignored': [
                        m.source_column
                        for m in self._results[file_key].column_mappings
                        if m.ignored
                    ]
                }
            }

        return {
            key: {
                'mappings': self.get_final_mappings(key),
                'corrections': self._corrections.get(key, {}),
                'ignored': [
                    m.source_column
                    for m in result.column_mappings
                    if m.ignored
                ]
            }
            for key, result in self._results.items()
        }

    def import_mappings(self, mapping_data: Dict):
        """Import previously saved mappings."""
        for file_key, data in mapping_data.items():
            if file_key in self._results:
                # Apply corrections
                corrections = data.get('corrections', {})
                self.apply_user_corrections(corrections, file_key)

                # Apply ignored columns
                for col in data.get('ignored', []):
                    self.ignore_column(col, file_key)


# Convenience function for quick validation
def validate_files(
    file_paths: List[Path],
    strand: Optional[str] = None
) -> Dict[str, FileValidationResult]:
    """
    Validate multiple files.

    Args:
        file_paths: List of file paths to validate
        strand: Optional strand to force

    Returns:
        Dict of file_key -> FileValidationResult
    """
    validator = PreFlightValidator()
    results = {}

    for path in file_paths:
        result = validator.validate_file(Path(path), strand=strand)
        key = f"{result.file_name}:{result.sheet_name or 'default'}"
        results[key] = result

    return results
