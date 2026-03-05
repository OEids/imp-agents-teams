
"""
IMP Planner - Agent Teams Web Application

A web interface for managing and running the S1, S2, S3 specialist agent teams.
Behind Ambitious MAT Finance Teams.

Brand: IMP Software (impsoftware.co.uk)
"""
from docx import Document

# PDF support
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

import streamlit as st
import pandas as pd
from pathlib import Path
import json
from datetime import datetime
import shutil
import time
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import threading

# Import config settings
from config.settings import (
    TEAMS as SETTINGS_TEAMS,
    get_data_directory,
    set_data_directory,
    get_templates_directory,
    set_templates_directory,
    get_reports_directory,
    set_reports_directory,
    load_user_config,
    refresh_paths,
    INTELLIGENCE_CONFIG,
    get_intelligence_config,
)

# Import Intelligence Module (optional)
try:
    from intelligence import InferenceEngine, ConfidenceLevel
    INFERENCE_AVAILABLE = True
except ImportError:
    INFERENCE_AVAILABLE = False

# Import PreFlightValidator (optional)
try:
    from teams.preflight_validator import PreFlightValidator, validate_files
    PREFLIGHT_AVAILABLE = True
except ImportError:
    PREFLIGHT_AVAILABLE = False

# Import Column Mapping Learner for persistent learning
try:
    from memory.column_mapping_learner import ColumnMappingLearner
    LEARNER_AVAILABLE = True
except ImportError:
    LEARNER_AVAILABLE = False
    ColumnMappingLearner = None

import os

# Page config
st.set_page_config(
    page_title="IMP Planner",
    page_icon="app brand/imp_logo_final.png",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =============================================================================
# IMP BRAND COLORS
# =============================================================================
IMP_COLORS = {
    "deep_purple": "#261342",
    "purple": "#6A0F8E",
    "light_purple": "#A093DB",
    "white": "#FFFFFF",
    "violet": "#6C63FF",
    "cyan": "#67D1FF",
    "dark_grey": "#B2B2B2",
    "light_grey": "#E5E5E5",
}

# Custom CSS for IMP branding
st.markdown(f"""
<style>
    /* Sidebar styling - subtle purple accent */
    [data-testid="stSidebar"] {{
        background: linear-gradient(180deg, {IMP_COLORS['white']} 0%, {IMP_COLORS['light_purple']}22 100%);
        border-right: 3px solid {IMP_COLORS['purple']};
    }}

    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3 {{
        color: {IMP_COLORS['deep_purple']} !important;
    }}

    [data-testid="stSidebar"] .stCaption {{
        color: {IMP_COLORS['purple']} !important;
    }}

    /* Button styling */
    .stButton > button[kind="primary"] {{
        background-color: {IMP_COLORS['purple']};
        border-color: {IMP_COLORS['purple']};
    }}

    .stButton > button[kind="primary"]:hover {{
        background-color: {IMP_COLORS['deep_purple']};
        border-color: {IMP_COLORS['deep_purple']};
    }}

    /* Tab styling */
    .stTabs [data-baseweb="tab-list"] {{
        gap: 8px;
    }}

    .stTabs [data-baseweb="tab"] {{
        background-color: {IMP_COLORS['light_grey']};
        border-radius: 4px;
        color: {IMP_COLORS['deep_purple']};
    }}

    .stTabs [aria-selected="true"] {{
        background-color: {IMP_COLORS['light_purple']} !important;
        color: {IMP_COLORS['deep_purple']} !important;
    }}

    /* Metric styling */
    [data-testid="stMetricValue"] {{
        color: {IMP_COLORS['deep_purple']};
    }}

    /* Expander styling */
    .streamlit-expanderHeader {{
        background-color: {IMP_COLORS['light_grey']};
        border-radius: 4px;
    }}

    /* Success/Info/Warning boxes */
    .stSuccess {{
        background-color: {IMP_COLORS['light_purple']}22;
        border-left-color: {IMP_COLORS['purple']};
    }}

    /* Links */
    a {{
        color: {IMP_COLORS['purple']} !important;
    }}

    /* Dividers */
    hr {{
        border-color: {IMP_COLORS['light_grey']};
    }}
</style>
""", unsafe_allow_html=True)

# Base paths
BASE_DIR = Path(__file__).parent

# Load paths from config
CUSTOMER_DATA_DIR = get_data_directory()
TEMPLATES_DIR = get_templates_directory()
REPORTS_DIR = get_reports_directory()

KNOWLEDGE_DIR = BASE_DIR / "knowledge"
SCOPE_DOCS_DIR = REPORTS_DIR / "scope_docs"
LC_LOCATIONS_FILE = BASE_DIR / "config" / "lc_locations.json"
PROJECTS_FILE = BASE_DIR / "config" / "projects.json"
LC_DOCS_DIR = REPORTS_DIR / "lc_docs"

# Ensure directories exist
REPORTS_DIR.mkdir(exist_ok=True)
for strand in ["S1", "S2", "S3"]:
    (CUSTOMER_DATA_DIR / strand).mkdir(parents=True, exist_ok=True)


# =============================================================================
# TEAM DEFINITIONS
# =============================================================================

TEAMS = {
    "S1": {
        "name": "Structure Team",
        "icon": "🏗️",
        "color": IMP_COLORS["purple"],
        "description": "Specialist Agent - Handles finance codes, schools, departments, and Chart of Accounts",
        "capabilities": [
            "🤖 Deep analysis of customer structure data",
            "📊 Auto Chart of Accounts extraction from customer files",
            "DFE COA mapping & finance code normalization",
            "School/department extraction with URN/LA lookup",
            "Validation & External Audit with scoring",
            "Auto-builds: FinanceCodes, Schools, Depts, Activity, Ledger",
        ],
        "template": "AA_New - Strand 1 Standard Workbook API",
        "knowledge_files": ["Strand 1 Process Notes.xlsx", "Strand 1 Training Notes.xlsx"],
    },
    "S2": {
        "name": "Staff Team",
        "icon": "👥",
        "color": IMP_COLORS["violet"],
        "description": "Specialist Agent - Processes staff contracts, pay scales, and personnel data",
        "capabilities": [
            "🤖 Deep analysis extracts ALL pay scales & points",
            "📊 Auto pay scale extraction from customer Excel files",
            "Pay scale setup (MPS, UPS, Leadership, NJC)",
            "Staff role classification (15 DFE groups)",
            "Contract building (Teaching FTE / Support Hours)",
            "Allowance extraction (TLR, SEN, etc.)",
            "Auto-builds: PayScales, StaffMembers, Contracts, EQWPatterns",
        ],
        "template": "AA_New - Strand 2 Standard Workbook API",
        "knowledge_files": ["Process Notes - Strand 2 (Structure).docx", "Process Notes - Strand 2 (Contracts).docx"],
    },
    "S3": {
        "name": "Financial Team",
        "icon": "💰",
        "color": IMP_COLORS["cyan"],
        "description": "Specialist Agent - Manages budgets, grants, funding, and pupil numbers",
        "capabilities": [
            "🤖 Deep analysis of budget & financial data",
            "📊 Auto pupil number extraction from census data",
            "Grant calculations (DFC, SCA, PE, UIFSM, Pupil Premium)",
            "Income/Expenditure line processing",
            "Validation & External Audit with scoring",
            "Auto-builds: Pupils, Funding, Calculators, Income, Expenditure",
        ],
        "template": "AA_New - Strand 3 Standard Workbook API",
        "knowledge_files": ["Strand 3 - Workbook - Phase 1.docx", "Strand 3 - Workbook - Phase 2.docx"],
    },
}


# =============================================================================
# SESSION STATE
# =============================================================================

if "processing_status" not in st.session_state:
    st.session_state.processing_status = {}
if "auto_process" not in st.session_state:
    st.session_state.auto_process = False
if "selected_team" not in st.session_state:
    st.session_state.selected_team = "S2"
# Pre-flight validation state
if "column_mappings" not in st.session_state:
    st.session_state.column_mappings = {}  # team_id -> mapping dict
if "mapping_validated" not in st.session_state:
    st.session_state.mapping_validated = {}  # team_id -> bool
if "validation_results" not in st.session_state:
    st.session_state.validation_results = {}  # team_id -> FileValidationResult dict
if "preflight_validator" not in st.session_state:
    st.session_state.preflight_validator = None
if "custom_mappings" not in st.session_state:
    st.session_state.custom_mappings = {}  # stores custom typed mappings
if "column_learner" not in st.session_state:
    if LEARNER_AVAILABLE:
        st.session_state.column_learner = ColumnMappingLearner()
    else:
        st.session_state.column_learner = None


# =============================================================================
# STANDARD FIELD NAMES FOR COLUMN MAPPING
# =============================================================================

# All valid S2 field names that users can map to
S2_FIELD_OPTIONS = [
    # Staff identification
    "payroll_number", "employee_id", "staff_code", "staff_id",
    # Names
    "surname", "forename", "first_name", "last_name", "name", "title",
    # Job/Role
    "job_title", "position", "role", "post", "staff_role_code", "staff_role_group",
    # School/Location
    "school_code", "school", "cost_centre", "department", "location",
    # Hours & FTE
    "weekly_hours", "ft_hours", "full_time_hours", "fte", "weekly_fte",
    "hours_per_week", "contracted_hours", "hours",
    # Pay scales
    "pay_scale", "pay_scale_type", "pay_scale_code", "scale", "grade",
    "scale_point", "current_scale_point", "scp", "point", "spine_point",
    # Salary
    "annual_salary", "salary", "actual_salary", "gross_salary", "rate",
    # Pension
    "pension", "pension_code", "pension_scheme",
    # Dates
    "service_start_date", "start_date", "contract_start", "hire_date",
    "end_date", "leave_date", "date_of_birth", "dob",
    # Contract
    "contract_ref", "contract_type", "reference", "contract_code",
    # Other
    "gender", "ni_number", "national_insurance",
    "eqw", "eqw_pattern", "weeks_worked", "weeks_paid",
    "finance_code", "fund_code", "nominal_code",
    # Allowances
    "allowance_type", "allowance_code", "tlr", "sen_allowance",
]

# All valid S1 field names
S1_FIELD_OPTIONS = [
    "finance_code", "fund_code", "activity_code", "ledger_code",
    "school_code", "school_name", "urn", "la_code",
    "department_code", "department_name",
    "cost_centre", "nominal_code", "description",
]

# All valid S3 field names
S3_FIELD_OPTIONS = [
    "school_code", "pupil_count", "fte_pupils",
    "grant_code", "grant_name", "grant_amount",
    "income_code", "expenditure_code", "budget_amount",
]

def get_field_options_for_strand(strand: str) -> list:
    """Get all valid field names for a strand."""
    if strand == "S1":
        return S1_FIELD_OPTIONS
    elif strand == "S2":
        return S2_FIELD_OPTIONS
    elif strand == "S3":
        return S3_FIELD_OPTIONS
    return []

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def get_customer_files(team_id: str) -> list:
    """Get list of customer data files for a team."""
    team_dir = CUSTOMER_DATA_DIR / team_id
    files = []
    if team_dir.exists():
        for ext in ["*.xlsx", "*.xlsm", "*.xls", "*.csv", "*.pdf", "*.docx", "*.doc", "*.png", "*.jpg", "*.jpeg"]:
            files.extend(team_dir.rglob(ext))
    return [f for f in files if not f.name.startswith("~$")]


def get_recent_reports(team_id: str, limit: int = 5) -> list:
    """Get recent report files for a team."""
    # Look for both old format and new specialist agent format
    reports = list(REPORTS_DIR.glob(f"{team_id}_*.xlsx"))
    reports += list(REPORTS_DIR.glob(f"{team_id}_complete_template_*.xlsx"))
    # Also check output folder
    output_dir = BASE_DIR / "output"
    if output_dir.exists():
        reports += list(output_dir.glob(f"{team_id}_*.xlsx"))
        reports += list(output_dir.glob(f"{team_id}_complete_template_*.xlsx"))
    # Remove duplicates and sort
    reports = list(set(reports))
    reports.sort(key=lambda f: f.stat().st_mtime, reverse=True)
    return reports[:limit]


def run_team_processing(team_id: str, column_mappings: dict = None) -> dict:
    """Run the data processing for a team using specialist agents.

    Args:
        team_id: The team/strand to process (S1, S2, S3)
        column_mappings: Optional dict of validated column mappings from pre-flight validation

    Returns:
        Processing result dictionary
    """
    customer_dir = CUSTOMER_DATA_DIR / team_id
    output_dir = REPORTS_DIR

    # Get validated mappings from session state if not provided
    if column_mappings is None:
        column_mappings = st.session_state.column_mappings.get(team_id, {})

    if team_id == "S1":
        from teams.s1_specialist import run_s1_specialist
        result = run_s1_specialist(customer_dir, output_dir, column_mappings=column_mappings)
        return {
            "success": result.get("success", False),
            "summary": result.get("summary", {}),
            "issues": result.get("issues", []),
            "assumptions": result.get("assumptions", []),
            "output_file": result.get("output_file"),
            "template_sheets": result.get("template_sheets", {})
        }
    elif team_id == "S2":
        from teams.s2_specialist import run_s2_specialist
        result = run_s2_specialist(customer_dir, output_dir, column_mappings=column_mappings)
        return {
            "success": result.get("success", False),
            "summary": result.get("summary", {}),
            "issues": result.get("issues", []),
            "assumptions": result.get("assumptions", []),
            "output_file": result.get("output_file"),
            "template_sheets": result.get("template_sheets", {}),
            "unclassified_data": result.get("unclassified_data", []),
            "created_role_codes": result.get("created_role_codes", []),
            "created_role_groups": result.get("created_role_groups", []),
            "skipped_staff": result.get("skipped_staff", []),
            "processing_log": result.get("processing_log", []),
        }
    elif team_id == "S3":
        from teams.s3_specialist import run_s3_specialist

        # Get template path if user uploaded one
        template_path = st.session_state.get("s3_template_path")

        # Get code mappings (customer codes -> template codes)
        code_mappings = st.session_state.get("s3_code_mappings")

        result = run_s3_specialist(
            customer_dir,
            output_dir,
            template_path=template_path,
            column_mappings=column_mappings,
            code_mappings=code_mappings
        )
        return {
            "success": result.get("success", False),
            "summary": result.get("summary", {}),
            "issues": result.get("issues", []),
            "assumptions": result.get("assumptions", []),
            "output_file": result.get("output_file"),
            "template_sheets": result.get("template_sheets", {}),
            "build_mode": "template" if template_path else "raw_data"
        }
    else:
        return {
            "success": False,
            "summary": {},
            "issues": [f"{team_id} processor not recognized"],
            "assumptions": [],
            "output_file": None,
            "template_sheets": {}
        }


def format_file_size(size_bytes: int) -> str:
    """Format file size in human-readable format."""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} TB"


def _split_lines(value: str) -> list:
    return [line.strip() for line in value.splitlines() if line.strip()]


def _k(prefix: str, name: str) -> str:
    return f"{prefix}_{name}"


def _extract_text_from_upload(uploaded_file) -> str:
    if not uploaded_file:
        return ""
    if isinstance(uploaded_file, Path):
        name = uploaded_file.name.lower()
        if name.endswith(".txt"):
            return uploaded_file.read_text(encoding="utf-8", errors="ignore")
        if name.endswith(".json"):
            try:
                payload = json.loads(uploaded_file.read_text(encoding="utf-8", errors="ignore"))
                return _extract_text_from_json(payload)
            except Exception:
                return ""
        if name.endswith(".docx"):
            doc = Document(str(uploaded_file))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if name.endswith(".pdf"):
            if not PDF_AVAILABLE:
                return ""
            try:
                text_parts = []
                with pdfplumber.open(str(uploaded_file)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                return "\n".join(text_parts)
            except Exception:
                return ""
        return ""

    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        return uploaded_file.getvalue().decode("utf-8", errors="ignore")
    if name.endswith(".json"):
        try:
            payload = json.loads(uploaded_file.getvalue().decode("utf-8", errors="ignore"))
            return _extract_text_from_json(payload)
        except Exception:
            return ""
    if name.endswith(".docx"):
        doc = Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if name.endswith(".pdf"):
        if not PDF_AVAILABLE:
            return ""
        try:
            import io
            text_parts = []
            with pdfplumber.open(io.BytesIO(uploaded_file.getvalue())) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except Exception:
            return ""
    return ""


def _extract_text_from_json(payload) -> str:
    """Best-effort extraction for transcript JSON shapes."""
    if payload is None:
        return ""
    # Common shapes: {"text": "..."} or {"transcript": "..."}
    if isinstance(payload, dict):
        for key in ["text", "transcript", "content", "body", "sentence"]:
            value = payload.get(key)
            if isinstance(value, str) and value.strip():
                return value
        # Fireflies-like: {"utterances":[{"speaker":"", "text":""}, ...]}
        for key in ["utterances", "segments", "results", "messages"]:
            items = payload.get(key)
            if isinstance(items, list):
                lines = []
                for item in items:
                    if isinstance(item, dict):
                        text = item.get("text") or item.get("content") or item.get("utterance") or item.get("sentence")
                        speaker = item.get("speaker") or item.get("name")
                        if text:
                            line = f"{speaker}: {text}" if speaker else str(text)
                            lines.append(line)
                if lines:
                    return "\n".join(lines)
        # Nested structures: try values that are lists
        for value in payload.values():
            if isinstance(value, list):
                lines = []
                for item in value:
                    if isinstance(item, dict):
                        text = item.get("text") or item.get("content") or item.get("utterance") or item.get("sentence")
                        if text:
                            lines.append(str(text))
                if lines:
                    return "\n".join(lines)
        return ""
    if isinstance(payload, list):
        lines = []
        for item in payload:
            if isinstance(item, dict):
                text = item.get("text") or item.get("content") or item.get("utterance") or item.get("sentence")
                speaker = item.get("speaker") or item.get("name")
                if text:
                    line = f"{speaker}: {text}" if speaker else str(text)
                    lines.append(line)
            elif isinstance(item, str):
                lines.append(item)
        return "\n".join(lines)
    if isinstance(payload, str):
        return payload
    return ""


def _list_scope_templates() -> list:
    templates = []
    if KNOWLEDGE_DIR.exists():
        for path in KNOWLEDGE_DIR.rglob("*.docx"):
            if "scope" in path.name.lower():
                templates.append(path)
    return sorted(templates)


def _load_projects() -> list:
    if PROJECTS_FILE.exists():
        try:
            data = json.loads(PROJECTS_FILE.read_text(encoding="utf-8"))
            projects = data.get("projects", [])
            return [p for p in projects if isinstance(p, str) and p.strip()]
        except Exception:
            return []
    return []


def _get_project_dir(project: str, strand_id: str) -> Path:
    return CUSTOMER_DATA_DIR / project / strand_id


def _get_uploaded_sheet_names(project: str, strand_id: str) -> set:
    uploaded = set()
    if not project:
        return uploaded
    folder = _get_project_dir(project, strand_id)
    if not folder.exists():
        return uploaded
    for file in folder.rglob("*"):
        if file.name.startswith("~$"):
            continue
        if file.suffix.lower() in [".xlsx", ".xlsm", ".xls"]:
            try:
                xl = pd.ExcelFile(file)
                for sheet in xl.sheet_names:
                    uploaded.add(sheet)
            except Exception:
                continue
    return uploaded


def _load_lc_locations() -> list:
    if LC_LOCATIONS_FILE.exists():
        try:
            data = json.loads(LC_LOCATIONS_FILE.read_text(encoding="utf-8"))
            locations = data.get("locations", [])
            return [loc for loc in locations if isinstance(loc, str) and loc.strip()]
        except Exception:
            return []
    return []


def _slugify_local(value: str) -> str:
    cleaned = []
    for ch in value.strip().lower():
        if ch.isalnum():
            cleaned.append(ch)
        elif ch in [" ", "-", "_"]:
            cleaned.append("-")
    slug = "".join(cleaned)
    while "--" in slug:
        slug = slug.replace("--", "-")
    return slug.strip("-") or "customer"


def _render_lc_documents(locations: list, customers: list, key_prefix: str = "lc"):
    st.subheader("Project Documentation")
    st.caption("Store and download project documents by customer and location.")

    if not locations:
        st.info("No locations configured. Edit config/lc_locations.json to add locations.")
        return

    customer_options = customers + ["(New customer)"]
    customer_choice = st.selectbox("Customer", customer_options, key=_k(key_prefix, "docs_customer"))
    if customer_choice == "(New customer)":
        customer_name = st.text_input("Customer name", key=_k(key_prefix, "docs_customer_name"))
    else:
        customer_name = customer_choice

    location = st.selectbox("Location", locations, key=_k(key_prefix, "docs_location"))
    if not customer_name.strip():
        st.info("Enter a customer name to manage documents.")
        return

    customer_dir = LC_DOCS_DIR / _slugify_local(customer_name)
    location_dir = customer_dir / _slugify_local(location)
    location_dir.mkdir(parents=True, exist_ok=True)

    uploaded = st.file_uploader(
        "Upload documents",
        type=None,
        accept_multiple_files=True,
        key=_k(key_prefix, f"docs_{location}")
    )
    if uploaded:
        for file in uploaded:
            out_path = location_dir / file.name
            with open(out_path, "wb") as f:
                f.write(file.getbuffer())
        st.success(f"Uploaded {len(uploaded)} file(s) for {customer_name} / {location}.")

    files = sorted(location_dir.glob("*"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not files:
        st.info("No documents uploaded for this location.")
        return

    for doc in files:
        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(f"**{doc.name}**")
            st.caption(datetime.fromtimestamp(doc.stat().st_mtime).strftime("%Y-%m-%d %H:%M"))
        with col2:
            with open(doc, "rb") as f:
                st.download_button(
                    "Download",
                    data=f.read(),
                    file_name=doc.name,
                    key=_k(key_prefix, f"dl_doc_{customer_name}_{location}_{doc.name}")
                )


# =============================================================================
# UI COMPONENTS
# =============================================================================

def render_sidebar():
    """Render the sidebar navigation."""
    with st.sidebar:
        # IMP Logo
        st.image("app brand/imp_logo_final.png", width=180)
        st.markdown(f"""
        <p style="color: {IMP_COLORS['purple']}; font-size: 0.9em; margin-top: -10px;">
        Behind Ambitious MAT Finance Teams
        </p>
        """, unsafe_allow_html=True)
        st.markdown("---")

        st.caption(f"Data folder: {CUSTOMER_DATA_DIR}")
        st.markdown("---")

        # Team selection
        st.subheader("Select Team")
        for team_id, team in TEAMS.items():
            col1, col2 = st.columns([1, 4])
            with col1:
                st.markdown(f"### {team['icon']}")
            with col2:
                if st.button(f"{team['name']}", key=f"btn_{team_id}", width="stretch"):
                    st.session_state.selected_team = team_id

        st.markdown("---")

        # Auto-processing toggle
        st.subheader("⚡ Auto-Processing")
        auto = st.toggle("Enable auto-processing", value=st.session_state.auto_process)
        st.session_state.auto_process = auto
        if auto:
            st.success("✅ ON - Auto-process uploads")
            # Show process all button when auto-processing is on
            if st.button("🚀 Process All Teams", width="stretch"):
                for tid in TEAMS:
                    files = get_customer_files(tid)
                    if files:
                        with st.spinner(f"Processing {tid}..."):
                            result = run_team_processing(tid)
                            st.session_state.processing_status[tid] = result
                st.success("All teams processed!")
                st.rerun()
        else:
            st.info("Manual mode")

        st.markdown("---")

        # Quick stats
        st.subheader("📊 Quick Stats")
        for team_id in TEAMS:
            files = get_customer_files(team_id)
            reports = get_recent_reports(team_id)
            st.metric(f"{team_id} Files", len(files), f"{len(reports)} reports")

        st.markdown("---")

        # App Sync Check
        st.subheader("🔄 App Sync")
        if st.button("Check Sync Status", width="stretch"):
            try:
                from teams.app_sync_agent import run_app_sync_check
                result = run_app_sync_check()
                if result["synced"]:
                    st.success("All teams synchronized!")
                else:
                    for tid, report in result["reports"].items():
                        if report["missing_in_app"]:
                            st.warning(f"{tid}: Missing {report['missing_in_app']}")
            except Exception as e:
                st.error(f"Sync check failed: {e}")

        st.markdown("---")

        # Cleanup Controls
        st.subheader("🧹 Cleanup")
        cleanup_col1, cleanup_col2 = st.columns(2)
        with cleanup_col1:
            if st.button("End of Day", width="stretch", help="Keep only latest file per strand"):
                try:
                    from teams.cleanup_agent import run_end_of_day_cleanup
                    result = run_end_of_day_cleanup()
                    if result["files_removed"] > 0:
                        st.success(f"Removed {result['files_removed']} files")
                    else:
                        st.info("No files to clean up")
                except Exception as e:
                    st.error(f"Cleanup failed: {e}")
        with cleanup_col2:
            if st.button("View Storage", width="stretch"):
                try:
                    from teams.cleanup_agent import get_storage_summary
                    summary = get_storage_summary()
                    st.caption(f"Total: {summary['total_files']} files ({summary['total_size']/1024:.1f} KB)")
                    for strand, data in summary["by_strand"].items():
                        st.caption(f"  {strand}: {data['count']} files")
                except Exception as e:
                    st.error(f"Error: {e}")

        # Remove Customer Data button
        if st.button("🗑️ Remove Customer Data", width="stretch", help="Remove all customer data files"):
            st.session_state.show_remove_confirm = True

        if st.session_state.get("show_remove_confirm", False):
            st.warning("⚠️ This will delete all customer data files!")
            remove_scope = st.radio(
                "Select scope:",
                ["Current team only", "All teams"],
                key="remove_scope",
                horizontal=True
            )
            col_confirm, col_cancel = st.columns(2)
            with col_confirm:
                if st.button("✅ Confirm Delete", width="stretch", type="primary"):
                    try:
                        files_removed = 0
                        if remove_scope == "Current team only":
                            team_dir = CUSTOMER_DATA_DIR / st.session_state.selected_team
                            if team_dir.exists():
                                for f in team_dir.rglob("*"):
                                    if f.is_file() and not f.name.startswith("~$"):
                                        f.unlink()
                                        files_removed += 1
                        else:
                            for strand in ["S1", "S2", "S3"]:
                                strand_dir = CUSTOMER_DATA_DIR / strand
                                if strand_dir.exists():
                                    for f in strand_dir.rglob("*"):
                                        if f.is_file() and not f.name.startswith("~$"):
                                            f.unlink()
                                            files_removed += 1
                        st.session_state.show_remove_confirm = False
                        st.success(f"Removed {files_removed} customer data file(s)")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error removing files: {e}")
            with col_cancel:
                if st.button("❌ Cancel", width="stretch"):
                    st.session_state.show_remove_confirm = False
                    st.rerun()

        st.markdown("---")

        # Settings Section
        st.subheader("⚙️ Settings")

        # Data Directory Setting
        current_data_dir = str(get_data_directory())
        st.caption("Data Directory:")
        new_data_dir = st.text_input(
            "Data Directory",
            value=current_data_dir,
            key="data_dir_input",
            label_visibility="collapsed",
            help="Path to customer data folder"
        )

        if new_data_dir != current_data_dir:
            if st.button("Save Data Directory", width="stretch"):
                if Path(new_data_dir).exists():
                    if set_data_directory(new_data_dir):
                        st.success("Data directory updated!")
                        st.rerun()  # Reloads module with new config
                    else:
                        st.error("Failed to save config")
                else:
                    st.warning("Directory does not exist")

        # Templates Directory Setting
        current_templates_dir = str(get_templates_directory())
        st.caption("Templates Directory:")
        new_templates_dir = st.text_input(
            "Templates Directory",
            value=current_templates_dir,
            key="templates_dir_input",
            label_visibility="collapsed",
            help="Path to template workbooks"
        )

        if new_templates_dir != current_templates_dir:
            if st.button("Save Templates Directory", width="stretch"):
                if Path(new_templates_dir).exists():
                    if set_templates_directory(new_templates_dir):
                        st.success("Templates directory updated!")
                        st.rerun()  # Reloads module with new config
                    else:
                        st.error("Failed to save config")
                else:
                    st.warning("Directory does not exist")

        # Show current config
        with st.expander("View Config"):
            config = load_user_config()
            st.json(config)

        # Intelligence Module Settings
        if INFERENCE_AVAILABLE:
            st.markdown("---")
            st.subheader("🧠 Intelligence")

            intel_config = get_intelligence_config()
            conf_thresholds = intel_config.get("confidence", {})

            st.caption(f"High: >{conf_thresholds.get('high_threshold', 0.9):.0%}")
            st.caption(f"Medium: >{conf_thresholds.get('medium_threshold', 0.7):.0%}")

            behavior = intel_config.get("behavior", {})
            if behavior.get("enable_learning"):
                st.success("Learning: ON")
            else:
                st.info("Learning: OFF")

            with st.expander("Intelligence Stats"):
                try:
                    from teams.expert_agents import get_inference_engine
                    engine = get_inference_engine()
                    if engine:
                        stats = engine.get_stats()
                        st.json(stats)
                    else:
                        st.info("InferenceEngine not initialized")
                except Exception as e:
                    st.error(f"Error: {e}")


def render_team_overview(team_id: str):
    """Render the team overview section."""
    team = TEAMS[team_id]

    # Header with IMP branding
    st.markdown(f"""
    <div style="background: linear-gradient(135deg, {IMP_COLORS['deep_purple']}15, {team['color']}20);
                padding: 20px; border-radius: 10px; border-left: 5px solid {team['color']};
                border-top: 1px solid {IMP_COLORS['light_grey']};">
        <h1 style="color: {IMP_COLORS['deep_purple']};">{team['icon']} {team['name']}</h1>
        <p style="font-size: 1.2em; color: {IMP_COLORS['dark_grey']};">{team['description']}</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("")

    # Capabilities
    st.subheader("🎯 Capabilities")
    for cap in team["capabilities"]:
        st.markdown(f"- {cap}")


def render_data_upload(team_id: str):
    """Render the data upload section."""
    st.subheader("📤 Upload Customer Data")

    # S3 Template Mode - Upload pre-populated template first
    if team_id == "S3":
        st.markdown("### 📋 Step 1: Upload Template (Required)")
        st.info("""
        **S3 Workflow:**
        1. **Upload Template** → Contains reference codes (Schools, Finance Codes, Depts)
        2. **Upload Raw Data** → Your customer budget files
        3. **Map Codes** → Match customer codes to template codes (in 'Validate & Map' tab)
        4. **Run S3** → Process data into template format
        """)

        with st.container():

            template_file = st.file_uploader(
                "Upload S3 Template Workbook",
                type=["xlsx", "xlsm"],
                key="s3_template_upload",
                help="Upload the pre-populated S3 template workbook. Customer data will be written into this template."
            )

            if template_file:
                # Save template to templates directory
                template_dir = TEMPLATES_DIR / "S3"
                template_dir.mkdir(parents=True, exist_ok=True)
                template_path = template_dir / template_file.name

                with open(template_path, "wb") as f:
                    f.write(template_file.getbuffer())

                st.session_state["s3_template_path"] = template_path
                st.success(f"✅ Template uploaded: {template_file.name}")
                st.info("🔧 **Template Mode ENABLED** - Customer data will be written into this template")

            # Show current template status
            current_template = st.session_state.get("s3_template_path")
            if current_template and Path(current_template).exists():
                st.success(f"📋 Active Template: {Path(current_template).name}")
                if st.button("Clear Template (use Raw Data mode)", key="clear_s3_template"):
                    st.session_state["s3_template_path"] = None
                    st.rerun()
            else:
                st.caption("No template selected - using Raw Data mode")

        st.markdown("---")
        st.markdown("### 📁 Step 2: Upload Raw Customer Data")

    # Show auto-processing status
    if st.session_state.auto_process:
        st.success("⚡ Auto-processing is **ON** - Files will be processed immediately after upload")

    upload_dir = CUSTOMER_DATA_DIR / team_id
    if team_id != "S3":
        st.caption(f"Upload to: {upload_dir}")

    # Use type=None to accept all files (Streamlit has MIME type issues with .docx on Windows)
    # Validate file types after upload instead
    ALLOWED_EXTENSIONS = {".xlsx", ".xlsm", ".xls", ".csv", ".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg"}

    uploaded_files = st.file_uploader(
        f"Upload {'raw budget files' if team_id == 'S3' else 'files'} for {team_id}",
        type=None,  # Accept all files - validate after upload (MIME type issues on Windows)
        accept_multiple_files=True,
        key=f"upload_{team_id}",
        help="Supported: Excel (.xlsx, .xls), CSV, PDF, Word (.docx, .doc), Images"
    )

    if uploaded_files:
        upload_dir.mkdir(parents=True, exist_ok=True)

        for uploaded_file in uploaded_files:
            # Validate file extension
            file_ext = Path(uploaded_file.name).suffix.lower()
            if file_ext not in ALLOWED_EXTENSIONS:
                st.error(f"❌ {uploaded_file.name}: Unsupported file type '{file_ext}'")
                continue

            file_path = upload_dir / uploaded_file.name
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            st.success(f"✅ Uploaded: {uploaded_file.name}")

        if st.session_state.auto_process:
            st.info("🔄 Auto-processing triggered...")
            with st.spinner(f"Running {team_id} Specialist Agent..."):
                result = run_team_processing(team_id)
                st.session_state.processing_status[team_id] = result
                if result.get("success"):
                    st.success("✅ Auto-processing completed!")
                else:
                    st.warning("⚠️ Auto-processing completed with issues")


def render_customer_files(team_id: str):
    """Render the customer files section."""
    st.subheader("📁 Customer Data Files")

    files = get_customer_files(team_id)
    base_path = CUSTOMER_DATA_DIR / team_id

    st.caption(f"Data folder: {base_path}")

    if not files:
        st.info(f"No customer data files found for {team_id}. Upload files above.")
        return

    # Group by subfolder
    file_groups = {}
    for f in files:
        try:
            rel_path = f.relative_to(base_path)
            folder = str(rel_path.parent) if rel_path.parent != Path(".") else "Root"
        except ValueError:
            folder = "Root"
        if folder not in file_groups:
            file_groups[folder] = []
        file_groups[folder].append(f)

    for folder, folder_files in file_groups.items():
        with st.expander(f"📂 {folder} ({len(folder_files)} files)", expanded=folder == "Root"):
            for f in folder_files:
                col1, col2, col3 = st.columns([3, 1, 1])
                with col1:
                    st.markdown(f"📄 **{f.name}**")
                with col2:
                    st.caption(format_file_size(f.stat().st_size))
                with col3:
                    st.caption(datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d"))


def render_processing(team_id: str):
    """Render the processing section."""
    st.subheader("⚙️ Process Data")

    # Show auto-processing status
    if st.session_state.auto_process:
        st.success("⚡ **Auto-Processing ENABLED** - Files will be processed automatically on upload")
    else:
        st.caption("💡 Enable auto-processing in sidebar to automatically process uploaded files")

    # Show specialist agent info
    agent_info = {
        "S1": "Structure Specialist - Finance codes, schools, departments, COA mapping",
        "S2": "Staff Specialist - Pay scales, staff roles, contracts, allowances",
        "S3": "Financial Specialist - Budgets, grants, pupil numbers, scenarios"
    }
    st.info(f"🤖 **{agent_info.get(team_id, 'Specialist Agent')}**")

    col1, col2, col3 = st.columns([2, 1, 1])

    with col1:
        if st.button(f"🚀 Run {team_id} Specialist Agent", type="primary", width="stretch"):
            with st.spinner(f"Running {team_id} Specialist Agent..."):
                result = run_team_processing(team_id)
                st.session_state.processing_status[team_id] = result

    with col2:
        if st.button("🔄 Refresh", width="stretch"):
            st.rerun()

    with col3:
        if st.button("🗑️ Clear Results", width="stretch"):
            if team_id in st.session_state.processing_status:
                del st.session_state.processing_status[team_id]
                st.rerun()

    # Show results
    if team_id in st.session_state.processing_status:
        result = st.session_state.processing_status[team_id]

        if result["success"]:
            st.success("✅ Processing completed successfully!")
        else:
            st.warning("⚠️ Processing completed with issues")

        # Summary metrics - team specific
        summary = result.get("summary", {})
        if summary:
            if team_id == "S2":
                cols = st.columns(4)
                cols[0].metric("Staff Members", summary.get("staff_members", 0))
                cols[1].metric("Staff Roles", summary.get("staff_roles", 0))
                cols[2].metric("Teaching Contracts", summary.get("teaching_contracts", 0))
                cols[3].metric("Support Contracts", summary.get("support_contracts", 0))

                cols2 = st.columns(4)
                cols2[0].metric("Pay Scales", summary.get("pay_scales", 0))
                cols2[1].metric("Pay Scale Points", summary.get("pay_scale_points", 0))
                audit_score = summary.get("audit_score", 0)
                cols2[2].metric("Audit Score", f"{audit_score:.1f}%")
                cols2[3].metric("Audit Passed", "YES" if summary.get("audit_passed", False) else "NO")

                # Show customer data load status
                customer_data_loaded = result.get("customer_data_loaded", False)
                if customer_data_loaded:
                    st.success("✅ Customer data successfully loaded and processed")
                else:
                    st.warning("⚠️ No customer data files found - using defaults")

                # Show audit status
                if summary.get("audit_passed", False):
                    st.success(f"✅ External Audit PASSED with score {audit_score:.1f}%")
                elif audit_score > 0:
                    st.warning(f"⚠️ External Audit FAILED with score {audit_score:.1f}%")

                # Show data source warnings
                data_warnings = result.get("data_source_warnings", [])
                if data_warnings:
                    with st.expander(f"📋 Data Source Warnings ({len(data_warnings)})", expanded=False):
                        for warning in data_warnings:
                            st.warning(warning)

                # Show detailed audit report for S2
                audit_data = result.get("audit", {})
                detailed_report = audit_data.get("detailed_report", {})
                if detailed_report:
                    issues = detailed_report.get("issues", [])
                    recommendations = detailed_report.get("recommendations", [])

                    if issues:
                        with st.expander(f"🔍 Audit Details - What's Missing & Why ({len(issues)} issues)", expanded=not summary.get("audit_passed", False)):
                            for issue in issues:
                                severity_icon = "🔴" if issue.get("severity") == "error" else "🟡" if issue.get("severity") == "warning" else "🔵"
                                st.markdown(f"### {severity_icon} {issue.get('check', 'Unknown Issue')}")
                                st.markdown(f"**Category:** {issue.get('category', '').replace('_', ' ').title()}")

                                col1, col2 = st.columns(2)
                                with col1:
                                    st.markdown("**What's Missing:**")
                                    st.info(issue.get("what_is_missing", "N/A"))
                                with col2:
                                    st.markdown("**Why It Matters:**")
                                    st.warning(issue.get("why_it_matters", "N/A"))

                                st.markdown("**How To Fix:**")
                                st.success(issue.get("how_to_fix", "N/A"))
                                st.markdown("---")

                    if recommendations:
                        with st.expander(f"💡 Recommendations ({len(recommendations)})", expanded=False):
                            for rec in recommendations:
                                priority_color = "🔴" if rec.get("priority") == "HIGH" else "🟡" if rec.get("priority") == "MEDIUM" else "🟢"
                                st.markdown(f"{priority_color} **[{rec.get('priority', '')}]** {rec.get('action', '')}")
                                st.caption(f"Reason: {rec.get('reason', '')}")

            elif team_id == "S1":
                cols = st.columns(4)
                cols[0].metric("Finance Codes", summary.get("finance_codes", 0))
                cols[1].metric("Schools", summary.get("schools", 0))
                cols[2].metric("Departments", summary.get("departments", 0))
                cols[3].metric("Validation Errors", summary.get("validation_errors", 0))

                cols2 = st.columns(4)
                cols2[0].metric("Validation Warnings", summary.get("validation_warnings", 0))
                audit_score = summary.get("audit_score", 0)
                cols2[1].metric("Audit Score", f"{audit_score:.1f}%")
                cols2[2].metric("Audit Passed", "YES" if summary.get("audit_passed", False) else "NO")
                cols2[3].empty()  # Placeholder

                # Show audit status
                if summary.get("audit_passed", False):
                    st.success(f"✅ External Audit PASSED with score {audit_score:.1f}%")
                elif audit_score > 0:
                    st.warning(f"⚠️ External Audit FAILED with score {audit_score:.1f}%")

                # Show detailed audit report if available
                audit_data = result.get("audit", {})
                detailed_report = audit_data.get("detailed_report", {})
                if detailed_report:
                    issues = detailed_report.get("issues", [])
                    recommendations = detailed_report.get("recommendations", [])

                    if issues:
                        with st.expander(f"🔍 Audit Details - What's Missing & Why ({len(issues)} issues)", expanded=not summary.get("audit_passed", False)):
                            for issue in issues:
                                severity_icon = "🔴" if issue.get("severity") == "error" else "🟡" if issue.get("severity") == "warning" else "🔵"
                                st.markdown(f"### {severity_icon} {issue.get('check', 'Unknown Issue')}")
                                st.markdown(f"**Category:** {issue.get('category', '').replace('_', ' ').title()}")

                                col1, col2 = st.columns(2)
                                with col1:
                                    st.markdown("**What's Missing:**")
                                    st.info(issue.get("what_is_missing", "N/A"))
                                with col2:
                                    st.markdown("**Why It Matters:**")
                                    st.warning(issue.get("why_it_matters", "N/A"))

                                st.markdown("**How To Fix:**")
                                st.success(issue.get("how_to_fix", "N/A"))

                                affected = issue.get("affected_records", [])
                                if affected:
                                    st.markdown(f"**Affected Records:** `{', '.join(affected[:10])}`" +
                                              (f" (+{len(affected)-10} more)" if len(affected) > 10 else ""))
                                st.markdown("---")

                    if recommendations:
                        with st.expander(f"💡 Recommendations ({len(recommendations)})", expanded=False):
                            for rec in recommendations:
                                priority_color = "🔴" if rec.get("priority") == "HIGH" else "🟡" if rec.get("priority") == "MEDIUM" else "🟢"
                                st.markdown(f"{priority_color} **[{rec.get('priority', '')}]** {rec.get('action', '')}")
                                st.caption(f"Reason: {rec.get('reason', '')}")

            elif team_id == "S3":
                # Show build mode
                build_mode = result.get("build_mode", "raw_data")
                if build_mode == "template":
                    st.info("📋 **Template Mode** - Data written into pre-populated template")
                else:
                    st.caption("📄 **Raw Data Mode** - New workbook created")

                cols = st.columns(4)
                cols[0].metric("Pupil Records", summary.get("pupils", 0))
                cols[1].metric("Grants", summary.get("grants", 0))
                cols[2].metric("Income Lines", summary.get("income_lines", 0))
                cols[3].metric("Expenditure Lines", summary.get("expenditure_lines", 0))

                cols2 = st.columns(4)
                schools = summary.get("schools", [])
                cols2[0].metric("Schools", len(schools) if isinstance(schools, list) else schools)
                audit_score = summary.get("audit_score", 0)
                cols2[1].metric("Audit Score", f"{audit_score:.1f}%")
                cols2[2].metric("Audit Passed", "YES" if summary.get("audit_passed", False) else "NO")
                cols2[3].empty()  # Placeholder

                # Show audit status
                if summary.get("audit_passed", False):
                    st.success(f"✅ External Audit PASSED with score {audit_score:.1f}%")
                elif audit_score > 0:
                    st.warning(f"⚠️ External Audit FAILED with score {audit_score:.1f}%")

                # Show detailed audit report for S3
                audit_data = result.get("audit", {})
                detailed_report = audit_data.get("detailed_report", {})
                if detailed_report:
                    issues = detailed_report.get("issues", [])
                    recommendations = detailed_report.get("recommendations", [])

                    if issues:
                        with st.expander(f"🔍 Audit Details - What's Missing & Why ({len(issues)} issues)", expanded=not summary.get("audit_passed", False)):
                            for issue in issues:
                                severity_icon = "🔴" if issue.get("severity") == "error" else "🟡" if issue.get("severity") == "warning" else "🔵"
                                st.markdown(f"### {severity_icon} {issue.get('check', 'Unknown Issue')}")
                                st.markdown(f"**Category:** {issue.get('category', '').replace('_', ' ').title()}")

                                col1, col2 = st.columns(2)
                                with col1:
                                    st.markdown("**What's Missing:**")
                                    st.info(issue.get("what_is_missing", "N/A"))
                                with col2:
                                    st.markdown("**Why It Matters:**")
                                    st.warning(issue.get("why_it_matters", "N/A"))

                                st.markdown("**How To Fix:**")
                                st.success(issue.get("how_to_fix", "N/A"))
                                st.markdown("---")

                    if recommendations:
                        with st.expander(f"💡 Recommendations ({len(recommendations)})", expanded=False):
                            for rec in recommendations:
                                priority_color = "🔴" if rec.get("priority") == "HIGH" else "🟡" if rec.get("priority") == "MEDIUM" else "🟢"
                                st.markdown(f"{priority_color} **[{rec.get('priority', '')}]** {rec.get('action', '')}")
                                st.caption(f"Reason: {rec.get('reason', '')}")

        # Template sheets built
        template_sheets = result.get("template_sheets", {})
        if template_sheets:
            with st.expander(f"📋 Template Sheets Built ({len(template_sheets)})", expanded=True):
                sheet_cols = st.columns(3)
                for idx, (sheet_name, df) in enumerate(template_sheets.items()):
                    col_idx = idx % 3
                    rows = len(df) if hasattr(df, '__len__') else 0
                    status = "✅" if rows > 0 else "⚪"
                    sheet_cols[col_idx].markdown(f"{status} **{sheet_name}**: {rows} rows")

        # S2-specific: Quick summary of processing details (full data in Excel output)
        if team_id == "S2":
            created_codes = result.get("created_role_codes", [])
            created_groups = result.get("created_role_groups", [])
            assumptions_list = result.get("assumptions", [])
            skipped_staff = result.get("skipped_staff", [])

            if created_codes or created_groups or assumptions_list or skipped_staff:
                with st.expander("📋 Processing Details (see output file for full data)", expanded=False):
                    col1, col2, col3, col4 = st.columns(4)
                    col1.metric("Role Codes Created", len(created_codes))
                    col2.metric("Role Groups Created", len(created_groups))
                    col3.metric("Assumptions Made", len(assumptions_list))
                    col4.metric("Staff Skipped", len(skipped_staff))

                    st.caption("Full details available in output Excel file sheets: _CreatedRoleCodes, _CreatedRoleGroups, _Assumptions, _SkippedStaff")

                    if skipped_staff:
                        st.warning(f"⚠️ {len(skipped_staff)} staff records were skipped - check _SkippedStaff sheet for details")

            # Show unclassified data that needs user mapping
            unclassified_data = result.get("unclassified_data", [])
            if unclassified_data:
                with st.expander(f"🔍 Unclassified Data - Needs Manual Mapping ({len(unclassified_data)} sheets)", expanded=True):
                    st.warning(f"**{len(unclassified_data)} data sheet(s) could not be auto-classified.** Please review and specify the data type.")

                    for idx, item in enumerate(unclassified_data):
                        st.markdown(f"**{item['file_name']}** / {item['sheet_name']} ({item['row_count']} rows)")
                        st.caption(f"Columns: {', '.join(item['columns'][:10])}{'...' if len(item['columns']) > 10 else ''}")

                        # Show sample data
                        if item.get('sample_data'):
                            # Pad arrays to same length to avoid DataFrame error
                            sample_items = list(item['sample_data'].items())[:5]
                            if sample_items:
                                max_len = max(len(v) for k, v in sample_items)
                                padded_data = {k: v + [None] * (max_len - len(v)) for k, v in sample_items}
                                sample_df = pd.DataFrame(padded_data)
                                st.dataframe(sample_df, use_container_width=True, hide_index=True)

                        # Let user specify what type of data this is
                        data_type = st.selectbox(
                            "What type of data is this?",
                            ["-- Select --", "Staff Contracts", "Pay Scales", "Allowances", "Not S2 Data (Ignore)"],
                            key=f"unclassified_type_{idx}"
                        )

                        if data_type != "-- Select --":
                            st.success(f"Marked as: {data_type}")
                            # Store in session state for reprocessing
                            if "unclassified_mappings" not in st.session_state:
                                st.session_state.unclassified_mappings = {}
                            st.session_state.unclassified_mappings[f"{item['file_name']}_{item['sheet_name']}"] = data_type

                        st.divider()

                    if st.button("Reprocess with Manual Classifications", key="reprocess_unclassified"):
                        st.info("Please re-run processing to apply your classifications")

        # Issues - show with full detail
        if result.get("issues"):
            with st.expander(f"⚠️ Issues ({len(result['issues'])})", expanded=True):
                for idx, issue in enumerate(result["issues"]):
                    # Check if it's a multi-line detailed error
                    if "\n" in str(issue):
                        # Show detailed errors in code block for readability
                        st.markdown(f"**Issue {idx + 1}:**")
                        st.code(issue, language=None)
                    else:
                        st.error(issue)

        # Processing Log (diagnostic info) - ALWAYS show for debugging
        processing_log = result.get("processing_log", [])
        with st.expander(f"📋 Processing Log ({len(processing_log)} entries)", expanded=True):
            if processing_log:
                for entry in processing_log:
                    st.text(entry)
            else:
                st.warning("No processing log entries - this may indicate the changes haven't been deployed yet")

        # Assumptions (for non-S2 teams, or as fallback)
        if team_id != "S2" and result.get("assumptions"):
            with st.expander(f"📝 Assumptions ({len(result['assumptions'])})"):
                for assumption in result["assumptions"]:
                    st.info(assumption)

        # Inference Details (if available)
        inference_result = result.get("inference_result")
        if inference_result:
            with st.expander("🧠 Inference Details", expanded=False):
                st.markdown("**Strand Detection Analysis**")

                # Confidence display
                confidence = inference_result.get("confidence", 0)
                conf_level = inference_result.get("confidence_level", "unknown")
                decision = inference_result.get("decision", "N/A")

                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Decision", decision)
                with col2:
                    st.metric("Confidence", f"{confidence:.0%}")
                with col3:
                    # Color-code confidence level
                    level_colors = {"high": "🟢", "medium": "🟡", "low": "🔴"}
                    st.metric("Level", f"{level_colors.get(conf_level, '⚪')} {conf_level.upper()}")

                # Reasoning
                reasoning = inference_result.get("reasoning", [])
                if reasoning:
                    st.markdown("**Reasoning:**")
                    for reason in reasoning:
                        st.markdown(f"- {reason}")

                # Alternatives
                alternatives = inference_result.get("alternatives", [])
                if alternatives:
                    st.markdown("**Alternatives Considered:**")
                    for alt in alternatives[:3]:
                        st.caption(f"• {alt.get('strand', 'N/A')}: {alt.get('confidence', 0):.0%}")

                # Review flag
                if inference_result.get("requires_review"):
                    st.warning("⚠️ This decision was flagged for review due to low confidence")

        # Reasoning Trail (for debugging)
        reasoning_trail = result.get("reasoning_trail")
        if reasoning_trail and st.checkbox("Show Full Reasoning Trail", key=f"trail_{team_id}"):
            with st.expander("🔍 Full Reasoning Trail (Debug)", expanded=False):
                st.json(reasoning_trail)

        # Download output
        output_file = result.get("output_file")
        if output_file and Path(output_file).exists():
            with open(output_file, "rb") as f:
                st.download_button(
                    label="📥 Download Template Data",
                    data=f.read(),
                    file_name=Path(output_file).name,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    type="primary"
                )


def render_reports(team_id: str):
    """Render the reports section."""
    st.subheader("📊 Recent Reports")

    reports = get_recent_reports(team_id, limit=10)

    if not reports:
        st.info("No reports generated yet. Run processing to create reports.")
        return

    for report in reports:
        col1, col2, col3 = st.columns([3, 1, 1])
        with col1:
            st.markdown(f"📄 **{report.name}**")
        with col2:
            st.caption(datetime.fromtimestamp(report.stat().st_mtime).strftime("%Y-%m-%d %H:%M"))
        with col3:
            with open(report, "rb") as f:
                st.download_button(
                    "⬇️",
                    data=f.read(),
                    file_name=report.name,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key=f"dl_{report.name}"
                )


def render_preview(team_id: str):
    """Render data preview section."""
    st.subheader("👁️ Data Preview")

    # Get latest report
    reports = get_recent_reports(team_id, limit=1)

    if not reports:
        st.info("No processed data to preview. Run processing first.")
        return

    report_file = reports[0]

    try:
        xl = pd.ExcelFile(report_file)

        selected_sheet = st.selectbox("Select sheet to preview", xl.sheet_names)

        if selected_sheet:
            df = pd.read_excel(xl, selected_sheet)

            # Filter options
            col1, col2 = st.columns(2)
            with col1:
                search = st.text_input("🔍 Search", placeholder="Filter rows...")
            with col2:
                max_rows = st.slider("Max rows", 10, 100, 25)

            # Apply filter
            if search:
                mask = df.astype(str).apply(lambda x: x.str.contains(search, case=False, na=False)).any(axis=1)
                df = df[mask]

            st.dataframe(df.head(max_rows), width="stretch")
            st.caption(f"Showing {min(len(df), max_rows)} of {len(df)} rows")

    except Exception as e:
        st.error(f"Error loading preview: {e}")


# =============================================================================
# VALIDATE & MAP TAB
# =============================================================================

def render_validate_and_map(team_id: str):
    """Render the Validate & Map tab for pre-flight column validation."""
    st.subheader("Validate & Map Columns")

    if not PREFLIGHT_AVAILABLE:
        st.warning("Pre-flight validation module not available. Install required dependencies.")
        return

    # Initialize validator if needed (with column learner for learned mappings)
    if st.session_state.preflight_validator is None:
        try:
            learner = st.session_state.get("column_learner")
            st.session_state.preflight_validator = PreFlightValidator(column_learner=learner)
        except Exception as e:
            st.error(f"Failed to initialize validator: {e}")
            return

    validator = st.session_state.preflight_validator

    # Show learned mappings info
    learner = st.session_state.get("column_learner")
    if learner and LEARNER_AVAILABLE:
        stats = learner.get_statistics()
        if stats['total_mappings'] > 0:
            with st.expander(f"Learned Mappings ({stats['total_mappings']} saved)", expanded=False):
                cols = st.columns(3)
                cols[0].metric("Total Learned", stats['total_mappings'])
                cols[1].metric("Times Applied", stats['total_applications'])
                cols[2].metric("Ignored Columns", stats['ignored_columns'])
                st.caption("Column mappings you've confirmed are automatically applied to new files.")

    # Step 1: File Selection
    st.markdown("### Step 1: Select Files to Validate")

    files = get_customer_files(team_id)

    if not files:
        st.info(f"No customer data files found for {team_id}. Upload files in the 'Upload & Files' tab first.")
        return

    # File checkboxes
    selected_files = []
    cols = st.columns(2)
    for idx, f in enumerate(files):  # Show all files
        col_idx = idx % 2
        with cols[col_idx]:
            if st.checkbox(f.name, key=f"validate_file_{f.name}_{team_id}", value=True):
                selected_files.append(f)

    if not selected_files:
        st.info("Select at least one file to validate.")
        return

    # Validate button
    if st.button("Analyze Selected Files", type="primary", width="stretch"):
        with st.spinner("Analyzing columns..."):
            results = {}
            for file_path in selected_files:
                try:
                    # Use validate_file_all_sheets to get all sheets/tables
                    file_results = validator.validate_file_all_sheets(file_path, strand=team_id)
                    for result in file_results:
                        key = f"{result.file_name}:{result.sheet_name or 'default'}"
                        results[key] = result
                except Exception as e:
                    error_type = type(e).__name__
                    st.error(f"Error validating {file_path.name}")
                    with st.expander("Error Details", expanded=True):
                        st.code(f"File: {file_path}\nType: {error_type}\nDetails: {str(e)}", language=None)

            st.session_state.validation_results[team_id] = results
            total_sheets = len(results)
            total_files = len(selected_files)
            st.success(f"Analyzed {total_sheets} sheet(s)/table(s) from {total_files} file(s)")

    # Step 2: Show Results
    if team_id in st.session_state.validation_results and st.session_state.validation_results[team_id]:
        results = st.session_state.validation_results[team_id]

        st.markdown("### Step 2: Column Analysis Results")

        # Group results by file name (consolidate multiple sheets into one entry per file)
        files_grouped = {}
        for file_key, result in results.items():
            file_name = result.file_name
            if file_name not in files_grouped:
                files_grouped[file_name] = []
            files_grouped[file_name].append(result)

        # Show one expander per file with aggregated metrics
        for file_name, file_results in files_grouped.items():
            # Aggregate metrics across all sheets/tables in this file
            total_rows = sum(r.row_count for r in file_results)
            total_cols = sum(r.column_count for r in file_results)
            total_matched = sum(r.mapping_summary['matched'] for r in file_results)
            total_review = sum(r.mapping_summary['review'] for r in file_results)
            total_unmapped = sum(r.mapping_summary['unmapped'] for r in file_results)
            total_corrected = sum(r.mapping_summary['corrected'] for r in file_results)
            total_ignored = sum(r.mapping_summary['ignored'] for r in file_results)

            # Get best strand detection
            best_strand = None
            best_confidence = 0
            for r in file_results:
                if r.detected_strand and r.strand_confidence > best_confidence:
                    best_strand = r.detected_strand
                    best_confidence = r.strand_confidence

            sheet_count = len(file_results)
            sheet_info = f", {sheet_count} sheets" if sheet_count > 1 else ""

            with st.expander(f"**{file_name}** ({total_rows} rows, {total_cols} columns{sheet_info})", expanded=True):
                # Summary metrics (aggregated)
                # Note: All mappings require user review - no auto-acceptance
                total_suggestions = total_matched + total_review  # Columns with suggested mappings
                cols = st.columns(4)
                cols[0].metric("Suggested", total_suggestions, help="Columns with suggested mappings (require confirmation)")
                cols[1].metric("Unmapped", total_unmapped, help="No suggestions - needs manual mapping")
                cols[2].metric("Corrected", total_corrected, help="User corrections applied")
                cols[3].metric("Ignored", total_ignored, help="Columns marked to ignore")

                # Strand detection
                if best_strand:
                    confidence_color = "green" if best_confidence >= 0.8 else "orange" if best_confidence >= 0.6 else "red"
                    st.markdown(f"**Detected Strand:** {best_strand} (:{confidence_color}[{best_confidence:.0%}])")

                # Errors (deduplicated across sheets) - show with full detail
                all_errors = []
                for r in file_results:
                    if hasattr(r, 'errors') and r.errors:
                        all_errors.extend(r.errors)
                if all_errors:
                    st.markdown("**Errors:**")
                    for error in all_errors:
                        # Show errors in expandable format for readability
                        with st.expander(f"Error: {error.split(chr(10))[0][:80]}...", expanded=True):
                            st.code(error, language=None)

                # Warnings (deduplicated across sheets)
                all_warnings = set()
                for r in file_results:
                    all_warnings.update(r.warnings or [])
                for warning in all_warnings:
                    st.warning(warning)

                # NOTE: All columns now require user review - no auto-accepted columns
                # Matched columns are included in the review section below

        # Collect and deduplicate columns needing review across all files
        all_review = {}  # source_column -> (mapping, file_keys)
        all_unmapped = {}  # source_column -> (mapping, file_keys)

        for file_key, result in results.items():
            for m in result.review_columns:
                if m.source_column and m.source_column.strip():
                    if m.source_column not in all_review:
                        all_review[m.source_column] = (m, [file_key])
                    else:
                        all_review[m.source_column][1].append(file_key)
            for m in result.unmapped_columns:
                if m.source_column and m.source_column.strip():
                    if m.source_column not in all_unmapped:
                        all_unmapped[m.source_column] = (m, [file_key])
                    else:
                        all_unmapped[m.source_column][1].append(file_key)

        # Show consolidated review columns (deduplicated)
        # All columns with suggested mappings require user confirmation
        if all_review:
            st.markdown("**Review Column Mappings** *(confirm or change each mapping)*")
            st.caption("All mappings require your confirmation. Use the dropdown to change any mapping regardless of confidence score.")
            for source_col, (mapping, file_keys) in all_review.items():
                col1, col2, col3, col4 = st.columns([3, 3, 2, 2])
                mapping_key = f"review_{source_col}"
                with col1:
                    st.text(source_col)
                with col2:
                    fuzzy_options = [mapping.mapped_to or ""] + [alt[0] for alt in mapping.alternatives]
                    fuzzy_options = [opt for opt in fuzzy_options if opt]
                    all_fields = get_field_options_for_strand(team_id)
                    standard_options = [f for f in all_fields if f not in fuzzy_options]
                    options = fuzzy_options.copy()
                    if standard_options:
                        options.append("── All Fields ──")
                        options.extend(standard_options)
                    options.append("-- Type custom --")
                    options.append("-- Ignore this column --")
                    options.append("-- Keep original --")

                    selected = st.selectbox("Map to", options, key=f"map_{mapping_key}", label_visibility="collapsed")

                    if selected == "-- Type custom --":
                        custom_value = st.text_input("Custom", key=f"custom_{mapping_key}", placeholder="Type column name...", label_visibility="collapsed")
                        if custom_value.strip():
                            st.session_state.custom_mappings[mapping_key] = custom_value.strip()
                    elif selected == "-- Ignore this column --":
                        st.session_state.custom_mappings[mapping_key] = "__IGNORE__"
                    elif selected == "-- Keep original --":
                        st.session_state.custom_mappings[mapping_key] = source_col
                    elif selected != "── All Fields ──" and selected:
                        st.session_state.custom_mappings[mapping_key] = selected
                with col3:
                    # Color-coded confidence indicator
                    conf = mapping.confidence
                    if conf >= 0.95:
                        st.caption(f":green[{conf:.0%}] (high)")
                    elif conf >= 0.8:
                        st.caption(f":orange[{conf:.0%}] (medium)")
                    else:
                        st.caption(f":red[{conf:.0%}] (low)")
                with col4:
                    if mapping.sample_values:
                        st.caption(f"e.g., {str(mapping.sample_values[0])[:20]}")

        # Show consolidated unmapped columns (deduplicated)
        if all_unmapped:
            st.markdown("**Unmapped Columns (needs manual mapping)**")
            for source_col, (mapping, file_keys) in all_unmapped.items():
                col1, col2, col3 = st.columns([3, 4, 3])
                mapping_key = f"unmap_{source_col}"
                with col1:
                    st.text(source_col)
                with col2:
                    suggestions = [alt[0] for alt in mapping.alternatives] if mapping.alternatives else []
                    all_fields = get_field_options_for_strand(team_id)
                    standard_options = [f for f in all_fields if f not in suggestions]
                    options = suggestions.copy()
                    if standard_options:
                        options.append("── All Fields ──")
                        options.extend(standard_options)
                    options.append("-- Type custom --")
                    options.append("-- Ignore this column --")
                    options.append("-- Keep original --")

                    selected = st.selectbox("Map to", options, key=f"select_{mapping_key}", label_visibility="collapsed", index=0)

                    if selected == "-- Type custom --":
                        custom_value = st.text_input("Custom", key=f"custom_{mapping_key}", placeholder="Type column name...", label_visibility="collapsed")
                        if custom_value.strip():
                            st.session_state.custom_mappings[mapping_key] = custom_value.strip()
                    elif selected == "-- Ignore this column --":
                        st.session_state.custom_mappings[mapping_key] = "__IGNORE__"
                    elif selected == "-- Keep original --":
                        st.session_state.custom_mappings[mapping_key] = source_col
                    elif selected != "── All Fields ──" and selected:
                        st.session_state.custom_mappings[mapping_key] = selected
                with col3:
                    if mapping.sample_values:
                        st.caption(f"e.g., {str(mapping.sample_values[0])[:20]}")

        # Step 3: Confirm & Proceed
        st.markdown("### Step 3: Confirm & Proceed")

        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("Confirm Mappings", type="primary", use_container_width=True):
                # Collect all final mappings using deduplicated keys
                all_mappings = {}
                custom_maps = st.session_state.get("custom_mappings", {})

                for file_key, result in results.items():
                    file_mappings = {}
                    for mapping in result.column_mappings:
                        source_col = mapping.source_column
                        # Check deduplicated keys first (review_ and unmap_ prefixes)
                        review_key = f"review_{source_col}"
                        unmap_key = f"unmap_{source_col}"

                        # Check for user selection in custom_maps
                        if review_key in custom_maps and custom_maps[review_key] != "__IGNORE__":
                            file_mappings[source_col] = custom_maps[review_key]
                        elif unmap_key in custom_maps and custom_maps[unmap_key] != "__IGNORE__":
                            file_mappings[source_col] = custom_maps[unmap_key]
                        elif mapping.user_override and mapping.user_override != "__IGNORE__":
                            # User explicitly overrode this mapping
                            file_mappings[source_col] = mapping.user_override
                        else:
                            # Check selectbox state directly for review columns
                            selectbox_key = f"map_{review_key}"
                            selectbox_value = st.session_state.get(selectbox_key)
                            if selectbox_value and selectbox_value not in ["── All Fields ──", "-- Type custom --", "-- Ignore this column --", "-- Keep original --"]:
                                file_mappings[source_col] = selectbox_value
                            elif selectbox_value == "-- Keep original --":
                                file_mappings[source_col] = source_col
                            elif mapping.mapped_to and mapping.status == 'review':
                                # Use suggested mapping if user didn't change it
                                file_mappings[source_col] = mapping.mapped_to

                    all_mappings[file_key] = file_mappings

                st.session_state.column_mappings[team_id] = all_mappings
                st.session_state.mapping_validated[team_id] = True
                st.success("Mappings confirmed! You can now proceed to Process tab.")

        with col2:
            # Save to Memory button - learns from user corrections
            learner = st.session_state.get("column_learner")
            custom_maps = st.session_state.get("custom_mappings", {})
            has_corrections = any(
                k.startswith("review_") or k.startswith("unmap_")
                for k in custom_maps.keys()
            )

            if st.button("Save to Memory", use_container_width=True, disabled=not has_corrections,
                        help="Save your column mappings so future files use them automatically"):
                if learner and has_corrections:
                    # Extract user corrections (both review and unmapped)
                    corrections = {}
                    for key, value in custom_maps.items():
                        if key.startswith("review_"):
                            source_col = key[7:]  # Remove "review_" prefix
                            corrections[source_col] = value
                        elif key.startswith("unmap_"):
                            source_col = key[6:]  # Remove "unmap_" prefix
                            corrections[source_col] = value

                    # Learn from corrections
                    learned_count = learner.learn_from_corrections(corrections, strand=team_id)
                    if learned_count > 0:
                        st.success(f"Learned {learned_count} column mappings for future use!")
                    else:
                        st.info("No new mappings to learn.")
                else:
                    st.warning("No corrections to save." if not has_corrections else "Learner not available.")

        with col3:
            if st.button("Clear & Re-analyze", use_container_width=True):
                if team_id in st.session_state.validation_results:
                    del st.session_state.validation_results[team_id]
                if team_id in st.session_state.column_mappings:
                    del st.session_state.column_mappings[team_id]
                if team_id in st.session_state.mapping_validated:
                    del st.session_state.mapping_validated[team_id]
                # Clear custom mappings for this team
                st.session_state.custom_mappings = {k: v for k, v in st.session_state.custom_mappings.items() if team_id not in k}
                # Force new validator instance
                st.session_state.preflight_validator = None
                st.rerun()

        # Show validation status
        if st.session_state.mapping_validated.get(team_id, False):
            st.success("Mappings have been validated and confirmed.")

            # Show final mappings summary
            with st.expander("View Final Mappings"):
                all_mappings = st.session_state.column_mappings.get(team_id, {})
                for file_key, mappings in all_mappings.items():
                    st.markdown(f"**{file_key}**")
                    if mappings:
                        mapping_df = pd.DataFrame([
                            {"Source": k, "Mapped To": v}
                            for k, v in mappings.items()
                        ])
                        st.dataframe(mapping_df, width="stretch", hide_index=True)
                    else:
                        st.caption("No mappings")

    # S3 Code Mapping Section
    if team_id == "S3":
        st.markdown("---")
        if st.session_state.mapping_validated.get(team_id, False):
            render_s3_code_mapping(team_id)
        else:
            st.markdown("### Step 3: Map Codes to Template")
            st.info("👆 **Complete Step 2 above first** - Confirm column mappings to unlock code mapping.")


def render_s3_code_mapping(team_id: str):
    """Render the S3 code mapping section (customer codes -> template codes)."""
    st.markdown("### Step 3: Map Codes to Template")

    template_path = st.session_state.get("s3_template_path")

    if not template_path or not Path(template_path).exists():
        st.info("📋 **Template Mode Required** - Upload a template in the 'Upload & Files' tab to enable code mapping.")
        st.caption("Without a template, customer codes will be used as-is.")
        return

    st.markdown(f"**Template:** {Path(template_path).name}")
    st.markdown("""
    This step maps your customer's codes to the template's codes:
    - **School codes** (e.g., "London Academy" → "SCH001")
    - **Finance codes** (e.g., "611100" → "I1201")
    - **Department codes** (e.g., "Teaching" → "TEACH")
    """)

    # Analyze codes button
    if st.button("🔍 Analyze Customer Codes", type="primary", key="analyze_codes"):
        with st.spinner("Analyzing customer data for codes..."):
            try:
                from teams.s3_code_mapper import run_code_mapping

                customer_dir = CUSTOMER_DATA_DIR / team_id
                column_mappings = st.session_state.column_mappings.get(team_id, {})

                result = run_code_mapping(
                    template_path=Path(template_path),
                    customer_data_dir=customer_dir,
                    column_mappings=column_mappings
                )

                st.session_state["s3_code_mapping_result"] = result
                st.success(f"Found {len(result.school_mappings)} schools, {len(result.finance_mappings)} finance codes")
            except Exception as e:
                st.error(f"Error analyzing codes: {e}")
                return

    # Show code mapping results
    result = st.session_state.get("s3_code_mapping_result")
    if result:
        # Warnings
        if result.warnings:
            for warning in result.warnings:
                st.warning(warning)

        # School Mappings
        if result.school_mappings:
            st.markdown("### School Mappings")
            review_schools = [m for m in result.school_mappings if m.requires_review]
            auto_schools = [m for m in result.school_mappings if not m.requires_review]

            if auto_schools:
                with st.expander(f"✅ Auto-matched Schools ({len(auto_schools)})", expanded=False):
                    school_df = pd.DataFrame([
                        {"Customer": m.customer_value, "Template": m.proposed_template_code,
                         "Description": m.proposed_template_description, "Confidence": f"{m.confidence:.0%}"}
                        for m in auto_schools
                    ])
                    st.dataframe(school_df, hide_index=True)

            if review_schools:
                st.markdown(f"**Review Required ({len(review_schools)} schools)**")
                for m in review_schools:
                    col1, col2, col3 = st.columns([2, 2, 1])
                    with col1:
                        st.text(f"{m.customer_value}")
                        if m.customer_description:
                            st.caption(m.customer_description)
                    with col2:
                        options = [m.proposed_template_code] if m.proposed_template_code else []
                        options += [alt['code'] for alt in m.alternatives if alt['code'] not in options]
                        options.append("-- Skip --")
                        selected = st.selectbox(
                            "Map to",
                            options,
                            key=f"school_map_{m.customer_value}",
                            label_visibility="collapsed"
                        )
                        st.session_state[f"approved_school_{m.customer_value}"] = selected if selected != "-- Skip --" else None
                    with col3:
                        st.caption(f"{m.confidence:.0%}")

        # Finance Code Mappings
        if result.finance_mappings:
            st.markdown("### Finance Code Mappings")
            review_finance = [m for m in result.finance_mappings if m.requires_review]
            auto_finance = [m for m in result.finance_mappings if not m.requires_review]

            if auto_finance:
                with st.expander(f"✅ Auto-matched Finance Codes ({len(auto_finance)})", expanded=False):
                    fin_df = pd.DataFrame([
                        {"Customer": m.customer_value, "Template": m.proposed_template_code,
                         "Description": m.proposed_template_description, "Confidence": f"{m.confidence:.0%}"}
                        for m in auto_finance
                    ])
                    st.dataframe(fin_df, hide_index=True)

            if review_finance:
                st.markdown(f"**Review Required ({len(review_finance)} codes)**")
                for m in review_finance[:20]:  # Limit to 20 for UI
                    col1, col2, col3 = st.columns([2, 2, 1])
                    with col1:
                        st.text(f"{m.customer_value}")
                        if m.customer_description:
                            st.caption(m.customer_description[:30])
                    with col2:
                        options = [m.proposed_template_code] if m.proposed_template_code else []
                        options += [alt['code'] for alt in m.alternatives if alt['code'] not in options]
                        options.append("-- Skip --")
                        selected = st.selectbox(
                            "Map to",
                            options,
                            key=f"finance_map_{m.customer_value}",
                            label_visibility="collapsed"
                        )
                        st.session_state[f"approved_finance_{m.customer_value}"] = selected if selected != "-- Skip --" else None
                    with col3:
                        st.caption(f"{m.confidence:.0%}")

                if len(review_finance) > 20:
                    st.caption(f"... and {len(review_finance) - 20} more")

        # Confirm button
        col1, col2 = st.columns(2)
        with col1:
            if st.button("✅ Confirm Code Mappings", type="primary", key="confirm_codes"):
                # Build code mappings from session state
                code_mappings = {
                    "school_mappings": {},
                    "finance_mappings": {},
                    "department_mappings": {}
                }

                # Auto-matched
                for m in result.school_mappings:
                    if not m.requires_review and m.proposed_template_code:
                        code_mappings["school_mappings"][m.customer_value] = m.proposed_template_code
                for m in result.finance_mappings:
                    if not m.requires_review and m.proposed_template_code:
                        code_mappings["finance_mappings"][m.customer_value] = m.proposed_template_code

                # User approved
                for m in result.school_mappings:
                    if m.requires_review:
                        approved = st.session_state.get(f"approved_school_{m.customer_value}")
                        if approved:
                            code_mappings["school_mappings"][m.customer_value] = approved
                for m in result.finance_mappings:
                    if m.requires_review:
                        approved = st.session_state.get(f"approved_finance_{m.customer_value}")
                        if approved:
                            code_mappings["finance_mappings"][m.customer_value] = approved

                st.session_state["s3_code_mappings"] = code_mappings
                st.success(f"✅ Code mappings confirmed! {len(code_mappings['school_mappings'])} schools, {len(code_mappings['finance_mappings'])} finance codes")

        with col2:
            if st.button("🔄 Clear & Re-analyze", key="clear_codes"):
                if "s3_code_mapping_result" in st.session_state:
                    del st.session_state["s3_code_mapping_result"]
                if "s3_code_mappings" in st.session_state:
                    del st.session_state["s3_code_mappings"]
                st.rerun()

        # Show confirmed status
        if st.session_state.get("s3_code_mappings"):
            st.success("✅ Code mappings are confirmed. Ready to process!")


# =============================================================================
# MAIN APP
# =============================================================================

def main():
    """Main application entry point."""
    render_sidebar()

    team_id = st.session_state.selected_team

    # Main content area
    render_team_overview(team_id)

    st.markdown("---")

    # Tabs for different sections
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📤 Upload & Files",
        "Validate & Map",
        "⚙️ Process",
        "📊 Reports",
        "👁️ Preview"
    ])

    with tab1:
        render_data_upload(team_id)
        st.markdown("---")
        render_customer_files(team_id)

    with tab2:
        render_validate_and_map(team_id)

    with tab3:
        render_processing(team_id)

    with tab4:
        render_reports(team_id)

    with tab5:
        render_preview(team_id)

    # Footer
    st.markdown("---")
    st.markdown(f"""
    <div style='text-align: center; padding: 20px; background: linear-gradient(135deg, {IMP_COLORS['deep_purple']}11, {IMP_COLORS['light_purple']}11); border-radius: 8px;'>
        <p style='color: {IMP_COLORS['deep_purple']}; margin-bottom: 5px;'>
            <strong>IMP Planner</strong> | Agent Teams v2.0
        </p>
        <p style='color: {IMP_COLORS['dark_grey']}; font-size: 0.85em; margin: 0;'>
            hello@impsoftware.co.uk | impsoftware.co.uk | 01392 573 620
        </p>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
