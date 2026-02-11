"""Read all knowledge files for each team."""
import pandas as pd
from pathlib import Path
import warnings
warnings.filterwarnings('ignore')

def safe_str(v, max_len=80):
    """Convert to ASCII-safe string."""
    if pd.isna(v):
        return None
    s = str(v).strip()
    if not s or s == 'nan':
        return None
    return s[:max_len].encode('ascii', 'replace').decode('ascii')

def read_excel_knowledge(folder, team_name):
    """Read Excel files from a knowledge folder."""
    print(f'\n{"="*70}')
    print(f'{team_name} KNOWLEDGE')
    print(f'{"="*70}')

    folder = Path(folder)
    for f in sorted(folder.glob('*.xlsx')):
        print(f'\n### FILE: {f.name} ###')
        try:
            xl = pd.ExcelFile(f)
            print(f'Sheets: {xl.sheet_names}')

            for sheet in xl.sheet_names:
                df = pd.read_excel(xl, sheet)
                print(f'\n--- [{sheet}] ({len(df)} rows) ---')

                # Print content
                for idx, row in df.head(40).iterrows():
                    vals = [safe_str(v) for v in row.values]
                    vals = [v for v in vals if v]
                    if vals:
                        line = ' | '.join(vals)
                        print(line[:200])
        except Exception as e:
            print(f'Error: {e}')

def read_docx_knowledge(folder, team_name):
    """Read Word documents from a knowledge folder."""
    print(f'\n{"="*70}')
    print(f'{team_name} KNOWLEDGE')
    print(f'{"="*70}')

    folder = Path(folder)

    try:
        from docx import Document
    except ImportError:
        print('python-docx not installed, trying alternative...')
        # Try using zipfile to read docx
        import zipfile
        import xml.etree.ElementTree as ET

        for f in sorted(folder.glob('*.docx')):
            print(f'\n### FILE: {f.name} ###')
            try:
                with zipfile.ZipFile(f, 'r') as z:
                    xml_content = z.read('word/document.xml')
                    tree = ET.fromstring(xml_content)

                    # Extract all text
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    paragraphs = tree.findall('.//w:p', ns)

                    for p in paragraphs[:100]:
                        texts = p.findall('.//w:t', ns)
                        para_text = ''.join([t.text or '' for t in texts])
                        if para_text.strip():
                            print(para_text.encode('ascii', 'replace').decode('ascii'))
            except Exception as e:
                print(f'Error: {e}')
        return

    for f in sorted(folder.glob('*.docx')):
        print(f'\n### FILE: {f.name} ###')
        try:
            doc = Document(f)
            for para in doc.paragraphs[:100]:
                if para.text.strip():
                    print(para.text.encode('ascii', 'replace').decode('ascii'))

            # Also read tables
            for table in doc.tables[:5]:
                print('\n[TABLE]')
                for row in table.rows[:20]:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        print(' | '.join(cells).encode('ascii', 'replace').decode('ascii')[:200])
        except Exception as e:
            print(f'Error: {e}')

if __name__ == '__main__':
    # Read S1 knowledge (Excel)
    read_excel_knowledge('knowledge/S1', 'S1 - STRUCTURE TEAM')

    # Read S2 knowledge (Word)
    read_docx_knowledge('knowledge/S2', 'S2 - STAFF TEAM')

    # Read S3 knowledge (Word)
    read_docx_knowledge('knowledge/S3', 'S3 - FINANCIAL TEAM')
